#!/usr/bin/env python3
"""Generate BRD Part 2: Section 4 - Data Model"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('/Users/n15318/Trust OMS/docs/CIM_BRD_Draft.docx')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1B3A5C')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E8F0FE')
                shading.set(qn('w:val'), 'clear')
                table.rows[r_idx + 1].cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()

# ============================================================
# SECTION 4: DATA MODEL
# ============================================================
doc.add_heading('4. Data Model', level=1)
doc.add_paragraph(
    'All entities follow TrustOMS conventions: snake_case table names, camelCase TypeScript exports, '
    'kebab-case entity keys. Every table includes created_at, updated_at, created_by, updated_by, '
    'is_deleted (soft delete), and version (optimistic locking) fields.'
)

# --- 4.1 meetings ---
doc.add_heading('4.1 meetings', level=2)
doc.add_paragraph('Stores all scheduled meetings/appointments. Entity key: meetings')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique meeting identifier'],
        ['meeting_ref', 'text', 'Yes', 'Auto: MTG-YYYYMMDD-XXXX', 'Human-readable reference'],
        ['subject', 'text', 'Yes', 'Max 200 chars', 'Meeting subject/title'],
        ['meeting_reason', 'text', 'Yes', 'Enum: CLIENT_CALL_NEW, CLIENT_CALL_EXISTING, BRANCH_VISIT, PORTFOLIO_REVIEW, SERVICE_REQUEST, CAMPAIGN_DETAILS, OTHERS', 'Purpose of meeting'],
        ['location', 'text', 'Yes', 'Max 500 chars', 'Physical or virtual location'],
        ['start_date', 'timestamp', 'Yes', 'Must be >= now for new meetings', 'Meeting start date/time'],
        ['end_date', 'timestamp', 'Yes', 'Must be > start_date', 'Meeting end date/time'],
        ['is_all_day', 'boolean', 'No', 'Default: false', 'If true, time fields ignored'],
        ['meeting_type', 'text', 'Yes', 'Enum: CIF, LEAD, PROSPECT, OTHERS', 'Target audience type'],
        ['mode_of_meeting', 'text', 'Yes', 'Enum: FACE_TO_FACE, IN_PERSON, IN_PERSON_OFFSHORE, TELEPHONE, TELEPHONE_OFFSHORE, VIDEO_CONFERENCE, VIDEO_CONFERENCE_OFFSHORE, OTHERS', 'How meeting is conducted'],
        ['client_id', 'text FK', 'Conditional', 'Required if meeting_type IN (CIF)', 'Link to clients table'],
        ['prospect_name', 'text', 'Conditional', 'Required if meeting_type = OTHERS', 'Free-text name for non-CIF'],
        ['contact_phone', 'text', 'No', 'Regex: +63XXXXXXXXXX or free-text', 'Contact phone number'],
        ['contact_email', 'text', 'No', 'Valid email format', 'Contact email address'],
        ['remarks', 'text', 'No', 'Max 2000 chars', 'Additional notes'],
        ['meeting_status', 'text', 'Yes', 'Enum: SCHEDULED, COMPLETED, CANCELLED, RESCHEDULED. Default: SCHEDULED', 'Current meeting status'],
        ['organizer_id', 'text FK', 'Yes', 'References users.id', 'RM who created the meeting'],
        ['branch_id', 'text FK', 'No', 'References branches.id', 'Branch context'],
        ['recurrence_rule', 'text', 'No', 'RRULE format or null', 'Recurring meeting pattern'],
        ['parent_meeting_id', 'integer FK', 'No', 'References meetings.id', 'For rescheduled meetings, links to original'],
        ['reminder_minutes', 'integer', 'No', 'Default: 30. Options: 5,10,15,30,60,120,1440', 'Minutes before meeting for reminder'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
        ['updated_at', 'timestamp', 'Yes', 'Auto-update on modify', 'Last modification time'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator user ID'],
        ['updated_by', 'text FK', 'Yes', 'References users.id', 'Last modifier user ID'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
        ['version', 'integer', 'Yes', 'Default: 1, increment on update', 'Optimistic locking'],
    ]
)

doc.add_heading('Sample Data: meetings', level=3)
add_table(doc,
    ['meeting_ref', 'subject', 'meeting_reason', 'start_date', 'meeting_type', 'mode', 'client_id', 'status', 'organizer_id'],
    [
        ['MTG-20260422-0001', 'Q2 Portfolio Review - Santos Family Trust', 'PORTFOLIO_REVIEW', '2026-04-25 10:00', 'CIF', 'FACE_TO_FACE', 'CL-00042', 'SCHEDULED', 'USR-RM-001'],
        ['MTG-20260422-0002', 'New UITF Pitch - Reyes Corp', 'CLIENT_CALL_NEW', '2026-04-25 14:30', 'PROSPECT', 'VIDEO_CONFERENCE', None, 'SCHEDULED', 'USR-RM-002'],
        ['MTG-20260422-0003', 'Branch Visit - Makati Main', 'BRANCH_VISIT', '2026-04-26 09:00', 'OTHERS', 'IN_PERSON', None, 'COMPLETED', 'USR-RM-001'],
    ]
)

# --- 4.2 meeting_attendees ---
doc.add_heading('4.2 meeting_attendees', level=2)
doc.add_paragraph('Junction table for meeting invitees. Entity key: meeting-attendees')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['meeting_id', 'integer FK', 'Yes', 'References meetings.id', 'Parent meeting'],
        ['user_id', 'text FK', 'Yes', 'References users.id', 'Invited user'],
        ['attendance_type', 'text', 'Yes', 'Enum: REQUIRED, OPTIONAL', 'Required or optional attendee'],
        ['rsvp_status', 'text', 'Yes', 'Enum: PENDING, ACCEPTED, DECLINED, TENTATIVE. Default: PENDING', 'Response status'],
        ['attended', 'boolean', 'No', 'Default: null. Set after meeting', 'Whether attendee actually attended'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
    ]
)

doc.add_heading('Sample Data: meeting_attendees', level=3)
add_table(doc,
    ['meeting_id', 'user_id', 'attendance_type', 'rsvp_status', 'attended'],
    [
        ['1', 'USR-RM-003', 'REQUIRED', 'ACCEPTED', 'true'],
        ['1', 'USR-SR-001', 'OPTIONAL', 'PENDING', None],
        ['2', 'USR-RM-001', 'REQUIRED', 'ACCEPTED', None],
    ]
)

# --- 4.3 call_reports ---
doc.add_heading('4.3 call_reports', level=2)
doc.add_paragraph('Core interaction tracking entity. Entity key: call-reports')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['report_ref', 'text', 'Yes', 'Auto: CR-YYYYMMDD-XXXX', 'Human-readable reference'],
        ['report_type', 'text', 'Yes', 'Enum: SCHEDULED, STANDALONE', 'Whether linked to meeting or ad-hoc'],
        ['meeting_id', 'integer FK', 'Conditional', 'Required if report_type=SCHEDULED. References meetings.id', 'Parent meeting (null for standalone)'],
        ['client_id', 'text FK', 'Conditional', 'Required if meeting_type IN (CIF)', 'Client this report is about'],
        ['subject', 'text', 'Yes', 'Max 200 chars. Pre-populated from meeting if scheduled', 'Report subject'],
        ['meeting_reason', 'text', 'Yes', 'Same enum as meetings.meeting_reason', 'Reason for interaction'],
        ['meeting_date', 'timestamp', 'Yes', 'Pre-populated from meeting start_date', 'When interaction occurred'],
        ['end_date', 'timestamp', 'No', 'Pre-populated from meeting end_date', 'When interaction ended'],
        ['location', 'text', 'Yes', 'Max 500 chars', 'Where interaction took place'],
        ['meeting_type', 'text', 'Yes', 'Enum: CIF, LEAD, PROSPECT, OTHERS', 'Audience type'],
        ['mode_of_meeting', 'text', 'Yes', 'Same enum as meetings.mode_of_meeting', 'Meeting mode'],
        ['client_relationship_status', 'text', 'Yes', 'Enum: CLIENT_ABROAD_EMAIL, CONFERENCE_CALL, NOT_CONTACTABLE, NOT_INTERESTED, ACTIVE_ENGAGED, FOLLOW_UP_NEEDED', 'Client engagement status after meeting'],
        ['state_of_mind', 'text', 'Yes', 'Enum: HAPPY, IRATE, SATISFIED, SENSITIVE, ULTRA_SENSITIVE', 'Client emotional state assessment'],
        ['person_met', 'text', 'No', 'Max 200 chars', 'Name of person met (for OTHERS type)'],
        ['summary_of_discussion', 'text', 'Yes', 'Min 50 chars, Max 5000 chars', 'Detailed narrative of meeting content'],
        ['products_discussed', 'jsonb', 'No', 'Array of {product_type, product_name, interest_level}', 'Products discussed during meeting'],
        ['contact_details', 'text', 'No', 'Auto-populated from client profile', 'Contact info used'],
        ['next_meeting_date', 'timestamp', 'No', 'Must be > meeting_date', 'Planned follow-up date'],
        ['next_meeting_time_start', 'time', 'No', 'HH:MM format', 'Follow-up start time'],
        ['next_meeting_time_end', 'time', 'No', 'Must be > start time', 'Follow-up end time'],
        ['linked_call_report_id', 'integer FK', 'No', 'References call_reports.id', 'Previous call report in chain'],
        ['report_status', 'text', 'Yes', 'Enum: DRAFT, SUBMITTED, UNDER_REVIEW, APPROVED, REJECTED, LATE_FILED. Default: DRAFT', 'Workflow status'],
        ['submitted_at', 'timestamp', 'No', 'Set when status changes to SUBMITTED', 'When report was submitted'],
        ['reviewed_at', 'timestamp', 'No', 'Set when approved/rejected', 'When supervisor reviewed'],
        ['reviewed_by', 'text FK', 'No', 'References users.id', 'Supervisor who reviewed'],
        ['rejection_reason', 'text', 'No', 'Required when status = REJECTED. Max 1000 chars', 'Why report was rejected'],
        ['quality_score', 'integer', 'No', 'Range: 1-5. Set by supervisor', 'Supervisor quality assessment'],
        ['is_late_filed', 'boolean', 'Yes', 'Default: false. True if filed >48h after meeting', 'Late filing flag'],
        ['filed_by', 'text FK', 'Yes', 'References users.id', 'RM who filed the report'],
        ['branch_id', 'text FK', 'No', 'References branches.id', 'Branch context'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
        ['updated_at', 'timestamp', 'Yes', 'Auto-update on modify', 'Last modification time'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator'],
        ['updated_by', 'text FK', 'Yes', 'References users.id', 'Last modifier'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
        ['version', 'integer', 'Yes', 'Default: 1', 'Optimistic locking'],
    ]
)

doc.add_heading('Sample Data: call_reports', level=3)
add_table(doc,
    ['report_ref', 'report_type', 'subject', 'client_id', 'state_of_mind', 'status', 'quality_score', 'filed_by'],
    [
        ['CR-20260422-0001', 'SCHEDULED', 'Q2 Review - Santos Trust', 'CL-00042', 'SATISFIED', 'APPROVED', '4', 'USR-RM-001'],
        ['CR-20260422-0002', 'STANDALONE', 'Walk-in Inquiry - New IMA', None, 'HAPPY', 'SUBMITTED', None, 'USR-BA-001'],
        ['CR-20260423-0001', 'SCHEDULED', 'UITF Pitch Follow-up - Reyes', 'CL-00108', 'SENSITIVE', 'DRAFT', None, 'USR-RM-002'],
    ]
)

# --- 4.4 call_report_attendees ---
doc.add_heading('4.4 call_report_attendees', level=2)
doc.add_paragraph('Tracks who actually attended the meeting recorded in a call report.')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['call_report_id', 'integer FK', 'Yes', 'References call_reports.id', 'Parent call report'],
        ['user_id', 'text FK', 'No', 'References users.id', 'Internal attendee (null if external)'],
        ['external_name', 'text', 'No', 'Max 200 chars. Required if user_id is null', 'External person name'],
        ['external_organization', 'text', 'No', 'Max 200 chars', 'External person org'],
        ['role_in_meeting', 'text', 'No', 'Free text: e.g. "Client PoA", "Co-RM"', 'Attendee role'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
    ]
)

doc.add_heading('Sample Data: call_report_attendees', level=3)
add_table(doc,
    ['call_report_id', 'user_id', 'external_name', 'external_organization', 'role_in_meeting'],
    [
        ['1', 'USR-RM-003', None, None, 'Co-RM'],
        ['1', None, 'Maria Santos', 'Santos Family Office', 'Client PoA'],
        ['2', 'USR-BA-001', None, None, 'Branch Associate'],
    ]
)

# --- 4.5 action_items ---
doc.add_heading('4.5 action_items', level=2)
doc.add_paragraph('Follow-up tasks created from call reports. Entity key: action-items')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['action_ref', 'text', 'Yes', 'Auto: ACT-YYYYMMDD-XXXX', 'Human-readable reference'],
        ['call_report_id', 'integer FK', 'Yes', 'References call_reports.id', 'Source call report'],
        ['title', 'text', 'Yes', 'Max 200 chars', 'Action item title'],
        ['description', 'text', 'No', 'Max 2000 chars', 'Detailed description'],
        ['assigned_to', 'text FK', 'Yes', 'References users.id', 'RM responsible'],
        ['assigned_by', 'text FK', 'Yes', 'References users.id', 'Who created the assignment'],
        ['due_date', 'date', 'Yes', 'Must be >= today', 'Deadline for completion'],
        ['priority', 'text', 'Yes', 'Enum: LOW, MEDIUM, HIGH, URGENT. Default: MEDIUM', 'Priority level'],
        ['action_status', 'text', 'Yes', 'Enum: OPEN, IN_PROGRESS, COMPLETED, OVERDUE, CANCELLED. Default: OPEN', 'Current status'],
        ['completed_at', 'timestamp', 'No', 'Set when status = COMPLETED', 'When action was completed'],
        ['completion_notes', 'text', 'No', 'Max 2000 chars. Required when completing', 'Notes on what was done'],
        ['client_id', 'text FK', 'No', 'References clients.id', 'Related client'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
        ['updated_at', 'timestamp', 'Yes', 'Auto-update', 'Last modification time'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator'],
        ['updated_by', 'text FK', 'Yes', 'References users.id', 'Last modifier'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
        ['version', 'integer', 'Yes', 'Default: 1', 'Optimistic locking'],
    ]
)

doc.add_heading('Sample Data: action_items', level=3)
add_table(doc,
    ['action_ref', 'title', 'assigned_to', 'due_date', 'priority', 'status', 'call_report_id'],
    [
        ['ACT-20260422-0001', 'Send updated IMA proposal to Santos Trust', 'USR-RM-001', '2026-04-28', 'HIGH', 'OPEN', '1'],
        ['ACT-20260422-0002', 'Prepare UITF fact sheet for Reyes Corp', 'USR-RM-002', '2026-04-30', 'MEDIUM', 'IN_PROGRESS', '3'],
        ['ACT-20260423-0001', 'Schedule KYC renewal for Santos Family', 'USR-RM-001', '2026-05-05', 'LOW', 'OPEN', '1'],
    ]
)

# --- 4.6 opportunities ---
doc.add_heading('4.6 opportunities', level=2)
doc.add_paragraph('Sales pipeline opportunities captured during client meetings. Entity key: opportunities')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['opportunity_ref', 'text', 'Yes', 'Auto: OPP-YYYYMMDD-XXXX', 'Human-readable reference'],
        ['name', 'text', 'Yes', 'Max 200 chars', 'Opportunity name/title'],
        ['description', 'text', 'No', 'Max 2000 chars', 'Detailed description'],
        ['call_report_id', 'integer FK', 'No', 'References call_reports.id', 'Source call report (null if bulk uploaded)'],
        ['client_id', 'text FK', 'No', 'References clients.id', 'Target client'],
        ['product_type', 'text', 'Yes', 'Enum: IMA_DIRECTED, IMA_DISCRETIONARY, UITF, PMT, PRE_NEED, EMPLOYEE_BENEFIT, ESCROW, INSURANCE, BONDS, EQUITIES, OTHER', 'Product category'],
        ['product_name', 'text', 'No', 'Max 200 chars', 'Specific product name'],
        ['expected_value', 'numeric(18,2)', 'Yes', 'Must be > 0. In PHP', 'Expected revenue/AUM'],
        ['currency', 'text', 'Yes', 'Default: PHP. ISO 4217 code', 'Currency of expected value'],
        ['probability', 'integer', 'Yes', 'Range: 0-100. Default: 50', 'Win probability percentage'],
        ['expected_close_date', 'date', 'Yes', 'Must be >= today', 'Target close date'],
        ['opportunity_status', 'text', 'Yes', 'Enum: OPEN, QUALIFIED, PROPOSAL, NEGOTIATION, WON, LOST, DEFERRED. Default: OPEN', 'Pipeline stage'],
        ['loss_reason', 'text', 'No', 'Required if status = LOST. Enum: PRICING, COMPETITOR, NO_BUDGET, TIMING, PRODUCT_FIT, OTHER', 'Why opportunity was lost'],
        ['won_date', 'date', 'No', 'Set when status = WON', 'Actual close date'],
        ['won_value', 'numeric(18,2)', 'No', 'Set when status = WON', 'Actual value realized'],
        ['owner_id', 'text FK', 'Yes', 'References users.id', 'RM who owns this opportunity'],
        ['branch_id', 'text FK', 'No', 'References branches.id', 'Branch context'],
        ['source', 'text', 'No', 'Enum: MEETING, CAMPAIGN, REFERRAL, INBOUND, BULK_UPLOAD', 'How opportunity was captured'],
        ['campaign_id', 'integer FK', 'No', 'References campaigns.id', 'Source campaign if applicable'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
        ['updated_at', 'timestamp', 'Yes', 'Auto-update', 'Last modification time'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator'],
        ['updated_by', 'text FK', 'Yes', 'References users.id', 'Last modifier'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
        ['version', 'integer', 'Yes', 'Default: 1', 'Optimistic locking'],
    ]
)

doc.add_heading('Sample Data: opportunities', level=3)
add_table(doc,
    ['opportunity_ref', 'name', 'product_type', 'expected_value', 'probability', 'status', 'owner_id', 'client_id'],
    [
        ['OPP-20260422-0001', 'Santos Trust - IMA Expansion', 'IMA_DIRECTED', '50000000.00', '75', 'PROPOSAL', 'USR-RM-001', 'CL-00042'],
        ['OPP-20260422-0002', 'Reyes Corp - New UITF Investment', 'UITF', '25000000.00', '40', 'QUALIFIED', 'USR-RM-002', None],
        ['OPP-20260423-0001', 'Garcia Family - PMT Setup', 'PMT', '100000000.00', '60', 'NEGOTIATION', 'USR-RM-001', 'CL-00089'],
    ]
)

# --- 4.7 expenses ---
doc.add_heading('4.7 expenses', level=2)
doc.add_paragraph('Client-meeting related expenses. Entity key: expenses')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['expense_ref', 'text', 'Yes', 'Auto: EXP-YYYYMMDD-XXXX', 'Human-readable reference'],
        ['call_report_id', 'integer FK', 'No', 'References call_reports.id', 'Linked call report'],
        ['meeting_id', 'integer FK', 'No', 'References meetings.id', 'Linked meeting'],
        ['expense_type', 'text', 'Yes', 'Enum: TRAVEL, MEALS, ENTERTAINMENT, ACCOMMODATION, TRANSPORTATION, COMMUNICATION, GIFTS, OTHER', 'Expense category'],
        ['amount', 'numeric(18,2)', 'Yes', 'Must be > 0', 'Expense amount'],
        ['currency', 'text', 'Yes', 'Default: PHP. ISO 4217', 'Expense currency'],
        ['expense_date', 'date', 'Yes', 'Must be <= today', 'When expense was incurred'],
        ['description', 'text', 'Yes', 'Max 500 chars', 'Expense description'],
        ['receipt_url', 'text', 'No', 'Valid URL to uploaded file', 'Receipt attachment'],
        ['expense_status', 'text', 'Yes', 'Enum: DRAFT, SUBMITTED, APPROVED, REJECTED. Default: DRAFT', 'Approval status'],
        ['approved_by', 'text FK', 'No', 'References users.id', 'Approver'],
        ['approved_at', 'timestamp', 'No', 'Set on approval', 'Approval timestamp'],
        ['rejection_reason', 'text', 'No', 'Required if REJECTED', 'Why rejected'],
        ['submitted_by', 'text FK', 'Yes', 'References users.id', 'RM who submitted'],
        ['branch_id', 'text FK', 'No', 'References branches.id', 'Branch context'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
        ['updated_at', 'timestamp', 'Yes', 'Auto-update', 'Last modification time'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator'],
        ['updated_by', 'text FK', 'Yes', 'References users.id', 'Last modifier'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
        ['version', 'integer', 'Yes', 'Default: 1', 'Optimistic locking'],
    ]
)

doc.add_heading('Sample Data: expenses', level=3)
add_table(doc,
    ['expense_ref', 'expense_type', 'amount', 'currency', 'expense_date', 'description', 'status', 'submitted_by'],
    [
        ['EXP-20260422-0001', 'MEALS', '3500.00', 'PHP', '2026-04-22', 'Client lunch - Santos Trust quarterly review', 'APPROVED', 'USR-RM-001'],
        ['EXP-20260423-0001', 'TRANSPORTATION', '1200.00', 'PHP', '2026-04-23', 'Grab to Makati client office', 'SUBMITTED', 'USR-RM-002'],
        ['EXP-20260423-0002', 'ENTERTAINMENT', '8500.00', 'PHP', '2026-04-23', 'Client dinner - Garcia Family PMT discussion', 'DRAFT', 'USR-RM-001'],
    ]
)

# --- 4.8 call_report_feedback ---
doc.add_heading('4.8 call_report_feedback', level=2)
doc.add_paragraph('Supervisor feedback and coaching comments on call reports.')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['call_report_id', 'integer FK', 'Yes', 'References call_reports.id', 'Target call report'],
        ['feedback_by', 'text FK', 'Yes', 'References users.id', 'Supervisor providing feedback'],
        ['feedback_type', 'text', 'Yes', 'Enum: GENERAL, COACHING, COMPLIANCE_FLAG, QUALITY_ISSUE', 'Category of feedback'],
        ['comment', 'text', 'Yes', 'Min 10 chars, Max 2000 chars', 'Feedback content'],
        ['is_private', 'boolean', 'Yes', 'Default: false. If true, only visible to RM and feedbacker', 'Private feedback flag'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation time'],
    ]
)

doc.add_heading('Sample Data: call_report_feedback', level=3)
add_table(doc,
    ['call_report_id', 'feedback_by', 'feedback_type', 'comment', 'is_private'],
    [
        ['1', 'USR-BH-001', 'GENERAL', 'Excellent detail on portfolio performance discussion. Good follow-up items identified.', 'false'],
        ['1', 'USR-BH-001', 'COACHING', 'Consider probing deeper on risk appetite changes given market volatility.', 'true'],
        ['2', 'USR-CCO-001', 'COMPLIANCE_FLAG', 'Walk-in meeting missing suitability assessment documentation. Please add.', 'false'],
    ]
)

# --- 4.9 tasks ---
doc.add_heading('4.9 tasks', level=2)
doc.add_paragraph('Personal and assigned task management for RMs. Entity key: tasks')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['task_ref', 'text', 'Yes', 'Auto: TSK-YYYYMMDD-XXXX', 'Human-readable reference'],
        ['title', 'text', 'Yes', 'Max 200 chars', 'Task title'],
        ['description', 'text', 'No', 'Max 2000 chars', 'Task details'],
        ['task_type', 'text', 'Yes', 'Enum: PERSONAL, ASSIGNED, SYSTEM_GENERATED. Default: PERSONAL', 'How task was created'],
        ['assigned_to', 'text FK', 'Yes', 'References users.id', 'Who must complete the task'],
        ['assigned_by', 'text FK', 'No', 'References users.id. Required if task_type=ASSIGNED', 'Who assigned the task'],
        ['due_date', 'date', 'Yes', 'Must be >= today', 'Task deadline'],
        ['due_time', 'time', 'No', 'HH:MM format', 'Specific deadline time'],
        ['priority', 'text', 'Yes', 'Enum: LOW, MEDIUM, HIGH, URGENT. Default: MEDIUM', 'Priority level'],
        ['task_status', 'text', 'Yes', 'Enum: PENDING, IN_PROGRESS, COMPLETED, OVERDUE, CANCELLED. Default: PENDING', 'Current status'],
        ['completed_at', 'timestamp', 'No', 'Set when COMPLETED', 'Completion timestamp'],
        ['reminder_date', 'date', 'No', 'Must be <= due_date', 'When to send reminder'],
        ['reminder_time', 'time', 'No', 'HH:MM', 'Reminder time'],
        ['related_entity_type', 'text', 'No', 'Enum: CALL_REPORT, MEETING, OPPORTUNITY, CLIENT', 'Linked entity type'],
        ['related_entity_id', 'text', 'No', 'ID of linked entity', 'Linked entity ID'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation'],
        ['updated_at', 'timestamp', 'Yes', 'Auto-update', 'Last modification'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator'],
        ['updated_by', 'text FK', 'Yes', 'References users.id', 'Last modifier'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
        ['version', 'integer', 'Yes', 'Default: 1', 'Optimistic locking'],
    ]
)

doc.add_heading('Sample Data: tasks', level=3)
add_table(doc,
    ['task_ref', 'title', 'task_type', 'assigned_to', 'due_date', 'priority', 'status'],
    [
        ['TSK-20260422-0001', 'Follow up Santos IMA expansion proposal', 'PERSONAL', 'USR-RM-001', '2026-04-28', 'HIGH', 'PENDING'],
        ['TSK-20260422-0002', 'Complete Q1 client review reports', 'ASSIGNED', 'USR-RM-002', '2026-04-30', 'MEDIUM', 'IN_PROGRESS'],
        ['TSK-20260423-0001', 'File overdue call report for MTG-20260420-0005', 'SYSTEM_GENERATED', 'USR-RM-001', '2026-04-24', 'URGENT', 'PENDING'],
    ]
)

# --- 4.10 campaigns ---
doc.add_heading('4.10 campaigns', level=2)
doc.add_paragraph('Marketing campaign definitions. Entity key: campaigns')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['campaign_code', 'text', 'Yes', 'Unique, Max 20 chars, uppercase alphanumeric', 'Campaign unique code'],
        ['name', 'text', 'Yes', 'Max 200 chars', 'Campaign name'],
        ['description', 'text', 'No', 'Max 2000 chars', 'Campaign description'],
        ['campaign_type', 'text', 'Yes', 'Enum: EVENT, PRODUCT_PROMOTION, CROSS_SELL, RETENTION, REACTIVATION', 'Campaign category'],
        ['target_segment', 'text', 'No', 'Max 500 chars', 'Target audience description'],
        ['budget', 'numeric(18,2)', 'No', 'Must be >= 0 if provided', 'Budget allocation'],
        ['currency', 'text', 'No', 'Default: PHP', 'Budget currency'],
        ['expected_roi', 'numeric(5,2)', 'No', 'Percentage', 'Expected ROI'],
        ['start_date', 'date', 'Yes', 'Must be >= today for new campaigns', 'Campaign start date'],
        ['end_date', 'date', 'Yes', 'Must be > start_date', 'Campaign end date'],
        ['campaign_status', 'text', 'Yes', 'Enum: DRAFT, ACTIVE, PAUSED, COMPLETED, CANCELLED. Default: DRAFT', 'Campaign status'],
        ['owner_id', 'text FK', 'Yes', 'References users.id', 'Campaign manager'],
        ['branch_id', 'text FK', 'No', 'References branches.id', 'Branch scope'],
        ['material_urls', 'jsonb', 'No', 'Array of {name, url, type}', 'Marketing materials'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation'],
        ['updated_at', 'timestamp', 'Yes', 'Auto-update', 'Last modification'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator'],
        ['updated_by', 'text FK', 'Yes', 'References users.id', 'Last modifier'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
        ['version', 'integer', 'Yes', 'Default: 1', 'Optimistic locking'],
    ]
)

doc.add_heading('Sample Data: campaigns', level=3)
add_table(doc,
    ['campaign_code', 'name', 'campaign_type', 'budget', 'start_date', 'end_date', 'status', 'owner_id'],
    [
        ['UITF-Q2-2026', 'Q2 UITF Growth Campaign', 'PRODUCT_PROMOTION', '500000.00', '2026-04-01', '2026-06-30', 'ACTIVE', 'USR-BH-001'],
        ['RET-HNW-2026', 'HNW Client Retention Program', 'RETENTION', '1000000.00', '2026-03-15', '2026-12-31', 'ACTIVE', 'USR-BH-001'],
        ['EVT-SEMINAR-05', 'May Investment Seminar', 'EVENT', '150000.00', '2026-05-15', '2026-05-15', 'DRAFT', 'USR-SR-001'],
    ]
)

# --- 4.11 lead_lists ---
doc.add_heading('4.11 lead_lists', level=2)
doc.add_paragraph('Filtered target lists for campaigns. Entity key: lead-lists')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['list_name', 'text', 'Yes', 'Max 200 chars', 'List name'],
        ['campaign_id', 'integer FK', 'No', 'References campaigns.id', 'Linked campaign'],
        ['filter_criteria', 'jsonb', 'No', 'JSON filter definition: {aum_min, aum_max, product_type, risk_profile, branch_id, tenure_years}', 'Auto-generation criteria'],
        ['list_type', 'text', 'Yes', 'Enum: AUTOMATIC, MANUAL, UPLOADED. Default: MANUAL', 'How list was created'],
        ['total_leads', 'integer', 'Yes', 'Default: 0. Auto-computed', 'Count of leads in list'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation'],
        ['created_by', 'text FK', 'Yes', 'References users.id', 'Creator'],
        ['is_deleted', 'boolean', 'Yes', 'Default: false', 'Soft delete flag'],
    ]
)

doc.add_heading('Sample Data: lead_lists', level=3)
add_table(doc,
    ['list_name', 'campaign_id', 'list_type', 'filter_criteria', 'total_leads'],
    [
        ['HNW Clients AUM>50M', '1', 'AUTOMATIC', '{"aum_min":50000000,"product_type":"UITF"}', '142'],
        ['Makati Branch Top Clients', '2', 'MANUAL', None, '35'],
        ['May Seminar Invitees', '3', 'UPLOADED', None, '200'],
    ]
)

# --- 4.12 lead_list_members ---
doc.add_heading('4.12 lead_list_members', level=2)
doc.add_paragraph('Individual leads within a lead list. Entity key: lead-list-members')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['lead_list_id', 'integer FK', 'Yes', 'References lead_lists.id', 'Parent list'],
        ['client_id', 'text FK', 'No', 'References clients.id', 'Existing client'],
        ['prospect_name', 'text', 'No', 'Required if client_id is null', 'Prospect name'],
        ['prospect_email', 'text', 'No', 'Valid email', 'Prospect email'],
        ['prospect_phone', 'text', 'No', 'Phone format', 'Prospect phone'],
        ['assigned_rm_id', 'text FK', 'No', 'References users.id', 'Assigned RM for contact'],
        ['response_status', 'text', 'Yes', 'Enum: NOT_CONTACTED, CONTACTED, INTERESTED, NOT_INTERESTED, NEED_MORE_INFO, CONVERTED, DO_NOT_CONTACT. Default: NOT_CONTACTED', 'Lead response'],
        ['response_date', 'timestamp', 'No', 'Set when response captured', 'When response was recorded'],
        ['response_notes', 'text', 'No', 'Max 1000 chars', 'Response details'],
        ['responded_by', 'text FK', 'No', 'References users.id', 'RM who captured response'],
        ['created_at', 'timestamp', 'Yes', 'Default: now()', 'Record creation'],
    ]
)

doc.add_heading('Sample Data: lead_list_members', level=3)
add_table(doc,
    ['lead_list_id', 'client_id', 'prospect_name', 'assigned_rm_id', 'response_status', 'response_date'],
    [
        ['1', 'CL-00042', None, 'USR-RM-001', 'INTERESTED', '2026-04-20'],
        ['1', 'CL-00089', None, 'USR-RM-001', 'NOT_CONTACTED', None],
        ['3', None, 'Juan Dela Cruz', 'USR-RM-002', 'CONVERTED', '2026-04-18'],
    ]
)

# --- 4.13 meeting_reminders ---
doc.add_heading('4.13 meeting_reminders', level=2)
doc.add_paragraph('Scheduled reminders for upcoming meetings.')
add_table(doc,
    ['Field', 'Type', 'Required', 'Validation / Default', 'Description'],
    [
        ['id', 'serial PK', 'Yes', 'Auto-increment', 'Unique identifier'],
        ['meeting_id', 'integer FK', 'Yes', 'References meetings.id', 'Target meeting'],
        ['user_id', 'text FK', 'Yes', 'References users.id', 'Recipient'],
        ['remind_at', 'timestamp', 'Yes', 'Computed: meeting.start_date - reminder_minutes', 'When to fire reminder'],
        ['reminder_sent', 'boolean', 'Yes', 'Default: false', 'Whether reminder was delivered'],
        ['sent_at', 'timestamp', 'No', 'Set when reminder fires', 'Actual delivery time'],
        ['channel', 'text', 'Yes', 'Enum: IN_APP, EMAIL, BOTH. Default: IN_APP', 'Delivery channel'],
    ]
)

doc.add_heading('Sample Data: meeting_reminders', level=3)
add_table(doc,
    ['meeting_id', 'user_id', 'remind_at', 'reminder_sent', 'channel'],
    [
        ['1', 'USR-RM-001', '2026-04-25 09:30:00', 'false', 'BOTH'],
        ['1', 'USR-RM-003', '2026-04-25 09:30:00', 'false', 'IN_APP'],
        ['2', 'USR-RM-002', '2026-04-25 14:00:00', 'false', 'EMAIL'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('Entity Relationship Summary:', style='Heading 3')
relationships = [
    'meetings 1:N meeting_attendees (one meeting has many attendees)',
    'meetings 1:N meeting_reminders (one meeting has many reminders)',
    'meetings 1:1 call_reports (one meeting produces at most one call report, via meeting_id)',
    'call_reports 1:N call_report_attendees (one report has many attendees)',
    'call_reports 1:N action_items (one report generates many action items)',
    'call_reports 1:N opportunities (one report can capture many opportunities)',
    'call_reports 1:N expenses (one report can have many expenses)',
    'call_reports 1:N call_report_feedback (one report receives many feedback entries)',
    'call_reports N:1 call_reports (linked_call_report_id creates interaction chains)',
    'campaigns 1:N lead_lists (one campaign has many lead lists)',
    'lead_lists 1:N lead_list_members (one list has many members)',
    'opportunities N:1 campaigns (opportunities can be sourced from campaigns)',
    'clients 1:N meetings (one client has many meetings via client_id)',
    'clients 1:N call_reports (one client has many call reports via client_id)',
    'users 1:N meetings (organizer_id — one RM organizes many meetings)',
    'users 1:N call_reports (filed_by — one RM files many call reports)',
    'users 1:N tasks (assigned_to — one user has many tasks)',
    'users 1:N opportunities (owner_id — one RM owns many opportunities)',
]
for r in relationships:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()
doc.save('/Users/n15318/Trust OMS/docs/CIM_BRD_Draft.docx')
print("Part 2 complete: Section 4 (Data Model) saved")
