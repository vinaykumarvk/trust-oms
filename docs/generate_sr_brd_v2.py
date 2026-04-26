#!/usr/bin/env python3
"""
Generate the Service Request / Task Management Module BRD v2.0 (FINAL)
Incorporates all 6 must-fix findings from adversarial evaluation + 6 Phase 2 items documented.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "CONFIDENTIAL — TrustOMS Service Request / Task Management BRD v2.0 FINAL"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style.font.size = Pt(8)
        hp.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style.font.size = Pt(8)


def styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF0F7"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)
    return table


def bullets(doc, items, style='List Bullet'):
    for item in items:
        doc.add_paragraph(item, style=style)


def bullets2(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet 2')


add_header_footer(doc)

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Service Request / Task Management Module")
run.bold = True; run.font.size = Pt(28); run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Business Requirements Document (BRD)")
run.font.size = Pt(18); run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x8C)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run("Version 2.0 FINAL — April 2026")
run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
sys_n = doc.add_paragraph()
sys_n.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sys_n.add_run("TrustOMS — Trust Banking Operations Management System")
run.font.size = Pt(12); run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("Incorporates findings from adversarial evaluation (6 must-fix items)")
run.font.size = Pt(10); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88); run.italic = True

doc.add_paragraph()
conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run("CONFIDENTIAL")
run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# DOCUMENT CONTROL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Document Control", level=1)
doc.add_heading("Version History", level=2)
styled_table(doc,
    ["Version", "Date", "Author", "Changes"],
    [
        ["0.1", "2026-04-20", "System Architect", "Initial draft based on reference implementations"],
        ["0.5", "2026-04-21", "System Architect", "Added field-level details, API specs, state diagrams"],
        ["1.0", "2026-04-22", "System Architect", "Complete v1 — all 14 sections"],
        ["2.0", "2026-04-23", "System Architect", "FINAL — Incorporated 6 must-fix items from adversarial evaluation: sequence-based IDs, status history table, DB-level filtering, authenticated user tracking, in-app notifications, RM reassignment"],
    ])

doc.add_paragraph()
doc.add_heading("Approvals", level=2)
styled_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [["Project Sponsor", "", "", ""], ["Business Analyst", "", "", ""], ["Technical Lead", "", "", ""], ["QA Lead", "", "", ""]])

doc.add_paragraph()
doc.add_heading("References", level=2)
styled_table(doc,
    ["#", "Document", "Source", "Relevance"],
    [
        ["1", "HSBC PWS — Journey of Customer App Service Request V10", "Intellect Design", "Full lifecycle, field-level specs, status transitions"],
        ["2", "Maybank DTC — Journey of Customer App Service Request V11", "Intellect Design", "Mobile UX patterns, push notifications, Incomplete flow"],
        ["3", "Jio BlackRock — Journey of Service Request Customer App V10", "Intellect Design", "Appointed dates, closure date configurability, field editability"],
        ["4", "Adversarial Evaluation Report — service-request-brd-evaluation.md", "Internal", "5-round structured debate identifying 12 gaps; 6 classified as must-fix"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc = [
    "1. Executive Summary",
    "2. Scope & Boundaries",
    "3. User Roles & Permissions",
    "4. Data Model",
    "   4.1 service_requests (Primary Table)",
    "   4.2 sr_status_history (Audit Log) — NEW IN V2",
    "   4.3 Enums",
    "   4.4 Indexes & Relationships",
    "   4.5 Sample Data",
    "5. Functional Requirements",
    "   5.1 Module A: Client Portal",
    "   5.2 Module B: Back-Office Workbench",
    "   5.3 Module C: Backend Service & API",
    "   5.4 Module D: In-App Notifications — NEW IN V2",
    "6. User Interface Requirements",
    "7. API & Integration Requirements",
    "8. Non-Functional Requirements",
    "9. Workflow & State Diagrams",
    "10. Notification & Communication Requirements",
    "11. Reporting & Analytics",
    "12. Migration & Launch Plan",
    "13. Glossary",
    "14. Appendices",
    "   Appendix A: Field Editability Matrix",
    "   Appendix B: V1 → V2 Change Log",
    "   Appendix C: Phase 2 Roadmap",
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

doc.add_heading("1.1 Project Overview", level=2)
doc.add_paragraph(
    "The Service Request / Task Management Module (SR-TM) is a core operational module within the "
    "TrustOMS platform that enables end-to-end lifecycle management of client service requests in a "
    "trust banking environment. The module serves two primary audiences: (1) trust banking clients who "
    "initiate, track, and manage service requests through the Client Portal, and (2) back-office staff "
    "(Relationship Managers, Tellers, Operations Officers) who process, verify, and resolve these requests "
    "through the Back-Office application. The system enforces SLA-driven timelines, dual-control "
    "verification workflows, document management, and full audit trails — all aligned with BSP (Bangko "
    "Sentral ng Pilipinas) regulatory requirements for Philippine trust banking operations."
)
doc.add_paragraph(
    "This v2 FINAL version incorporates six must-fix findings from a structured adversarial evaluation: "
    "(1) concurrency-safe request ID generation via PostgreSQL SEQUENCE, (2) an append-only status history "
    "audit table for BSP examination readiness, (3) database-level query filtering replacing in-memory filtering, "
    "(4) authenticated user identity tracking in all audit fields, (5) in-app notification badges for client "
    "action items, and (6) RM reassignment capability for operational continuity."
)

doc.add_heading("1.2 Business Objectives", level=2)
bullets(doc, [
    "Digitize the service request lifecycle — eliminate paper-based SR forms and manual tracking spreadsheets, reducing processing time by 60%.",
    "Enforce SLA compliance — auto-compute closure dates based on priority (HIGH=3 days, MEDIUM=5 days, LOW=7 days) and surface overdue SRs in real-time dashboards.",
    "Enable client self-service — allow clients to create, modify, track, and close service requests through the Client Portal without visiting a branch.",
    "Implement dual-control verification — require RM preparation and Teller verification before completing any SR, ensuring segregation of duties per BSP Circular 871.",
    "Provide operational visibility — give operations managers a real-time KPI dashboard with status distribution, SLA breach counts, and aging analysis.",
    "Maintain complete audit trail — log every status change with actor identity and timestamp in an append-only audit table for BSP examination readiness.",
    "Support operational continuity — enable RM reassignment when staff are on leave or reassigned, preventing SR orphaning.",
])

doc.add_heading("1.3 Target Users & Pain Points", level=2)
styled_table(doc,
    ["User", "Current Pain Point", "How SR-TM Solves It"],
    [
        ["Trust Banking Client", "Must visit branch or call RM; no visibility into request status", "Self-service portal with real-time tracking, SLA visibility, in-app notification badges"],
        ["Relationship Manager (RM)", "Tracks SRs in spreadsheets; SLA breaches discovered too late", "Structured SR queue with status filters, one-click send-for-verification, overdue SLA alerts"],
        ["Teller / Operations", "Receives SR instructions verbally; no standardized verification", "Dedicated verification queue with complete/incomplete/reject actions and notes"],
        ["Operations Manager", "No consolidated view; manual SLA tracking; RM absences cause orphaned SRs", "KPI dashboard, aging analysis, RM reassignment capability"],
        ["Compliance Officer", "Difficult to reconstruct SR processing history for BSP examinations", "Append-only status history table with full timeline and actor identity"],
    ])

doc.add_heading("1.4 Success Metrics (KPIs)", level=2)
styled_table(doc,
    ["KPI", "Target", "Measurement Method"],
    [
        ["Average SR Processing Time", "≤ SLA days per priority level", "Median days from creation to COMPLETED status"],
        ["SLA Compliance Rate", "≥ 95% of SRs completed within SLA", "Count of on-time completions / total completions"],
        ["Client Self-Service Adoption", "≥ 70% of SRs created via Client Portal", "Portal-created SRs / total SRs"],
        ["First-Pass Completion Rate", "≥ 85% without INCOMPLETE cycle", "Direct READY_FOR_TELLER → COMPLETED count"],
        ["Digital Document Upload Rate", "≥ 80% of SRs have digital documents", "SRs with documents[] not empty / total SRs"],
        ["API Response Time (P95)", "≤ 500ms for list endpoints", "Database-level filtered query latency"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 2: SCOPE & BOUNDARIES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. Scope & Boundaries", level=1)

doc.add_heading("2.1 In Scope (Phase 1)", level=2)
bullets(doc, [
    "Client Portal: Service Request list view with search, status filters, and pagination",
    "Client Portal: Create New Service Request form (SR type, priority, details, remarks, document upload)",
    "Client Portal: Service Request detail view with status-aware field editability",
    "Client Portal: Close Request and Resubmit for Verification actions",
    "Client Portal: In-app notification badges showing count of SRs requiring client action [V2 NEW]",
    "Client Portal: Mobile-responsive design (card-based mobile, table-based desktop)",
    "Back-Office: Service Request Workbench with KPI summary cards and filterable data table",
    "Back-Office: Send for Verification workflow (RM fills branch, unit, date)",
    "Back-Office: Complete / Mark Incomplete / Reject actions (Teller)",
    "Back-Office: RM Reassignment capability (Operations Manager can reassign SRs) [V2 NEW]",
    "Back-Office: SR detail update (remarks, documents, closure date)",
    "Backend: RESTful API endpoints for all CRUD and status transition operations",
    "Backend: Concurrency-safe request ID generation via PostgreSQL SEQUENCE [V2 NEW]",
    "Backend: Database-level query filtering with Drizzle ORM dynamic WHERE clauses [V2 NEW]",
    "Backend: Authenticated user identity tracked in all created_by/updated_by fields [V2 NEW]",
    "Backend: Role-based access control (client portal scoped to client_id, back-office protected)",
    "Backend: SLA-driven closure date computation",
    "Database: service_requests table with all fields, indexes, and relationships",
    "Database: sr_status_history append-only audit log table [V2 NEW]",
    "Database: Enum types for status, priority, and SR type",
])

doc.add_heading("2.2 Out of Scope (Documented Phase 2 Deferrals)", level=2)
bullets(doc, [
    "Business-day SLA computation with market calendar integration (currently uses calendar days)",
    "Configurable SR Type table replacing fixed enum (dynamic type creation via admin UI)",
    "Duplicate SR detection (warn if same client has open SR of same type)",
    "Priority change workflow with SLA recalculation",
    "Document versioning (tracking upload history per SR)",
    "Bulk operations (bulk close, bulk reassign, bulk export)",
    "Push notification / email delivery (notification events defined; delivery deferred)",
    "SLA auto-escalation to supervisor on breach",
    "Client satisfaction surveys post-SR completion",
    "Integration with external CRM systems",
])

doc.add_heading("2.3 Assumptions", level=2)
bullets(doc, [
    "Users are authenticated via JWT in httpOnly cookies (existing auth module).",
    "Client IDs exist in the system; the SR module references but does not create clients.",
    "Upload service at /api/v1/uploads handles file storage; SR stores only file references.",
    "SR types are a fixed enum set; adding types requires a schema migration (Phase 2: lookup table).",
    "Dates stored in UTC; displayed in user's local timezone by frontend.",
    "Back-office app accessed from managed corporate devices on internal networks.",
    "Client portal accessible from modern browsers (Chrome 90+, Safari 14+, Firefox 88+, Edge 90+).",
    "Each client has a mapped RM in the clients table; auto-assignment uses this mapping.",
])

doc.add_heading("2.4 Constraints", level=2)
bullets(doc, [
    "Regulatory: All SR state changes must produce an append-only audit record per BSP Circular 871.",
    "Technical: Backend uses Express.js + Drizzle ORM + PostgreSQL stack.",
    "Technical: Frontend uses React + React Router + TanStack Query + shadcn/ui.",
    "Technical: Document uploads limited to PDF format, maximum 10MB total per SR.",
    "Security: Back-office endpoints protected by requireBackOfficeRole() middleware.",
    "Security: Client portal endpoints scoped to authenticated client_id from JWT token.",
    "Performance: List endpoints must return within 500ms for up to 10,000 SRs (database-level filtering).",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 3: USER ROLES & PERMISSIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. User Roles & Permissions", level=1)

doc.add_heading("3.1 Role Definitions", level=2)
styled_table(doc,
    ["Role", "Description", "Application"],
    [
        ["Client (Customer)", "Trust banking client who initiates and tracks their own service requests", "Client Portal"],
        ["Relationship Manager (RM)", "Bank officer assigned to manage a portfolio of clients; prepares SRs for teller verification", "Back-Office"],
        ["Teller", "Operations staff who performs final verification and completes/rejects service requests", "Back-Office"],
        ["Operations Manager", "Supervisor who monitors SR pipeline, SLA compliance, and team performance; can reassign SRs [V2]", "Back-Office"],
        ["Compliance Officer", "Reviews audit trails and status history for regulatory examination readiness", "Back-Office"],
        ["System Administrator", "Manages system configuration, user accounts, and SR type definitions", "Back-Office"],
    ])

doc.add_paragraph()
doc.add_heading("3.2 Permissions Matrix", level=2)
styled_table(doc,
    ["Action", "Client", "RM", "Teller", "Ops Mgr", "Compliance", "Admin"],
    [
        ["Create Service Request", "Yes (own)", "Yes (any)", "No", "No", "No", "Yes"],
        ["View SR List (own)", "Yes", "—", "—", "—", "—", "—"],
        ["View SR List (all)", "No", "Yes", "Yes", "Yes", "Yes (read)", "Yes"],
        ["View SR Detail", "Yes (own)", "Yes", "Yes", "Yes", "Yes (read)", "Yes"],
        ["Edit SR Fields (non-terminal)", "Yes (own)", "Yes", "No", "No", "No", "Yes"],
        ["Close Request", "Yes (own, open)", "Yes", "No", "No", "No", "Yes"],
        ["Send for Verification", "No", "Yes", "No", "No", "No", "Yes"],
        ["Complete Request", "No", "No", "Yes", "No", "No", "Yes"],
        ["Mark Incomplete", "No", "No", "Yes", "No", "No", "Yes"],
        ["Reject Request", "No", "No", "Yes", "No", "No", "Yes"],
        ["Resubmit for Verification", "Yes (own)", "Yes", "No", "No", "No", "Yes"],
        ["Reassign RM [V2 NEW]", "No", "No", "No", "Yes", "No", "Yes"],
        ["View KPI Summary", "No", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["View Status History [V2]", "Yes (own)", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["Export SR Data", "No", "No", "No", "Yes", "Yes", "Yes"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 4: DATA MODEL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. Data Model", level=1)

# 4.1 service_requests
doc.add_heading("4.1 Entity: service_requests", level=2)
doc.add_paragraph("Primary table storing all service request records with full lifecycle tracking.")

styled_table(doc,
    ["#", "Field Name", "Data Type", "Required", "Default", "Validation / Notes"],
    [
        ["1", "id", "serial (PK)", "Auto", "Auto-increment", "Internal primary key"],
        ["2", "request_id", "varchar(20)", "Yes", "Auto (SEQUENCE)", "Format: SR-YYYY-NNNNNN, UNIQUE constraint. Generated via PostgreSQL SEQUENCE sr_request_id_seq [V2: was COUNT-based]"],
        ["3", "client_id", "varchar(50)", "Yes", "—", "FK → clients table"],
        ["4", "sr_type", "sr_type_enum", "Yes", "—", "Must be valid enum value"],
        ["5", "sr_details", "text", "No", "null", "Max 2000 characters"],
        ["6", "priority", "sr_priority_enum", "Yes", "'MEDIUM'", "LOW | MEDIUM | HIGH"],
        ["7", "sr_status", "sr_status_enum", "Yes", "'APPROVED'", "Current lifecycle status"],
        ["8", "request_date", "timestamp", "Yes", "now()", "Cannot be future date"],
        ["9", "closure_date", "timestamp", "Yes", "Computed", "request_date + SLA days; must be >= request_date"],
        ["10", "actual_closure_date", "timestamp", "No", "null", "Set on COMPLETED/REJECTED/CLOSED"],
        ["11", "appointed_start_date", "timestamp", "No", "null", "Scheduled service delivery start"],
        ["12", "appointed_end_date", "timestamp", "No", "null", "Must be >= appointed_start_date"],
        ["13", "remarks", "text", "No", "null", "Max 1000 characters"],
        ["14", "closure_reason", "text", "No", "null", "Required when → CLOSED"],
        ["15", "rejection_reason", "text", "No", "null", "Required when → REJECTED"],
        ["16", "verification_notes", "text", "No", "null", "Required when → INCOMPLETE"],
        ["17", "documents", "jsonb", "No", "[]", "Array of {filename, url, uploadedAt}"],
        ["18", "assigned_rm_id", "integer", "No", "null", "FK → users (RM role); auto-populated from client's mapped RM [V2: explicit]"],
        ["19", "service_branch", "varchar(100)", "No", "null", "Branch code handling the request"],
        ["20", "resolution_unit", "varchar(100)", "No", "null", "Team responsible for resolution"],
        ["21", "sales_date", "timestamp", "No", "null", "Sales date recorded by RM"],
        ["22", "teller_id", "integer", "No", "null", "FK → users (Teller role)"],
        ["23", "created_by", "varchar(100)", "Yes", "—", "Authenticated user ID from JWT [V2: was 'system']"],
        ["24", "updated_by", "varchar(100)", "Yes", "—", "Authenticated user ID from JWT [V2: was 'system']"],
        ["25", "created_at", "timestamp", "Yes", "now()", "Record creation timestamp"],
        ["26", "updated_at", "timestamp", "Yes", "now()", "Auto-set on every update"],
        ["27", "is_deleted", "boolean", "Yes", "false", "Soft delete flag"],
    ])

doc.add_paragraph()

# 4.2 sr_status_history — V2 NEW
doc.add_heading("4.2 Entity: sr_status_history [V2 NEW]", level=2)
doc.add_paragraph(
    "Append-only audit log table recording every status transition. This table is INSERT-only — "
    "records are never updated or deleted. Required for BSP examination readiness per BSP Circular 871 "
    "and MORB Part 4 Section 16 (complete records of all trust operation activities)."
)

styled_table(doc,
    ["#", "Field Name", "Data Type", "Required", "Default", "Validation / Notes"],
    [
        ["1", "id", "serial (PK)", "Auto", "Auto-increment", "Internal primary key"],
        ["2", "service_request_id", "integer", "Yes", "—", "FK → service_requests.id; indexed"],
        ["3", "from_status", "varchar(30)", "No", "null", "null for initial creation event"],
        ["4", "to_status", "varchar(30)", "Yes", "—", "Target status after transition"],
        ["5", "changed_by", "varchar(100)", "Yes", "—", "Authenticated user ID who triggered the change"],
        ["6", "changed_at", "timestamp", "Yes", "now()", "Timestamp of the status change"],
        ["7", "notes", "text", "No", "null", "Context: verification_notes, closure_reason, rejection_reason, or general notes"],
        ["8", "action", "varchar(50)", "Yes", "—", "Human-readable action: 'created', 'sent_for_verification', 'completed', 'marked_incomplete', 'rejected', 'closed', 'resubmitted', 'reassigned'"],
    ])

doc.add_paragraph()
doc.add_heading("4.2.1 sr_status_history Sample Data", level=3)
styled_table(doc,
    ["id", "service_request_id", "from_status", "to_status", "changed_by", "changed_at", "action", "notes"],
    [
        ["1", "1", "null", "APPROVED", "CLT-001", "2026-04-20 09:30", "created", "null"],
        ["2", "1", "APPROVED", "READY_FOR_TELLER", "RM-015", "2026-04-20 14:00", "sent_for_verification", "null"],
        ["3", "1", "READY_FOR_TELLER", "INCOMPLETE", "TLR-042", "2026-04-21 10:15", "marked_incomplete", "Missing KYC update form"],
        ["4", "1", "INCOMPLETE", "READY_FOR_TELLER", "CLT-001", "2026-04-21 16:30", "resubmitted", "Uploaded corrected KYC form"],
        ["5", "1", "READY_FOR_TELLER", "COMPLETED", "TLR-042", "2026-04-22 09:00", "completed", "null"],
    ])

doc.add_paragraph()

# 4.3 Enums
doc.add_heading("4.3 Enums", level=2)

doc.add_heading("4.3.1 sr_status_enum", level=3)
styled_table(doc,
    ["Value", "Description", "Terminal?"],
    [
        ["NEW", "Initial status (auto-transitions to APPROVED on create)", "No"],
        ["APPROVED", "Auto-approved; visible to client and RM for processing", "No"],
        ["READY_FOR_TELLER", "RM has sent for teller verification", "No"],
        ["COMPLETED", "Teller verified and completed", "Yes"],
        ["INCOMPLETE", "Teller returned with notes; action required", "No"],
        ["REJECTED", "Teller rejected with reason", "Yes"],
        ["CLOSED", "Manually closed by client or RM", "Yes"],
    ])

doc.add_paragraph()
doc.add_heading("4.3.2 sr_priority_enum", level=3)
styled_table(doc,
    ["Value", "SLA Days", "Description"],
    [
        ["HIGH", "3", "Urgent — expedited processing"],
        ["MEDIUM", "5", "Standard priority (default)"],
        ["LOW", "7", "Low-urgency — extended window"],
    ])

doc.add_paragraph()
doc.add_heading("4.3.3 sr_type_enum", level=3)
styled_table(doc,
    ["Value", "Display Label", "Description"],
    [
        ["REVIEW_PORTFOLIO", "Review Portfolio", "Client requests a portfolio review meeting"],
        ["MULTIPLE_MANDATE_REGISTRATION", "Multiple Mandate Registration", "Register additional mandates on trust account"],
        ["NOMINEE_UPDATION", "Nominee Update", "Update nominee/beneficiary details"],
        ["ACCOUNT_CLOSURE", "Account Closure", "Request closure of a trust account"],
        ["STATEMENT_REQUEST", "Statement Request", "Request for specific account statements"],
        ["ADDRESS_CHANGE", "Address Change", "Update registered address"],
        ["BENEFICIARY_UPDATE", "Beneficiary Update", "Add, remove, or modify trust beneficiaries"],
        ["GENERAL_INQUIRY", "General Inquiry", "General inquiry not covered by specific types"],
    ])

doc.add_paragraph()

# 4.4 Indexes & Relationships
doc.add_heading("4.4 Indexes & Relationships", level=2)

doc.add_heading("4.4.1 Indexes", level=3)
styled_table(doc,
    ["Index Name", "Table", "Column(s)", "Type", "Purpose"],
    [
        ["idx_sr_request_id", "service_requests", "request_id", "Unique", "Fast lookup by human-readable ID"],
        ["idx_sr_client_id", "service_requests", "client_id", "Non-unique", "Filter by client"],
        ["idx_sr_status", "service_requests", "sr_status", "Non-unique", "Filter by status"],
        ["idx_sr_priority", "service_requests", "priority", "Non-unique", "Filter by priority"],
        ["idx_sr_created_at", "service_requests", "created_at", "Non-unique", "Order by date"],
        ["idx_sr_assigned_rm", "service_requests", "assigned_rm_id", "Non-unique", "Filter by RM"],
        ["idx_srh_sr_id", "sr_status_history", "service_request_id", "Non-unique", "Join history to SR [V2]"],
        ["idx_srh_changed_at", "sr_status_history", "changed_at", "Non-unique", "Order history [V2]"],
    ])

doc.add_paragraph()
doc.add_heading("4.4.2 Relationships", level=3)
styled_table(doc,
    ["From", "To", "Type", "FK Column"],
    [
        ["service_requests", "clients", "Many-to-One", "client_id"],
        ["service_requests", "users (RM)", "Many-to-One", "assigned_rm_id"],
        ["service_requests", "users (Teller)", "Many-to-One", "teller_id"],
        ["sr_status_history", "service_requests", "Many-to-One", "service_request_id [V2]"],
    ])

doc.add_paragraph()

# 4.5 Sample Data
doc.add_heading("4.5 Sample Data: service_requests", level=2)
styled_table(doc,
    ["request_id", "client_id", "sr_type", "priority", "sr_status", "request_date", "closure_date", "assigned_rm_id", "remarks"],
    [
        ["SR-2026-000001", "CLT-001", "REVIEW_PORTFOLIO", "HIGH", "COMPLETED", "2026-04-20 09:30", "2026-04-23 09:30", "15", "Review equity allocation"],
        ["SR-2026-000002", "CLT-002", "NOMINEE_UPDATION", "MEDIUM", "READY_FOR_TELLER", "2026-04-18 14:15", "2026-04-23 14:15", "15", "Update nominee to spouse"],
        ["SR-2026-000003", "CLT-001", "STATEMENT_REQUEST", "LOW", "APPROVED", "2026-04-22 11:00", "2026-04-29 11:00", "15", "Q1 2026 statement for tax filing"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 5: FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. Functional Requirements", level=1)

# ── Module A: Client Portal ──
doc.add_heading("5.1 Module A: Client Portal", level=2)

doc.add_heading("FR-001: View Service Request List", level=3)
doc.add_paragraph("As a Client, I want to view a list of all my service requests so that I can track their status and take action on pending items.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-001.1: List displays: Request ID, SR Type, Priority (badge), Request Date, Closure Date, Request Age, Status (badge).",
    "AC-001.2: Search bar filters by Request ID (partial match) and Priority text.",
    "AC-001.3: Status tabs: All (default), Approved, In Progress, Completed, Rejected. Counts shown on tabs.",
    "AC-001.4: Desktop: data table; Mobile (< 768px): stacked cards.",
    "AC-001.5: Clicking a row navigates to /service-requests/:id.",
    "AC-001.6: '+ New Request' button in header links to /service-requests/new.",
    "AC-001.7: Empty state: 'No service requests found' with create link.",
    "AC-001.8: Pagination when results exceed 25 per page.",
    "AC-001.9: [V2] Notification badge on nav item showing count of INCOMPLETE SRs requiring action.",
])
doc.add_paragraph("Business Rules:")
bullets2(doc, [
    "BR-001.1: Only SRs for authenticated client_id (derived from JWT, not from request parameter) are returned.",
    "BR-001.2: Soft-deleted SRs (is_deleted=true) never displayed.",
    "BR-001.3: [V2] Filtering uses database-level WHERE clauses, not in-memory filtering.",
])

doc.add_heading("FR-002: Create New Service Request", level=3)
doc.add_paragraph("As a Client, I want to create a new service request so that I can request banking services without visiting a branch.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-002.1: SR Type dropdown lists all sr_type_enum values with display labels. Mandatory.",
    "AC-002.2: Service Request Details textarea (max 2000 chars).",
    "AC-002.3: Priority toggle: Low / Medium (default) / High.",
    "AC-002.4: Closure Date auto-computed and displayed read-only.",
    "AC-002.5: Remarks textarea (max 1000 chars, optional).",
    "AC-002.6: Document Upload: PDF only, max 10MB total.",
    "AC-002.7: Submit disabled until SR Type selected.",
    "AC-002.8: Success dialog: '<<Request ID>> is successfully submitted for processing.'",
    "AC-002.9: After dismiss, navigate to SR list.",
    "AC-002.10: Created SR has status APPROVED.",
    "AC-002.11: [V2] created_by set to authenticated user ID from JWT (not 'system').",
    "AC-002.12: [V2] assigned_rm_id auto-populated from client's mapped RM.",
    "AC-002.13: [V2] An sr_status_history record is inserted: {from_status: null, to_status: 'APPROVED', action: 'created'}.",
])
doc.add_paragraph("Business Rules:")
bullets2(doc, [
    "BR-002.1: [V2] Request ID generated via PostgreSQL SEQUENCE: SELECT nextval('sr_request_id_seq'), formatted as SR-YYYY-NNNNNN.",
    "BR-002.2: Closure Date = request_date + SLA days (HIGH=3, MEDIUM=5, LOW=7).",
    "BR-002.3: client_id derived from authenticated JWT token, not from request body.",
])

doc.add_heading("FR-003: View/Modify SR in APPROVED Status", level=3)
doc.add_paragraph("As a Client, I want to view and modify my SR in APPROVED status to add details before the RM processes it.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-003.1: Header: Request ID + Status badge (green 'Approved').",
    "AC-003.2: Read-only: Request ID, SR Type, Request Date, Closure Date, Priority, Age, Status.",
    "AC-003.3: Editable: Service Request Details, Remarks.",
    "AC-003.4: Documents: show uploaded files, option to add more.",
    "AC-003.5: 'Save' persists edits with confirmation toast.",
    "AC-003.6: 'Close Request' opens closure reason dialog (mandatory textarea).",
    "AC-003.7: After close: status → CLOSED, all fields read-only.",
    "AC-003.8: [V2] All saves record updated_by = authenticated user ID.",
    "AC-003.9: [V2] Status history entry inserted on close: {from: 'APPROVED', to: 'CLOSED', action: 'closed'}.",
])

doc.add_heading("FR-004: View/Modify SR in READY_FOR_TELLER Status", level=3)
doc.add_paragraph("As a Client, I want to view my SR during teller verification to track progress.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-004.1: All fields displayed; additional RM-filled fields visible (branch, unit, dates).",
    "AC-004.2: Closure Date editable only if SR Type config allows.",
    "AC-004.3: Remarks editable.",
    "AC-004.4: 'Close Request' available.",
    "AC-004.5: Cannot modify SR Type, Priority, or Details.",
])

doc.add_heading("FR-005: View/Modify SR in INCOMPLETE Status", level=3)
doc.add_paragraph("As a Client, I want to address teller feedback and resubmit my SR.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-005.1: Verification notes from teller displayed in yellow alert.",
    "AC-005.2: Editable: Details, Remarks.",
    "AC-005.3: Document re-upload available.",
    "AC-005.4: 'Save' persists (status stays INCOMPLETE).",
    "AC-005.5: 'Re-send for Verification' transitions → READY_FOR_TELLER.",
    "AC-005.6: 'Close Request' available.",
    "AC-005.7: [V2] Status history on resubmit: {from: 'INCOMPLETE', to: 'READY_FOR_TELLER', action: 'resubmitted'}.",
])

doc.add_heading("FR-006: View SR in Terminal Status", level=3)
doc.add_paragraph("As a Client, I want to view completed/rejected/closed SRs for my records.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-006.1: All fields read-only.",
    "AC-006.2: REJECTED: red alert with rejection_reason.",
    "AC-006.3: CLOSED: yellow alert with closure_reason.",
    "AC-006.4: COMPLETED: actual_closure_date displayed.",
    "AC-006.5: [V2] Status history timeline available showing all transitions.",
])

doc.add_page_break()

# ── Module B: Back-Office ──
doc.add_heading("5.2 Module B: Back-Office Workbench", level=2)

doc.add_heading("FR-007: SR Workbench KPI Dashboard", level=3)
doc.add_paragraph("As an Operations Manager, I want a summary dashboard to monitor the SR pipeline and SLA breaches.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-007.1: KPI cards: Total, Approved, Ready for Teller, Completed, Overdue SLA.",
    "AC-007.2: Overdue SLA card highlighted red/amber when count > 0.",
    "AC-007.3: Refresh on page load; manual refresh button.",
    "AC-007.4: Loading skeleton while fetching.",
])

doc.add_heading("FR-008: SR Workbench List View", level=3)
doc.add_paragraph("As an RM, I want to view and filter all SRs to prioritize my work queue.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-008.1: Columns: Request ID, Client, SR Type, Priority, Request Date, Closure Date, Age, Status, Actions.",
    "AC-008.2: Status tabs: All, New, Approved, Ready for Teller, Incomplete, Completed, Rejected, Closed.",
    "AC-008.3: Search across Request ID, Client Name, SR Type.",
    "AC-008.4: Pagination (10/25/50/100 per page).",
    "AC-008.5: Context-specific action buttons per row.",
    "AC-008.6: SLA-breached rows highlighted.",
    "AC-008.7: [V2] Database-level filtering; no in-memory filtering.",
])

doc.add_heading("FR-009: Send for Verification (RM)", level=3)
doc.add_paragraph("As an RM, I want to prepare and send an SR for teller verification.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-009.1: Available only for APPROVED status.",
    "AC-009.2: Dialog: Service Branch (mandatory), Resolution Unit, Sales Date.",
    "AC-009.3: Transitions APPROVED → READY_FOR_TELLER.",
    "AC-009.4: [V2] Status history: {action: 'sent_for_verification', changed_by: RM user ID}.",
    "AC-009.5: [V2] updated_by = authenticated RM user ID.",
])

doc.add_heading("FR-010: Complete Request (Teller)", level=3)
doc.add_paragraph("As a Teller, I want to mark an SR as completed after verification.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-010.1: Available only for READY_FOR_TELLER status.",
    "AC-010.2: Confirmation dialog before completing.",
    "AC-010.3: Transitions → COMPLETED; sets actual_closure_date and teller_id.",
    "AC-010.4: [V2] Status history: {action: 'completed', changed_by: Teller user ID}.",
])

doc.add_heading("FR-011: Mark Incomplete (Teller)", level=3)
doc.add_paragraph("As a Teller, I want to return an SR with notes explaining what needs correction.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-011.1: Available only for READY_FOR_TELLER status.",
    "AC-011.2: Dialog with mandatory verification notes (min 10 chars).",
    "AC-011.3: Transitions → INCOMPLETE; saves teller_id and verification_notes.",
    "AC-011.4: [V2] Status history: {action: 'marked_incomplete', notes: verification_notes}.",
])

doc.add_heading("FR-012: Reject Request (Teller)", level=3)
doc.add_paragraph("As a Teller, I want to reject an SR with a documented reason.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-012.1: Available only for READY_FOR_TELLER status.",
    "AC-012.2: Dialog with mandatory rejection reason (min 10 chars).",
    "AC-012.3: Warning: 'This action cannot be undone.'",
    "AC-012.4: Transitions → REJECTED; sets rejection_reason and actual_closure_date.",
    "AC-012.5: [V2] Status history: {action: 'rejected', notes: rejection_reason}.",
])

doc.add_heading("FR-013: Update SR Details (RM)", level=3)
doc.add_paragraph("As an RM, I want to update SR details before sending for verification.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-013.1: Editable: Details, Remarks, Documents, Closure Date.",
    "AC-013.2: Cannot modify terminal-status SRs (COMPLETED/REJECTED/CLOSED).",
    "AC-013.3: [V2] updated_by = authenticated RM user ID.",
])

# FR-014: RM Reassignment — V2 NEW
doc.add_heading("FR-014: Reassign RM [V2 NEW]", level=3)
doc.add_paragraph("As an Operations Manager, I want to reassign a service request to a different RM so that SRs are not orphaned when an RM is on leave or reassigned.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-014.1: Available for SRs in non-terminal statuses (APPROVED, READY_FOR_TELLER, INCOMPLETE).",
    "AC-014.2: Reassign dialog: dropdown of active users with RM role, filtered by branch (optional).",
    "AC-014.3: Only Operations Manager and Admin roles can perform reassignment.",
    "AC-014.4: assigned_rm_id is updated; updated_by set to the Ops Manager's user ID.",
    "AC-014.5: Status history entry: {action: 'reassigned', notes: 'Reassigned from RM-015 to RM-022', changed_by: Ops Manager ID}.",
    "AC-014.6: Success toast: '<<Request ID>> reassigned to <<RM Name>>.'",
    "AC-014.7: The SR status does NOT change during reassignment — only the assigned RM changes.",
])

doc.add_page_break()

# ── Module C: Backend ──
doc.add_heading("5.3 Module C: Backend Service & API", level=2)

doc.add_heading("FR-015: Concurrency-Safe Request ID Generation [V2 NEW]", level=3)
doc.add_paragraph("As the System, I want to generate unique request IDs safely under concurrent load so that no duplicate IDs are created.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-015.1: A PostgreSQL SEQUENCE named sr_request_id_seq is created.",
    "AC-015.2: ID generation uses SELECT nextval('sr_request_id_seq'), not COUNT(*).",
    "AC-015.3: Format: SR-YYYY-NNNNNN where NNNNNN is the sequence value zero-padded to 6 digits.",
    "AC-015.4: YYYY is the current year at creation time.",
    "AC-015.5: The UNIQUE constraint on request_id provides a final safety net.",
    "AC-015.6: Concurrent creation of 100 SRs results in 100 unique request_ids with no errors.",
])

doc.add_heading("FR-016: SLA Closure Date Computation", level=3)
doc.add_paragraph("As the System, I want to compute SLA closure dates so that every SR has a target completion date.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-016.1: closure_date = request_date + SLA_DAYS[priority].",
    "AC-016.2: SLA_DAYS: HIGH=3, MEDIUM=5, LOW=7 (calendar days for Phase 1).",
    "AC-016.3: Unknown priority defaults to 5 days.",
    "AC-016.4: Closure date set at creation, not auto-recalculated on priority change.",
    "AC-016.5: Phase 2: Business-day computation with market calendar integration.",
])

doc.add_heading("FR-017: Status Transition Validation", level=3)
doc.add_paragraph("As the System, I want to enforce valid status transitions so that SRs cannot enter invalid states.")
styled_table(doc,
    ["From Status", "To Status", "Triggered By", "Required Fields"],
    [
        ["NEW", "APPROVED", "System (auto)", "None"],
        ["APPROVED", "READY_FOR_TELLER", "RM", "service_branch (mandatory)"],
        ["APPROVED", "CLOSED", "Client / RM", "closure_reason"],
        ["READY_FOR_TELLER", "COMPLETED", "Teller", "teller_id"],
        ["READY_FOR_TELLER", "INCOMPLETE", "Teller", "teller_id, verification_notes"],
        ["READY_FOR_TELLER", "REJECTED", "Teller", "rejection_reason"],
        ["READY_FOR_TELLER", "CLOSED", "Client / RM", "closure_reason"],
        ["INCOMPLETE", "READY_FOR_TELLER", "Client / RM", "None"],
        ["INCOMPLETE", "CLOSED", "Client / RM", "closure_reason"],
    ])
doc.add_paragraph("Rules: Invalid transitions return HTTP 400. Terminal statuses cannot transition. Every transition inserts an sr_status_history record [V2].")

doc.add_heading("FR-018: Database-Level Filtering [V2 NEW]", level=3)
doc.add_paragraph("As the System, I want to filter service requests at the database level so that list queries perform efficiently at scale.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-018.1: All filter parameters (client_id, status, priority, search) are applied as SQL WHERE clauses.",
    "AC-018.2: Pagination uses SQL LIMIT/OFFSET, not in-memory array slicing.",
    "AC-018.3: Search uses SQL ILIKE for partial matching on request_id, sr_type, sr_details.",
    "AC-018.4: Total count uses a separate COUNT(*) query with the same WHERE conditions.",
    "AC-018.5: P95 response time ≤ 500ms with 10,000 records and combined filters.",
])

doc.add_heading("FR-019: Authenticated User Tracking [V2 NEW]", level=3)
doc.add_paragraph("As the System, I want to record the actual authenticated user identity in all audit fields so that the audit trail is accurate.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-019.1: created_by is set to the authenticated user's ID extracted from the JWT token.",
    "AC-019.2: updated_by is set to the authenticated user's ID on every update operation.",
    "AC-019.3: sr_status_history.changed_by is set to the authenticated user's ID.",
    "AC-019.4: No service method defaults to 'system' — the caller must always provide the user ID.",
    "AC-019.5: API route handlers extract user ID from req.user (populated by auth middleware) and pass to service methods.",
])

doc.add_page_break()

# ── Module D: In-App Notifications — V2 NEW ──
doc.add_heading("5.4 Module D: In-App Notifications [V2 NEW]", level=2)

doc.add_heading("FR-020: Client Portal Action-Required Badge", level=3)
doc.add_paragraph("As a Client, I want to see a badge on the Service Requests navigation item showing how many SRs need my attention so that I don't miss items requiring action.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-020.1: The 'Service Requests' nav item in the Client Portal sidebar shows a numeric badge.",
    "AC-020.2: Badge count = number of SRs in INCOMPLETE status for the authenticated client.",
    "AC-020.3: Badge is hidden (not shown as '0') when count is zero.",
    "AC-020.4: Badge data is fetched on initial page load and refreshed every 60 seconds via polling.",
    "AC-020.5: API endpoint: GET /api/v1/client-portal/service-requests/action-count/:clientId returning {count: N}.",
    "AC-020.6: Badge uses a red circle with white text, consistent with other notification badges in the app.",
])

doc.add_heading("FR-021: View Status History Timeline", level=3)
doc.add_paragraph("As a Client or Back-Office user, I want to see the complete status history of an SR so that I can understand its full processing timeline.")
doc.add_paragraph("Acceptance Criteria:")
bullets2(doc, [
    "AC-021.1: SR detail page includes a 'History' section (collapsible, default collapsed).",
    "AC-021.2: History displays a vertical timeline of all status changes from sr_status_history.",
    "AC-021.3: Each entry shows: timestamp, action label, from → to status, changed_by name, and notes (if any).",
    "AC-021.4: Entries ordered chronologically (oldest first).",
    "AC-021.5: API endpoint: GET /api/v1/service-requests/:id/history returning array of history records.",
    "AC-021.6: Client portal version shows history for own SRs only.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 6: USER INTERFACE REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. User Interface Requirements", level=1)

doc.add_heading("6.1 Client Portal Screens", level=2)

doc.add_heading("6.1.1 Service Request List (/service-requests)", level=3)
bullets(doc, [
    "Layout: Full-width page; header (title + '+ New Request' button), search bar, tab strip, content area.",
    "Search: Text input with debounce (300ms), placeholder 'Search by Request ID or Priority'.",
    "Tabs: All | Approved | In Progress | Completed | Rejected. Active tab underlined.",
    "Desktop (≥ 768px): Data table; Mobile (< 768px): Stacked cards.",
    "Empty: Centered icon + 'No service requests found' + create link.",
    "Loading: Skeleton rows/cards. Error: Red alert + 'Retry' button.",
    "Pagination: Page buttons + Previous/Next; 'Showing 1-25 of 47'.",
    "[V2] Nav badge: Red circle badge on 'Service Requests' sidebar item showing INCOMPLETE count.",
])

doc.add_heading("6.1.2 Create Service Request (/service-requests/new)", level=3)
bullets(doc, [
    "Layout: Centered form card (max-width 600px) with back arrow.",
    "Fields: SR Type (dropdown, mandatory), Details (textarea), Priority (segmented), Closure Date (read-only), Remarks (textarea), Status (read-only 'New'), Document Upload (PDF, 10MB).",
    "Submit: Full-width 'Send Service Request' button; disabled until SR Type selected; spinner during submit.",
    "Success: Modal with check icon, Request ID, success message, 'OK' button.",
])

doc.add_heading("6.1.3 Service Request Detail (/service-requests/:id)", level=3)
bullets(doc, [
    "Layout: Card detail view with back nav. Header: Request ID + Status badge.",
    "Info Grid: SR Type + Closure Date (row 1), Request Age + Priority (row 2).",
    "Sections: Details, Remarks, Documents, Verification Notes (INCOMPLETE), Rejection/Closure Reason (terminal).",
    "Actions: Status-dependent buttons (Save, Close Request, Re-send for Verification).",
    "Close Dialog: AlertDialog with mandatory closure reason textarea.",
    "[V2] History Section: Collapsible vertical timeline from sr_status_history.",
    "[V2] All edits set updated_by from authenticated user.",
])

doc.add_heading("6.2 Back-Office Screens", level=2)

doc.add_heading("6.2.1 SR Workbench (/operations/service-requests)", level=3)
bullets(doc, [
    "Layout: KPI cards row → tab strip → content area.",
    "KPI Cards: Total (blue), Approved (green), Ready for Teller (amber), Completed (teal), Overdue SLA (red).",
    "Tabs: All | New | Approved | Ready for Teller | Incomplete | Completed | Rejected | Closed.",
    "Table: Request ID, Client, SR Type, Priority, Request Date, Closure Date, Age, Status, Actions.",
    "Actions: Send for Verification (APPROVED), Complete/Incomplete/Reject (READY_FOR_TELLER).",
    "[V2] Reassign RM button: Available for non-terminal SRs; dropdown of RM users.",
    "Dialogs: Send for Verification (branch/unit/date), Mark Incomplete (notes), Reject (reason), Reassign RM (dropdown) [V2].",
    "Row highlighting: Amber tint for SLA-breached open SRs.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 7: API & INTEGRATION REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. API & Integration Requirements", level=1)

doc.add_heading("7.1 Authentication & Error Format", level=2)
doc.add_paragraph("Authentication: JWT in httpOnly cookies. Client portal scoped to client_id from JWT. Back-office: requireBackOfficeRole() middleware.")
doc.add_paragraph("Standard Error Response:")
doc.add_paragraph('{ "error": { "code": "VALIDATION_ERROR|NOT_FOUND|FORBIDDEN|INVALID_TRANSITION|INTERNAL_ERROR", "message": "...", "field": "..." } }')

doc.add_heading("7.2 Client Portal Endpoints", level=2)
styled_table(doc,
    ["Method", "Path", "Description", "Auth"],
    [
        ["GET", "/api/v1/client-portal/service-requests/:clientId", "List SRs for client (DB-filtered) [V2]", "Client (own)"],
        ["POST", "/api/v1/client-portal/service-requests", "Create new SR (sequence ID) [V2]", "Client"],
        ["GET", "/api/v1/client-portal/service-requests/detail/:id", "Get SR detail + history [V2]", "Client (own)"],
        ["PUT", "/api/v1/client-portal/service-requests/:id", "Update SR fields", "Client (own)"],
        ["PUT", "/api/v1/client-portal/service-requests/:id/close", "Close SR", "Client (own)"],
        ["PUT", "/api/v1/client-portal/service-requests/:id/resubmit", "Resubmit from INCOMPLETE", "Client (own)"],
        ["GET", "/api/v1/client-portal/service-requests/action-count/:clientId [V2]", "Count SRs needing action", "Client (own)"],
    ])

doc.add_paragraph()
doc.add_heading("7.3 Back-Office Endpoints", level=2)
styled_table(doc,
    ["Method", "Path", "Description", "Auth"],
    [
        ["GET", "/api/v1/service-requests", "List all SRs (DB-filtered) [V2]", "Back-office"],
        ["GET", "/api/v1/service-requests/summary", "KPI summary", "Back-office"],
        ["GET", "/api/v1/service-requests/:id", "Get SR detail", "Back-office"],
        ["GET", "/api/v1/service-requests/:id/history [V2]", "Get status history timeline", "Back-office"],
        ["PUT", "/api/v1/service-requests/:id", "Update SR fields", "Back-office (RM)"],
        ["PUT", "/api/v1/service-requests/:id/send-for-verification", "RM → teller", "Back-office (RM)"],
        ["PUT", "/api/v1/service-requests/:id/complete", "Teller completes", "Back-office (Teller)"],
        ["PUT", "/api/v1/service-requests/:id/incomplete", "Teller → incomplete", "Back-office (Teller)"],
        ["PUT", "/api/v1/service-requests/:id/reject", "Teller rejects", "Back-office (Teller)"],
        ["PUT", "/api/v1/service-requests/:id/reassign [V2]", "Ops Mgr reassigns RM", "Back-office (Ops Mgr)"],
    ])

doc.add_paragraph()
doc.add_heading("7.4 Key Request/Response Examples", level=2)

doc.add_heading("7.4.1 Create SR (POST)", level=3)
doc.add_paragraph("Request:")
doc.add_paragraph('{\n  "sr_type": "REVIEW_PORTFOLIO",\n  "sr_details": "Review equity allocation",\n  "priority": "HIGH",\n  "remarks": "Prefer Makati branch meeting",\n  "documents": ["portfolio_summary.pdf"]\n}')
doc.add_paragraph("Note: client_id and created_by derived from JWT token, not from request body [V2].")
doc.add_paragraph("Response (201):")
doc.add_paragraph('{\n  "id": 1,\n  "request_id": "SR-2026-000001",\n  "sr_status": "APPROVED",\n  "closure_date": "2026-04-26T09:30:00.000Z",\n  "assigned_rm_id": 15,\n  "created_by": "CLT-001"\n}')

doc.add_heading("7.4.2 Reassign RM (PUT) [V2 NEW]", level=3)
doc.add_paragraph("Request: PUT /api/v1/service-requests/1/reassign")
doc.add_paragraph('{\n  "assigned_rm_id": 22\n}')
doc.add_paragraph("Response (200):")
doc.add_paragraph('{\n  "id": 1,\n  "request_id": "SR-2026-000001",\n  "assigned_rm_id": 22,\n  "updated_by": "OPS-MGR-003"\n}')

doc.add_heading("7.4.3 Get Status History (GET) [V2 NEW]", level=3)
doc.add_paragraph("GET /api/v1/service-requests/1/history")
doc.add_paragraph("Response (200):")
doc.add_paragraph(
    '[\n'
    '  {"id": 1, "from_status": null, "to_status": "APPROVED", "action": "created", "changed_by": "CLT-001", "changed_at": "2026-04-20T09:30:00Z", "notes": null},\n'
    '  {"id": 2, "from_status": "APPROVED", "to_status": "READY_FOR_TELLER", "action": "sent_for_verification", "changed_by": "RM-015", "changed_at": "2026-04-20T14:00:00Z", "notes": null},\n'
    '  {"id": 3, "from_status": "READY_FOR_TELLER", "to_status": "COMPLETED", "action": "completed", "changed_by": "TLR-042", "changed_at": "2026-04-21T10:00:00Z", "notes": null}\n'
    ']'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 8: NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("8. Non-Functional Requirements", level=1)

doc.add_heading("8.1 Performance", level=2)
styled_table(doc,
    ["Metric", "Target", "Notes"],
    [
        ["List API P95", "≤ 500ms", "Database-level filtering with indexes [V2]"],
        ["Detail API P95", "≤ 200ms", "Single record by ID"],
        ["Create API P95", "≤ 1000ms", "Includes SEQUENCE nextval + history insert [V2]"],
        ["History API P95", "≤ 300ms", "Indexed by service_request_id [V2]"],
        ["UI page load", "≤ 2 seconds", "Lazy-loaded components"],
        ["Concurrent users", "100+", "Back-office users simultaneously"],
    ])

doc.add_heading("8.2 Security", level=2)
bullets(doc, [
    "Auth: JWT in httpOnly cookies.",
    "Authorization: Client scoped to client_id from JWT [V2]; back-office via requireBackOfficeRole().",
    "Input: Server-side sanitization; Drizzle ORM parameterized queries.",
    "CSRF: SameSite cookie. XSS: React default escaping.",
    "Uploads: PDF only, 10MB max, malware scan by upload service.",
    "Audit: Append-only sr_status_history table; never updated or deleted [V2].",
    "Encryption: PostgreSQL TDE at rest; HTTPS/TLS 1.2+ in transit.",
    "[V2] All audit fields (created_by, updated_by, changed_by) use authenticated user ID, never 'system'.",
])

doc.add_heading("8.3 Other NFRs", level=2)
bullets(doc, [
    "Scalability: Designed for 50,000 SRs/year. Database-level filtering ensures linear query performance [V2].",
    "Availability: 99.5% during business hours (Mon-Fri, 8AM-6PM PHT).",
    "Accessibility: WCAG 2.1 AA. Labels, keyboard nav, screen reader support.",
    "Browsers: Chrome/Firefox/Safari/Edge latest 2 versions; mobile Chrome/Safari.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 9: WORKFLOW & STATE DIAGRAMS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("9. Workflow & State Diagrams", level=1)

doc.add_heading("9.1 State Transition Table", level=2)
styled_table(doc,
    ["#", "Current", "Action", "Next", "Actor", "Side Effects"],
    [
        ["T1", "—", "Create", "APPROVED", "Client/RM", "Generate request_id (SEQUENCE) [V2]; compute closure_date; set assigned_rm_id from client mapping [V2]; insert status_history [V2]"],
        ["T2", "APPROVED", "Send for Verification", "READY_FOR_TELLER", "RM", "Set branch, unit, date; insert status_history [V2]"],
        ["T3", "APPROVED", "Close", "CLOSED", "Client/RM", "Set closure_reason, actual_closure_date; insert status_history [V2]"],
        ["T4", "READY_FOR_TELLER", "Complete", "COMPLETED", "Teller", "Set teller_id, actual_closure_date; insert status_history [V2]"],
        ["T5", "READY_FOR_TELLER", "Mark Incomplete", "INCOMPLETE", "Teller", "Set teller_id, verification_notes; insert status_history [V2]"],
        ["T6", "READY_FOR_TELLER", "Reject", "REJECTED", "Teller", "Set rejection_reason, actual_closure_date; insert status_history [V2]"],
        ["T7", "READY_FOR_TELLER", "Close", "CLOSED", "Client/RM", "Set closure_reason, actual_closure_date; insert status_history [V2]"],
        ["T8", "INCOMPLETE", "Resubmit", "READY_FOR_TELLER", "Client/RM", "Update remarks, docs; insert status_history [V2]"],
        ["T9", "INCOMPLETE", "Close", "CLOSED", "Client/RM", "Set closure_reason, actual_closure_date; insert status_history [V2]"],
        ["T10", "Any non-terminal", "Reassign RM [V2]", "(unchanged)", "Ops Mgr", "Update assigned_rm_id; insert status_history with action='reassigned' [V2]"],
    ])

doc.add_heading("9.2 Happy Path Workflow", level=2)
steps = [
    "1. Client creates SR via portal → status: APPROVED, request_id generated via SEQUENCE [V2].",
    "2. assigned_rm_id auto-populated from client's mapped RM [V2].",
    "3. Status history: {action: 'created', changed_by: client ID} [V2].",
    "4. RM sees SR in workbench, contacts client, collects docs.",
    "5. RM clicks 'Send for Verification' (branch=MAKATI-001, unit=Trust Ops-Equity).",
    "6. Status history: {action: 'sent_for_verification', changed_by: RM ID} [V2].",
    "7. Teller reviews, clicks 'Complete'.",
    "8. Status history: {action: 'completed', changed_by: Teller ID} [V2].",
    "9. Client sees COMPLETED status; can view full history timeline [V2].",
]
for s in steps:
    doc.add_paragraph(s)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 10: NOTIFICATIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("10. Notification & Communication Requirements", level=1)
doc.add_paragraph("Phase 1 delivers in-app notification badges [V2]. Full push/email delivery is Phase 2.")

styled_table(doc,
    ["#", "Event", "Phase 1 (In-App)", "Phase 2 (Email/Push)", "Recipient", "Template"],
    [
        ["N1", "SR Created", "Badge increment", "Email", "Client", '"{request_id} submitted. Closure: {closure_date}."'],
        ["N2", "Sent for Verification", "—", "In-app", "Client", '"{request_id} is being processed."'],
        ["N3", "Completed", "—", "Email", "Client", '"{request_id} completed successfully."'],
        ["N4", "Marked Incomplete", "Badge increment [V2]", "Email", "Client", '"{request_id} needs attention: {notes}"'],
        ["N5", "Rejected", "—", "Email", "Client", '"{request_id} rejected: {reason}"'],
        ["N6", "Closed", "—", "In-app", "Client/RM", '"{request_id} closed: {reason}"'],
        ["N7", "SLA Warning (1 day)", "—", "Email", "RM/Ops Mgr", '"{request_id} SLA breach in 24h."'],
        ["N8", "SLA Breached", "—", "Email", "RM/Ops Mgr", '"{request_id} SLA breached. Age: {age}d."'],
        ["N9", "Resubmitted", "—", "In-app", "Teller", '"{request_id} resubmitted."'],
        ["N10", "RM Reassigned [V2]", "—", "In-app", "New RM", '"{request_id} assigned to you."'],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 11: REPORTING
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("11. Reporting & Analytics", level=1)

doc.add_heading("11.1 Phase 1: KPI Summary Dashboard", level=2)
styled_table(doc,
    ["Metric", "Calculation", "Audience"],
    [
        ["Total SRs", "COUNT(*) WHERE is_deleted=false", "RM, Ops Mgr"],
        ["By Status", "COUNT(*) GROUP BY sr_status", "RM, Ops Mgr"],
        ["Overdue SLA", "COUNT(*) WHERE status IN (open) AND now() > closure_date", "Ops Mgr"],
    ])

doc.add_heading("11.2 Phase 2: Extended Reports", level=2)
bullets(doc, [
    "SR Volume Trends (weekly/monthly chart)",
    "SR Type Distribution (pie chart)",
    "Branch Performance Comparison",
    "RM Workload Distribution",
    "Aging Analysis (histogram by age bucket)",
    "SLA Compliance Rate over time",
    "Status History Analytics (average time in each status)",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 12: MIGRATION & LAUNCH
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("12. Migration & Launch Plan", level=1)

doc.add_heading("12.1 Data Migration", level=2)
doc.add_paragraph("Greenfield module — no legacy data migration. Existing paper/spreadsheet SRs can be manually entered post-launch.")

doc.add_heading("12.2 Phased Rollout", level=2)
styled_table(doc,
    ["Phase", "Scope"],
    [
        ["Phase 1 (MVP)", "Full lifecycle + 6 must-fix items from adversarial evaluation (this BRD)"],
        ["Phase 2", "Business-day SLA, configurable SR types, duplicate detection, priority changes, doc versioning, bulk ops, push/email notifications, SLA escalation"],
        ["Phase 3", "Analytics/reporting, CRM integration, client satisfaction surveys"],
    ])

doc.add_heading("12.3 Go-Live Checklist", level=2)
bullets(doc, [
    "DB migration: service_requests table, sr_status_history table, sr_request_id_seq SEQUENCE, enums, indexes [V2].",
    "Backend: All service methods use authenticated user ID (no 'system' defaults) [V2].",
    "Backend: Request ID uses SEQUENCE-based generation [V2].",
    "Backend: List queries use database-level WHERE filtering [V2].",
    "Backend: Every status transition inserts sr_status_history record [V2].",
    "Client Portal: List, Create, Detail views render on desktop and mobile.",
    "Client Portal: Notification badge shows INCOMPLETE count [V2].",
    "Back-Office: KPI cards, tab filters, all action dialogs functional.",
    "Back-Office: RM reassignment dialog works for Ops Manager role [V2].",
    "Security: Client endpoints scoped to JWT client_id; back-office role-protected.",
    "Performance: List endpoint < 500ms with 1000+ records.",
    "Audit: Status history verified for all transitions end-to-end.",
])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 13: GLOSSARY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("13. Glossary", level=1)
styled_table(doc,
    ["Term", "Definition"],
    [
        ["SR", "Service Request — a client-initiated request for banking services"],
        ["SLA", "Service Level Agreement — agreed timeframe for completing a service request"],
        ["RM", "Relationship Manager — bank officer managing a client portfolio"],
        ["Teller", "Operations staff performing SR verification and processing"],
        ["BSP", "Bangko Sentral ng Pilipinas — Philippine central bank and financial regulator"],
        ["MORB", "Manual of Regulations for Banks — BSP regulatory framework"],
        ["KPI", "Key Performance Indicator"],
        ["Closure Date", "SLA-computed target date for SR completion"],
        ["Request Age", "Calendar days since SR creation"],
        ["Dual Control", "Security principle: RM prepares, Teller verifies"],
        ["SEQUENCE", "PostgreSQL object for generating unique sequential numbers [V2]"],
        ["Append-Only", "Insert-only table pattern ensuring audit records are never modified [V2]"],
        ["JWT", "JSON Web Token — authentication token format"],
        ["JSONB", "PostgreSQL binary JSON type for document references"],
        ["Drizzle ORM", "TypeScript ORM for database operations"],
        ["shadcn/ui", "React component library for UI elements"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 14: APPENDICES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("14. Appendices", level=1)

doc.add_heading("14.1 Appendix A: Field Editability Matrix by Status", level=2)
styled_table(doc,
    ["Field", "Initiate", "Approved", "Ready for Teller", "Incomplete", "Terminal"],
    [
        ["SR Type", "Select (M)", "Display", "Display", "Display", "Display"],
        ["SR Details", "Input", "Editable", "Display", "Editable", "Display"],
        ["Priority", "Select (M)", "Display", "Display", "Display", "Display"],
        ["Request Date", "Input (M)", "Display", "Display", "Display", "Display"],
        ["Closure Date", "Computed", "Computed", "Editable*", "Editable*", "Display"],
        ["Remarks", "Input", "Editable", "Editable", "Editable", "Display"],
        ["Documents", "Upload", "Upload", "Display", "Upload", "Display"],
        ["Status", "Display", "Display", "Display", "Display", "Display"],
    ])
doc.add_paragraph("(M)=Mandatory. *=Editable only if SR Type allows closure date changes. Terminal=COMPLETED/REJECTED/CLOSED.")

doc.add_paragraph()
doc.add_heading("14.2 Appendix B: V1 → V2 Change Log", level=2)
doc.add_paragraph("Summary of all changes made from BRD v1.0 to v2.0 FINAL based on adversarial evaluation findings:")
styled_table(doc,
    ["#", "Finding", "Severity", "Change Made"],
    [
        ["1", "Request ID race condition (COUNT-based generation)", "HIGH", "Replaced with PostgreSQL SEQUENCE (sr_request_id_seq). FR-015 added."],
        ["2", "Missing status change audit history", "HIGH", "Added sr_status_history append-only table (Section 4.2). FR-021 added. All transitions insert history record."],
        ["3", "In-memory filtering not scalable", "MEDIUM", "Specified database-level WHERE filtering. FR-018 added. BR-001.3 updated."],
        ["4", "updated_by hardcoded to 'system'", "MEDIUM", "All audit fields use authenticated user ID from JWT. FR-019 added."],
        ["5", "No client notification mechanism", "MEDIUM", "Added in-app notification badges for INCOMPLETE SRs. FR-020 added. API endpoint added."],
        ["6", "No RM reassignment capability", "MEDIUM", "Added reassign action for Ops Manager. FR-014 added. API endpoint added. Status history tracks reassignment."],
    ])

doc.add_paragraph()
doc.add_heading("14.3 Appendix C: Phase 2 Roadmap", level=2)
styled_table(doc,
    ["#", "Item", "Priority", "Rationale"],
    [
        ["P2-1", "Business-day SLA with market calendar", "HIGH", "Calendar-day SLA inflates breach counts on weekends/holidays"],
        ["P2-2", "Configurable SR Type table (replacing enum)", "HIGH", "Eliminates schema migrations for new SR types"],
        ["P2-3", "Duplicate SR detection", "MEDIUM", "Prevents queue clutter from identical requests"],
        ["P2-4", "Priority change with SLA recalculation", "MEDIUM", "Handles urgency changes post-creation"],
        ["P2-5", "Document versioning", "MEDIUM", "Tracks original vs. corrected document uploads"],
        ["P2-6", "Bulk operations (close, reassign, export)", "LOW", "Batch processing for operations efficiency"],
        ["P2-7", "Push/email notification delivery", "HIGH", "Full notification infrastructure for all 10 events"],
        ["P2-8", "SLA auto-escalation to supervisor", "MEDIUM", "Automated escalation on breach"],
    ])

doc.add_paragraph()
doc.add_heading("14.4 Appendix D: Badge Color Coding", level=2)
styled_table(doc,
    ["Element", "Value", "Color", "Hex"],
    [
        ["Priority", "HIGH", "Red", "#EF4444"],
        ["Priority", "MEDIUM", "Amber", "#F59E0B"],
        ["Priority", "LOW", "Green", "#10B981"],
        ["Status", "NEW", "Blue", "#3B82F6"],
        ["Status", "APPROVED", "Green", "#10B981"],
        ["Status", "READY_FOR_TELLER", "Amber", "#F59E0B"],
        ["Status", "COMPLETED", "Teal", "#14B8A6"],
        ["Status", "INCOMPLETE", "Orange", "#F97316"],
        ["Status", "REJECTED", "Red", "#EF4444"],
        ["Status", "CLOSED", "Gray", "#6B7280"],
    ])

doc.add_paragraph()
doc.add_heading("14.5 Appendix E: Regulatory References", level=2)
styled_table(doc,
    ["Regulation", "Relevance", "How This BRD Complies"],
    [
        ["BSP Circular 871 (Internal Controls)", "Audit Trail", "sr_status_history table logs every transition with user ID and timestamp [V2]"],
        ["BSP Circular 808 (IT Risk Management)", "Data Security", "Encrypted at rest/transit; role-based access; JWT auth"],
        ["Data Privacy Act 2012 (RA 10173)", "Data Protection", "Soft delete preserves data; access scoped to client_id"],
        ["MORB Part 4 Sec 16 (Trust Operations)", "Complete Records", "Append-only history; 5-year retention; no record deletion [V2]"],
    ])


# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
out = "/Users/n15318/Trust OMS/docs/ServiceRequest_TaskManagement_BRD_v2_FINAL.docx"
doc.save(out)
print(f"BRD v2 FINAL generated: {out}")
