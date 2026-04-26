#!/usr/bin/env python3
"""Generate BRD Part 4: Sections 6-10"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document('/Users/n15318/Trust OMS/docs/CIM_BRD_Draft.docx')

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1B3A5C')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E8F0FE')
                shading.set(qn('w:val'), 'clear')
                table.rows[r_idx + 1].cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()

# ============================================================
# SECTION 6: USER INTERFACE REQUIREMENTS
# ============================================================
doc.add_heading('6. User Interface Requirements', level=1)
doc.add_paragraph('All screens follow TrustOMS UI patterns: Radix UI components, Tailwind CSS, dark mode support, collapsible sidebar (256px/64px), and responsive design down to 768px tablet width.')

doc.add_heading('6.1 CRM Calendar Page', level=2)
doc.add_paragraph('Path: /crm/calendar | Navigation: CRM > Calendar')
items = [
    'Layout: Full-width content area. Top bar: view toggle buttons (Day|Week|Month|All), date navigation (< Today >), date display, "+ New Meeting" button, "Standalone Report" button.',
    'Month View (default): CSS Grid calendar. 7 columns (Mon-Sun). 6 rows. Day cells: day number top-left, meeting count badge top-right, up to 3 meeting pills (colored by reason). Click day -> Day view. Click meeting pill -> meeting popover.',
    'Week View: 8-column grid (time labels + 7 days). Rows for each hour 7AM-8PM. Meeting blocks positioned absolutely based on time. Color-coded by meeting_reason. Drag to reschedule (with confirmation).',
    'Day View: Single column timeline. Each meeting as a Card component: avatar, subject, time, client, location, mode icon, status badge. Action buttons: File Report (if completed), Reschedule, Cancel. Empty slots have dashed borders.',
    'All View: OpsDataTable component (reuses existing CRUD table). Columns: Date, Subject, Client, Type, Mode, Status, Location. Search bar, status filter chips, date range picker, pagination.',
    'Responsive: On tablet (768px-1024px), month view reduces to 5-day work week. Below 768px, default to Day view with swipe navigation.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 Meeting Create/Edit Modal', level=2)
doc.add_paragraph('Path: Modal overlay on Calendar page | Triggered by: "+ New Meeting" button or meeting edit')
items = [
    'Layout: Sheet (slide-from-right) component, 600px wide. Two-section accordion: "Meeting Details" (expanded by default) and "Participants" (collapsed).',
    'Meeting Details section: Subject (text input), Meeting Reason (select dropdown), Location (text input), Start Date/Time (date+time pickers side-by-side), End Date/Time (date+time pickers), All Day (switch toggle), Meeting Type (radio group: CIF|Lead|Prospect|Others), Mode (select dropdown).',
    'Conditional field: When Type=CIF, show "Client" combobox with search. When Type=OTHERS, show "Contact Name" text input.',
    'Participants section: Required Attendees (multi-select combobox searching users in same hierarchy). Optional Attendees (multi-select combobox searching all users). Selected users shown as removable chips.',
    'Footer: "Cancel" (ghost button) and "Schedule Meeting" (primary button). Form validation on submit. Loading spinner during API call.',
    'Dark mode: All form controls support dark theme. Background: slate-900, borders: slate-700, text: slate-100.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.3 Call Report Form Page', level=2)
doc.add_paragraph('Path: /crm/call-reports/new?meetingId=X or /crm/call-reports/new (standalone) | Navigation: CRM > Call Reports > New')
items = [
    'Layout: Multi-step form (stepper at top). Steps: 1. Meeting Particulars, 2. Meeting Information, 3. Products Discussed, 4. Action Items, 5. Opportunities, 6. Review & Submit.',
    'Step 1 - Meeting Particulars: If scheduled, read-only display of meeting data in a summary card. If standalone, editable form matching meeting create fields.',
    'Step 2 - Meeting Information: Client Status (select), State of Mind (radio group with emoji icons: Happy/Satisfied/Sensitive/Irate/Ultra-Sensitive), Summary of Discussion (rich text area, min 50 chars with char counter), Next Meeting fields (date, start time, end time), Person Met (text if Others).',
    'Step 3 - Products Discussed: Dynamic table with add-row button. Columns: Product Type (select from trust_product_types), Product Name (text), Interest Level (select: High/Medium/Low). Delete row button per row.',
    'Step 4 - Action Items: Dynamic card list. Each card: Title (text), Description (textarea), Assigned To (user search), Due Date (date picker), Priority (select). Add/Remove buttons.',
    'Step 5 - Opportunities: Dynamic card list. Each card: Name (text), Product Type (select), Expected Value (currency input), Probability (slider), Expected Close Date (date). Weighted value displayed.',
    'Step 6 - Review: Read-only summary of all sections. "Save as Draft" (secondary) and "Submit for Review" (primary) buttons. Expense entry link: "Log an expense for this meeting?"',
    'Responsive: Steps collapse to vertical accordion on mobile. Form fields stack single-column below 768px.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.4 Call Report Approval Queue', level=2)
doc.add_paragraph('Path: /crm/call-reports/approval-queue | Navigation: CRM > Approval Queue (visible to BO_HEAD only)')
items = [
    'Layout: Card list (default) or table view (toggle). Filter bar: status chips (Submitted, Under Review), late-filed toggle, RM filter (combobox), date range.',
    'Card layout: Each card shows: report_ref, RM avatar+name, client name, submission date, SLA indicator (green/yellow/red circle), late-filed badge (amber), state_of_mind emoji.',
    'Expanded card: Full call report content (scrollable). Bottom action bar: Quality Score (1-5 stars), "Approve" (green), "Request Info" (yellow), "Reject" (red with reason modal).',
    'SLA indicator: Green = <24h since submission, Yellow = 24-48h, Red = >48h.',
    'Table view: Columns: Ref, RM, Client, Submitted Date, SLA, Late Filed, Status, Actions. Sort by any column.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.5 Call Report List Page', level=2)
doc.add_paragraph('Path: /crm/call-reports | Navigation: CRM > Call Reports')
items = [
    'Layout: OpsDataTable with tabs: My Reports | Team Reports (SENIOR_RM/BO_HEAD) | All Reports (BO_HEAD/COMPLIANCE).',
    'Columns: Report Ref, Type (Scheduled/Standalone badge), Date, Client, Subject, Status badge (color-coded), Quality Score (stars), Filed By.',
    'Filters: Status (multi-select chips), Type, Date Range, Client search, RM (for team/all tabs).',
    'Row click: navigates to /crm/call-reports/:id detail page.',
    'Actions: Export CSV (BO_HEAD+), "+ New Standalone" button, "+ Scheduled" button.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.6 Opportunity Pipeline Page', level=2)
doc.add_paragraph('Path: /crm/opportunities | Navigation: CRM > Opportunities')
items = [
    'Layout: Toggle between Kanban board and Table view. Summary bar at top: Total Pipeline Value, Weighted Value, Win Rate, Average Deal Size.',
    'Kanban: 7 columns (Open, Qualified, Proposal, Negotiation, Won, Lost, Deferred). Cards show: name, client, value (formatted currency), probability bar, close date, RM avatar.',
    'Table: Columns: Ref, Name, Client, Product, Value, Probability, Close Date, Status, Owner. Standard OpsDataTable with search, filters, pagination.',
    'Filters: Owner RM, Branch, Product Type, Value Range (min/max sliders), Date Range.',
    'Detail panel (slide-out sheet): Full opportunity detail with edit capability, linked call report reference, status history timeline.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.7 Task Manager Page', level=2)
doc.add_paragraph('Path: /crm/tasks | Navigation: CRM > Tasks')
items = [
    'Layout: Three-column kanban: To Do | In Progress | Completed. Filter bar: priority filter, due date range, entity type link filter.',
    'Task cards: Title, due date (with overdue badge), priority badge (color: red=urgent, orange=high, blue=medium, grey=low), linked entity chip, assigned by (for ASSIGNED type).',
    'Drag-and-drop between columns updates task_status.',
    'Quick add: "+" button at top of To Do column opens inline form (title + due date only, expand for full form).',
    'Team view (SENIOR_RM/BO_HEAD): Shows all team tasks grouped by assignee with activity summary.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.8 Campaign Management Page', level=2)
doc.add_paragraph('Path: /crm/campaigns | Navigation: CRM > Campaigns (BO_HEAD+)')
items = [
    'Layout: Campaign list as cards. Each card: name, code, type badge, date range, status badge, progress bar (contacted/total), budget display.',
    'Campaign detail page (/crm/campaigns/:id): Tabs - Overview | Lead Lists | Responses | Performance.',
    'Overview tab: Campaign metadata, material downloads, edit button.',
    'Lead Lists tab: List of linked lead lists with member count, creation method badge. Create new list button.',
    'Responses tab: Table of leads with response status, contact info, assigned RM, response date. Inline response capture.',
    'Performance tab: Dashboard with metrics cards, response pie chart, timeline chart, RM leaderboard.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.9 CRM Dashboard (RM Home)', level=2)
doc.add_paragraph('Path: /crm/dashboard | Navigation: CRM > Dashboard (default CRM landing page)')
items = [
    'Layout: 4-quadrant grid. Top-left: Today\'s Meetings (compact card list). Top-right: Pending Action Items (list with overdue count). Bottom-left: Pipeline Summary (mini funnel chart). Bottom-right: Recent Call Reports (last 5 with status).',
    'KPI bar at top: Meetings This Week, Reports Filed This Week, Pending Approvals, Overdue Items.',
    'Quick actions: "+ Meeting", "+ Standalone Report", "+ Task" buttons in top-right.',
    'For BO_HEAD: Additional "Team Overview" section showing RM activity grid (heatmap of meetings/reports per RM per day for past 2 weeks).',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.10 Navigation Structure', level=2)
doc.add_paragraph('New sidebar section "CRM" added to the existing back-office navigation:')
add_table(doc,
    ['Menu Item', 'Path', 'Icon', 'Visible To'],
    [
        ['CRM Dashboard', '/crm/dashboard', 'LayoutDashboard', 'RM, SENIOR_RM, BO_HEAD, TRUST_BIZ_HEAD'],
        ['Calendar', '/crm/calendar', 'Calendar', 'RM, SENIOR_RM, BO_HEAD, BRANCH_ASSOC'],
        ['Call Reports', '/crm/call-reports', 'FileText', 'RM, SENIOR_RM, BO_HEAD, COMPLIANCE, CCO'],
        ['Approval Queue', '/crm/call-reports/approval-queue', 'CheckSquare', 'BO_HEAD'],
        ['Opportunities', '/crm/opportunities', 'TrendingUp', 'RM, SENIOR_RM, BO_HEAD, TRUST_BIZ_HEAD'],
        ['Tasks', '/crm/tasks', 'ListChecks', 'RM, SENIOR_RM, BO_HEAD'],
        ['Expenses', '/crm/expenses', 'Receipt', 'RM, SENIOR_RM, BO_HEAD'],
        ['Campaigns', '/crm/campaigns', 'Megaphone', 'BO_HEAD, TRUST_BIZ_HEAD, SYSTEM_ADMIN'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 7: API & INTEGRATION REQUIREMENTS
# ============================================================
doc.add_heading('7. API & Integration Requirements', level=1)

doc.add_heading('7.1 Authentication & Authorization', level=2)
doc.add_paragraph('All CIM API endpoints use the existing TrustOMS authentication system: JWT tokens stored in httpOnly cookies. Authorization uses role-based middleware guards. A new CRM-specific guard will be introduced:')
items = [
    'requireCRMRole(): Allows RELATIONSHIP_MANAGER, SENIOR_RM, BO_HEAD, BRANCH_ASSOCIATE, TRUST_BUSINESS_HEAD',
    'requireCRMSupervisor(): Allows BO_HEAD, SENIOR_RM (for team-scoped operations)',
    'requireCRMAdmin(): Allows BO_HEAD, SYSTEM_ADMIN (for campaign and configuration management)',
    'Data scoping middleware: RM sees own data, SENIOR_RM sees team data, BO_HEAD sees branch data, TRUST_BIZ_HEAD sees all data',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7.2 Standardized Error Response', level=2)
doc.add_paragraph('All CIM endpoints use the existing TrustOMS error response format:')
doc.add_paragraph('{\n  "error": {\n    "code": "VALIDATION_ERROR",\n    "message": "Summary of discussion must be at least 50 characters",\n    "field": "summary_of_discussion",\n    "correlation_id": "req-uuid-here"\n  }\n}')
doc.add_paragraph('Error codes: VALIDATION_ERROR, NOT_FOUND, FORBIDDEN, CONFLICT (duplicate), PRECONDITION_FAILED (optimistic lock), RATE_LIMITED, INTERNAL_ERROR')

doc.add_heading('7.3 API Endpoint Catalog', level=2)

add_table(doc,
    ['Method', 'Path', 'Auth Guard', 'Description'],
    [
        ['GET', '/api/v1/meetings', 'requireCRMRole()', 'List meetings (paginated, filtered by user scope)'],
        ['GET', '/api/v1/meetings/:id', 'requireCRMRole()', 'Get meeting detail'],
        ['POST', '/api/v1/meetings', 'requireCRMRole()', 'Create meeting'],
        ['PUT', '/api/v1/meetings/:id', 'requireCRMRole()', 'Update meeting'],
        ['POST', '/api/v1/meetings/:id/reschedule', 'requireCRMRole()', 'Reschedule meeting (creates new, marks old)'],
        ['POST', '/api/v1/meetings/:id/cancel', 'requireCRMRole()', 'Cancel meeting'],
        ['DELETE', '/api/v1/meetings/:id', 'requireCRMRole()', 'Soft delete meeting'],
        ['GET', '/api/v1/meetings/:parentId/attendees', 'requireCRMRole()', 'List meeting attendees'],
        ['POST', '/api/v1/meetings/:parentId/attendees', 'requireCRMRole()', 'Add attendee'],
        ['DELETE', '/api/v1/meetings/:parentId/attendees/:id', 'requireCRMRole()', 'Remove attendee'],
        ['GET', '/api/v1/call-reports', 'requireCRMRole()', 'List call reports (scoped)'],
        ['GET', '/api/v1/call-reports/:id', 'requireCRMRole()', 'Get call report detail'],
        ['POST', '/api/v1/call-reports', 'requireCRMRole()', 'Create call report'],
        ['PUT', '/api/v1/call-reports/:id', 'requireCRMRole()', 'Update call report (draft only)'],
        ['POST', '/api/v1/call-reports/:id/submit', 'requireCRMRole()', 'Submit for review'],
        ['POST', '/api/v1/call-reports/:id/approve', 'requireCRMSupervisor()', 'Approve call report'],
        ['POST', '/api/v1/call-reports/:id/reject', 'requireCRMSupervisor()', 'Reject call report'],
        ['GET', '/api/v1/call-reports/:parentId/feedback', 'requireCRMRole()', 'List feedback for report'],
        ['POST', '/api/v1/call-reports/:parentId/feedback', 'requireCRMSupervisor()', 'Add feedback'],
        ['GET', '/api/v1/call-reports/:parentId/action-items', 'requireCRMRole()', 'List action items for report'],
        ['POST', '/api/v1/call-reports/:parentId/action-items', 'requireCRMRole()', 'Create action item'],
        ['PUT', '/api/v1/action-items/:id', 'requireCRMRole()', 'Update action item status'],
        ['GET', '/api/v1/action-items', 'requireCRMRole()', 'List all action items (scoped)'],
        ['GET', '/api/v1/opportunities', 'requireCRMRole()', 'List opportunities (scoped)'],
        ['POST', '/api/v1/opportunities', 'requireCRMRole()', 'Create opportunity'],
        ['PUT', '/api/v1/opportunities/:id', 'requireCRMRole()', 'Update opportunity'],
        ['POST', '/api/v1/opportunities/bulk-upload', 'requireCRMRole()', 'Bulk CSV upload'],
        ['GET', '/api/v1/tasks', 'requireCRMRole()', 'List tasks (scoped)'],
        ['POST', '/api/v1/tasks', 'requireCRMRole()', 'Create task'],
        ['PUT', '/api/v1/tasks/:id', 'requireCRMRole()', 'Update task'],
        ['GET', '/api/v1/expenses', 'requireCRMRole()', 'List expenses (scoped)'],
        ['POST', '/api/v1/expenses', 'requireCRMRole()', 'Create expense'],
        ['PUT', '/api/v1/expenses/:id', 'requireCRMRole()', 'Update expense'],
        ['POST', '/api/v1/expenses/:id/approve', 'requireCRMSupervisor()', 'Approve expense'],
        ['GET', '/api/v1/campaigns', 'requireCRMAdmin()', 'List campaigns'],
        ['POST', '/api/v1/campaigns', 'requireCRMAdmin()', 'Create campaign'],
        ['PUT', '/api/v1/campaigns/:id', 'requireCRMAdmin()', 'Update campaign'],
        ['GET', '/api/v1/campaigns/:parentId/lead-lists', 'requireCRMAdmin()', 'List lead lists'],
        ['POST', '/api/v1/campaigns/:parentId/lead-lists', 'requireCRMAdmin()', 'Create lead list'],
        ['GET', '/api/v1/lead-lists/:parentId/members', 'requireCRMRole()', 'List lead list members'],
        ['PUT', '/api/v1/lead-list-members/:id/response', 'requireCRMRole()', 'Capture lead response'],
        ['GET', '/api/v1/crm/dashboard', 'requireCRMRole()', 'CRM dashboard aggregated data'],
        ['GET', '/api/v1/crm/reports/rm-productivity', 'requireCRMSupervisor()', 'RM productivity report'],
        ['GET', '/api/v1/crm/reports/campaign-performance', 'requireCRMAdmin()', 'Campaign performance report'],
    ]
)

doc.add_heading('7.4 Request/Response Examples', level=2)

doc.add_heading('7.4.1 POST /api/v1/meetings (Create Meeting)', level=3)
doc.add_paragraph('Request Body:')
doc.add_paragraph('''{
  "subject": "Q2 Portfolio Review - Santos Family Trust",
  "meeting_reason": "PORTFOLIO_REVIEW",
  "location": "BGC Office - Meeting Room 3A",
  "start_date": "2026-04-25T10:00:00+08:00",
  "end_date": "2026-04-25T11:30:00+08:00",
  "is_all_day": false,
  "meeting_type": "CIF",
  "mode_of_meeting": "FACE_TO_FACE",
  "client_id": "CL-00042",
  "remarks": "Bring Q1 performance report and IMA expansion proposal",
  "reminder_minutes": 30,
  "attendees": [
    { "user_id": "USR-RM-003", "attendance_type": "REQUIRED" },
    { "user_id": "USR-SR-001", "attendance_type": "OPTIONAL" }
  ]
}''')
doc.add_paragraph('Success Response (201):')
doc.add_paragraph('''{
  "data": {
    "id": 1,
    "meeting_ref": "MTG-20260422-0001",
    "subject": "Q2 Portfolio Review - Santos Family Trust",
    "meeting_status": "SCHEDULED",
    "organizer_id": "USR-RM-001",
    "created_at": "2026-04-22T09:15:00+08:00"
  }
}''')

doc.add_heading('7.4.2 POST /api/v1/call-reports (Create Call Report)', level=3)
doc.add_paragraph('Request Body:')
doc.add_paragraph('''{
  "report_type": "SCHEDULED",
  "meeting_id": 1,
  "client_relationship_status": "ACTIVE_ENGAGED",
  "state_of_mind": "SATISFIED",
  "summary_of_discussion": "Reviewed Q1 portfolio performance (+8.2% vs benchmark +6.1%). Client expressed satisfaction with equity allocation but requested exploring fixed-income opportunities for risk diversification. Discussed IMA expansion to include peso bond allocation of PHP 20M. Client agreed to review proposal by end of month.",
  "products_discussed": [
    { "product_type": "IMA_DIRECTED", "product_name": "Peso Bond Fund", "interest_level": "HIGH" },
    { "product_type": "UITF", "product_name": "Growth Equity Fund", "interest_level": "MEDIUM" }
  ],
  "next_meeting_date": "2026-05-10T14:00:00+08:00",
  "next_meeting_time_start": "14:00",
  "next_meeting_time_end": "15:00",
  "action_items": [
    {
      "title": "Prepare peso bond allocation proposal",
      "description": "Include 3 options: conservative, moderate, aggressive bond allocations",
      "assigned_to": "USR-RM-001",
      "due_date": "2026-04-30",
      "priority": "HIGH"
    }
  ],
  "opportunities": [
    {
      "name": "Santos Trust - IMA Expansion (Bonds)",
      "product_type": "IMA_DIRECTED",
      "expected_value": 20000000,
      "currency": "PHP",
      "probability": 75,
      "expected_close_date": "2026-05-31"
    }
  ],
  "attendees": [
    { "user_id": "USR-RM-003", "role_in_meeting": "Co-RM" },
    { "external_name": "Maria Santos", "external_organization": "Santos Family Office", "role_in_meeting": "Client PoA" }
  ]
}''')
doc.add_paragraph('Success Response (201):')
doc.add_paragraph('''{
  "data": {
    "id": 1,
    "report_ref": "CR-20260425-0001",
    "report_type": "SCHEDULED",
    "report_status": "DRAFT",
    "filed_by": "USR-RM-001",
    "is_late_filed": false,
    "created_at": "2026-04-25T12:30:00+08:00",
    "action_items_created": 1,
    "opportunities_created": 1,
    "next_meeting_created": { "meeting_ref": "MTG-20260510-0001" }
  }
}''')

doc.add_heading('7.4.3 GET /api/v1/call-reports?status=SUBMITTED&page=1 (List for Approval)', level=3)
doc.add_paragraph('Response (200):')
doc.add_paragraph('''{
  "data": [
    {
      "id": 1,
      "report_ref": "CR-20260425-0001",
      "report_type": "SCHEDULED",
      "subject": "Q2 Portfolio Review - Santos Family Trust",
      "client_name": "Santos Family Trust",
      "meeting_date": "2026-04-25T10:00:00+08:00",
      "state_of_mind": "SATISFIED",
      "report_status": "SUBMITTED",
      "submitted_at": "2026-04-25T12:35:00+08:00",
      "is_late_filed": false,
      "filed_by": "USR-RM-001",
      "filed_by_name": "Juan Cruz",
      "quality_score": null
    }
  ],
  "total": 12,
  "page": 1,
  "pageSize": 25
}''')

doc.add_heading('7.5 Integration with Existing TrustOMS Modules', level=2)
add_table(doc,
    ['Integration Point', 'Direction', 'Description'],
    [
        ['Client 360 Dashboard', 'CIM -> Client 360', 'New "Interactions" tab on client detail page showing call report timeline, upcoming meetings, and open action items for the client'],
        ['clients table', 'Client 360 -> CIM', 'CIM reads client_id, legal_name, contact info, risk_profile, client_status from existing clients table via FK'],
        ['users table', 'Auth -> CIM', 'CIM reads user_id, name, role, branch from existing users table for attendees, organizers, assignees'],
        ['branches table', 'Reference -> CIM', 'CIM reads branch_id, branch_name for filtering and scoping'],
        ['market_calendar table', 'Reference -> CIM', 'Calendar view shows holidays from market_calendar; meeting scheduling warns if date is a holiday'],
        ['trust_product_types table', 'Reference -> CIM', 'Products Discussed and Opportunity product_type use trust_product_types for dropdown options'],
        ['notification system', 'CIM -> Notifications', 'CIM dispatches events to existing notificationLog table via notificationService.dispatch()'],
        ['audit system', 'CIM -> Audit', 'All CIM mutations logged via existing hash-chained auditEvents / auditRecords tables'],
        ['entity_registry', 'CIM -> Config', 'CIM entities (meetings, call-reports, etc.) registered in entity_registry for dynamic field configuration'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 8: NON-FUNCTIONAL REQUIREMENTS
# ============================================================
doc.add_heading('8. Non-Functional Requirements', level=1)

doc.add_heading('8.1 Performance', level=2)
add_table(doc,
    ['Metric', 'Target', 'Measurement'],
    [
        ['Calendar month view load', '< 2 seconds for 100 meetings', 'Time from page load to last meeting pill rendered'],
        ['Call report form load', '< 1.5 seconds including pre-population', 'Time from navigation to form interactive'],
        ['Call report list (10K records)', '< 3 seconds with pagination', 'Time for /api/v1/call-reports?page=1 response'],
        ['Opportunity kanban load', '< 2 seconds for 500 opportunities', 'Time from page load to all cards rendered'],
        ['Search/filter response', '< 500ms for ILIKE search', 'API response time for search queries'],
        ['Concurrent users', '500 RMs + 50 supervisors simultaneously', 'Load test: 550 concurrent sessions without degradation'],
        ['Meeting creation throughput', '50 meetings/minute', 'Sustained creation rate during peak (quarter-end)'],
    ]
)

doc.add_heading('8.2 Security', level=2)
items = [
    'Authentication: JWT tokens in httpOnly secure cookies (existing TrustOMS pattern). Token expiry: 8 hours. Refresh token: 24 hours.',
    'Authorization: Role-based middleware guards on all endpoints. Data scoping: RMs see own data only; supervisors see team; executives see all.',
    'PII Classification: summary_of_discussion = Sensitive-PII. client contact fields = PII. All PII access logged via logDataAccess() middleware.',
    'Data Encryption: At rest via Supabase/PostgreSQL TDE. In transit via TLS 1.3.',
    'OWASP: Input validation (Zod schemas), parameterized queries (Drizzle ORM), CSRF protection (double-submit cookie), XSS prevention (React auto-escaping + CSP headers).',
    'Rate Limiting: 100 requests/minute per user for mutation endpoints. 300 requests/minute for read endpoints.',
    'Audit Trail: Hash-chained audit logging for all create/update/delete operations. Immutable once written.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8.3 Scalability', level=2)
items = [
    'Expected growth: 20% YoY increase in RMs (from 500 to 600 in year 2). Proportional increase in meetings and call reports.',
    'Database: PostgreSQL with proper indexes on frequently queried columns (organizer_id, client_id, start_date, report_status, filed_by). Partitioning on meetings and call_reports by year if table exceeds 1M rows.',
    'Caching: React Query client-side cache (5 min stale time for list queries). No server-side cache for v1.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8.4 Availability & Backup', level=2)
items = [
    'Uptime target: 99.5% (allows ~43 hours downtime per year for maintenance).',
    'Backup: Supabase automated daily backups. Point-in-time recovery enabled (RPO: 1 hour, RTO: 4 hours).',
    'Disaster recovery: Supabase handles infrastructure-level DR. Application is stateless (can be redeployed in minutes).',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8.5 Accessibility', level=2)
items = [
    'WCAG 2.1 AA compliance required for all CIM pages.',
    'Keyboard navigation: All interactive elements focusable via Tab. Calendar day/week views navigable via arrow keys.',
    'Screen reader: ARIA labels on all form controls, calendar cells, and status badges.',
    'Color contrast: Minimum 4.5:1 ratio for body text, 3:1 for large text. Status colors supplemented with icons/text labels.',
    'Dark mode: Full support via Tailwind dark: variants. Tested in both light and dark themes.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('8.6 Browser & Device Support', level=2)
add_table(doc,
    ['Platform', 'Supported Versions', 'Testing Priority'],
    [
        ['Chrome', '90+', 'Primary'],
        ['Firefox', '90+', 'Secondary'],
        ['Safari', '15+', 'Secondary'],
        ['Edge', '90+', 'Secondary'],
        ['iPad/Tablet', '768px+ responsive', 'Primary (RM field use)'],
        ['Mobile phone', '375px+ responsive', 'Read-only views only'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 9: WORKFLOW & STATE DIAGRAMS
# ============================================================
doc.add_heading('9. Workflow & State Diagrams', level=1)

doc.add_heading('9.1 Meeting Lifecycle', level=2)
add_table(doc,
    ['Current State', 'Action', 'Next State', 'Triggered By', 'Side Effects'],
    [
        ['(new)', 'Create meeting', 'SCHEDULED', 'RM', 'Attendee notifications sent; reminders scheduled; audit log'],
        ['SCHEDULED', 'Meeting end_date passes', 'COMPLETED', 'System (cron)', 'Enables "File Call Report" button; task created if no report filed in 48h'],
        ['SCHEDULED', 'Reschedule', 'RESCHEDULED', 'Organizer', 'New meeting created; attendees notified of new date; audit log'],
        ['SCHEDULED', 'Cancel', 'CANCELLED', 'Organizer/Supervisor', 'Attendees notified; reminders cancelled; reason recorded; audit log'],
        ['COMPLETED', 'Call report filed', 'COMPLETED', 'RM', 'Meeting linked to call_report via meeting_id; no status change'],
        ['RESCHEDULED', 'No further action', 'RESCHEDULED (terminal)', 'N/A', 'Remains in history; new meeting continues the lifecycle'],
        ['CANCELLED', 'No further action', 'CANCELLED (terminal)', 'N/A', 'Remains in history for audit'],
    ]
)

doc.add_heading('9.2 Call Report Lifecycle', level=2)
add_table(doc,
    ['Current State', 'Action', 'Next State', 'Triggered By', 'Side Effects'],
    [
        ['(new)', 'Create / Save as Draft', 'DRAFT', 'RM', 'Report created; action items/opportunities NOT yet created (only on submit)'],
        ['DRAFT', 'Submit for Review', 'SUBMITTED', 'RM', 'Action items created; opportunities created; next meeting created; notification to supervisor; audit log; is_late_filed computed'],
        ['DRAFT', 'Edit', 'DRAFT', 'RM', 'Fields updated; no notifications'],
        ['SUBMITTED', 'Supervisor opens review', 'UNDER_REVIEW', 'Supervisor', 'Notification to RM: "Your report is being reviewed"'],
        ['SUBMITTED', 'Approve directly', 'APPROVED', 'Supervisor', 'Report locked; quality_score set; notification to RM; audit log'],
        ['UNDER_REVIEW', 'Approve', 'APPROVED', 'Supervisor', 'Report locked; quality_score set; notification to RM; audit log'],
        ['UNDER_REVIEW', 'Request info', 'UNDER_REVIEW', 'Supervisor', 'Notification to RM with question; feedback entry created'],
        ['SUBMITTED', 'Reject', 'REJECTED', 'Supervisor', 'rejection_reason recorded; notification to RM; report unlocked for editing'],
        ['UNDER_REVIEW', 'Reject', 'REJECTED', 'Supervisor', 'Same as above'],
        ['REJECTED', 'RM edits and resubmits', 'SUBMITTED', 'RM', 'Resets review cycle; supervisor re-notified'],
        ['APPROVED', 'Feedback added', 'APPROVED', 'Supervisor', 'No status change; feedback entry created; notification to RM'],
    ]
)

doc.add_heading('9.3 Action Item Lifecycle', level=2)
add_table(doc,
    ['Current State', 'Action', 'Next State', 'Triggered By', 'Side Effects'],
    [
        ['(new)', 'Created from call report submit', 'OPEN', 'System', 'Notification to assignee; appears in task list'],
        ['OPEN', 'Start working', 'IN_PROGRESS', 'Assignee', 'Audit log'],
        ['OPEN', 'Due date passes', 'OVERDUE', 'System (cron)', 'Notification to assignee + supervisor; visual badge change'],
        ['IN_PROGRESS', 'Complete', 'COMPLETED', 'Assignee', 'completion_notes required; completed_at set; notification to reporter; audit log'],
        ['IN_PROGRESS', 'Due date passes', 'OVERDUE', 'System (cron)', 'Same as OPEN -> OVERDUE'],
        ['OVERDUE', 'Complete', 'COMPLETED', 'Assignee', 'Same as IN_PROGRESS -> COMPLETED'],
        ['OPEN/IN_PROGRESS', 'Cancel', 'CANCELLED', 'Assignee/Supervisor', 'Reason required; audit log'],
    ]
)

doc.add_heading('9.4 Opportunity Lifecycle', level=2)
add_table(doc,
    ['Current State', 'Action', 'Next State', 'Triggered By', 'Side Effects'],
    [
        ['(new)', 'Created from call report or bulk upload', 'OPEN', 'RM/System', 'Appears on pipeline board'],
        ['OPEN', 'Qualify lead', 'QUALIFIED', 'RM', 'Audit log'],
        ['QUALIFIED', 'Send proposal', 'PROPOSAL', 'RM', 'Audit log'],
        ['PROPOSAL', 'Enter negotiation', 'NEGOTIATION', 'RM', 'Audit log'],
        ['NEGOTIATION', 'Win deal', 'WON', 'RM', 'won_value + won_date required; notification to supervisor; audit log'],
        ['NEGOTIATION', 'Lose deal', 'LOST', 'RM', 'loss_reason required; audit log'],
        ['OPEN/QUALIFIED/PROPOSAL', 'Defer', 'DEFERRED', 'RM', 'Audit log; can be reopened later'],
        ['DEFERRED', 'Reopen', 'OPEN', 'RM', 'Audit log'],
        ['Any except WON/LOST', 'Lose deal', 'LOST', 'RM', 'loss_reason required'],
    ]
)

doc.add_heading('9.5 Expense Lifecycle', level=2)
add_table(doc,
    ['Current State', 'Action', 'Next State', 'Triggered By', 'Side Effects'],
    [
        ['(new)', 'Create', 'DRAFT', 'RM', 'Expense saved'],
        ['DRAFT', 'Submit', 'SUBMITTED', 'RM', 'Notification to supervisor; appears in approval queue'],
        ['SUBMITTED', 'Approve', 'APPROVED', 'Supervisor', 'Notification to RM; audit log'],
        ['SUBMITTED', 'Reject', 'REJECTED', 'Supervisor', 'rejection_reason; notification to RM; RM can edit and resubmit'],
        ['REJECTED', 'Resubmit', 'SUBMITTED', 'RM', 'Resets approval cycle'],
    ]
)

doc.add_heading('9.6 Campaign Lifecycle', level=2)
add_table(doc,
    ['Current State', 'Action', 'Next State', 'Triggered By', 'Side Effects'],
    [
        ['(new)', 'Create', 'DRAFT', 'Supervisor', 'Campaign created; no lead outreach yet'],
        ['DRAFT', 'Activate', 'ACTIVE', 'Supervisor', 'Requires at least 1 lead list; lead assignments visible to RMs'],
        ['ACTIVE', 'Pause', 'PAUSED', 'Supervisor', 'Leads remain assigned but hidden from RM response capture'],
        ['PAUSED', 'Resume', 'ACTIVE', 'Supervisor', 'Leads visible again'],
        ['ACTIVE', 'End date passes', 'COMPLETED', 'System', 'Final metrics computed; no further responses accepted'],
        ['DRAFT/ACTIVE/PAUSED', 'Cancel', 'CANCELLED', 'Supervisor', 'All lead assignments removed; audit log'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 10: NOTIFICATION & COMMUNICATION REQUIREMENTS
# ============================================================
doc.add_heading('10. Notification & Communication Requirements', level=1)

doc.add_paragraph('All notifications use the existing TrustOMS notificationService.dispatch() system. Channels: IN_APP (always) and EMAIL (configurable per notification type).')

doc.add_heading('10.1 Notification Catalog', level=2)
add_table(doc,
    ['Event', 'Channel', 'Recipient(s)', 'Trigger', 'Message Template'],
    [
        ['Meeting Invitation', 'IN_APP + EMAIL', 'All attendees', 'Meeting created', 'You have been invited to "[subject]" on [date] at [time] by [organizer_name].'],
        ['Meeting Reschedule', 'IN_APP + EMAIL', 'All attendees', 'Meeting rescheduled', 'Meeting "[subject]" has been rescheduled from [old_date] to [new_date] by [organizer_name].'],
        ['Meeting Cancellation', 'IN_APP + EMAIL', 'All attendees', 'Meeting cancelled', 'Meeting "[subject]" on [date] has been cancelled by [organizer_name]. Reason: [cancel_reason].'],
        ['Meeting Reminder', 'IN_APP + EMAIL', 'Organizer + attendees', 'reminder_minutes before start', 'Reminder: "[subject]" starts in [N] minutes at [location].'],
        ['Call Report Submitted', 'IN_APP', 'BO_HEAD (supervisor)', 'RM submits report', 'New call report [report_ref] submitted by [rm_name] for review.'],
        ['Call Report Approved', 'IN_APP', 'Filing RM', 'Supervisor approves', 'Your call report [report_ref] has been approved by [supervisor_name]. Quality score: [score]/5.'],
        ['Call Report Rejected', 'IN_APP + EMAIL', 'Filing RM', 'Supervisor rejects', 'Your call report [report_ref] requires revision. Reason: [rejection_reason].'],
        ['Call Report Feedback', 'IN_APP', 'Filing RM', 'Supervisor adds feedback', 'New [feedback_type] feedback on call report [report_ref] from [supervisor_name].'],
        ['Call Report Overdue', 'IN_APP + EMAIL', 'RM + BO_HEAD', '48h after completed meeting, no report filed', 'Meeting "[subject]" on [date] has no call report filed. Please submit within [N] hours.'],
        ['Review SLA Warning', 'IN_APP', 'BO_HEAD', 'Report SUBMITTED >48h without review', 'Call report [report_ref] has been pending review for [hours] hours. SLA at risk.'],
        ['Action Item Assigned', 'IN_APP', 'Assignee', 'Action item created', 'New action item "[title]" assigned by [assigner_name]. Due: [due_date].'],
        ['Action Item Overdue', 'IN_APP + EMAIL', 'Assignee + supervisor', 'Due date passed', 'Action item "[title]" is overdue. Was due on [due_date].'],
        ['Task Assigned', 'IN_APP', 'Assignee', 'Supervisor assigns task', 'New task "[title]" assigned by [supervisor_name]. Due: [due_date].'],
        ['Task Reminder', 'IN_APP', 'Assignee', 'reminder_date reached', 'Reminder: Task "[title]" is due on [due_date].'],
        ['Expense Approved', 'IN_APP', 'RM', 'Supervisor approves', 'Your expense [expense_ref] for [amount] [currency] has been approved.'],
        ['Expense Rejected', 'IN_APP', 'RM', 'Supervisor rejects', 'Your expense [expense_ref] was rejected. Reason: [rejection_reason].'],
        ['Opportunity Won', 'IN_APP', 'RM + Supervisor', 'Opportunity status -> WON', 'Opportunity "[name]" won! Value: [won_value] [currency].'],
        ['Campaign Activated', 'IN_APP', 'Assigned RMs', 'Campaign status -> ACTIVE', 'Campaign "[name]" is now active. You have [N] leads to contact.'],
    ]
)

doc.add_heading('10.2 Notification Preferences', level=2)
doc.add_paragraph(
    'Users cannot opt out of critical notifications (meeting invitations, call report rejection, action item overdue). '
    'Users can configure EMAIL delivery preference for non-critical notifications via a settings page. '
    'Default: all notifications sent via IN_APP; EMAIL enabled only for: meeting invitations, meeting reminders, '
    'call report rejection, and action item overdue.'
)

doc.add_page_break()
doc.save('/Users/n15318/Trust OMS/docs/CIM_BRD_Draft.docx')
print("Part 4 complete: Sections 6-10 saved")
