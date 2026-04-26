#!/usr/bin/env python3
"""
Generate Calendar & Call Report Management BRD as a .docx file.
Uses python-docx to create a professionally formatted Business Requirements Document.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Color constants ──────────────────────────────────────────────────
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
LIGHT_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
HEADER_BG = "1B3A5C"
SUBHEADER_BG = "D6E4F0"

doc = Document()

# ── Styles ───────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size, color in [('Heading 1', 18, DARK_BLUE), ('Heading 2', 14, DARK_BLUE), ('Heading 3', 12, DARK_BLUE)]:
    s = doc.styles[level]
    s.font.name = 'Arial'
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

# ── Helper functions ─────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        set_cell_shading(cell, HEADER_BG)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")
    doc.add_paragraph()
    return table

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    return p

# ── Page setup ───────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

# ── Header & Footer ─────────────────────────────────────────────────
header = section.header
hp = header.paragraphs[0]
hp.text = "Trust OMS — Calendar & Call Report Management BRD\tConfidential"
hp.style.font.size = Pt(8)
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.style.font.size = Pt(8)

# ── Title Page ───────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tp.add_run("BUSINESS REQUIREMENTS DOCUMENT")
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE
run.bold = True

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = tp2.add_run("Calendar & Call Report Management")
run2.font.size = Pt(22)
run2.font.color.rgb = DARK_BLUE

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = tp3.add_run("Trust OMS Platform")
run3.font.size = Pt(16)
run3.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

tp4 = doc.add_paragraph()
tp4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = tp4.add_run(f"Version 1.0  |  {datetime.date.today().strftime('%B %d, %Y')}")
run4.font.size = Pt(12)
run4.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

tp5 = doc.add_paragraph()
tp5.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = tp5.add_run("CONFIDENTIAL")
run5.font.size = Pt(14)
run5.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
run5.bold = True

doc.add_page_break()

# ── Document Control ─────────────────────────────────────────────────
doc.add_heading("Document Control", level=1)
add_styled_table(doc,
    ["Version", "Date", "Author", "Changes"],
    [
        ["1.0", datetime.date.today().strftime("%Y-%m-%d"), "Trust OMS Team", "Initial BRD — synthesized from SBI, HSBC, Jio BlackRock reference implementations"],
    ])

doc.add_heading("Approvals", level=2)
add_styled_table(doc,
    ["Name", "Role", "Signature", "Date"],
    [
        ["", "Product Owner", "", ""],
        ["", "Technical Architect", "", ""],
        ["", "QA Lead", "", ""],
    ])

doc.add_page_break()

# ── TABLE OF CONTENTS placeholder ────────────────────────────────────
doc.add_heading("Table of Contents", level=1)
doc.add_paragraph("[Auto-generated TOC — update field after opening in Word]")
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

doc.add_heading("1.1 Project Name", level=2)
doc.add_paragraph("Trust OMS — Calendar & Call Report Management Module")

doc.add_heading("1.2 Project Description", level=2)
doc.add_paragraph(
    "This module provides Relationship Managers (RMs) in wealth management and private banking "
    "with a unified platform to schedule and track client meetings via a multi-view calendar, "
    "file detailed call reports (both for scheduled meetings and standalone/ad-hoc interactions), "
    "capture sales opportunities and expenses linked to meetings, manage feedback, and maintain "
    "a chronological conversation history per client. A configurable approval workflow ensures "
    "call reports filed after a threshold period (default: 5 business days) are routed to the "
    "RM's supervisor for review before finalization. The module integrates with the existing "
    "Trust OMS CRM, portfolio, and compliance subsystems."
)

doc.add_heading("1.3 Business Objectives", level=2)
objectives = [
    "Centralize all RM-client interaction tracking in a single calendar-driven interface, replacing fragmented spreadsheets and email-based logs.",
    "Enforce timely call-report filing through the 5-day threshold rule with mandatory supervisor approval for late submissions.",
    "Capture sales opportunities at the point of client conversation, increasing pipeline visibility and conversion tracking.",
    "Provide supervisors with real-time dashboards showing RM activity, meeting coverage, opportunity pipeline, and compliance adherence.",
    "Generate auditable conversation histories per client relationship for regulatory reporting (MiFID II, RBI KYC, MAS TBO).",
]
for o in objectives:
    add_bullet(doc, o)

doc.add_heading("1.4 Target Users & Pain Points", level=2)
add_styled_table(doc,
    ["User Role", "Pain Point Addressed"],
    [
        ["Relationship Manager (RM)", "Manual meeting tracking, no unified view of upcoming/past meetings; call reports filed in disconnected systems; missed opportunity capture."],
        ["RM Supervisor / Branch Manager", "No visibility into RM meeting activity; cannot enforce call-report discipline; late or missing call reports go undetected."],
        ["Compliance Officer", "Incomplete audit trail of client interactions; difficulty proving suitability conversations occurred."],
        ["Operations / MIS", "Manual extraction of meeting/call-report data for management reporting."],
    ])

doc.add_heading("1.5 Success Metrics (KPIs)", level=2)
add_styled_table(doc,
    ["KPI", "Target", "Measurement Method"],
    [
        ["Call-report filing rate", "≥ 95% of meetings have a call report within 5 business days", "Automated report: meetings without call reports after 5 days"],
        ["Average call-report turnaround", "≤ 2 business days from meeting date", "Median(call_report.created_at − meeting.end_date_time)"],
        ["Opportunity capture rate", "≥ 40% of call reports have at least one linked opportunity", "Count(call_reports with opportunities) / Count(all call_reports)"],
        ["Supervisor approval turnaround", "≤ 1 business day from submission", "Median(approval.decided_at − call_report.submitted_for_approval_at)"],
        ["RM calendar adoption", "≥ 90% of RMs schedule ≥ 4 meetings/week via platform", "Weekly active scheduler count / total RM count"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 2: SCOPE & BOUNDARIES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("2. Scope & Boundaries", level=1)

doc.add_heading("2.1 In Scope", level=2)
in_scope = [
    "Multi-view calendar (All Activities, Month, Week, Day) with color-coded meeting status.",
    "Schedule a Meeting — with required/optional invitees, meeting reason, mode, location.",
    "Edit / Cancel / Reschedule meeting.",
    "File Call Report for a scheduled meeting (pre-populated fields from meeting).",
    "Standalone (ad-hoc) Call Report for unscheduled interactions.",
    "5-day threshold rule: if call report filed > 5 business days after meeting, route to supervisor approval.",
    "Supervisor approval workspace: claim, view, authorize, reject call reports.",
    "Opportunity management: add, view, modify, upload (bulk via XLS/XLSX), auto-expiry.",
    "Expense capture: add, view, modify with conditional fields for Conveyance type.",
    "Feedback capture from Calendar and Customer Dashboard.",
    "Conversation History: chronological timeline of all interactions per client.",
    "Action Item tracking linked to call reports.",
    "Attachment upload (PDF, DOCX, XLSX, images) on call reports.",
    "Notification system: email + in-app for meeting reminders, approval requests, overdue call reports.",
    "Reporting dashboards: RM activity, meeting coverage, opportunity pipeline, expense summary.",
    "Role-based access control for RM, Supervisor, Compliance Officer, MIS/Ops.",
    "REST API layer for all operations.",
    "Integration with existing Trust OMS CRM (Customer, Lead, Prospect entities).",
]
for s in in_scope:
    add_bullet(doc, s)

doc.add_heading("2.2 Out of Scope", level=2)
out_scope = [
    "Video conferencing integration (Zoom/Teams embedded) — meetings link to external tools via URL only.",
    "Automated speech-to-text transcription of calls.",
    "Mobile native app — responsive web only for v1.",
    "Direct calendar sync with Outlook/Google Calendar (planned for v2).",
    "Expense approval workflow (expenses are recorded for MIS; approval is outside this module).",
    "Client-facing meeting booking portal (client portal integration planned for v2).",
    "AI-powered meeting summary generation.",
]
for s in out_scope:
    add_bullet(doc, s)

doc.add_heading("2.3 Assumptions", level=2)
assumptions = [
    "RMs have been onboarded in the Trust OMS user directory with correct branch/hierarchy assignments.",
    "Customer, Lead, and Prospect master data exists in the CRM module and is accessible via internal API.",
    "Business-day calendar (excluding public holidays) is maintained in a shared configuration table.",
    "The 5-day threshold is configurable at the organization level (default = 5 business days).",
    "File attachments are stored in an S3-compatible object store; only metadata is persisted in the database.",
    "Email/SMTP service is available for notification delivery.",
    "The front end is React/TypeScript; the back end is Node.js/Express with PostgreSQL.",
]
for a in assumptions:
    add_bullet(doc, a)

doc.add_heading("2.4 Constraints", level=2)
constraints = [
    "Must comply with data-residency regulations (data stored in-country).",
    "Call reports and conversation history are immutable once approved (append-only audit log).",
    "Maximum file attachment size: 10 MB per file, 50 MB total per call report.",
    "Calendar view must render ≤ 2 seconds for a month with up to 200 meetings.",
    "All timestamps stored in UTC; displayed in the user's configured timezone.",
]
for c in constraints:
    add_bullet(doc, c)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 3: USER ROLES & PERMISSIONS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("3. User Roles & Permissions", level=1)

doc.add_heading("3.1 Role Definitions", level=2)
add_styled_table(doc,
    ["Role", "Description"],
    [
        ["Relationship Manager (RM)", "Front-line advisor managing client relationships. Schedules meetings, files call reports, captures opportunities and expenses."],
        ["RM Supervisor", "Manages a team of RMs. Reviews and approves/rejects late call reports. Views team-level dashboards and activity reports."],
        ["Branch Manager", "Oversees all RMs and supervisors in a branch. Read-only access to all data within the branch."],
        ["Compliance Officer", "Audits call reports and conversation histories for regulatory adherence. Read-only access across the organization."],
        ["MIS / Operations", "Generates reports and extracts data for management information. Read-only access to aggregated data."],
        ["System Administrator", "Configures system parameters (5-day threshold, meeting reasons, expense types). Manages user roles."],
    ])

doc.add_heading("3.2 Permissions Matrix", level=2)
add_styled_table(doc,
    ["Feature / Action", "RM", "Supervisor", "Branch Mgr", "Compliance", "MIS/Ops", "Admin"],
    [
        ["View own calendar", "Yes", "Yes", "Yes", "No", "No", "No"],
        ["View team calendar", "No", "Yes", "Yes", "No", "No", "No"],
        ["Schedule meeting", "Yes", "Yes", "No", "No", "No", "No"],
        ["Edit/Cancel own meeting", "Yes", "Yes", "No", "No", "No", "No"],
        ["File call report (own meetings)", "Yes", "Yes", "No", "No", "No", "No"],
        ["File standalone call report", "Yes", "Yes", "No", "No", "No", "No"],
        ["View own call reports", "Yes", "Yes", "Yes", "Yes", "Yes", "No"],
        ["View team call reports", "No", "Yes", "Yes", "Yes", "Yes", "No"],
        ["Approve/Reject call reports", "No", "Yes", "No", "No", "No", "No"],
        ["Add opportunity", "Yes", "Yes", "No", "No", "No", "No"],
        ["View opportunities (own)", "Yes", "Yes", "Yes", "Yes", "Yes", "No"],
        ["View opportunities (team)", "No", "Yes", "Yes", "Yes", "Yes", "No"],
        ["Bulk upload opportunities", "Yes", "Yes", "No", "No", "No", "No"],
        ["Add expense", "Yes", "Yes", "No", "No", "No", "No"],
        ["View expenses (own)", "Yes", "Yes", "Yes", "No", "Yes", "No"],
        ["View expenses (team)", "No", "Yes", "Yes", "No", "Yes", "No"],
        ["Submit feedback", "Yes", "Yes", "No", "No", "No", "No"],
        ["View conversation history", "Yes", "Yes", "Yes", "Yes", "No", "No"],
        ["View RM activity dashboard", "Own only", "Team", "Branch", "Org", "Org", "No"],
        ["Configure system parameters", "No", "No", "No", "No", "No", "Yes"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 4: DATA MODEL
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("4. Data Model", level=1)
doc.add_paragraph("All entities use UUID primary keys. Timestamps are stored in UTC. Soft-delete is implemented via a deleted_at nullable timestamp column.")

# ── 4.1 Meeting ──────────────────────────────────────────────────────
doc.add_heading("4.1 Meeting", level=2)
doc.add_paragraph("Represents a scheduled calendar event between an RM and one or more clients/leads/prospects.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["subject", "VARCHAR(255)", "Yes", "Min 3 chars, max 255", "—"],
        ["meeting_reason", "ENUM", "Yes", "Values: client_call_new, client_call_existing, branch_visit, portfolio_review, service_request, campaign_details, others", "—"],
        ["meeting_reason_other", "VARCHAR(255)", "Conditional", "Required if meeting_reason = 'others'", "NULL"],
        ["location", "VARCHAR(500)", "Yes", "Free text", "—"],
        ["start_date_time", "TIMESTAMPTZ", "Yes", "Must be ≥ current time (for new meetings)", "—"],
        ["end_date_time", "TIMESTAMPTZ", "Yes", "Must be > start_date_time", "—"],
        ["is_all_day", "BOOLEAN", "No", "If true, time portion is ignored (00:00–23:59)", "false"],
        ["meeting_type", "ENUM", "Yes", "Values: cif, lead, prospect, others", "—"],
        ["mode_of_meeting", "ENUM", "Yes", "Values: face_to_face, in_person, in_person_offshore, telephone, telephone_offshore, video_conference, video_conference_offshore, others", "—"],
        ["relationship_id", "UUID", "Conditional", "FK → customers/leads/prospects depending on meeting_type. Required unless meeting_type = 'others'", "NULL"],
        ["relationship_name", "VARCHAR(255)", "Yes", "Denormalized from relationship entity for display", "—"],
        ["contact_phone", "VARCHAR(20)", "No", "Auto-populated from relationship; editable. Format: +<country><number>", "NULL"],
        ["contact_email", "VARCHAR(255)", "No", "Auto-populated from relationship; editable. RFC 5322 format", "NULL"],
        ["remarks", "TEXT", "No", "Max 2000 chars", "NULL"],
        ["status", "ENUM", "Yes", "Values: scheduled, completed, cancelled, no_show", "scheduled"],
        ["call_report_status", "ENUM", "Yes", "Values: pending, filed, overdue", "pending"],
        ["created_by", "UUID", "Yes", "FK → users.id (the RM)", "Current user"],
        ["branch_id", "UUID", "Yes", "FK → branches.id", "Current user's branch"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set on insert", "NOW()"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set on update", "NOW()"],
        ["deleted_at", "TIMESTAMPTZ", "No", "Set on soft-delete", "NULL"],
    ])

doc.add_paragraph("Relationships: Meeting 1→N MeetingInvitee, Meeting 1→0..1 CallReport, Meeting N→1 User (created_by).")

doc.add_heading("Sample Data — Meeting", level=3)
add_styled_table(doc,
    ["id", "subject", "meeting_reason", "start_date_time", "status", "call_report_status", "created_by"],
    [
        ["a1b2c3d4-...", "Q1 Portfolio Review – Sharma Family", "portfolio_review", "2026-04-25 10:00 UTC", "scheduled", "pending", "rm-user-001"],
        ["e5f6g7h8-...", "New Investment Discussion – Patel Trust", "client_call_new", "2026-04-22 14:30 UTC", "completed", "filed", "rm-user-001"],
        ["i9j0k1l2-...", "Branch Visit – Andheri West", "branch_visit", "2026-04-20 09:00 UTC", "completed", "overdue", "rm-user-002"],
    ])

# ── 4.2 MeetingInvitee ──────────────────────────────────────────────
doc.add_heading("4.2 MeetingInvitee", level=2)
doc.add_paragraph("Join table linking meetings to internal users invited as attendees.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["meeting_id", "UUID", "Yes", "FK → meetings.id", "—"],
        ["user_id", "UUID", "Yes", "FK → users.id", "—"],
        ["invitee_type", "ENUM", "Yes", "Values: required, optional", "—"],
        ["rsvp_status", "ENUM", "No", "Values: pending, accepted, declined, tentative", "pending"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_paragraph("Relationships: MeetingInvitee N→1 Meeting, MeetingInvitee N→1 User.")

doc.add_heading("Sample Data — MeetingInvitee", level=3)
add_styled_table(doc,
    ["id", "meeting_id", "user_id", "invitee_type", "rsvp_status"],
    [
        ["inv-001", "a1b2c3d4-...", "rm-user-001", "required", "accepted"],
        ["inv-002", "a1b2c3d4-...", "sup-user-010", "optional", "pending"],
    ])

# ── 4.3 CallReport ──────────────────────────────────────────────────
doc.add_heading("4.3 CallReport", level=2)
doc.add_paragraph("Captures the outcome and details of an RM-client interaction. Can be linked to a scheduled meeting or filed as a standalone report.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["meeting_id", "UUID", "No", "FK → meetings.id. NULL for standalone reports", "NULL"],
        ["report_type", "ENUM", "Yes", "Values: scheduled, standalone", "—"],
        ["subject", "VARCHAR(255)", "Yes", "Pre-populated from meeting if linked; editable for standalone", "—"],
        ["meeting_reason", "ENUM", "Yes", "Same enum as Meeting.meeting_reason", "—"],
        ["meeting_reason_other", "VARCHAR(255)", "Conditional", "Required if meeting_reason = 'others'", "NULL"],
        ["start_date_time", "TIMESTAMPTZ", "Yes", "Display only if from meeting; editable for standalone", "—"],
        ["end_date_time", "TIMESTAMPTZ", "Yes", "Must be > start_date_time", "—"],
        ["location", "VARCHAR(500)", "Yes", "Pre-populated from meeting if linked", "—"],
        ["meeting_type", "ENUM", "Yes", "Values: cif, lead, prospect, others", "—"],
        ["mode_of_meeting", "ENUM", "Yes", "Same enum as Meeting.mode_of_meeting", "—"],
        ["relationship_id", "UUID", "Conditional", "FK → customers/leads/prospects", "NULL"],
        ["relationship_name", "VARCHAR(255)", "Yes", "Denormalized for display", "—"],
        ["person_met", "VARCHAR(255)", "Yes", "Name of the individual met (may differ from relationship_name)", "—"],
        ["client_status", "ENUM", "No", "Values: client_abroad_email, conference_call, not_contactable, not_interested", "NULL"],
        ["state_of_mind", "ENUM", "No", "Values: happy, irate, satisfied, sensitive, ultra_sensitive", "NULL"],
        ["summary_of_discussion", "TEXT", "Yes", "Min 20 chars, max 5000 chars", "—"],
        ["remarks", "TEXT", "No", "Max 2000 chars", "NULL"],
        ["next_meeting_start", "TIMESTAMPTZ", "No", "If provided, auto-creates a new meeting", "NULL"],
        ["next_meeting_end", "TIMESTAMPTZ", "No", "Required if next_meeting_start is provided", "NULL"],
        ["status", "ENUM", "Yes", "Values: draft, pending, completed, pending_approval, approved, rejected", "draft"],
        ["filed_date", "TIMESTAMPTZ", "No", "Set when status transitions from draft to pending/completed", "NULL"],
        ["days_since_meeting", "INTEGER", "No", "Computed: business_days(meeting.end_date_time, filed_date)", "NULL"],
        ["requires_approval", "BOOLEAN", "No", "True if days_since_meeting > threshold (default 5)", "false"],
        ["approval_submitted_at", "TIMESTAMPTZ", "No", "When submitted for supervisor approval", "NULL"],
        ["created_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["branch_id", "UUID", "Yes", "FK → branches.id", "Current user's branch"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set on update", "NOW()"],
        ["deleted_at", "TIMESTAMPTZ", "No", "Soft-delete", "NULL"],
    ])

doc.add_paragraph("Relationships: CallReport 0..1→1 Meeting, CallReport 1→N Opportunity, CallReport 1→N Expense, CallReport 1→N ActionItem, CallReport 1→N Attachment, CallReport 0..1→1 CallReportApproval.")

doc.add_heading("Sample Data — CallReport", level=3)
add_styled_table(doc,
    ["id", "meeting_id", "report_type", "subject", "status", "days_since_meeting", "requires_approval"],
    [
        ["cr-001", "e5f6g7h8-...", "scheduled", "Q1 Review – Patel Trust", "completed", "1", "false"],
        ["cr-002", "i9j0k1l2-...", "scheduled", "Branch Visit – Andheri", "pending_approval", "7", "true"],
        ["cr-003", "NULL", "standalone", "Ad-hoc call – Kapoor HNI", "completed", "0", "false"],
    ])

# ── 4.4 CallReportApproval ───────────────────────────────────────────
doc.add_heading("4.4 CallReportApproval", level=2)
doc.add_paragraph("Tracks the supervisor review workflow for call reports that exceed the filing threshold.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["call_report_id", "UUID", "Yes", "FK → call_reports.id. Unique constraint.", "—"],
        ["supervisor_id", "UUID", "No", "FK → users.id. Set when supervisor claims the report.", "NULL"],
        ["status", "ENUM", "Yes", "Values: pending, claimed, approved, rejected", "pending"],
        ["claimed_at", "TIMESTAMPTZ", "No", "Set when supervisor claims", "NULL"],
        ["decided_at", "TIMESTAMPTZ", "No", "Set when approved or rejected", "NULL"],
        ["reviewer_comments", "TEXT", "No", "Mandatory on rejection; optional on approval. Max 2000 chars.", "NULL"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_heading("Sample Data — CallReportApproval", level=3)
add_styled_table(doc,
    ["id", "call_report_id", "supervisor_id", "status", "reviewer_comments"],
    [
        ["apr-001", "cr-002", "sup-user-010", "claimed", "NULL"],
        ["apr-002", "cr-005", "sup-user-010", "rejected", "Summary too brief; add details about client risk appetite discussion."],
    ])

# ── 4.5 Opportunity ─────────────────────────────────────────────────
doc.add_heading("4.5 Opportunity", level=2)
doc.add_paragraph("A sales opportunity identified during a client interaction, linked to a call report.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["call_report_id", "UUID", "No", "FK → call_reports.id", "NULL"],
        ["opportunity_type", "ENUM", "Yes", "Values: cif, lead, prospect", "—"],
        ["relationship_id", "UUID", "Yes", "FK → customers/leads/prospects", "—"],
        ["relationship_name", "VARCHAR(255)", "Yes", "Denormalized", "—"],
        ["sub_product", "VARCHAR(255)", "Yes", "Free text or dropdown from product master", "—"],
        ["investment_mode", "ENUM", "No", "Values: lumpsum, periodic", "NULL"],
        ["opportunity_discovered", "DECIMAL(18,2)", "Yes", "Amount in base currency (INR). Must be > 0.", "—"],
        ["currency", "VARCHAR(3)", "Yes", "ISO 4217 code", "INR"],
        ["opportunity_date", "DATE", "Yes", "Cannot be in the future", "TODAY"],
        ["description", "TEXT", "No", "Max 2000 chars", "NULL"],
        ["due_date", "DATE", "Yes", "Must be ≥ opportunity_date", "—"],
        ["opportunity_closed", "DECIMAL(18,2)", "No", "Amount closed/won. Must be ≤ opportunity_discovered.", "NULL"],
        ["status", "ENUM", "Yes", "Values: open, closed, aborted, expired", "open"],
        ["stage", "ENUM", "Yes", "Values: interested, to_be_approached, won_doc_in_process, won_doc_completed, declined, not_proceeding", "interested"],
        ["created_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["deleted_at", "TIMESTAMPTZ", "No", "Soft-delete", "NULL"],
    ])

doc.add_paragraph("Business rule: A nightly batch job sets status = 'expired' for all opportunities where due_date < current_business_date AND status = 'open'.")
doc.add_paragraph("Relationships: Opportunity N→1 CallReport, Opportunity N→1 User (created_by).")

doc.add_heading("Sample Data — Opportunity", level=3)
add_styled_table(doc,
    ["id", "relationship_name", "sub_product", "opportunity_discovered", "due_date", "status", "stage"],
    [
        ["opp-001", "Sharma Family Trust", "Equity MF – Large Cap", "5000000.00", "2026-06-30", "open", "interested"],
        ["opp-002", "Patel Industries", "Fixed Deposit – 1Y", "10000000.00", "2026-05-15", "closed", "won_doc_completed"],
        ["opp-003", "Kapoor HNI", "PMS – Balanced", "25000000.00", "2026-04-20", "expired", "to_be_approached"],
    ])

# ── 4.6 Expense ──────────────────────────────────────────────────────
doc.add_heading("4.6 Expense", level=2)
doc.add_paragraph("Expenses incurred by RMs in the course of client interactions.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["expense_type", "ENUM", "Yes", "Values: client_entertainment, conveyance, festive_events, others", "—"],
        ["expense_date", "DATE", "Yes", "Cannot be in the future", "TODAY"],
        ["amount", "DECIMAL(12,2)", "Yes", "Must be > 0. Currency is base currency (INR).", "—"],
        ["currency", "VARCHAR(3)", "Yes", "ISO 4217", "INR"],
        ["purpose", "VARCHAR(500)", "Yes", "Min 5 chars", "—"],
        ["entity_type", "ENUM", "Yes", "Values: cif, lead, prospect", "—"],
        ["relationship_id", "UUID", "Yes", "FK → customers/leads/prospects", "—"],
        ["relationship_name", "VARCHAR(255)", "Yes", "Denormalized", "—"],
        ["call_report_id", "UUID", "No", "FK → call_reports.id. Strongly recommended for conveyance.", "NULL"],
        ["from_place", "VARCHAR(255)", "Conditional", "Required if expense_type = 'conveyance'", "NULL"],
        ["to_place", "VARCHAR(255)", "Conditional", "Required if expense_type = 'conveyance'", "NULL"],
        ["transport_mode", "ENUM", "Conditional", "Required if expense_type = 'conveyance'. Values: car, taxi, bus, train, flight, auto, others", "NULL"],
        ["distance_km", "DECIMAL(8,2)", "Conditional", "Required if expense_type = 'conveyance'. Must be > 0.", "NULL"],
        ["created_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["deleted_at", "TIMESTAMPTZ", "No", "Soft-delete", "NULL"],
    ])

doc.add_heading("Sample Data — Expense", level=3)
add_styled_table(doc,
    ["id", "expense_type", "expense_date", "amount", "purpose", "from_place", "to_place"],
    [
        ["exp-001", "client_entertainment", "2026-04-22", "3500.00", "Lunch with Sharma Family Trust client", "NULL", "NULL"],
        ["exp-002", "conveyance", "2026-04-21", "850.00", "Travel to Andheri branch for client meet", "BKC Office", "Andheri West Branch"],
        ["exp-003", "festive_events", "2026-03-25", "12000.00", "Holi client appreciation event", "NULL", "NULL"],
    ])

# ── 4.7 Feedback ─────────────────────────────────────────────────────
doc.add_heading("4.7 Feedback", level=2)
doc.add_paragraph("Feedback recorded by an RM about a client interaction, captured from Calendar or Customer Dashboard.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["meeting_id", "UUID", "No", "FK → meetings.id", "NULL"],
        ["call_report_id", "UUID", "No", "FK → call_reports.id", "NULL"],
        ["relationship_id", "UUID", "Yes", "FK → customers/leads/prospects", "—"],
        ["feedback_text", "TEXT", "Yes", "Min 10 chars, max 3000 chars", "—"],
        ["sentiment", "ENUM", "No", "Values: positive, neutral, negative", "NULL"],
        ["source", "ENUM", "Yes", "Values: calendar, customer_dashboard, manual", "—"],
        ["created_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_heading("Sample Data — Feedback", level=3)
add_styled_table(doc,
    ["id", "relationship_id", "feedback_text", "sentiment", "source"],
    [
        ["fb-001", "rel-sharma", "Client very happy with portfolio performance, interested in increasing equity allocation.", "positive", "calendar"],
        ["fb-002", "rel-kapoor", "Client frustrated with delayed statement delivery. Escalated to ops.", "negative", "customer_dashboard"],
    ])

# ── 4.8 ConversationHistory ─────────────────────────────────────────
doc.add_heading("4.8 ConversationHistory", level=2)
doc.add_paragraph("Immutable, append-only timeline of all interactions with a client relationship. Entries are auto-generated when a call report is filed, feedback is submitted, or a meeting status changes.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["relationship_id", "UUID", "Yes", "FK → customers/leads/prospects", "—"],
        ["interaction_type", "ENUM", "Yes", "Values: meeting_scheduled, meeting_completed, meeting_cancelled, call_report_filed, feedback_submitted, opportunity_created, opportunity_updated, action_item_completed", "—"],
        ["interaction_date", "TIMESTAMPTZ", "Yes", "Date/time of the actual interaction", "—"],
        ["summary", "TEXT", "Yes", "Auto-generated summary. Max 1000 chars.", "—"],
        ["reference_type", "VARCHAR(50)", "Yes", "Entity type: meeting, call_report, feedback, opportunity, action_item", "—"],
        ["reference_id", "UUID", "Yes", "FK to the referenced entity", "—"],
        ["created_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_paragraph("This entity is INSERT-ONLY. No updates or deletes are permitted.")

doc.add_heading("Sample Data — ConversationHistory", level=3)
add_styled_table(doc,
    ["id", "relationship_id", "interaction_type", "interaction_date", "summary"],
    [
        ["ch-001", "rel-sharma", "meeting_completed", "2026-04-22 14:30 UTC", "Q1 Portfolio Review completed. Discussed rebalancing equity allocation."],
        ["ch-002", "rel-sharma", "call_report_filed", "2026-04-22 16:00 UTC", "Call report filed for Q1 Portfolio Review meeting."],
        ["ch-003", "rel-sharma", "opportunity_created", "2026-04-22 16:05 UTC", "New opportunity: Equity MF Large Cap, INR 50,00,000."],
    ])

# ── 4.9 ActionItem ──────────────────────────────────────────────────
doc.add_heading("4.9 ActionItem", level=2)
doc.add_paragraph("Follow-up tasks generated from call reports.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["call_report_id", "UUID", "Yes", "FK → call_reports.id", "—"],
        ["title", "VARCHAR(255)", "Yes", "Min 5 chars", "—"],
        ["description", "TEXT", "No", "Max 2000 chars", "NULL"],
        ["assigned_to", "UUID", "Yes", "FK → users.id", "Current user"],
        ["due_date", "DATE", "Yes", "Must be ≥ today", "—"],
        ["priority", "ENUM", "Yes", "Values: low, medium, high, critical", "medium"],
        ["status", "ENUM", "Yes", "Values: open, in_progress, completed, cancelled", "open"],
        ["completed_at", "TIMESTAMPTZ", "No", "Set when status → completed", "NULL"],
        ["created_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_heading("Sample Data — ActionItem", level=3)
add_styled_table(doc,
    ["id", "call_report_id", "title", "assigned_to", "due_date", "priority", "status"],
    [
        ["ai-001", "cr-001", "Send updated portfolio statement to client", "rm-user-001", "2026-04-25", "high", "open"],
        ["ai-002", "cr-001", "Schedule follow-up call for MF redemption", "rm-user-001", "2026-04-28", "medium", "open"],
    ])

# ── 4.10 Attachment ──────────────────────────────────────────────────
doc.add_heading("4.10 Attachment", level=2)
doc.add_paragraph("File attachments linked to call reports.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["call_report_id", "UUID", "Yes", "FK → call_reports.id", "—"],
        ["file_name", "VARCHAR(255)", "Yes", "Original file name", "—"],
        ["file_type", "VARCHAR(50)", "Yes", "MIME type. Allowed: application/pdf, application/vnd.openxmlformats-officedocument.*, image/jpeg, image/png", "—"],
        ["file_size_bytes", "INTEGER", "Yes", "Max 10,485,760 (10 MB)", "—"],
        ["storage_path", "VARCHAR(1000)", "Yes", "S3 object key", "—"],
        ["uploaded_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_heading("Sample Data — Attachment", level=3)
add_styled_table(doc,
    ["id", "call_report_id", "file_name", "file_type", "file_size_bytes"],
    [
        ["att-001", "cr-001", "portfolio_summary_q1.pdf", "application/pdf", "245760"],
        ["att-002", "cr-002", "meeting_notes_scan.jpg", "image/jpeg", "1048576"],
    ])

# ── 4.11 OpportunityUpload ──────────────────────────────────────────
doc.add_heading("4.11 OpportunityUpload", level=2)
doc.add_paragraph("Tracks bulk opportunity uploads via XLS/XLSX file.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["file_name", "VARCHAR(255)", "Yes", "Original file name (.xls or .xlsx only)", "—"],
        ["file_size_bytes", "INTEGER", "Yes", "Max 5,242,880 (5 MB)", "—"],
        ["storage_path", "VARCHAR(1000)", "Yes", "S3 object key", "—"],
        ["total_rows", "INTEGER", "No", "Parsed after upload", "NULL"],
        ["success_count", "INTEGER", "No", "Rows successfully imported", "NULL"],
        ["error_count", "INTEGER", "No", "Rows with validation errors", "NULL"],
        ["error_log", "JSONB", "No", "Array of {row, field, error} objects", "NULL"],
        ["status", "ENUM", "Yes", "Values: uploaded, processing, completed, failed", "uploaded"],
        ["uploaded_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
        ["completed_at", "TIMESTAMPTZ", "No", "Set when processing finishes", "NULL"],
    ])

doc.add_heading("Sample Data — OpportunityUpload", level=3)
add_styled_table(doc,
    ["id", "file_name", "total_rows", "success_count", "error_count", "status"],
    [
        ["ou-001", "opportunities_apr_batch.xlsx", "150", "148", "2", "completed"],
        ["ou-002", "leads_pipeline.xlsx", "50", "0", "0", "processing"],
    ])

# ── 4.12 SystemConfig ───────────────────────────────────────────────
doc.add_heading("4.12 SystemConfig", level=2)
doc.add_paragraph("Key-value configuration store for module-level parameters.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["config_key", "VARCHAR(100)", "Yes", "Unique. e.g., 'call_report_threshold_days'", "—"],
        ["config_value", "VARCHAR(500)", "Yes", "Stored as string, parsed by application", "—"],
        ["description", "VARCHAR(500)", "No", "Human-readable description", "NULL"],
        ["updated_by", "UUID", "Yes", "FK → users.id", "—"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_heading("Sample Data — SystemConfig", level=3)
add_styled_table(doc,
    ["config_key", "config_value", "description"],
    [
        ["call_report_threshold_days", "5", "Number of business days after which call report requires supervisor approval"],
        ["max_attachment_size_bytes", "10485760", "Maximum size per attachment in bytes (10 MB)"],
        ["opportunity_auto_expiry_enabled", "true", "Enable nightly batch to auto-expire overdue opportunities"],
    ])

# ── Entity Relationship Summary ──────────────────────────────────────
doc.add_heading("4.13 Entity Relationship Summary", level=2)
relationships = [
    "User 1→N Meeting (created_by)",
    "Meeting 1→N MeetingInvitee",
    "Meeting 1→0..1 CallReport",
    "CallReport 1→0..1 CallReportApproval",
    "CallReport 1→N Opportunity",
    "CallReport 1→N Expense",
    "CallReport 1→N ActionItem",
    "CallReport 1→N Attachment",
    "CallReport 1→N Feedback",
    "Customer/Lead/Prospect 1→N Meeting (relationship_id)",
    "Customer/Lead/Prospect 1→N ConversationHistory (relationship_id)",
    "OpportunityUpload 1→N Opportunity (via bulk import)",
    "User 1→N CallReportApproval (supervisor_id)",
]
for r in relationships:
    add_bullet(doc, r)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 5: FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("5. Functional Requirements", level=1)

frs = [
    {
        "id": "FR-001",
        "name": "Calendar Multi-View Display",
        "desc": "The system shall display the RM's calendar in four views: All Activities (list), Month (grid), Week (7-column grid), and Day (hourly slots). Each meeting is color-coded by status: Scheduled (blue), Completed (green), Cancelled (grey), No Show (red), Overdue Call Report (orange).",
        "story": "As an RM, I want to view my scheduled meetings in multiple calendar formats so that I can quickly see my upcoming and past interactions.",
        "criteria": [
            "AC-001-1: Month view displays all meetings for the selected month with color-coded status dots.",
            "AC-001-2: Week view shows 7 columns (Mon–Sun) with meetings as time blocks.",
            "AC-001-3: Day view shows hourly slots (8 AM–8 PM default) with meeting cards.",
            "AC-001-4: All Activities view shows a paginated table (20 rows/page) sortable by date, subject, status.",
            "AC-001-5: Clicking a meeting card/row opens the meeting detail panel.",
            "AC-001-6: A status legend is visible at the top of every view.",
        ],
        "rules": [
            "BR-001-1: Default view on page load is Week view centered on the current week.",
            "BR-001-2: Meetings in the past 90 days and future 365 days are loaded; older meetings via pagination.",
            "BR-001-3: If a meeting is > threshold days old and has call_report_status = 'pending', it is colored orange (overdue).",
        ],
        "ui": "View toggle (tabs or segmented control) at top. Filter panel: date range, meeting reason, status. Calendar grid fills remaining viewport. Mobile: month view defaults to list mode.",
        "edge": "If RM has > 10 meetings in a single day, Day view scrolls vertically. If no meetings exist for a period, display 'No meetings scheduled' placeholder.",
    },
    {
        "id": "FR-002",
        "name": "Schedule a Meeting",
        "desc": "The system shall allow RMs to schedule a new meeting by filling in subject, meeting reason, location, date/time, meeting type, mode, relationship, contact details, remarks, and invitees.",
        "story": "As an RM, I want to schedule a meeting with a client so that I can track and prepare for the interaction.",
        "criteria": [
            "AC-002-1: 'Schedule Meeting' button is visible on all calendar views.",
            "AC-002-2: A modal/drawer form opens with all required fields.",
            "AC-002-3: Meeting reason dropdown contains exactly: Client Call – New, Client Call – Existing, Branch Visit, Portfolio Review, Service Request, Campaign Details, Others.",
            "AC-002-4: Selecting 'Others' reveals a free-text 'Specify Reason' field (mandatory).",
            "AC-002-5: Meeting type dropdown: CIF, Lead, Prospect, Others.",
            "AC-002-6: Selecting CIF/Lead/Prospect enables a type-ahead search on the relationship field; selecting Others allows free text.",
            "AC-002-7: Contact phone and email auto-populate from the selected relationship but are editable.",
            "AC-002-8: 'All Day' checkbox disables time pickers and sets start=00:00, end=23:59.",
            "AC-002-9: At least one Required Invitee must be added (can be the RM themselves).",
            "AC-002-10: On save, meeting appears on the calendar immediately.",
        ],
        "rules": [
            "BR-002-1: Start date/time must be in the future (≥ current time) for new meetings.",
            "BR-002-2: End date/time must be after start date/time.",
            "BR-002-3: Meeting duration must be ≤ 8 hours unless is_all_day = true.",
            "BR-002-4: Duplicate detection: warn (not block) if another meeting exists within ±30 minutes for the same RM.",
        ],
        "ui": "Modal form with two columns on desktop (left: meeting details, right: invitees). Required fields marked with red asterisk. Save and Cancel buttons at bottom. Success: toast notification + modal closes. Error: inline field-level errors.",
        "edge": "If the selected relationship has no phone/email on file, show a warning banner: 'Contact details not available. Please enter manually.' If the RM's calendar has a conflict, show a yellow warning but allow save.",
    },
    {
        "id": "FR-003",
        "name": "Edit / Cancel / Reschedule Meeting",
        "desc": "The system shall allow RMs to edit meeting details, cancel a meeting, or reschedule it to a new date/time. Cancelled meetings cannot be re-opened.",
        "story": "As an RM, I want to modify or cancel a scheduled meeting so that my calendar accurately reflects my plans.",
        "criteria": [
            "AC-003-1: Edit button is visible on meetings with status = 'scheduled'.",
            "AC-003-2: All editable fields can be modified (subject, reason, location, date/time, invitees, remarks).",
            "AC-003-3: Cancel button prompts a confirmation dialog: 'Are you sure you want to cancel this meeting? This cannot be undone.'",
            "AC-003-4: Cancelled meetings show as greyed-out on the calendar with a strikethrough on the subject.",
            "AC-003-5: Reschedule changes the date/time and keeps all other fields; a ConversationHistory entry is created noting the reschedule.",
        ],
        "rules": [
            "BR-003-1: Only the meeting creator or a Supervisor can edit/cancel.",
            "BR-003-2: Meetings with status = 'completed' or 'cancelled' cannot be edited.",
            "BR-003-3: If a call report has already been filed for this meeting, the meeting cannot be cancelled.",
        ],
        "ui": "Edit opens the same form as FR-002 pre-populated. Cancel shows a confirmation modal. Reschedule is an edit action that changes date/time fields.",
        "edge": "If all invitees decline after scheduling, meeting remains active (RM's discretion to cancel).",
    },
    {
        "id": "FR-004",
        "name": "File Call Report (Scheduled Meeting)",
        "desc": "After a meeting is completed, the RM files a call report. Key meeting fields (subject, reason, date/time, location, type, mode, relationship) are pre-populated from the meeting and shown as read-only. The RM adds: person met, client status, state of mind, summary of discussion, and optional next meeting details.",
        "story": "As an RM, I want to file a call report for a completed meeting so that the interaction is documented for compliance and follow-up.",
        "criteria": [
            "AC-004-1: 'File Call Report' button appears on meetings with status = 'completed' and call_report_status = 'pending'.",
            "AC-004-2: Pre-populated fields from the meeting are displayed as read-only.",
            "AC-004-3: Summary of Discussion is mandatory, min 20 characters.",
            "AC-004-4: If days_since_meeting > threshold, a yellow banner warns: 'This call report exceeds the X-day filing window and will require supervisor approval.'",
            "AC-004-5: On submit, if requires_approval = true, status is set to 'pending_approval' and a CallReportApproval record is created.",
            "AC-004-6: On submit, if requires_approval = false, status is set to 'completed'.",
            "AC-004-7: A ConversationHistory entry is auto-created.",
            "AC-004-8: If next meeting date/time is provided, a new Meeting record is auto-created.",
        ],
        "rules": [
            "BR-004-1: days_since_meeting is calculated as business days between meeting.end_date_time and current date (excluding weekends and holidays from the business calendar).",
            "BR-004-2: The threshold defaults to 5 business days and is read from SystemConfig.",
            "BR-004-3: A meeting can have exactly one call report (1:1 relationship).",
            "BR-004-4: Filing a call report sets the parent meeting's call_report_status to 'filed'.",
        ],
        "ui": "Full-page form divided into sections: Meeting Info (read-only), Interaction Details, Opportunities, Action Items, Attachments, Next Steps. Submit and Save Draft buttons.",
        "edge": "If the RM navigates away with unsaved changes, show a 'Discard changes?' confirmation. If the meeting was cancelled after the RM opened the call report form, show an error on submit.",
    },
    {
        "id": "FR-005",
        "name": "Standalone Call Report",
        "desc": "The system shall allow RMs to file a call report without a pre-existing meeting, for ad-hoc phone calls, walk-ins, or informal interactions. All fields are editable (no pre-population).",
        "story": "As an RM, I want to file a standalone call report for an unscheduled interaction so that all client touchpoints are captured.",
        "criteria": [
            "AC-005-1: 'New Call Report' button is accessible from the calendar toolbar and the call reports list page.",
            "AC-005-2: All fields are editable (no read-only pre-population).",
            "AC-005-3: report_type is automatically set to 'standalone'.",
            "AC-005-4: The 5-day threshold rule does NOT apply (requires_approval = false).",
            "AC-005-5: A ConversationHistory entry is auto-created on submit.",
        ],
        "rules": [
            "BR-005-1: Standalone reports do not create a Meeting record.",
            "BR-005-2: Start/end date/time for standalone reports can be in the past.",
            "BR-005-3: Standalone reports are directly set to 'completed' on submit.",
        ],
        "ui": "Same form layout as FR-004 but all fields are editable. Mode of meeting defaults to 'telephone'.",
        "edge": "If the RM selects a date in the far past (> 90 days), show a warning but allow submission.",
    },
    {
        "id": "FR-006",
        "name": "Call Report Approval Workflow (5-Day Rule)",
        "desc": "When a call report is filed more than the configured threshold (default: 5 business days) after the meeting, it enters a supervisor approval queue. The supervisor can claim, review, approve, or reject the report.",
        "story": "As a Supervisor, I want to review late-filed call reports so that I can ensure interaction records are accurate and timely.",
        "criteria": [
            "AC-006-1: Supervisor dashboard shows two sections: 'Pending Approval Listing' (unclaimed) and 'My Approvals' (claimed by this supervisor).",
            "AC-006-2: Supervisor can 'Claim' a report from the pending listing, which locks it from other supervisors.",
            "AC-006-3: Claimed report opens in a read-only detail view with Approve and Reject buttons.",
            "AC-006-4: Approve sets call_report.status = 'approved' and approval.status = 'approved'.",
            "AC-006-5: Reject requires reviewer_comments (min 20 chars); sets call_report.status = 'rejected' and approval.status = 'rejected'.",
            "AC-006-6: Rejected reports are returned to the RM's 'Rejected Call Reports' queue for revision.",
            "AC-006-7: When a rejected report is re-submitted, a new approval cycle starts (new CallReportApproval record).",
        ],
        "rules": [
            "BR-006-1: Only supervisors in the same branch or hierarchy as the RM can claim reports.",
            "BR-006-2: A supervisor can claim a maximum of 20 uncompleted reports at a time.",
            "BR-006-3: If a claimed report is not actioned within 2 business days, it is auto-unclaimed and returned to the pending pool.",
            "BR-006-4: The RM is notified via email and in-app when their report is approved or rejected.",
        ],
        "ui": "Supervisor workspace with tabs: Pending (table with Claim button per row), My Approvals (table with View/Approve/Reject). Detail view: left panel = call report details, right panel = approval actions + comments textarea.",
        "edge": "If the RM updates the call report while it is claimed by a supervisor, the supervisor sees a 'Report updated by RM' banner and must re-review.",
    },
    {
        "id": "FR-007",
        "name": "Opportunity Capture",
        "desc": "The system shall allow RMs to capture sales opportunities linked to a call report. Each opportunity records the product, amount, stage, and expected close date.",
        "story": "As an RM, I want to record a sales opportunity during a call report so that my pipeline is visible to management.",
        "criteria": [
            "AC-007-1: 'Add Opportunity' button is available within the call report form and on the Opportunities list page.",
            "AC-007-2: Opportunity form captures: type, relationship, sub-product, investment mode, discovered amount, date, description, due date.",
            "AC-007-3: Stage dropdown: Interested, To be approached, Won – Doc in Process, Won – Doc completed, Declined, Not Proceeding.",
            "AC-007-4: Status auto-set to 'open' on creation.",
            "AC-007-5: Editing an opportunity updates the ConversationHistory.",
        ],
        "rules": [
            "BR-007-1: opportunity_discovered must be > 0.",
            "BR-007-2: opportunity_closed must be ≤ opportunity_discovered.",
            "BR-007-3: When stage is set to 'Won – Doc completed', opportunity_closed becomes mandatory.",
            "BR-007-4: Due date must be ≥ opportunity_date.",
        ],
        "ui": "Inline form within call report (accordion section). Standalone page: filterable data table with Add/Edit/View actions.",
        "edge": "If opportunity_discovered is entered without a currency, default to INR. If duplicate sub_product + relationship within 30 days, show a warning.",
    },
    {
        "id": "FR-008",
        "name": "Opportunity Bulk Upload",
        "desc": "RMs can bulk-upload opportunities via an XLS/XLSX file. The system validates each row and imports valid records while reporting errors.",
        "story": "As an RM, I want to upload multiple opportunities from a spreadsheet so that I can quickly populate my pipeline from offline notes.",
        "criteria": [
            "AC-008-1: Upload button accepts only .xls and .xlsx files up to 5 MB.",
            "AC-008-2: A downloadable template file is available with column headers matching the Opportunity entity fields.",
            "AC-008-3: After upload, system shows a processing indicator.",
            "AC-008-4: On completion, a summary is displayed: X rows imported, Y rows failed.",
            "AC-008-5: Failed rows are downloadable as an error log with row number, field, and error description.",
        ],
        "rules": [
            "BR-008-1: Maximum 500 rows per upload.",
            "BR-008-2: Each row is validated against the same rules as the single-create form.",
            "BR-008-3: Valid rows are imported even if some rows fail (partial success).",
        ],
        "ui": "Upload dialog with drag-and-drop zone. Template download link. Results table showing success/error summary with error log download.",
        "edge": "If file has 0 valid rows, show error: 'No valid records found. Please check the template format.' If file is corrupt/unreadable, show 'Unable to parse file. Please use the provided template.'",
    },
    {
        "id": "FR-009",
        "name": "Opportunity Auto-Expiry",
        "desc": "A nightly batch job automatically expires open opportunities whose due date has passed.",
        "story": "As an RM Supervisor, I want overdue opportunities to auto-expire so that the pipeline reflects only active pursuits.",
        "criteria": [
            "AC-009-1: Batch runs daily at a configurable time (default: 01:00 UTC).",
            "AC-009-2: All opportunities with status = 'open' and due_date < current_business_date are set to status = 'expired'.",
            "AC-009-3: A batch summary email is sent to the RM listing expired opportunities.",
            "AC-009-4: Expired opportunities appear with a distinct visual indicator (e.g., strikethrough or muted color).",
        ],
        "rules": [
            "BR-009-1: Only 'open' status opportunities are expired; 'closed' and 'aborted' are unaffected.",
            "BR-009-2: Auto-expiry can be disabled via SystemConfig ('opportunity_auto_expiry_enabled').",
        ],
        "ui": "No direct UI trigger; results visible in the Opportunities list. Expired items show in a separate 'Expired' tab or with a red badge.",
        "edge": "If the batch fails mid-run, it should be idempotent — re-running processes remaining records without double-processing.",
    },
    {
        "id": "FR-010",
        "name": "Expense Capture",
        "desc": "RMs can record expenses incurred during client interactions. Expense types include Client Entertainment, Conveyance, Festive Events, and Others. Conveyance expenses require additional fields: from/to place, transport mode, and distance.",
        "story": "As an RM, I want to log expenses related to client meetings so that I have a record for reimbursement and MIS reporting.",
        "criteria": [
            "AC-010-1: 'Add Expense' is available from the call report form and the Expenses list page.",
            "AC-010-2: Expense type dropdown: Client Entertainment, Conveyance, Festive Events, Others.",
            "AC-010-3: Selecting 'Conveyance' reveals fields: From Place, To Place, Transport Mode, Distance (km).",
            "AC-010-4: Transport mode dropdown: Car, Taxi, Bus, Train, Flight, Auto, Others.",
            "AC-010-5: Expense date cannot be in the future.",
            "AC-010-6: Amount must be > 0.",
            "AC-010-7: Expenses list page supports filters: expense type, date range.",
        ],
        "rules": [
            "BR-010-1: From Place, To Place, Transport Mode, and Distance are mandatory when expense_type = 'conveyance'.",
            "BR-010-2: Call Report linkage is optional but recommended for conveyance.",
            "BR-010-3: Distance must be > 0 km when required.",
        ],
        "ui": "Inline form within call report (accordion section). Standalone Expenses list page with data table, filters, and Add button.",
        "edge": "If the RM changes expense type from Conveyance to another type, the conditional fields are cleared.",
    },
    {
        "id": "FR-011",
        "name": "Feedback Capture",
        "desc": "RMs can record qualitative feedback about a client from the Calendar (linked to a meeting) or the Customer Dashboard (general feedback).",
        "story": "As an RM, I want to capture client feedback so that the relationship team has context on client sentiment.",
        "criteria": [
            "AC-011-1: 'Feedback' button is available on meeting detail view and customer dashboard.",
            "AC-011-2: Feedback form has: free-text feedback (min 10 chars), optional sentiment selector (Positive/Neutral/Negative).",
            "AC-011-3: Source is auto-set based on entry point (calendar vs customer_dashboard).",
            "AC-011-4: On submit, a ConversationHistory entry is created.",
        ],
        "rules": [
            "BR-011-1: Feedback is append-only — cannot be edited or deleted once submitted.",
            "BR-011-2: Multiple feedbacks can exist per meeting or per relationship.",
        ],
        "ui": "Modal dialog with textarea and sentiment radio buttons. Character counter showing min/max. Submit and Cancel.",
        "edge": "If feedback is submitted for a relationship that has been marked inactive/closed, show a warning but allow submission.",
    },
    {
        "id": "FR-012",
        "name": "Conversation History",
        "desc": "The system maintains and displays a chronological, immutable timeline of all interactions with a client/relationship. Entries are auto-generated by the system on key events.",
        "story": "As an RM, I want to view the full interaction history with a client so that I can prepare for meetings and demonstrate suitability compliance.",
        "criteria": [
            "AC-012-1: Conversation History is accessible from the Customer Dashboard and the Calendar (per-client filter).",
            "AC-012-2: Timeline displays entries in reverse chronological order with: date, interaction type icon, summary text, link to source entity.",
            "AC-012-3: Filter by interaction type, date range.",
            "AC-012-4: Entries are immutable — no edit/delete UI is provided.",
            "AC-012-5: Lazy-loads in pages of 20 entries.",
        ],
        "rules": [
            "BR-012-1: Auto-created on: meeting status change, call report filed, feedback submitted, opportunity created/updated, action item completed.",
            "BR-012-2: Summary is auto-generated: '{interaction_type} — {entity subject/title}. {first 100 chars of detail}'.",
        ],
        "ui": "Vertical timeline component with icons per type. Each entry is a card with date, type badge, summary, and 'View Details' link. Infinite scroll.",
        "edge": "If a referenced entity is soft-deleted, the conversation history entry remains with a 'Source record deleted' note.",
    },
    {
        "id": "FR-013",
        "name": "Action Item Management",
        "desc": "RMs can create follow-up action items linked to call reports, assign them, set due dates and priorities, and track completion.",
        "story": "As an RM, I want to create and track action items from a meeting so that follow-ups are not forgotten.",
        "criteria": [
            "AC-013-1: Action items can be added inline within the call report form.",
            "AC-013-2: Each action item has: title, description, assigned user, due date, priority.",
            "AC-013-3: Action items appear in a 'My Action Items' dashboard widget with filters: status, priority, due date.",
            "AC-013-4: Marking an action item as completed sets completed_at and creates a ConversationHistory entry.",
            "AC-013-5: Overdue action items (due_date < today AND status != completed) are highlighted in red.",
        ],
        "rules": [
            "BR-013-1: Action items can only be assigned to users within the same branch/hierarchy.",
            "BR-013-2: Completed action items cannot be re-opened.",
        ],
        "ui": "Inline repeater within call report form. Standalone 'My Action Items' page with kanban or table view. Quick-complete checkbox.",
        "edge": "If the assigned user is deactivated, action item remains and shows 'Assigned user inactive — please reassign.'",
    },
    {
        "id": "FR-014",
        "name": "Attachment Management",
        "desc": "RMs can upload file attachments to call reports for supporting documentation (meeting notes, presentations, signed documents).",
        "story": "As an RM, I want to attach documents to a call report so that supporting evidence is linked to the interaction record.",
        "criteria": [
            "AC-014-1: Upload zone accepts PDF, DOCX, XLSX, JPEG, PNG.",
            "AC-014-2: Max file size: 10 MB per file.",
            "AC-014-3: Max total per call report: 50 MB.",
            "AC-014-4: Files can be previewed (PDF/images inline; others download).",
            "AC-014-5: Files can be deleted only while the call report is in draft status.",
        ],
        "rules": [
            "BR-014-1: Files are virus-scanned before storage (integration with antivirus API).",
            "BR-014-2: Once call report status = 'completed' or 'approved', attachments are immutable.",
        ],
        "ui": "Drag-and-drop upload zone with file list showing name, size, type, and delete/preview icons.",
        "edge": "If upload fails mid-stream, show retry option. If total exceeds 50 MB, block additional uploads with error message.",
    },
    {
        "id": "FR-015",
        "name": "Call Reports List & Search",
        "desc": "The system provides a paginated, filterable list of all call reports accessible to the current user based on their role.",
        "story": "As an RM, I want to search and filter my past call reports so that I can quickly find a specific interaction record.",
        "criteria": [
            "AC-015-1: Data table with columns: Date, Subject, Client, Type, Status, Filed By.",
            "AC-015-2: Filters: date range, status, meeting reason, relationship.",
            "AC-015-3: Full-text search on subject and summary_of_discussion.",
            "AC-015-4: Export to CSV/Excel.",
            "AC-015-5: Clicking a row opens the call report detail view.",
        ],
        "rules": [
            "BR-015-1: RM sees own reports. Supervisor sees team reports. Branch Manager sees branch reports. Compliance/MIS sees all.",
        ],
        "ui": "Standard data table with filter panel (collapsible sidebar or top bar). Pagination: 20 rows/page. Export button in toolbar.",
        "edge": "If search returns 0 results, show 'No call reports match your filters' with a 'Clear filters' link.",
    },
    {
        "id": "FR-016",
        "name": "Meeting Notifications & Reminders",
        "desc": "The system sends automated notifications for meeting reminders, overdue call reports, and approval status changes.",
        "story": "As an RM, I want to receive reminders before meetings and alerts for overdue call reports so that I stay on top of my schedule.",
        "criteria": [
            "AC-016-1: Meeting reminder: 24 hours before, 1 hour before (configurable).",
            "AC-016-2: Overdue call report alert: sent daily at 9 AM local time for meetings > threshold days without a call report.",
            "AC-016-3: Approval submitted: in-app + email to the supervisors in the RM's hierarchy.",
            "AC-016-4: Approval decision: in-app + email to the RM.",
            "AC-016-5: Users can configure notification preferences (enable/disable per channel per event type).",
        ],
        "rules": [
            "BR-016-1: Notification preferences are stored per user.",
            "BR-016-2: Critical notifications (compliance-related) cannot be disabled.",
        ],
        "ui": "Bell icon in header with badge count. Notification center slide-out panel. Preferences page under Settings.",
        "edge": "If email delivery fails, retry up to 3 times with exponential backoff. Log failures for admin review.",
    },
    {
        "id": "FR-017",
        "name": "Supervisor Team Dashboard",
        "desc": "Supervisors see a dashboard showing their team's meeting activity, call report compliance, opportunity pipeline, and overdue items.",
        "story": "As a Supervisor, I want a dashboard showing my team's activity so that I can monitor productivity and compliance.",
        "criteria": [
            "AC-017-1: Widgets: Meetings This Week (bar chart by RM), Call Report Filing Rate (gauge), Opportunity Pipeline (funnel chart), Overdue Items (count cards).",
            "AC-017-2: Drill-down from any widget to the underlying data table.",
            "AC-017-3: Date range selector (default: current month).",
            "AC-017-4: RM filter to focus on a specific team member.",
        ],
        "rules": [
            "BR-017-1: Dashboard data refreshes every 5 minutes (or on manual refresh).",
            "BR-017-2: Branch Manager sees branch-level aggregates; Supervisor sees team-level.",
        ],
        "ui": "Grid layout with responsive cards. Charts use the shared charting library. Loading skeleton during data fetch.",
        "edge": "If a team has no data for the selected period, show 'No activity recorded' per widget.",
    },
]

for fr in frs:
    doc.add_heading(f"{fr['id']}: {fr['name']}", level=2)
    doc.add_paragraph(fr["desc"])
    p_story = doc.add_paragraph()
    run_label = p_story.add_run("User Story: ")
    run_label.bold = True
    p_story.add_run(fr["story"])

    doc.add_heading("Acceptance Criteria", level=3)
    for ac in fr["criteria"]:
        add_bullet(doc, ac)

    doc.add_heading("Business Rules", level=3)
    for br in fr["rules"]:
        add_bullet(doc, br)

    doc.add_heading("UI Behavior", level=3)
    doc.add_paragraph(fr["ui"])

    doc.add_heading("Edge Cases & Error Handling", level=3)
    doc.add_paragraph(fr["edge"])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 6: USER INTERFACE REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("6. User Interface Requirements", level=1)

screens = [
    {
        "name": "Calendar Main View",
        "purpose": "Primary landing page for RMs to view and manage their meeting schedule.",
        "layout": "Top bar: View toggle (All/Month/Week/Day), date navigation (< Today >), status legend, 'Schedule Meeting' button, filter icon. Main area: Calendar grid (month/week/day) or data table (all). Right sidebar (collapsible): selected meeting detail preview.",
        "components": "Calendar grid (react-big-calendar or custom), meeting cards with color-coded left border, filter panel (date range, reason, status dropdowns), data table with sortable columns, pagination.",
        "nav": "From: Main navigation sidebar. To: Meeting Detail, File Call Report, Schedule Meeting modal.",
        "responsive": "Desktop: full calendar grid. Tablet: week/day only. Mobile: list view (All Activities) as default; month shows dots per day.",
    },
    {
        "name": "Schedule Meeting Modal",
        "purpose": "Form to create or edit a meeting.",
        "layout": "Two-column layout. Left: Subject, Meeting Reason, Location, Start/End DateTime, All Day toggle, Mode of Meeting, Meeting Type, Relationship search, Contact details, Remarks. Right: Required Invitees (multi-select), Optional Invitees (multi-select). Bottom: Save, Cancel buttons.",
        "components": "Text inputs, select dropdowns, date-time pickers, toggle switch, type-ahead search with debounce, multi-select user picker with hierarchy tree, textarea.",
        "nav": "From: Calendar Main View (button). To: Calendar Main View (on save/cancel).",
        "responsive": "Desktop: two-column. Mobile: single-column stacked, invitee section collapsed by default.",
    },
    {
        "name": "Call Report Form",
        "purpose": "File a call report for a scheduled meeting or create a standalone report.",
        "layout": "Full page, vertical sections: 1) Meeting Info (read-only if linked), 2) Interaction Details (person met, client status, state of mind, summary), 3) Opportunities (repeater/accordion), 4) Expenses (repeater/accordion), 5) Action Items (repeater/accordion), 6) Attachments (upload zone), 7) Next Steps (next meeting date/time). Sticky footer: Save Draft, Submit.",
        "components": "Read-only display cards, text inputs, textareas with char counters, select dropdowns, repeater rows with add/remove, drag-and-drop file upload, date-time pickers.",
        "nav": "From: Calendar (File Call Report), Call Reports List. To: Calendar, Call Reports List.",
        "responsive": "Single-column on all devices. Accordion sections collapsed by default on mobile.",
    },
    {
        "name": "Supervisor Approval Workspace",
        "purpose": "Supervisors review, claim, approve, or reject late-filed call reports.",
        "layout": "Tabs: 'Pending Approval' (unclaimed), 'My Approvals' (claimed). Each tab: data table with columns: RM Name, Client, Subject, Filed Date, Days Late, Status, Actions (Claim/View/Approve/Reject). Detail panel: full call report view + approval action panel with comments textarea.",
        "components": "Tabbed interface, data tables with action buttons, detail slide-over panel, textarea for reviewer comments, Approve/Reject buttons.",
        "nav": "From: Main navigation sidebar (Supervisor menu). To: Call Report Detail (view).",
        "responsive": "Desktop: table + side panel. Mobile: table collapses to card list; detail opens as full-screen.",
    },
    {
        "name": "Opportunities List",
        "purpose": "View, filter, and manage all opportunities.",
        "layout": "Top bar: 'Add Opportunity' button, 'Upload' button, filters (status, stage, date range, relationship). Main: data table with columns: Date, Relationship, Sub Product, Amount Discovered, Amount Closed, Stage, Status, Actions.",
        "components": "Data table with sorting, pagination, filters, export button, upload modal.",
        "nav": "From: Main navigation sidebar. To: Opportunity detail/edit, Upload modal.",
        "responsive": "Desktop: full table. Mobile: card layout with key fields.",
    },
    {
        "name": "Expenses List",
        "purpose": "View, filter, and manage all expenses.",
        "layout": "Top bar: 'Add Expense' button, filters (expense type, date range). Main: data table with columns: Date, Type, Amount, Purpose, Client, Call Report, Actions.",
        "components": "Data table, filter dropdowns, add/edit modal.",
        "nav": "From: Main navigation sidebar. To: Expense add/edit modal.",
        "responsive": "Desktop: full table. Mobile: card layout.",
    },
    {
        "name": "Conversation History",
        "purpose": "View chronological interaction timeline for a specific client.",
        "layout": "Left: client info summary card. Main: vertical timeline with date separators, interaction type icons, summary cards with 'View Details' links. Top: filter bar (interaction type, date range).",
        "components": "Timeline component, filter bar, info card, lazy-load scroll.",
        "nav": "From: Customer Dashboard, Calendar (client filter). To: Meeting detail, Call Report detail, Opportunity detail.",
        "responsive": "Single-column timeline on all devices.",
    },
    {
        "name": "Notification Preferences",
        "purpose": "Users configure which notifications they receive and via which channels.",
        "layout": "Table: rows = event types (meeting reminder, overdue alert, approval request, approval decision), columns = channels (email, in-app, SMS). Checkboxes at intersections. 'Save' button.",
        "components": "Grid of checkboxes, save button, info tooltip per event type.",
        "nav": "From: Settings menu. To: Settings menu.",
        "responsive": "Horizontal scroll on mobile for the grid.",
    },
]

for s in screens:
    doc.add_heading(s["name"], level=2)
    p = doc.add_paragraph()
    run = p.add_run("Purpose: ")
    run.bold = True
    p.add_run(s["purpose"])

    p2 = doc.add_paragraph()
    run2 = p2.add_run("Layout: ")
    run2.bold = True
    p2.add_run(s["layout"])

    p3 = doc.add_paragraph()
    run3 = p3.add_run("Key Components: ")
    run3.bold = True
    p3.add_run(s["components"])

    p4 = doc.add_paragraph()
    run4 = p4.add_run("Navigation: ")
    run4.bold = True
    p4.add_run(s["nav"])

    p5 = doc.add_paragraph()
    run5 = p5.add_run("Responsive Behavior: ")
    run5.bold = True
    p5.add_run(s["responsive"])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 7: API & INTEGRATION REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("7. API & Integration Requirements", level=1)

doc.add_heading("7.1 Authentication", level=2)
doc.add_paragraph("All API endpoints require a valid JWT token passed in the Authorization header as 'Bearer <token>'. Tokens are issued by the Trust OMS authentication service (existing module). Token expiry: 8 hours. Refresh tokens: 30 days.")

doc.add_heading("7.2 Standardized Error Response", level=2)
doc.add_paragraph("All endpoints return errors in this format:")
doc.add_paragraph(
    '{\n'
    '  "error": {\n'
    '    "code": "VALIDATION_ERROR",\n'
    '    "message": "Summary of discussion is required and must be at least 20 characters.",\n'
    '    "field": "summary_of_discussion",\n'
    '    "details": []\n'
    '  }\n'
    '}'
)
doc.add_paragraph("Common error codes: VALIDATION_ERROR, NOT_FOUND, UNAUTHORIZED, FORBIDDEN, CONFLICT, INTERNAL_ERROR.")

doc.add_heading("7.3 API Endpoints", level=2)

endpoints = [
    ["GET", "/api/v1/meetings", "List meetings with filters (date range, status, reason). Paginated.", "200, 401"],
    ["POST", "/api/v1/meetings", "Create a new meeting.", "201, 400, 401"],
    ["GET", "/api/v1/meetings/:id", "Get meeting details by ID.", "200, 404, 401"],
    ["PUT", "/api/v1/meetings/:id", "Update meeting details.", "200, 400, 404, 401, 403"],
    ["PATCH", "/api/v1/meetings/:id/cancel", "Cancel a meeting.", "200, 400, 404, 401, 403"],
    ["GET", "/api/v1/call-reports", "List call reports with filters. Paginated.", "200, 401"],
    ["POST", "/api/v1/call-reports", "Create a new call report (scheduled or standalone).", "201, 400, 401"],
    ["GET", "/api/v1/call-reports/:id", "Get call report details.", "200, 404, 401"],
    ["PUT", "/api/v1/call-reports/:id", "Update a draft call report.", "200, 400, 404, 401, 403"],
    ["PATCH", "/api/v1/call-reports/:id/submit", "Submit a call report (draft → pending/completed/pending_approval).", "200, 400, 404, 401"],
    ["GET", "/api/v1/approvals", "List approval queue (supervisor). Filters: status, branch.", "200, 401, 403"],
    ["PATCH", "/api/v1/approvals/:id/claim", "Claim a pending approval.", "200, 400, 404, 401, 403"],
    ["PATCH", "/api/v1/approvals/:id/approve", "Approve a claimed call report.", "200, 400, 404, 401, 403"],
    ["PATCH", "/api/v1/approvals/:id/reject", "Reject a claimed call report. Body: { reviewer_comments }.", "200, 400, 404, 401, 403"],
    ["GET", "/api/v1/opportunities", "List opportunities with filters. Paginated.", "200, 401"],
    ["POST", "/api/v1/opportunities", "Create an opportunity.", "201, 400, 401"],
    ["PUT", "/api/v1/opportunities/:id", "Update an opportunity.", "200, 400, 404, 401, 403"],
    ["POST", "/api/v1/opportunities/upload", "Bulk upload opportunities (multipart/form-data).", "202, 400, 401"],
    ["GET", "/api/v1/opportunities/uploads/:id", "Get upload processing status.", "200, 404, 401"],
    ["GET", "/api/v1/expenses", "List expenses with filters. Paginated.", "200, 401"],
    ["POST", "/api/v1/expenses", "Create an expense.", "201, 400, 401"],
    ["PUT", "/api/v1/expenses/:id", "Update an expense.", "200, 400, 404, 401, 403"],
    ["POST", "/api/v1/feedback", "Submit feedback.", "201, 400, 401"],
    ["GET", "/api/v1/conversation-history/:relationshipId", "Get conversation history for a relationship. Paginated.", "200, 404, 401"],
    ["GET", "/api/v1/action-items", "List action items with filters.", "200, 401"],
    ["POST", "/api/v1/action-items", "Create an action item.", "201, 400, 401"],
    ["PATCH", "/api/v1/action-items/:id/complete", "Mark action item as completed.", "200, 404, 401"],
    ["POST", "/api/v1/call-reports/:id/attachments", "Upload attachment (multipart/form-data).", "201, 400, 413, 401"],
    ["DELETE", "/api/v1/call-reports/:id/attachments/:attachmentId", "Delete an attachment (draft only).", "204, 400, 404, 401, 403"],
    ["GET", "/api/v1/dashboard/supervisor", "Get supervisor dashboard data.", "200, 401, 403"],
    ["GET", "/api/v1/notifications", "List notifications for current user. Paginated.", "200, 401"],
    ["PATCH", "/api/v1/notifications/:id/read", "Mark notification as read.", "200, 404, 401"],
    ["GET", "/api/v1/notification-preferences", "Get notification preferences.", "200, 401"],
    ["PUT", "/api/v1/notification-preferences", "Update notification preferences.", "200, 400, 401"],
    ["GET", "/api/v1/system-config", "Get all config values (admin only).", "200, 401, 403"],
    ["PUT", "/api/v1/system-config/:key", "Update a config value (admin only).", "200, 400, 404, 401, 403"],
]

add_styled_table(doc,
    ["Method", "Endpoint", "Description", "Status Codes"],
    endpoints)

doc.add_heading("7.4 Request/Response Examples", level=2)

doc.add_heading("POST /api/v1/meetings — Create Meeting", level=3)
doc.add_paragraph("Request Body:")
doc.add_paragraph(
    '{\n'
    '  "subject": "Q2 Portfolio Review – Sharma Family",\n'
    '  "meeting_reason": "portfolio_review",\n'
    '  "location": "BKC Office, Conference Room 3",\n'
    '  "start_date_time": "2026-05-01T10:00:00Z",\n'
    '  "end_date_time": "2026-05-01T11:30:00Z",\n'
    '  "is_all_day": false,\n'
    '  "meeting_type": "cif",\n'
    '  "mode_of_meeting": "in_person",\n'
    '  "relationship_id": "rel-sharma-001",\n'
    '  "contact_phone": "+919876543210",\n'
    '  "contact_email": "sharma@example.com",\n'
    '  "remarks": "Bring updated portfolio statement",\n'
    '  "invitees": [\n'
    '    { "user_id": "rm-user-001", "invitee_type": "required" },\n'
    '    { "user_id": "sup-user-010", "invitee_type": "optional" }\n'
    '  ]\n'
    '}'
)

doc.add_paragraph("Success Response (201):")
doc.add_paragraph(
    '{\n'
    '  "data": {\n'
    '    "id": "mtg-abc123",\n'
    '    "subject": "Q2 Portfolio Review – Sharma Family",\n'
    '    "status": "scheduled",\n'
    '    "call_report_status": "pending",\n'
    '    "start_date_time": "2026-05-01T10:00:00Z",\n'
    '    "end_date_time": "2026-05-01T11:30:00Z",\n'
    '    "created_by": "rm-user-001",\n'
    '    "created_at": "2026-04-22T09:15:00Z",\n'
    '    "invitees": [\n'
    '      { "user_id": "rm-user-001", "invitee_type": "required", "rsvp_status": "accepted" },\n'
    '      { "user_id": "sup-user-010", "invitee_type": "optional", "rsvp_status": "pending" }\n'
    '    ]\n'
    '  }\n'
    '}'
)

doc.add_heading("POST /api/v1/call-reports — Create Call Report", level=3)
doc.add_paragraph("Request Body:")
doc.add_paragraph(
    '{\n'
    '  "meeting_id": "mtg-abc123",\n'
    '  "report_type": "scheduled",\n'
    '  "person_met": "Mr. Rajesh Sharma",\n'
    '  "client_status": null,\n'
    '  "state_of_mind": "happy",\n'
    '  "summary_of_discussion": "Reviewed Q1 portfolio performance. Client satisfied with 12% returns. Discussed increasing equity allocation from 40% to 55%. Client interested in Large Cap MF for additional INR 50L investment. Agreed to share product recommendations by next week.",\n'
    '  "remarks": null,\n'
    '  "next_meeting_start": "2026-05-08T10:00:00Z",\n'
    '  "next_meeting_end": "2026-05-08T11:00:00Z",\n'
    '  "opportunities": [\n'
    '    {\n'
    '      "opportunity_type": "cif",\n'
    '      "relationship_id": "rel-sharma-001",\n'
    '      "sub_product": "Equity MF – Large Cap",\n'
    '      "investment_mode": "lumpsum",\n'
    '      "opportunity_discovered": 5000000.00,\n'
    '      "opportunity_date": "2026-05-01",\n'
    '      "due_date": "2026-06-30",\n'
    '      "stage": "interested"\n'
    '    }\n'
    '  ],\n'
    '  "action_items": [\n'
    '    {\n'
    '      "title": "Send Large Cap MF product recommendations",\n'
    '      "assigned_to": "rm-user-001",\n'
    '      "due_date": "2026-05-07",\n'
    '      "priority": "high"\n'
    '    }\n'
    '  ]\n'
    '}'
)

doc.add_heading("GET /api/v1/meetings — List Meetings", level=3)
doc.add_paragraph("Query Parameters: ?start_date=2026-04-01&end_date=2026-04-30&status=scheduled&page=1&page_size=20")
doc.add_paragraph("Success Response (200):")
doc.add_paragraph(
    '{\n'
    '  "data": [ { "id": "mtg-abc123", "subject": "...", "status": "scheduled", ... } ],\n'
    '  "pagination": {\n'
    '    "page": 1,\n'
    '    "page_size": 20,\n'
    '    "total_count": 47,\n'
    '    "total_pages": 3\n'
    '  }\n'
    '}'
)

doc.add_heading("PATCH /api/v1/approvals/:id/reject — Reject Call Report", level=3)
doc.add_paragraph("Request Body:")
doc.add_paragraph(
    '{\n'
    '  "reviewer_comments": "Summary lacks detail on risk appetite discussion. Please add specifics about the client\'s current risk tolerance and how the proposed allocation change aligns with their investment policy statement."\n'
    '}'
)

doc.add_heading("7.5 External Integrations", level=2)
add_styled_table(doc,
    ["Service", "Purpose", "Protocol", "Notes"],
    [
        ["Trust OMS CRM", "Customer/Lead/Prospect master data lookup", "Internal REST API", "Existing; used for relationship type-ahead search"],
        ["Email (SMTP)", "Send notification emails", "SMTP / SES API", "Templates stored in notification_templates table"],
        ["S3-compatible Object Store", "File attachment storage", "S3 API", "Presigned URLs for upload/download"],
        ["Antivirus Scanner", "Scan uploaded files", "REST API / ClamAV", "Block file if infected"],
        ["Business Calendar Service", "Business-day calculations (holidays)", "Internal REST API", "Used for 5-day threshold calculation"],
    ])

doc.add_heading("7.6 Rate Limiting", level=2)
add_styled_table(doc,
    ["Endpoint Group", "Rate Limit", "Window"],
    [
        ["Read endpoints (GET)", "200 requests", "Per minute per user"],
        ["Write endpoints (POST/PUT/PATCH)", "50 requests", "Per minute per user"],
        ["File upload", "10 requests", "Per minute per user"],
        ["Bulk upload", "2 requests", "Per hour per user"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 8: NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("8. Non-Functional Requirements", level=1)

doc.add_heading("8.1 Performance", level=2)
add_styled_table(doc,
    ["Metric", "Target"],
    [
        ["Calendar month-view load time", "≤ 2 seconds for up to 200 meetings"],
        ["Call report form load (with pre-population)", "≤ 1.5 seconds"],
        ["API response time (read endpoints)", "P95 ≤ 300 ms"],
        ["API response time (write endpoints)", "P95 ≤ 500 ms"],
        ["Full-text search (call reports)", "P95 ≤ 1 second"],
        ["Concurrent users per instance", "≥ 500"],
        ["Bulk upload (500 rows)", "≤ 30 seconds end-to-end"],
    ])

doc.add_heading("8.2 Security", level=2)
security = [
    "Authentication: JWT tokens with httpOnly cookie storage (no localStorage).",
    "Authorization: Role-based access control (RBAC) enforced at API gateway and service layers.",
    "Data encryption: TLS 1.2+ in transit; AES-256 at rest for database and object storage.",
    "Input validation: Server-side validation on all endpoints; parameterized queries (no raw SQL).",
    "OWASP Top 10 compliance: XSS prevention (React auto-escaping + CSP headers), CSRF tokens, SQL injection prevention, rate limiting.",
    "Audit trail: All create/update/delete operations logged with user ID, timestamp, before/after values.",
    "Session management: 8-hour token expiry, sliding refresh, forced logout on password change.",
    "File upload security: MIME type validation, file extension whitelist, antivirus scan, max size enforcement.",
]
for s in security:
    add_bullet(doc, s)

doc.add_heading("8.3 Scalability", level=2)
doc.add_paragraph("Horizontal scaling: stateless API servers behind a load balancer. Database: read replicas for reporting queries. Object storage: S3 with CDN for attachment downloads. Expected growth: 50% YoY increase in meetings/call reports.")

doc.add_heading("8.4 Availability", level=2)
doc.add_paragraph("Target uptime: 99.9% (excluding planned maintenance windows). Planned maintenance: max 1 hour/month during off-hours. Database failover: automatic within 30 seconds.")

doc.add_heading("8.5 Data Backup & Recovery", level=2)
add_styled_table(doc,
    ["Parameter", "Value"],
    [
        ["Backup frequency", "Full daily, incremental every 6 hours"],
        ["Recovery Point Objective (RPO)", "6 hours"],
        ["Recovery Time Objective (RTO)", "1 hour"],
        ["Backup retention", "30 days for daily, 7 days for incremental"],
        ["Backup location", "Cross-region replication"],
    ])

doc.add_heading("8.6 Accessibility", level=2)
doc.add_paragraph("WCAG 2.1 Level AA compliance. All interactive elements keyboard-navigable. ARIA labels on custom components. Minimum contrast ratio: 4.5:1 for text, 3:1 for large text. Screen reader compatibility tested with NVDA and VoiceOver.")

doc.add_heading("8.7 Browser & Device Support", level=2)
add_styled_table(doc,
    ["Browser/Device", "Minimum Version"],
    [
        ["Chrome", "Last 2 major versions"],
        ["Firefox", "Last 2 major versions"],
        ["Safari", "Last 2 major versions"],
        ["Edge", "Last 2 major versions"],
        ["Mobile (responsive web)", "iOS Safari 15+, Chrome Android 100+"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 9: WORKFLOW & STATE DIAGRAMS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("9. Workflow & State Diagrams", level=1)

doc.add_heading("9.1 Meeting Lifecycle", level=2)
add_styled_table(doc,
    ["Current State", "Action", "Next State", "Side Effects"],
    [
        ["(New)", "RM creates meeting", "Scheduled", "ConversationHistory entry; invitee notifications sent"],
        ["Scheduled", "Meeting time passes + RM marks complete", "Completed", "ConversationHistory entry; call_report_status stays 'pending'"],
        ["Scheduled", "RM cancels meeting", "Cancelled", "ConversationHistory entry; invitee cancellation notifications"],
        ["Scheduled", "RM edits meeting", "Scheduled", "Updated_at set; if date changed, ConversationHistory 'rescheduled' entry"],
        ["Completed", "RM files call report within threshold", "Completed (call_report_status → filed)", "CallReport created; ConversationHistory entry"],
        ["Completed", "Threshold days pass without call report", "Completed (call_report_status → overdue)", "Overdue notification sent to RM"],
        ["Scheduled", "Meeting time passes + no action from RM", "No Show (after configurable grace period)", "Notification to RM and supervisor"],
    ])

doc.add_heading("9.2 Call Report Lifecycle", level=2)
add_styled_table(doc,
    ["Current State", "Action", "Next State", "Side Effects"],
    [
        ["(New)", "RM opens call report form", "Draft", "Record created in DB"],
        ["Draft", "RM saves progress", "Draft", "Updated_at set"],
        ["Draft", "RM submits (within threshold)", "Completed", "ConversationHistory entry; meeting.call_report_status → filed"],
        ["Draft", "RM submits (exceeds threshold)", "Pending Approval", "CallReportApproval created; supervisor notified; ConversationHistory entry"],
        ["Pending Approval", "Supervisor claims", "Pending Approval (approval.status → claimed)", "Locked from other supervisors"],
        ["Pending Approval", "Supervisor approves", "Approved", "RM notified; ConversationHistory entry"],
        ["Pending Approval", "Supervisor rejects", "Rejected", "RM notified with comments; ConversationHistory entry"],
        ["Rejected", "RM revises and re-submits", "Pending Approval", "New CallReportApproval record; supervisor notified"],
    ])

doc.add_heading("9.3 Opportunity Lifecycle", level=2)
add_styled_table(doc,
    ["Current State", "Action", "Next State", "Side Effects"],
    [
        ["(New)", "RM creates opportunity", "Open (stage: Interested)", "ConversationHistory entry"],
        ["Open", "RM updates stage", "Open (new stage)", "ConversationHistory entry"],
        ["Open", "Stage → Won – Doc completed", "Closed", "opportunity_closed becomes mandatory; ConversationHistory entry"],
        ["Open", "Stage → Declined / Not Proceeding", "Aborted", "ConversationHistory entry"],
        ["Open", "due_date < current_business_date (batch)", "Expired", "Batch notification to RM"],
    ])

doc.add_heading("9.4 Approval Lifecycle (Supervisor)", level=2)
add_styled_table(doc,
    ["Current State", "Action", "Next State", "Side Effects"],
    [
        ["Pending", "Supervisor claims", "Claimed", "Locked from other supervisors; claimed_at set"],
        ["Claimed", "Supervisor approves", "Approved", "decided_at set; call_report.status → approved; RM notified"],
        ["Claimed", "Supervisor rejects", "Rejected", "decided_at set; reviewer_comments saved; call_report.status → rejected; RM notified"],
        ["Claimed", "2 business days pass without action", "Pending", "Auto-unclaimed; supervisor notified"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 10: NOTIFICATION & COMMUNICATION
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("10. Notification & Communication Requirements", level=1)

add_styled_table(doc,
    ["Event", "Channel", "Recipient", "Trigger Condition", "Message Template"],
    [
        ["Meeting Reminder (24h)", "Email + In-App", "RM + Invitees", "24 hours before meeting.start_date_time", "Subject: Reminder — {meeting.subject} tomorrow at {time}. Body: Meeting with {relationship_name} at {location}. [View Details]"],
        ["Meeting Reminder (1h)", "In-App + Push", "RM + Invitees", "1 hour before meeting.start_date_time", "Your meeting '{subject}' starts in 1 hour at {location}."],
        ["Overdue Call Report", "Email + In-App", "RM", "Daily 9 AM if meeting > threshold days AND call_report_status = 'pending'", "Subject: Overdue call report for {meeting.subject}. Body: Your meeting on {date} is {N} days overdue for a call report. File now to avoid supervisor review."],
        ["Approval Required", "Email + In-App", "Supervisors in branch", "Call report submitted with requires_approval = true", "Subject: Call report approval required — {rm_name}. Body: {rm_name} has filed a call report {days_late} days late for meeting '{subject}'. [Review Now]"],
        ["Approval Decision", "Email + In-App", "RM", "Supervisor approves or rejects", "Subject: Call report {approved/rejected} — {subject}. Body: Your call report has been {decision}. {If rejected: Reviewer comments: {comments}. Please revise and resubmit.}"],
        ["Meeting Cancelled", "Email + In-App", "Invitees", "Meeting status → cancelled", "Subject: Meeting cancelled — {subject}. Body: The meeting on {date} at {location} has been cancelled by {created_by.name}."],
        ["Action Item Assigned", "In-App", "Assigned user", "Action item created", "{creator_name} assigned you an action item: '{title}' due {due_date}."],
        ["Action Item Overdue", "Email + In-App", "Assigned user", "Daily 9 AM if due_date < today AND status != completed", "Action item '{title}' is overdue. Due date was {due_date}. [Mark Complete]"],
        ["Opportunity Expired", "Email", "RM", "Batch job expires opportunity", "Subject: Opportunity expired — {sub_product} for {relationship_name}. Body: Amount: {amount}. The opportunity was not closed by the due date {due_date}."],
        ["Bulk Upload Complete", "In-App", "RM", "OpportunityUpload.status → completed", "Your opportunity upload is complete. {success_count} imported, {error_count} failed. [View Details]"],
    ])

doc.add_heading("10.1 Notification Preferences", level=2)
doc.add_paragraph(
    "Users can opt out of non-critical notifications (meeting reminders, action item alerts) per channel. "
    "Critical notifications — overdue call report alerts, approval requests, and approval decisions — cannot be disabled. "
    "Preferences are stored in a user_notification_preferences table with columns: user_id, event_type, channel, enabled (boolean)."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 11: REPORTING & ANALYTICS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("11. Reporting & Analytics", level=1)

add_styled_table(doc,
    ["Report Name", "Audience", "Data Sources", "Key Metrics", "Filters", "Refresh"],
    [
        ["RM Activity Summary", "Supervisor, Branch Manager", "meetings, call_reports", "Meetings per RM (weekly/monthly), call report filing rate, avg turnaround time", "Date range, RM, branch", "Real-time"],
        ["Call Report Compliance", "Compliance Officer", "call_reports, call_report_approvals", "Filing rate, overdue count, approval rate, rejection reasons", "Date range, branch, status", "Daily"],
        ["Opportunity Pipeline", "Supervisor, Branch Manager", "opportunities", "Total discovered amount, pipeline by stage (funnel), conversion rate, average deal size", "Date range, RM, status, stage", "Real-time"],
        ["Expense Summary", "MIS/Ops, Branch Manager", "expenses", "Total expense by type, per-RM expense trends, top expense categories", "Date range, expense type, RM", "Weekly"],
        ["Meeting Coverage Heatmap", "Branch Manager", "meetings", "Meetings per client (color intensity), clients with no meetings in 90 days", "Date range, branch", "Weekly"],
        ["Supervisor Approval Metrics", "Branch Manager", "call_report_approvals", "Avg approval turnaround, pending queue size, approval vs rejection ratio", "Date range, supervisor", "Real-time"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 12: MIGRATION & LAUNCH PLAN
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("12. Migration & Launch Plan", level=1)

doc.add_heading("12.1 Data Migration", level=2)
doc.add_paragraph("If replacing an existing system (spreadsheets, legacy CRM), the following data must be migrated:")
migrations = [
    "Historical meetings: Import from CSV with mapping to the Meeting entity schema.",
    "Existing call reports: Map to CallReport entity; set status = 'completed' and requires_approval = false for all historical records.",
    "Client relationships: Already in Trust OMS CRM; no migration needed.",
    "Opportunity data: Import via the bulk upload feature (FR-008).",
]
for m in migrations:
    add_bullet(doc, m)

doc.add_heading("12.2 Phased Rollout", level=2)
add_styled_table(doc,
    ["Phase", "Features", "Timeline", "Notes"],
    [
        ["Phase 1 (MVP)", "Calendar views, schedule meeting, file call report (scheduled + standalone), basic call report list", "Sprint 1–3", "Core functionality for RM daily workflow"],
        ["Phase 2", "5-day rule + supervisor approval workflow, conversation history, action items, attachments", "Sprint 4–5", "Compliance and workflow features"],
        ["Phase 3", "Opportunity management (CRUD + bulk upload + auto-expiry), expense capture", "Sprint 6–7", "Sales pipeline and expense tracking"],
        ["Phase 4", "Notification system, reporting dashboards, notification preferences", "Sprint 8–9", "Automation and analytics"],
        ["Phase 5", "Performance optimization, mobile responsiveness polish, accessibility audit", "Sprint 10", "Production hardening"],
    ])

doc.add_heading("12.3 Go-Live Checklist", level=2)
checklist = [
    "All Phase 1–4 features pass UAT with ≥ 95% test case pass rate.",
    "Performance benchmarks met (Section 8.1).",
    "Security audit completed; all critical/high findings resolved.",
    "Data migration validated: imported records match source count ± 0.",
    "Notification templates reviewed and approved by stakeholders.",
    "User training completed for RM, Supervisor, and Compliance roles.",
    "Rollback plan documented and tested.",
    "Monitoring and alerting configured (error rates, response times, queue depths).",
]
for c in checklist:
    add_bullet(doc, c)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 13: GLOSSARY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("13. Glossary", level=1)

add_styled_table(doc,
    ["Term", "Definition"],
    [
        ["RM", "Relationship Manager — front-line wealth management advisor who manages client relationships."],
        ["Call Report", "A structured record of an RM-client interaction, documenting the discussion, outcomes, and follow-ups."],
        ["Standalone Call Report", "A call report filed without a linked scheduled meeting, for ad-hoc interactions."],
        ["5-Day Threshold Rule", "Business rule requiring supervisor approval if a call report is filed more than 5 business days after the meeting."],
        ["CIF", "Customer Information File — a unique identifier for an existing bank customer."],
        ["Lead", "A potential customer who has been identified but not yet qualified."],
        ["Prospect", "A qualified lead who has shown interest and is being actively pursued."],
        ["State of Mind", "RM's assessment of the client's emotional state during the interaction (Happy, Irate, Satisfied, Sensitive, Ultra-Sensitive)."],
        ["Opportunity", "A potential sales deal identified during a client interaction, with an expected value and close date."],
        ["Pipeline", "The aggregate of all open opportunities, typically visualized as a funnel by stage."],
        ["Conversation History", "A chronological, immutable record of all interactions with a specific client relationship."],
        ["Action Item", "A follow-up task generated from a call report, assigned to a user with a due date."],
        ["Business Day", "A working day excluding weekends and public holidays as defined in the business calendar."],
        ["ConversationHistory Entry", "An immutable log record auto-created when a significant event occurs (meeting, call report, feedback, opportunity)."],
        ["RSVP", "Response to a meeting invitation: Accepted, Declined, Tentative, or Pending."],
        ["MIS", "Management Information Systems — the team responsible for generating business reports and analytics."],
        ["Trust OMS", "Trust Order Management System — the broader platform that this module integrates with."],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# SECTION 14: APPENDICES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("14. Appendices", level=1)

doc.add_heading("14.1 Meeting Reason Values", level=2)
add_styled_table(doc,
    ["Code", "Display Label"],
    [
        ["client_call_new", "Client Call – New"],
        ["client_call_existing", "Client Call – Existing"],
        ["branch_visit", "Branch Visit"],
        ["portfolio_review", "Portfolio Review"],
        ["service_request", "Service Request"],
        ["campaign_details", "Campaign Details"],
        ["others", "Others (requires free-text specification)"],
    ])

doc.add_heading("14.2 Meeting Mode Values", level=2)
add_styled_table(doc,
    ["Code", "Display Label"],
    [
        ["face_to_face", "Face-to-Face"],
        ["in_person", "In-Person"],
        ["in_person_offshore", "In-Person (Off-Shore)"],
        ["telephone", "Telephone"],
        ["telephone_offshore", "Telephone (Off-Shore)"],
        ["video_conference", "Video Conference"],
        ["video_conference_offshore", "Video Conference (Off-Shore)"],
        ["others", "Others"],
    ])

doc.add_heading("14.3 Calendar Color Legend", level=2)
add_styled_table(doc,
    ["Status", "Color", "Hex Code"],
    [
        ["Scheduled", "Blue", "#3B82F6"],
        ["Completed", "Green", "#10B981"],
        ["Cancelled", "Grey", "#9CA3AF"],
        ["No Show", "Red", "#EF4444"],
        ["Overdue Call Report", "Orange", "#F97316"],
    ])

doc.add_heading("14.4 Opportunity Stage Progression", level=2)
doc.add_paragraph("Typical (but not enforced) stage progression:")
add_styled_table(doc,
    ["Order", "Stage", "Description"],
    [
        ["1", "Interested", "Client has expressed interest in a product/service."],
        ["2", "To be approached", "RM plans to approach client with a specific offering."],
        ["3", "Won – Doc in Process", "Client has agreed; documentation is being prepared."],
        ["4", "Won – Doc completed", "All documentation completed; deal closed."],
        ["—", "Declined", "Client declined the offering."],
        ["—", "Not Proceeding", "RM/bank decided not to proceed."],
    ])

doc.add_heading("14.5 Opportunity Bulk Upload Template Columns", level=2)
add_styled_table(doc,
    ["Column", "Data Type", "Required", "Notes"],
    [
        ["opportunity_type", "String", "Yes", "cif, lead, or prospect"],
        ["relationship_id", "String", "Yes", "Must exist in CRM"],
        ["sub_product", "String", "Yes", "Product name"],
        ["investment_mode", "String", "No", "lumpsum or periodic"],
        ["opportunity_discovered", "Number", "Yes", "> 0"],
        ["currency", "String", "No", "Default: INR"],
        ["opportunity_date", "Date (YYYY-MM-DD)", "Yes", "≤ today"],
        ["due_date", "Date (YYYY-MM-DD)", "Yes", "≥ opportunity_date"],
        ["description", "String", "No", "Max 2000 chars"],
        ["stage", "String", "No", "Default: interested"],
    ])

doc.add_heading("14.6 Reference Documents", level=2)
refs = [
    "SBI — Journey of RM from Customer Meeting to Closure of Action Items V10 (42 pages)",
    "HSBC — Journey of RM from Customer Meeting to Closure of Action Items V11 (40+ pages)",
    "Jio BlackRock — Journey of RM from Customer Meeting to Closure of Action Items V11 (40+ pages)",
    "SBI — Opportunity Management V10 (16 pages)",
    "SBI — Expense Capture V10 (13 pages)",
]
for r in refs:
    add_bullet(doc, r)

# ── Save ─────────────────────────────────────────────────────────────
output_path = "/Users/n15318/Trust OMS/docs/Calendar_CallReport_Management_BRD_v1.docx"
doc.save(output_path)
print(f"BRD saved to: {output_path}")
