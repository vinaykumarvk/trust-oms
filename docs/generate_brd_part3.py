#!/usr/bin/env python3
"""Generate BRD Part 3: Section 5 - Functional Requirements"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document('/Users/n15318/Trust OMS/docs/CIM_BRD_Draft.docx')

def add_fr(doc, fr_id, name, desc, story, criteria, rules, ui_notes, edge_cases):
    doc.add_heading(f'{fr_id}: {name}', level=3)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    r = p.add_run('User Story: ')
    r.bold = True
    p.add_run(story)
    
    p = doc.add_paragraph()
    r = p.add_run('Acceptance Criteria:')
    r.bold = True
    for c in criteria:
        doc.add_paragraph(c, style='List Bullet')
    
    p = doc.add_paragraph()
    r = p.add_run('Business Rules:')
    r.bold = True
    for rule in rules:
        doc.add_paragraph(rule, style='List Bullet')
    
    p = doc.add_paragraph()
    r = p.add_run('UI Behavior:')
    r.bold = True
    for note in ui_notes:
        doc.add_paragraph(note, style='List Bullet')
    
    p = doc.add_paragraph()
    r = p.add_run('Edge Cases & Error Handling:')
    r.bold = True
    for ec in edge_cases:
        doc.add_paragraph(ec, style='List Bullet')

doc.add_heading('5. Functional Requirements', level=1)

# --- Module A: Calendar Management ---
doc.add_heading('5.1 Calendar Management', level=2)

add_fr(doc, 'FR-001', 'Calendar View - Month',
    'Display a monthly calendar grid showing all meetings for the selected month. Each day cell shows meeting count and first 2-3 meeting subjects. Clicking a day navigates to day view. The current day is highlighted. Navigation arrows allow moving between months.',
    'As a Relationship Manager, I want to see my monthly calendar at a glance so that I can plan my schedule and identify open slots for new client meetings.',
    [
        'AC-001: Calendar displays 42 cells (6 weeks x 7 days) with correct day numbers for the selected month',
        'AC-002: Each day cell shows a count badge (e.g., "3 meetings") and the first 2 meeting subjects truncated to 30 chars',
        'AC-003: Clicking a day cell navigates to Day view for that date',
        'AC-004: Left/Right arrows navigate to previous/next month; clicking month/year header opens month picker',
        'AC-005: Meetings from the logged-in user (as organizer or attendee) are displayed; team meetings shown if SENIOR_RM/BO_HEAD',
    ],
    [
        'BR-001: Default view is Month for current month on page load',
        'BR-002: Weekend days (Sat/Sun) are visually distinguished with lighter background',
        'BR-003: Public holidays from market_calendar table are shown as badges on relevant days',
        'BR-004: Maximum 500 meetings fetched per month view query; paginate if exceeded',
    ],
    [
        'Month view loads within 2 seconds for 100 meetings',
        'Clicking a meeting subject opens a popover with meeting details and "View Full" link',
        'Empty days show no badge; hovering shows "No meetings scheduled"',
        'Current day has a blue circle highlight on the day number',
    ],
    [
        'If API returns error, show toast "Unable to load calendar. Please try again." with retry button',
        'If user has no meetings for the month, show empty state: "No meetings scheduled this month. Click + to create one."',
        'If date range spans DST change, times are correctly displayed in local timezone',
    ]
)

add_fr(doc, 'FR-002', 'Calendar View - Week',
    'Display a weekly view with 7 columns (Mon-Sun) and hourly time slots (7:00 AM to 8:00 PM). Meetings appear as colored blocks spanning their duration. Overlapping meetings are displayed side-by-side within the same time slot.',
    'As a Relationship Manager, I want to see my weekly schedule in a timeline format so that I can identify time conflicts and available slots.',
    [
        'AC-006: Week view shows 7 columns with 13 hourly rows (7 AM to 8 PM)',
        'AC-007: Meeting blocks are colored by meeting_reason (e.g., blue for Portfolio Review, green for Client Call)',
        'AC-008: Overlapping meetings render side-by-side (50% width each for 2 overlaps, 33% for 3)',
        'AC-009: Clicking an empty time slot opens the "Create Meeting" form pre-filled with that date/time',
        'AC-010: Dragging a meeting block to a new time slot triggers reschedule confirmation',
    ],
    [
        'BR-005: Week starts on Monday per Philippine business convention',
        'BR-006: All-day meetings appear in a separate row above the time grid',
        'BR-007: Cancelled meetings are shown with strikethrough text and 50% opacity',
    ],
    [
        'Hover on meeting block shows tooltip: subject, time, location, client name',
        'Right-click on meeting block shows context menu: View, Edit, Reschedule, Cancel',
        'Week navigation via left/right arrows; "Today" button returns to current week',
    ],
    [
        'If more than 5 meetings overlap in one slot, show "+N more" overflow indicator',
        'Meetings starting before 7 AM or ending after 8 PM are clipped with scroll indicator',
    ]
)

add_fr(doc, 'FR-003', 'Calendar View - Day',
    'Display a single day timeline with hourly slots showing all meetings. Each meeting card shows subject, time, client name, location, and status badge. Action buttons for each meeting: File Call Report, Reschedule, Cancel.',
    'As a Relationship Manager, I want to see my daily schedule in detail so that I can prepare for each meeting and quickly file call reports afterward.',
    [
        'AC-011: Day view shows hourly slots from 7:00 AM to 8:00 PM with meetings as cards',
        'AC-012: Each meeting card displays: subject, time range, client name, location, mode icon, status badge',
        'AC-013: Completed meetings show "File Call Report" button if no call report exists yet',
        'AC-014: Meeting cards show attendee avatars (first 3 + "+N" overflow)',
        'AC-015: Day navigation via left/right arrows and date picker',
    ],
    [
        'BR-008: Meetings past their end_date automatically transition to COMPLETED status on next page load',
        'BR-009: "File Call Report" button is only visible for meetings where status = COMPLETED and no linked call_report exists',
    ],
    [
        'Clicking "File Call Report" navigates to call report form pre-populated with meeting data',
        'Empty slots show a dashed border with "+" icon on hover for quick meeting creation',
    ],
    [
        'If a meeting spans midnight, it appears on both days with a continuation indicator',
        'If the day has no meetings, show illustration with "Your day is clear! Schedule a meeting?"',
    ]
)

add_fr(doc, 'FR-004', 'Calendar View - All (List)',
    'Display all meetings as a sortable, filterable table/list without date constraint. Supports search, status filter, meeting type filter, and pagination. Useful for finding past meetings and reviewing history.',
    'As a Relationship Manager, I want to search and filter all my meetings across any date range so that I can find specific past interactions quickly.',
    [
        'AC-016: Table columns: Date, Subject, Client, Type, Mode, Status, Location, Attendees count',
        'AC-017: Search bar filters on subject, client name, and location (ILIKE)',
        'AC-018: Filter dropdowns for: status (multi-select), meeting_reason, meeting_type, date range picker',
        'AC-019: Default sort: start_date descending (newest first)',
        'AC-020: Pagination with 25 records per page, total count displayed',
    ],
    [
        'BR-010: SENIOR_RM sees own meetings + team members meetings. BO_HEAD sees all branch meetings.',
        'BR-011: Deleted (cancelled) meetings are hidden by default but visible with "Show cancelled" toggle',
    ],
    [
        'Clicking a row expands an inline detail panel or navigates to meeting detail page',
        'Bulk actions: Export to CSV (BO_HEAD and above)',
    ],
    [
        'If search returns no results, show "No meetings match your filters" with clear filters link',
        'Export CSV with >10,000 rows triggers async generation with notification on completion',
    ]
)

add_fr(doc, 'FR-005', 'Create Meeting',
    'Multi-field form to schedule a new meeting. Fields are organized in two sections: Meeting Details (subject, reason, location, dates, type, mode) and Participants (client/prospect selection, attendees). Form validation ensures all mandatory fields are filled before submission.',
    'As a Relationship Manager, I want to schedule a new client meeting with all relevant details so that it appears on my calendar and invitees are notified.',
    [
        'AC-021: Form displays all fields defined in the meetings data model with correct input types',
        'AC-022: Meeting_type selection drives conditional fields: CIF shows client search, LEAD/PROSPECT shows prospect search, OTHERS shows free-text name field',
        'AC-023: Client search uses combobox with autocomplete against clients table (legal_name, client_id)',
        'AC-024: When client is selected, contact_phone and contact_email auto-populate from client profile',
        'AC-025: Attendee selection allows multi-select from users table with type toggle (Required/Optional)',
        'AC-026: On successful submission: meeting created, attendees added, reminders scheduled, notification sent to all attendees',
    ],
    [
        'BR-012: If is_all_day is true, start_time and end_time fields are hidden and set to 00:00/23:59',
        'BR-013: End date/time must be after start date/time; form shows inline error if violated',
        'BR-014: Meeting subject auto-generates suggestion: "[meeting_reason] - [client_name]" but is editable',
        'BR-015: Overlapping meetings are ALLOWED but a warning toast is shown: "You have another meeting at this time"',
        'BR-016: meeting_ref auto-generated server-side on creation',
    ],
    [
        'Submit button shows loading spinner; disabled to prevent double-submission',
        'On success: toast "Meeting scheduled successfully", redirect to calendar day view for meeting date',
        'On validation error: inline red text under each invalid field; form does not submit',
        'Cancel button shows confirmation if any field has been modified: "Discard changes?"',
    ],
    [
        'If client_id references a SUSPENDED or CLOSED client, show warning: "Client status is [status]. Proceed?"',
        'If selected date is a market holiday, show info: "Note: [date] is a holiday ([holiday_name])"',
        'If more than 20 attendees are added, show warning: "Large meeting. Consider using a conference link."',
        'Network failure during submission: show error toast with "Retry" button; form data preserved',
    ]
)

add_fr(doc, 'FR-006', 'Reschedule Meeting',
    'Allow the meeting organizer to change the date, time, or location of a scheduled meeting. All attendees receive a notification about the change. The original meeting is marked as RESCHEDULED and a new meeting record is created with parent_meeting_id linking to the original.',
    'As a Relationship Manager, I want to reschedule a meeting when plans change so that all attendees are informed of the new date/time.',
    [
        'AC-027: Reschedule option available for meetings with status = SCHEDULED only',
        'AC-028: Reschedule form pre-fills current values; only date, time, and location are editable',
        'AC-029: On submit: original meeting status changes to RESCHEDULED, new meeting created with parent_meeting_id reference',
        'AC-030: All attendees receive in-app + email notification: "Meeting [subject] has been rescheduled from [old_date] to [new_date]"',
        'AC-031: Reschedule history is visible on the meeting detail page showing original and all rescheduled versions',
    ],
    [
        'BR-017: Only the organizer_id user can reschedule. Others see the option greyed out with tooltip "Only the organizer can reschedule"',
        'BR-018: Rescheduling to a past date is not allowed',
    ],
    [
        'Confirmation dialog: "Reschedule meeting to [new_date]? All [N] attendees will be notified."',
        'After rescheduling, calendar view refreshes to show updated meeting',
    ],
    [
        'If notification delivery fails for some attendees, log the failure but complete the reschedule; show warning: "Meeting rescheduled but [N] attendees could not be notified"',
    ]
)

add_fr(doc, 'FR-007', 'Cancel Meeting',
    'Allow the meeting organizer to cancel a scheduled meeting. All attendees receive a cancellation notification. The meeting status changes to CANCELLED but the record is retained for audit purposes (not soft-deleted).',
    'As a Relationship Manager, I want to cancel a meeting that is no longer needed so that attendees are freed up and the calendar is accurate.',
    [
        'AC-032: Cancel option available for meetings with status = SCHEDULED or RESCHEDULED only',
        'AC-033: Cancel requires a reason (text field, min 10 chars)',
        'AC-034: On confirmation: meeting status = CANCELLED, cancellation notification sent to all attendees',
        'AC-035: Cancelled meetings appear with strikethrough styling on calendar views',
    ],
    [
        'BR-019: Only the organizer can cancel. SENIOR_RM and BO_HEAD can also cancel meetings for their team.',
        'BR-020: Cancelling a meeting with an existing call report is blocked: "Cannot cancel — a call report has been filed for this meeting"',
    ],
    [
        'Confirmation dialog with reason field: "Are you sure you want to cancel this meeting? This cannot be undone."',
    ],
    [
        'If meeting is within 1 hour, show urgent warning: "This meeting starts in less than 1 hour. Cancel anyway?"',
    ]
)

doc.add_page_break()

# --- Module B: Call Report Management ---
doc.add_heading('5.2 Call Report Management', level=2)

add_fr(doc, 'FR-008', 'File Scheduled Call Report',
    'Create a call report for a previously scheduled meeting. Form is pre-populated with meeting data (subject, reason, dates, location, type, mode, client, attendees). RM completes the meeting information section: client status, state of mind, discussion summary, products discussed, next meeting details, and action items.',
    'As a Relationship Manager, I want to file a call report for a completed meeting so that the interaction is documented and available for supervisor review and future reference.',
    [
        'AC-036: Call report form is accessible from: (1) Day view "File Call Report" button, (2) Meeting detail page, (3) Call Reports list "New Scheduled" button',
        'AC-037: Meeting Particulars section is pre-populated and read-only (subject, reason, dates, location, type, mode)',
        'AC-038: Meeting Information section requires: client_relationship_status, state_of_mind, summary_of_discussion (min 50 chars)',
        'AC-039: Products Discussed allows adding multiple rows: {product_type dropdown, product_name text, interest_level: HIGH/MEDIUM/LOW}',
        'AC-040: Next Meeting section allows scheduling follow-up: date, start time, end time (creates new meeting record on submit)',
        'AC-041: Attendees section shows pre-populated attendees with option to add external attendees (name, organization, role)',
    ],
    [
        'BR-021: Call report can only be filed after the meeting end_date has passed. If meeting is still in future, button shows "Available after meeting"',
        'BR-022: If call report is filed more than 48 hours after meeting_date, is_late_filed = true and status requires supervisor approval before becoming APPROVED',
        'BR-023: If call report is filed more than 5 business days after meeting, system generates a SYSTEM_GENERATED task for BO_HEAD to review',
        'BR-024: report_ref is auto-generated server-side: CR-YYYYMMDD-XXXX',
        'BR-025: A meeting can have at most one call report. Attempting to create a second shows error: "A call report already exists for this meeting"',
    ],
    [
        'Form has "Save as Draft" and "Submit" buttons. Draft saves without validation; Submit enforces all mandatory fields',
        'On Submit: toast "Call report submitted for review", status changes to SUBMITTED',
        'Progress indicator shows: "Meeting Particulars > Meeting Info > Products > Next Meeting > Attendees > Review"',
    ],
    [
        'If the linked meeting was cancelled after the call report was drafted, show warning: "The linked meeting has been cancelled. You can still submit this report as a standalone."',
        'If client_id references a recently merged client, show info banner: "Client profile was updated on [date]"',
    ]
)

add_fr(doc, 'FR-009', 'File Standalone Call Report',
    'Create a call report for an unscheduled/ad-hoc client interaction (walk-in, unexpected phone call, informal meeting). All fields must be manually entered since there is no linked meeting. Accessible from Customer Dashboard or Calendar action menu.',
    'As a Relationship Manager, I want to file a call report for an unplanned client interaction so that every touchpoint is captured regardless of whether it was pre-scheduled.',
    [
        'AC-042: Standalone call report form shows all fields as editable (no pre-population)',
        'AC-043: Meeting Particulars requires: subject, meeting_reason, meeting_date (defaults to today), location, meeting_type, mode_of_meeting',
        'AC-044: Client selection is required for CIF type; shows combobox search against clients table',
        'AC-045: All Meeting Information fields are mandatory: client_relationship_status, state_of_mind, summary_of_discussion',
        'AC-046: report_type is set to STANDALONE automatically',
    ],
    [
        'BR-026: Standalone reports do not require a linked meeting_id',
        'BR-027: meeting_date for standalone reports can be backdated up to 30 days. Beyond 30 days, BO_HEAD approval is required.',
        'BR-028: BRANCH_ASSOCIATE can file standalone reports only (no scheduled reports)',
    ],
    [
        'Accessible from: Calendar top-right "+ Standalone" button, Client 360 page "Log Interaction" button',
        'Same form layout and validation as scheduled, but all fields editable',
    ],
    [
        'If client is selected but their assigned RM is different from the filing user, show info: "Note: You are not the assigned RM for this client. The assigned RM ([name]) will be notified."',
    ]
)

add_fr(doc, 'FR-010', 'Call Report Approval Workflow',
    'Submitted call reports enter a review queue visible to supervisors (BO_HEAD). Supervisors can approve, reject (with reason), or request revisions. Approved reports are locked from further editing. Rejected reports return to DRAFT status for RM correction.',
    'As a Supervisor (BO_HEAD), I want to review and approve submitted call reports so that I can ensure quality, compliance, and provide coaching feedback to my team.',
    [
        'AC-047: Approval queue page shows all SUBMITTED and UNDER_REVIEW call reports for the supervisor\'s team',
        'AC-048: Queue is sortable by: submitted_at, RM name, client name, is_late_filed flag',
        'AC-049: Late-filed reports are highlighted with amber badge and sorted to top by default',
        'AC-050: Supervisor can: Approve (status -> APPROVED), Reject (status -> REJECTED, requires rejection_reason min 20 chars), Request Info (status -> UNDER_REVIEW, sends notification to RM)',
        'AC-051: Supervisor can assign quality_score (1-5 stars) on approval',
        'AC-052: SLA indicator shows time since submission: green (<24h), yellow (24-48h), red (>48h)',
    ],
    [
        'BR-029: Only BO_HEAD role can approve/reject call reports. SENIOR_RM can provide feedback but not change status.',
        'BR-030: Approved reports are immutable — no further edits allowed. Only metadata (feedback, quality_score) can be added.',
        'BR-031: Rejected reports revert to DRAFT status and the RM receives notification: "Call report [ref] requires revision: [rejection_reason]"',
        'BR-032: If a report remains in SUBMITTED status for >48 hours without review, system sends reminder to BO_HEAD',
    ],
    [
        'Approval queue uses card layout: each card shows report_ref, RM name, client, submission date, late flag, SLA indicator',
        'Clicking a card expands full report detail inline (no page navigation)',
        'Approve/Reject buttons at bottom of expanded card with quality score stars',
    ],
    [
        'If supervisor tries to approve their own call report (if they also file reports), block with: "You cannot approve your own call report. Another supervisor must review."',
        'If call report was submitted for a client the supervisor does not manage, show info: "This client belongs to [other branch]. Cross-branch approval."',
    ]
)

add_fr(doc, 'FR-011', 'Call Report Feedback',
    'Supervisors and senior RMs can add feedback comments to call reports at any status. Feedback supports general comments, coaching notes (private), compliance flags, and quality observations. All feedback is timestamped and attributed.',
    'As a Supervisor, I want to provide coaching feedback on my team\'s call reports so that RMs improve their client interaction documentation and skills.',
    [
        'AC-053: Feedback section appears on every call report detail view as a chronological thread',
        'AC-054: Feedback form has: type dropdown (General, Coaching, Compliance Flag, Quality Issue), comment text area, is_private checkbox',
        'AC-055: Private feedback (is_private=true) is only visible to the feedback author and the RM who filed the report',
        'AC-056: COMPLIANCE_FLAG type feedback sends notification to COMPLIANCE_OFFICER role',
        'AC-057: Feedback thread shows avatar, name, role, timestamp, and type badge for each entry',
    ],
    [
        'BR-033: Feedback can be added at any call report status (including APPROVED)',
        'BR-034: Feedback cannot be edited or deleted after creation (immutable audit trail)',
        'BR-035: RM receives notification for each new feedback: "New [type] feedback on call report [ref] from [supervisor_name]"',
    ],
    [
        'Feedback appears in a chat-like thread below the call report details',
        'Supervisor feedback has blue background; compliance feedback has red border',
    ],
    [
        'If feedback text exceeds 2000 chars, truncate with "Read more" expansion',
    ]
)

add_fr(doc, 'FR-012', 'Call Report Linked Chain',
    'Call reports can be linked to previous call reports to create a chronological interaction chain for a client. When filing a new report, the RM can select a previous report to link to, establishing a parent-child relationship visible as a timeline.',
    'As a Relationship Manager, I want to link call reports in a chain so that I can track the progression of a client relationship across multiple interactions.',
    [
        'AC-058: Call report form includes "Link to Previous Report" optional field with combobox searching existing call reports for the same client',
        'AC-059: Client 360 page shows "Interaction Timeline" tab displaying all linked call reports as a vertical timeline',
        'AC-060: Timeline shows: date, subject, state_of_mind badge, status badge, summary preview (first 100 chars)',
        'AC-061: Clicking a timeline entry navigates to the full call report detail',
    ],
    [
        'BR-036: linked_call_report_id must reference a call report for the same client_id. Cross-client linking is not allowed.',
        'BR-037: Circular linking is prevented: A->B->C->A is blocked by the system',
    ],
    [
        'Timeline uses vertical line with circular nodes; nodes colored by state_of_mind (green=Happy, red=Irate, yellow=Sensitive)',
    ],
    [
        'If a linked parent report is soft-deleted, show "[Deleted Report]" placeholder in the chain with info icon',
    ]
)

doc.add_page_break()

# --- Module C: Action Items ---
doc.add_heading('5.3 Action Item Management', level=2)

add_fr(doc, 'FR-013', 'Create Action Items from Call Report',
    'During call report filing, RMs can add multiple action items with title, description, assignee, due date, and priority. Action items are created as child records of the call report and appear in the assignee\'s task list.',
    'As a Relationship Manager, I want to create follow-up action items while filing a call report so that important tasks are not forgotten and are trackable.',
    [
        'AC-062: Action Items section in call report form allows adding 0-20 items',
        'AC-063: Each item requires: title (max 200 chars), assigned_to (user search), due_date, priority dropdown',
        'AC-064: Description is optional but supports max 2000 chars',
        'AC-065: On call report submission, action items are created with status OPEN and notifications sent to assignees',
        'AC-066: Action items are visible on: call report detail, assignee task list, client 360 interaction timeline',
    ],
    [
        'BR-038: assigned_to defaults to the RM filing the report but can be changed to any user in the same branch or team',
        'BR-039: due_date must be >= call report filing date. System suggests due_date = filing_date + 5 business days',
        'BR-040: When an action item\'s due_date passes without completion, status auto-transitions to OVERDUE and notification sent to assignee + supervisor',
    ],
    [
        'Add Item button appends a new row to the action items table within the form',
        'Each row has a delete (X) button to remove before submission',
    ],
    [
        'If assignee user is INACTIVE or SUSPENDED, show error: "User [name] is not active. Select a different assignee."',
    ]
)

add_fr(doc, 'FR-014', 'Action Item Dashboard',
    'Centralized view of all action items assigned to the logged-in user, with filters for status, priority, due date range, and source call report. Overdue items are highlighted. Supports status updates and completion notes.',
    'As a Relationship Manager, I want to see all my pending action items in one place so that I can prioritize and track my follow-up tasks.',
    [
        'AC-067: Dashboard shows table with columns: Action Ref, Title, Client, Due Date, Priority, Status, Source Report',
        'AC-068: Default filter: status IN (OPEN, IN_PROGRESS, OVERDUE); sorted by due_date ascending',
        'AC-069: Overdue items have red background row and "OVERDUE" badge',
        'AC-070: Clicking an action item opens detail panel with: full description, source call report link, status update controls',
        'AC-071: Status update allows: OPEN -> IN_PROGRESS -> COMPLETED (with mandatory completion_notes)',
    ],
    [
        'BR-041: SENIOR_RM and BO_HEAD can see action items assigned to their team members',
        'BR-042: Completed action items move to "Completed" tab, visible for 90 days before archiving',
    ],
    [
        'Quick status update via dropdown on each row (no need to open detail panel)',
        'Batch select + "Mark as Complete" for completing multiple items',
    ],
    [
        'If the source call report was rejected, action items show info: "Source report under revision"',
    ]
)

doc.add_page_break()

# --- Module D: Opportunity Management ---
doc.add_heading('5.4 Opportunity Management', level=2)

add_fr(doc, 'FR-015', 'Capture Opportunity from Call Report',
    'During call report filing, RMs can capture sales opportunities identified during the client meeting. Each opportunity includes product type, expected value, probability, and expected close date. Opportunities feed into the pipeline dashboard.',
    'As a Relationship Manager, I want to capture sales opportunities during client meetings so that my pipeline is visible to management and I can track deals to closure.',
    [
        'AC-072: Opportunities section in call report form allows adding 0-10 opportunities',
        'AC-073: Each opportunity requires: name, product_type, expected_value, currency, probability (slider 0-100%), expected_close_date',
        'AC-074: Optional fields: description, product_name (specific product)',
        'AC-075: source field auto-set to MEETING; campaign_id available if meeting was for a campaign',
        'AC-076: Opportunities are created on call report submission with status OPEN',
    ],
    [
        'BR-043: expected_value must be > 0. Currency defaults to PHP.',
        'BR-044: probability defaults to 50%. UI shows a color-coded slider: red (0-25), yellow (26-50), light-green (51-75), green (76-100)',
        'BR-045: expected_close_date must be >= today',
    ],
    [
        'Opportunity cards within call report form show estimated weighted value (expected_value * probability/100)',
        'Product type dropdown populated from trust_product_types reference table',
    ],
    [
        'If expected_value exceeds PHP 100,000,000, show confirmation: "High-value opportunity. Please verify the amount."',
    ]
)

add_fr(doc, 'FR-016', 'Opportunity Pipeline Dashboard',
    'Visual funnel/kanban board showing opportunities by stage (Open -> Qualified -> Proposal -> Negotiation -> Won/Lost). Supports filtering by RM, branch, product type, and date range. Shows aggregated pipeline value.',
    'As a Supervisor/Trust Business Head, I want to see the opportunity pipeline across my team so that I can forecast revenue and coach RMs on deal progression.',
    [
        'AC-077: Kanban board with columns for each opportunity_status stage',
        'AC-078: Each card shows: opportunity name, client, expected value, probability, expected close date, owner RM',
        'AC-079: Drag-and-drop between columns updates opportunity_status',
        'AC-080: Pipeline summary bar shows: total count per stage, total weighted value, weighted pipeline value',
        'AC-081: Filters: owner RM, branch, product_type, date range, value range (min/max)',
    ],
    [
        'BR-046: Moving to WON status requires won_value and won_date fields (modal prompt)',
        'BR-047: Moving to LOST status requires loss_reason selection',
        'BR-048: Only the opportunity owner or BO_HEAD can change status',
    ],
    [
        'Kanban cards show probability as a mini progress bar',
        'Won column cards have green border; Lost have red border',
        'Summary shows conversion rate: Won / (Won + Lost) for selected period',
    ],
    [
        'If pipeline has >200 opportunities, disable drag-and-drop and switch to table view for performance',
    ]
)

add_fr(doc, 'FR-017', 'Bulk Upload Opportunities',
    'Upload opportunities in bulk via CSV file. System validates each row, reports errors per row, and creates valid opportunities. Template CSV is downloadable from the UI.',
    'As a Relationship Manager, I want to upload multiple opportunities at once from a spreadsheet so that I can efficiently capture pipeline from events or campaigns.',
    [
        'AC-082: Upload page provides downloadable CSV template with all required columns and sample data',
        'AC-083: Upload accepts .csv files up to 5MB (max 500 rows per upload)',
        'AC-084: Validation runs per row: required fields, enum values, numeric ranges, date formats',
        'AC-085: Results page shows: success count, error count, error details per row (row number + field + error message)',
        'AC-086: Successfully validated rows are inserted; failed rows are not. User can fix and re-upload failed rows.',
    ],
    [
        'BR-049: CSV columns: name, product_type, expected_value, currency, probability, expected_close_date, client_id (optional), description (optional)',
        'BR-050: source is auto-set to BULK_UPLOAD for all imported opportunities',
        'BR-051: Duplicate detection: if name + client_id + expected_close_date matches an existing opportunity, mark as potential duplicate and skip (user decides)',
    ],
    [
        'Upload area supports drag-and-drop or file picker',
        'Progress bar during processing; results table is downloadable as CSV',
    ],
    [
        'If CSV has >500 rows, reject with: "Maximum 500 rows per upload. Please split your file."',
        'If CSV encoding is not UTF-8, attempt auto-detection with warning',
    ]
)

doc.add_page_break()

# --- Module E: Task Management ---
doc.add_heading('5.5 Task Management', level=2)

add_fr(doc, 'FR-018', 'Personal Task Creation',
    'RMs can create personal to-do tasks with title, description, due date, priority, and optional reminders. Tasks appear on a centralized task list and can be linked to clients, meetings, or call reports.',
    'As a Relationship Manager, I want to create personal tasks so that I can manage my daily workload and remember important follow-ups.',
    [
        'AC-087: Task creation form: title (required), description (optional), due_date (required), due_time (optional), priority (required), reminder_date/time (optional)',
        'AC-088: Optional entity linking: select entity type (Client, Meeting, Call Report, Opportunity) and search for specific entity',
        'AC-089: task_type set to PERSONAL automatically',
        'AC-090: Created task appears immediately in "My Tasks" list',
    ],
    [
        'BR-052: Personal tasks are visible only to the creator unless task_type = ASSIGNED',
        'BR-053: Reminder triggers in-app notification at specified date/time',
    ],
    ['Quick-add task from top nav bar via "+" button; opens minimal form in modal'],
    ['If due_date is today, show amber "Due Today" badge. If past, show red "Overdue" badge.']
)

add_fr(doc, 'FR-019', 'Supervisor Task Assignment',
    'Supervisors (BO_HEAD, SENIOR_RM) can assign tasks to team members. Assigned tasks appear in the assignee\'s task list with an "Assigned" badge and the assigner\'s name.',
    'As a Supervisor, I want to assign tasks to my team members so that I can delegate work and track their completion.',
    [
        'AC-091: Assign Task form adds: assigned_to (user search within team), task fields same as personal',
        'AC-092: task_type set to ASSIGNED; assigned_by set to supervisor user ID',
        'AC-093: Assignee receives notification: "New task assigned by [supervisor]: [title]"',
        'AC-094: Supervisor can see assigned tasks in "Team Tasks" tab with status tracking',
    ],
    [
        'BR-054: Supervisors can only assign to users within their branch/team hierarchy',
        'BR-055: Assigned tasks cannot be deleted by the assignee — only marked COMPLETED or CANCELLED (with reason)',
    ],
    ['Team Tasks tab shows: assignee name, title, due date, status, with filter by assignee'],
    ['If assigned user is on leave (future feature), show warning: "User is out of office until [date]"']
)

doc.add_page_break()

# --- Module F: Expense Management ---
doc.add_heading('5.6 Expense Management', level=2)

add_fr(doc, 'FR-020', 'Log Client Meeting Expense',
    'RMs can log expenses related to client meetings including travel, meals, entertainment, and accommodation. Expenses can be linked to specific call reports or meetings. Receipt uploads are supported.',
    'As a Relationship Manager, I want to log client meeting expenses so that I can track spending and submit for reimbursement approval.',
    [
        'AC-095: Expense form: expense_type (dropdown), amount (numeric), currency (default PHP), expense_date, description, receipt upload (image/PDF up to 5MB)',
        'AC-096: Optional linking to call_report_id or meeting_id via dropdown',
        'AC-097: Save as Draft or Submit for Approval',
        'AC-098: Submitted expenses appear in supervisor\'s approval queue',
    ],
    [
        'BR-056: expense_date cannot be more than 30 days in the past',
        'BR-057: Expenses > PHP 10,000 require receipt attachment (receipt_url mandatory)',
    ],
    [
        'Receipt upload shows thumbnail preview for images; PDF icon for documents',
        'Currency selector shows PHP, USD, EUR, GBP, SGD, HKD, JPY',
    ],
    [
        'If file upload fails, expense is saved without receipt; user can attach later from edit view',
    ]
)

doc.add_page_break()

# --- Module G: Campaign Management ---
doc.add_heading('5.7 Campaign Management', level=2)

add_fr(doc, 'FR-021', 'Create Campaign',
    'Supervisors create marketing campaigns with code, name, type, target segment, budget, dates, and marketing materials. Campaigns are the top-level container for organizing lead lists and tracking responses.',
    'As a Supervisor (BO_HEAD), I want to create marketing campaigns so that I can organize targeted outreach to clients and prospects and measure results.',
    [
        'AC-099: Campaign form: campaign_code (auto-suggest from type+quarter), name, description, campaign_type, target_segment, budget, currency, expected_roi, start_date, end_date',
        'AC-100: Material upload section: multiple file attachments (PDF, images) with name labels',
        'AC-101: Campaign created with status DRAFT; can be activated when at least one lead list is attached',
        'AC-102: Campaign dashboard shows all campaigns with status, dates, lead count, response rate',
    ],
    [
        'BR-058: campaign_code must be unique. System suggests: [TYPE_ABBREVIATION]-[QUARTER]-[YEAR] (e.g., PP-Q2-2026)',
        'BR-059: Only BO_HEAD and SYSTEM_ADMIN can create campaigns',
        'BR-060: Budget is informational only — no financial integration in v1',
    ],
    ['Campaign cards on dashboard show progress bar: (contacted / total leads) as percentage'],
    ['If start_date is in the past when activating, show warning: "Campaign start date has passed. Activate anyway?"']
)

add_fr(doc, 'FR-022', 'Manage Lead Lists',
    'Create filtered target lists for campaigns using automatic filter rules (AUM, product type, risk profile), manual selection, or bulk CSV upload. Lists can be merged and members can be assigned to specific RMs.',
    'As a Supervisor, I want to create targeted lead lists so that the right clients/prospects are contacted for each campaign.',
    [
        'AC-103: Three list creation modes: Automatic (filter builder), Manual (search + select), Upload (CSV)',
        'AC-104: Automatic filter builder: AUM range (min/max), product_type, risk_profile, branch, tenure_years, client_status',
        'AC-105: Preview shows matching count before list creation; "Generate List" creates the list',
        'AC-106: Manual mode: search clients/prospects, select checkbox, add to list',
        'AC-107: Upload mode: CSV with client_id or prospect details, same validation as opportunity upload',
        'AC-108: Merge function: select 2+ lists, create a new merged list with de-duplication',
    ],
    [
        'BR-061: Lead lists can be shared across campaigns by linking lead_list to campaign_id',
        'BR-062: Automatic lists are regenerated on demand (not auto-refreshed)',
        'BR-063: RM assignment on lead_list_members is optional; if not set, any RM in the branch can capture responses',
    ],
    ['Filter builder uses visual chips for active filters; removing a chip updates the preview count'],
    ['If automatic filter returns 0 results, show: "No matches found. Try broadening your criteria."']
)

add_fr(doc, 'FR-023', 'Capture Lead Response',
    'RMs update lead response status when contacting campaign targets. Responses include Interested, Not Interested, Need More Info, Converted, and Do Not Contact. Response data feeds campaign performance metrics.',
    'As a Relationship Manager, I want to record client responses to campaign outreach so that the campaign team can track effectiveness and I can focus on interested leads.',
    [
        'AC-109: RM sees assigned campaign leads in "My Campaigns" section with contact details and response status',
        'AC-110: Response capture: select status from dropdown, add response_notes (optional), save',
        'AC-111: CONVERTED status triggers opportunity creation prompt: "Create an opportunity for this lead?"',
        'AC-112: DO_NOT_CONTACT status flags the client/prospect to be excluded from future campaigns',
        'AC-113: Response capture sets response_date = now() and responded_by = current user',
    ],
    [
        'BR-064: Only the assigned RM or any RM in the same branch (if no assignment) can capture responses',
        'BR-065: Response status can be updated multiple times; only the latest is used for metrics',
    ],
    ['Lead list shows as a card grid with status color coding: grey=not contacted, blue=contacted, green=interested, red=not interested'],
    ['If lead is a client with active service request, show info badge: "Active SR: [reference]"']
)

add_fr(doc, 'FR-024', 'Campaign Performance Dashboard',
    'Analytics dashboard showing campaign metrics: total leads, contacted count, response breakdown, conversion rate, cost per conversion, and ROI. Filterable by campaign, date range, and branch.',
    'As a Trust Business Head, I want to see campaign performance metrics so that I can evaluate ROI and make data-driven decisions about future campaigns.',
    [
        'AC-114: Dashboard cards: Total Leads, Contacted %, Interested %, Converted %, Cost per Lead, ROI',
        'AC-115: Response breakdown pie chart: Interested / Not Interested / Need More Info / Converted / Not Contacted',
        'AC-116: Trend line chart: responses over time (daily/weekly granularity)',
        'AC-117: RM leaderboard: table showing RMs ranked by response capture rate',
        'AC-118: Export dashboard data to CSV',
    ],
    [
        'BR-066: ROI = (total_won_opportunity_value_from_campaign - budget) / budget * 100',
        'BR-067: Cost per lead = budget / total_leads',
    ],
    ['Dashboard auto-refreshes every 5 minutes. Manual refresh button available.'],
    ['If campaign has no responses yet, show: "No responses captured yet. Campaign is [X days] old."']
)

doc.add_page_break()
doc.save('/Users/n15318/Trust OMS/docs/CIM_BRD_Draft.docx')
print("Part 3 complete: Section 5 (Functional Requirements) saved")
