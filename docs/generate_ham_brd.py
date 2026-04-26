#!/usr/bin/env python3
"""Generate Handover & Assignment Management BRD as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Arial'
        set_cell_shading(cell, '1F3864')
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)
        run.font.name = 'Arial'
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)
        run.font.name = 'Arial'
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(47, 84, 150)
        run.font.name = 'Arial'
    return p

def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.bold = bold
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    return p

def numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    return p

# ── Main Document ────────────────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

header = doc.sections[0].header
hp = header.paragraphs[0]
hp.text = "CONFIDENTIAL — Trust OMS Handover & Assignment Management BRD"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in hp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Page ")
run.font.size = Pt(8)
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
fp._p.append(fld)

# ── Title Page ───────────────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Business Requirements Document")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)
run.font.name = 'Arial'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Handover & Assignment Management")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(47, 84, 150)
run.font.name = 'Arial'

proj = doc.add_paragraph()
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = proj.add_run("Trust OMS — Wealth Management Platform")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(89, 89, 89)
run.font.name = 'Arial'

for _ in range(4):
    doc.add_paragraph()

info_data = [
    ["Document ID", "BRD-TRUSTOMS-HAM-2026-001"],
    ["Version", "2.0 (Post-Adversarial Review)"],
    ["Date", "2026-04-22"],
    ["Status", "Final"],
    ["Classification", "Confidential"],
    ["Author", "Trust OMS Platform Team"],
]
add_styled_table(doc, ["Property", "Value"], info_data)

doc.add_page_break()

# ── Table of Contents ────────────────────────────────────────────────────

h1(doc, "Table of Contents")
para(doc, "[Table of Contents — auto-generated in Word: Insert > Table of Contents]")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "1. Executive Summary")

h2(doc, "1.1 Project Name")
para(doc, "Trust OMS — Handover & Assignment Management Module (HAM)")

h2(doc, "1.2 Project Description")
para(doc, "The Handover & Assignment Management module is a critical relationship management component of the Trust OMS wealth management platform. It enables Operations users to permanently transfer (Handover) or temporarily reassign (Delegate) Leads, Prospects, and Clients from one Relationship Manager (RM) to another. Handover is used when an RM leaves the organization, retires, transfers branches, or when a client requests a change of RM. Delegation is used for temporary absences such as vacation, medical leave, training, or short-term branch coverage. The module enforces maker-checker authorization workflows for all permanent handovers, provides auto-authorization for temporary delegations, supports bulk upload operations for mass transfers, and maintains a comprehensive audit trail for regulatory compliance. The design is validated against three reference implementations: Maybank Asset Management, HSBC Private Wealth Solutions, and Jio BlackRock (JBR).")

h2(doc, "1.3 Business Objectives")
objectives = [
    "Enable seamless permanent transfer of Leads, Prospects, and Clients between Relationship Managers with full audit trail and maker-checker authorization to ensure operational integrity.",
    "Provide temporary delegation capability for RM absences with auto-authorization and automatic expiry to ensure continuous client coverage without service gaps.",
    "Support bulk handover operations via CSV file upload for mass RM transitions such as branch restructuring, team reorganization, or RM exits.",
    "Enforce scrutiny checklist completion for Client Handovers to ensure all client commitments, call reports, and contact details are validated before transfer.",
    "Provide a comprehensive Handover Dashboard with pending authorization queue, recent transfers, delegation calendar, and AUM impact assessment.",
    "Maintain complete audit trail of all handover and delegation actions for regulatory reporting and compliance with fiduciary obligations in trust management.",
]
for o in objectives:
    bullet(doc, o)

h2(doc, "1.4 Target Users and Pain Points")
add_styled_table(doc, ["User Role", "Pain Points Addressed"], [
    ["Operations Maker", "Manual spreadsheet-based RM reassignments are error-prone; no systematic tracking of handover reasons; no scrutiny checklist enforcement; bulk transfers require manual database updates."],
    ["Operations Checker", "No centralized authorization queue for handover approvals; cannot view full handover context (client details, AUM impact, checklist status) before approving."],
    ["Relationship Manager (RM)", "No visibility into newly assigned clients post-handover; no notification when clients are transferred in or out; unclear delegation coverage during absences."],
    ["Branch Manager / Supervisor", "No dashboard view of RM capacity, delegation coverage, or handover pipeline; cannot assess AUM impact of RM departures."],
    ["Compliance Officer", "No systematic audit trail for RM-client reassignments; cannot verify scrutiny checklist completion; difficult to report on regulatory fiduciary transfer requirements."],
])

h2(doc, "1.5 Success Metrics (KPIs)")
add_styled_table(doc, ["KPI", "Target", "Measurement Method"], [
    ["Handover processing time", "<30 min per handover (vs 4h manual)", "System timestamp tracking"],
    ["Authorization turnaround", "100% authorizations within 4 business hours", "Approval queue timestamps"],
    ["Scrutiny checklist compliance", "100% of client handovers have completed checklists", "Checklist completion audit"],
    ["Delegation coverage", "Zero client service gaps during RM absences", "Delegation coverage report"],
    ["Bulk upload success rate", ">98% rows processed successfully per upload", "Upload processing logs"],
    ["Audit trail completeness", "100% of transfers have full audit records", "Audit log validation"],
    ["System availability", "99.9% uptime during business hours", "Infrastructure monitoring"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2: SCOPE & BOUNDARIES
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "2. Scope & Boundaries")

h2(doc, "2.1 In Scope")
in_scope = [
    "Lead Handover — Permanent transfer of leads between RMs with mandatory handover reason and maker-checker authorization.",
    "Prospect Handover — Permanent transfer of prospects between RMs with mandatory handover reason and maker-checker authorization.",
    "Client Handover — Permanent transfer of clients between RMs with mandatory handover reason, scrutiny checklist, and maker-checker authorization.",
    "Bulk Client Handover — CSV file upload for mass client transfers with preview/dry-run, supervisor notification, and per-row audit logging (compensating controls in lieu of full maker-checker).",
    "Lead Delegation — Temporary reassignment of leads with date range and auto-authorization.",
    "Prospect Delegation — Temporary reassignment of prospects with date range and auto-authorization.",
    "Client Delegation — Temporary reassignment of clients with date range, scrutiny checklist, and auto-authorization.",
    "Handover Authorization Workflow — Operations Checker approve/reject with reason capture.",
    "Delegation Auto-Expiry — Automatic revert of delegated relationships when delegation period ends.",
    "Handover History & Audit Trail — Complete log of all transfers with timestamps, actors, and reasons.",
    "Handover Dashboard — Pending authorizations, recent transfers, delegation calendar, AUM impact widget.",
    "Notification System — Email/in-app notifications for handover initiation, authorization, rejection, and delegation start/end.",
    "Portfolio Impact Assessment — AUM summary display for clients being handed over.",
    "Pending Activities Warning — Alert for open orders, service requests, or upcoming maturity dates before handover.",
    "Filterable Data Tables — All handover/delegation grids support column-level filtering, sorting, and pagination.",
    "Navigation from Relationship Menu — Handover and Delegation accessible from Relationship menu in back-office portal.",
    "Batch Authorization — Batch authorize/reject multiple pending handover requests in a single action.",
    "Delegation Extension — Extend active delegations by up to 90 additional days with supervisor approval.",
    "Concurrent Conflict Resolution — Automatic delegation termination when permanent handover is authorized for same entity.",
    "Data Archival — Year-based partitioning of audit logs with cold storage archival after 2 years.",
]
for item in in_scope:
    bullet(doc, item)

h2(doc, "2.2 Out of Scope")
out_scope = [
    "RM self-assignment — RMs cannot assign clients to themselves; only Operations users can initiate handovers.",
    "Cross-entity handover — Transfers between different legal entities (e.g., Maybank to HSBC) are not supported.",
    "Client-initiated RM change requests via client portal — Client portal displays current RM but does not allow self-service RM change.",
    "Family Grouping management — Separate module; not included in HAM.",
    "FID Grouping management — Separate module; not included in HAM.",
    "RM onboarding/offboarding — HR processes are handled outside this module.",
    "Performance attribution — Revenue/AUM attribution during transition periods handled by reporting module.",
]
for item in out_scope:
    bullet(doc, item)

h2(doc, "2.3 Assumptions")
assumptions = [
    "Lead, Prospect, and Client master records already exist in the system from onboarding modules.",
    "RM and Supervisor master records are maintained in User Maintenance and are available for dropdown selection.",
    "The Operations portal already has role-based authentication supporting Maker and Checker roles.",
    "Branch and Location master data is pre-configured in the system.",
    "Email service is available for sending notification emails.",
    "The system operates in a multi-branch, single-entity deployment model.",
]
for item in assumptions:
    bullet(doc, item)

h2(doc, "2.4 Constraints")
constraints = [
    "Handover can only be initiated by users with Operations Maker role.",
    "Handover authorization can only be performed by users with Operations Checker role.",
    "Delegation is limited to RMs under the same supervisor or within the same branch.",
    "Bulk upload file size limit: 10 MB, maximum 5,000 rows per file.",
    "A client/prospect/lead can only have one active handover request at a time (no concurrent pending handovers for same entity).",
    "Delegation period cannot exceed 90 calendar days.",
]
for item in constraints:
    bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3: USER ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "3. User Roles & Permissions")

h2(doc, "3.1 Role Definitions")
add_styled_table(doc, ["Role", "Description"], [
    ["Operations Maker", "Back-office operations user who initiates handover and delegation requests. Can create, filter, select, and save handover/delegation records."],
    ["Operations Checker", "Back-office operations user who authorizes or rejects handover requests. Cannot initiate handovers. A Checker cannot authorize their own Maker submissions."],
    ["Branch Manager", "Supervisory role with read access to handover dashboard and delegation calendar for their branch. Cannot initiate or authorize handovers directly."],
    ["System Administrator", "Configures handover scrutiny checklist items, upload templates, notification templates, and delegation period limits. No transactional access."],
    ["Relationship Manager", "Read-only access to their own incoming/outgoing handover history. Receives notifications. Cannot initiate handovers."],
])

h2(doc, "3.2 Permissions Matrix")
add_styled_table(doc, ["Permission", "Ops Maker", "Ops Checker", "Branch Mgr", "Sys Admin", "RM"], [
    ["View Handover Screen", "Yes", "Yes", "Yes (read)", "No", "No"],
    ["Initiate Lead Handover", "Yes", "No", "No", "No", "No"],
    ["Initiate Prospect Handover", "Yes", "No", "No", "No", "No"],
    ["Initiate Client Handover", "Yes", "No", "No", "No", "No"],
    ["Upload Bulk Handover CSV", "Yes", "No", "No", "No", "No"],
    ["Authorize/Reject Handover", "No", "Yes", "No", "No", "No"],
    ["Initiate Lead Delegation", "Yes", "No", "No", "No", "No"],
    ["Initiate Prospect Delegation", "Yes", "No", "No", "No", "No"],
    ["Initiate Client Delegation", "Yes", "No", "No", "No", "No"],
    ["View Handover Dashboard", "Yes", "Yes", "Yes", "No", "No"],
    ["View Delegation Calendar", "Yes", "Yes", "Yes", "No", "No"],
    ["View Handover History", "Yes", "Yes", "Yes", "No", "Own only"],
    ["Configure Checklist Items", "No", "No", "No", "Yes", "No"],
    ["Configure Upload Templates", "No", "No", "No", "Yes", "No"],
    ["Receive Handover Notifications", "Yes", "Yes", "No", "No", "Yes"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4: DATA MODEL
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "4. Data Model")

# --- 4.1 Handover Request ---
h2(doc, "4.1 HandoverRequest")
para(doc, "Central entity tracking every handover (permanent transfer) request. One request covers one or more Leads/Prospects/Clients being transferred from an outgoing RM to an incoming RM.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["handover_type", "ENUM", "Yes", "'LEAD' | 'PROSPECT' | 'CLIENT'", "—"],
    ["status", "ENUM", "Yes", "'PENDING' | 'AUTHORIZED' | 'REJECTED'", "PENDING"],
    ["version", "INTEGER", "Yes", "Auto-incremented on update; starts at 1", "1"],
    ["outgoing_rm_id", "VARCHAR(50)", "Yes", "Must exist in RM master", "—"],
    ["outgoing_rm_name", "VARCHAR(200)", "Yes", "Populated from RM master", "—"],
    ["incoming_rm_id", "VARCHAR(50)", "Yes", "Must exist in RM master; != outgoing_rm_id", "—"],
    ["incoming_rm_name", "VARCHAR(200)", "Yes", "Populated from RM master", "—"],
    ["incoming_srm_id", "VARCHAR(50)", "No", "Auto-populated from incoming RM's supervisor", "—"],
    ["incoming_srm_name", "VARCHAR(200)", "No", "Populated from supervisor master", "—"],
    ["branch_code", "VARCHAR(20)", "Yes", "Must exist in Branch master", "—"],
    ["handover_reason", "TEXT", "Yes", "Min 5 chars, max 500 chars", "—"],
    ["reject_reason", "TEXT", "No", "Required when status = REJECTED; min 5 chars", "—"],
    ["initiated_by", "VARCHAR(50)", "Yes", "Maker user ID", "Current user"],
    ["authorized_by", "VARCHAR(50)", "No", "Checker user ID; set on authorize/reject", "—"],
    ["authorized_at", "TIMESTAMP", "No", "Set on authorize/reject", "—"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "Auto-updated", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "handover_type", "status", "outgoing_rm_id", "incoming_rm_id", "handover_reason", "initiated_by"], [
    ["a1b2c3d4", "CLIENT", "PENDING", "BASERM2", "WFRM6", "RM leaving organization", "OPSMAKER1"],
    ["e5f6g7h8", "LEAD", "AUTHORIZED", "BASERM1", "BASERM2", "Branch restructuring", "OPSMAKER2"],
    ["i9j0k1l2", "PROSPECT", "REJECTED", "WL001", "WL002", "Client request", "OPSMAKER1"],
])

# --- 4.2 HandoverItem ---
h2(doc, "4.2 HandoverItem")
para(doc, "Each individual Lead/Prospect/Client included in a handover request. Multiple items can belong to one HandoverRequest.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["handover_request_id", "UUID", "Yes", "FK → HandoverRequest.id", "—"],
    ["entity_type", "ENUM", "Yes", "'LEAD' | 'PROSPECT' | 'CLIENT'", "—"],
    ["entity_id", "VARCHAR(50)", "Yes", "Must exist in corresponding master table", "—"],
    ["entity_name", "VARCHAR(200)", "Yes", "Populated from master", "—"],
    ["branch_code", "VARCHAR(20)", "Yes", "From entity master", "—"],
    ["cust_id", "VARCHAR(50)", "No", "Only for CLIENT type; customer ID", "—"],
    ["outgoing_rm_id", "VARCHAR(50)", "Yes", "Current RM of the entity", "—"],
    ["outgoing_rm_name", "VARCHAR(200)", "Yes", "Populated from RM master", "—"],
    ["referring_rm_id", "VARCHAR(50)", "No", "For CLIENT type; referring RM if applicable", "—"],
    ["branch_rm_id", "VARCHAR(50)", "No", "For CLIENT type; branch-level RM", "—"],
    ["aum_amount", "DECIMAL(18,2)", "No", "Auto-populated from portfolio for CLIENT type", "0.00"],
    ["pending_orders_count", "INTEGER", "No", "Count of open orders; auto-computed", "0"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "handover_request_id", "entity_type", "entity_id", "entity_name", "branch_code", "aum_amount"], [
    ["x1y2z3", "a1b2c3d4", "CLIENT", "080820231", "Ricky Narayan", "Pragathi Nagar", "2500000.00"],
    ["x4y5z6", "a1b2c3d4", "CLIENT", "100100231", "KTMR004 Kumar", "Pragathi Nagar", "1800000.00"],
    ["x7y8z9", "e5f6g7h8", "LEAD", "2306479000059034", "Ricky Narayan", "BR001", "0.00"],
])

# --- 4.3 ScrutinyChecklist ---
h2(doc, "4.3 ScrutinyChecklist")
para(doc, "Checklist items that must be completed for Client Handovers. Each item tracks validation status and optional remarks.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["handover_request_id", "UUID", "Yes", "FK → HandoverRequest.id", "—"],
    ["checklist_item_id", "UUID", "Yes", "FK → ChecklistItemConfig.id", "—"],
    ["validation_label", "VARCHAR(300)", "Yes", "Display text of checklist item", "—"],
    ["status", "ENUM", "Yes", "'NOT_STARTED' | 'WORK_IN_PROGRESS' | 'COMPLETED'", "NOT_STARTED"],
    ["remarks", "TEXT", "No", "Free-text remarks; max 500 chars", "—"],
    ["updated_by", "VARCHAR(50)", "No", "User who last updated", "—"],
    ["updated_at", "TIMESTAMP", "Yes", "Auto-updated", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "handover_request_id", "validation_label", "status", "remarks"], [
    ["sc001", "a1b2c3d4", "Assurance/Promises given to the client", "COMPLETED", "All promises documented"],
    ["sc002", "a1b2c3d4", "Call Report", "COMPLETED", "Final call completed 20-Apr-2026"],
    ["sc003", "a1b2c3d4", "Client not contactable", "COMPLETED", "Client reachable via email"],
    ["sc004", "a1b2c3d4", "Customer key contact person details should be updated", "WORK_IN_PROGRESS", "Awaiting alternate contact"],
])

# --- 4.4 ChecklistItemConfig ---
h2(doc, "4.4 ChecklistItemConfig")
para(doc, "Admin-configurable list of scrutiny checklist items. These define what validations must be completed for client handovers.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["label", "VARCHAR(300)", "Yes", "Unique label text", "—"],
    ["description", "TEXT", "No", "Detailed description for operations users", "—"],
    ["is_mandatory", "BOOLEAN", "Yes", "If true, status must be COMPLETED before save", "true"],
    ["applies_to", "ENUM", "Yes", "'HANDOVER_ONLY' | 'DELEGATION_ONLY' | 'BOTH'", "HANDOVER_ONLY"],
    ["display_order", "INTEGER", "Yes", "Sort order in UI", "0"],
    ["is_active", "BOOLEAN", "Yes", "Soft-delete flag", "true"],
    ["created_by", "VARCHAR(50)", "Yes", "Admin user ID", "—"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "Auto-updated", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "label", "is_mandatory", "display_order", "is_active"], [
    ["cfg001", "Assurance/Promises given to the client", "true", "1", "true"],
    ["cfg002", "Call Report", "true", "2", "true"],
    ["cfg003", "Client not contactable", "false", "3", "true"],
    ["cfg004", "Customer key contact person details should be updated", "true", "4", "true"],
])

# --- 4.5 DelegationRequest ---
h2(doc, "4.5 DelegationRequest")
para(doc, "Tracks temporary delegation of Leads/Prospects/Clients from one RM to another. Delegations are auto-authorized and auto-expire.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["delegation_type", "ENUM", "Yes", "'LEAD' | 'PROSPECT' | 'CLIENT'", "—"],
    ["status", "ENUM", "Yes", "'ACTIVE' | 'EXPIRED' | 'CANCELLED' | 'EARLY_TERMINATED'", "ACTIVE"],
    ["extended_from_id", "UUID", "No", "FK → DelegationRequest.id if this is an extension of a prior delegation", "—"],
    ["extension_count", "INTEGER", "Yes", "Number of times this delegation has been extended (max 1)", "0"],
    ["outgoing_rm_id", "VARCHAR(50)", "Yes", "Must exist in RM master", "—"],
    ["outgoing_rm_name", "VARCHAR(200)", "Yes", "Populated from RM master", "—"],
    ["delegate_rm_id", "VARCHAR(50)", "Yes", "Must be under same supervisor or branch; != outgoing_rm_id", "—"],
    ["delegate_rm_name", "VARCHAR(200)", "Yes", "Populated from RM master", "—"],
    ["delegate_srm_id", "VARCHAR(50)", "No", "Auto-populated from delegate RM's supervisor", "—"],
    ["branch_code", "VARCHAR(20)", "Yes", "Must exist in Branch master", "—"],
    ["delegation_reason", "TEXT", "Yes", "Min 5 chars, max 500 chars", "—"],
    ["start_date", "DATE", "Yes", ">= today; start_date < end_date", "—"],
    ["end_date", "DATE", "Yes", "<= start_date + 90 days", "—"],
    ["auto_revert_completed", "BOOLEAN", "Yes", "Set to true when auto-revert runs", "false"],
    ["initiated_by", "VARCHAR(50)", "Yes", "Maker user ID", "Current user"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "Auto-updated", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "delegation_type", "status", "outgoing_rm_id", "delegate_rm_id", "start_date", "end_date", "delegation_reason"], [
    ["d1e2f3", "CLIENT", "ACTIVE", "BASERM2", "BASERM3", "2026-04-25", "2026-05-10", "Annual leave"],
    ["d4e5f6", "LEAD", "EXPIRED", "WL001", "WL002", "2026-03-01", "2026-03-15", "Training program"],
    ["d7e8f9", "PROSPECT", "CANCELLED", "BASERM1", "BASERM4", "2026-04-20", "2026-05-20", "Medical leave"],
])

# --- 4.6 DelegationItem ---
h2(doc, "4.6 DelegationItem")
para(doc, "Individual Lead/Prospect/Client included in a delegation request.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["delegation_request_id", "UUID", "Yes", "FK → DelegationRequest.id", "—"],
    ["entity_type", "ENUM", "Yes", "'LEAD' | 'PROSPECT' | 'CLIENT'", "—"],
    ["entity_id", "VARCHAR(50)", "Yes", "Must exist in corresponding master", "—"],
    ["entity_name", "VARCHAR(200)", "Yes", "Populated from master", "—"],
    ["branch_code", "VARCHAR(20)", "Yes", "From entity master", "—"],
    ["original_rm_id", "VARCHAR(50)", "Yes", "RM before delegation", "—"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "delegation_request_id", "entity_type", "entity_id", "entity_name", "branch_code"], [
    ["di01", "d1e2f3", "CLIENT", "080820231", "Ricky Narayan", "Pragathi Nagar"],
    ["di02", "d1e2f3", "CLIENT", "210820231", "Sidhu Narayan", "Pragathi Nagar"],
    ["di03", "d4e5f6", "LEAD", "23053811059463", "NonINDLeas", "BR001"],
])

# --- 4.7 BulkUploadLog ---
h2(doc, "4.7 BulkUploadLog")
para(doc, "Tracks bulk CSV uploads for client handover processing.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["upload_type", "ENUM", "Yes", "'CLIENT_HANDOVER'", "—"],
    ["file_name", "VARCHAR(255)", "Yes", "Original file name", "—"],
    ["file_size_bytes", "INTEGER", "Yes", "Max 10485760 (10MB)", "—"],
    ["total_rows", "INTEGER", "Yes", "Max 5000", "—"],
    ["success_count", "INTEGER", "Yes", "Rows processed successfully", "0"],
    ["error_count", "INTEGER", "Yes", "Rows that failed processing", "0"],
    ["status", "ENUM", "Yes", "'PROCESSING' | 'COMPLETED' | 'FAILED'", "PROCESSING"],
    ["error_details", "JSON", "No", "Array of {row, field, error} objects", "[]"],
    ["uploaded_by", "VARCHAR(50)", "Yes", "User who uploaded the file", "—"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set", "NOW()"],
    ["completed_at", "TIMESTAMP", "No", "Set when processing finishes", "—"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "file_name", "total_rows", "success_count", "error_count", "status", "uploaded_by"], [
    ["bu01", "client_handover_apr2026.csv", "150", "148", "2", "COMPLETED", "OPSMAKER1"],
    ["bu02", "branch_transfer_batch.csv", "500", "0", "0", "PROCESSING", "OPSMAKER2"],
])

# --- 4.8 HandoverAuditLog ---
h2(doc, "4.8 HandoverAuditLog")
para(doc, "Immutable audit trail for all handover and delegation actions.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["action_type", "ENUM", "Yes", "'HANDOVER_INITIATED' | 'HANDOVER_AUTHORIZED' | 'HANDOVER_REJECTED' | 'HANDOVER_AMENDED' | 'DELEGATION_CREATED' | 'DELEGATION_EXPIRED' | 'DELEGATION_CANCELLED' | 'DELEGATION_EARLY_TERMINATED' | 'DELEGATION_EXTENSION_REQUESTED' | 'DELEGATION_EXTENDED' | 'BULK_UPLOAD' | 'BULK_UPLOAD_PREVIEW' | 'BATCH_AUTHORIZE' | 'BATCH_REJECT' | 'CHECKLIST_UPDATED'", "—"],
    ["reference_type", "ENUM", "Yes", "'HANDOVER_REQUEST' | 'DELEGATION_REQUEST' | 'BULK_UPLOAD'", "—"],
    ["reference_id", "UUID", "Yes", "FK to corresponding request/upload ID", "—"],
    ["entity_type", "ENUM", "No", "'LEAD' | 'PROSPECT' | 'CLIENT'", "—"],
    ["entity_id", "VARCHAR(50)", "No", "ID of affected Lead/Prospect/Client", "—"],
    ["from_rm_id", "VARCHAR(50)", "No", "Outgoing RM", "—"],
    ["to_rm_id", "VARCHAR(50)", "No", "Incoming RM", "—"],
    ["action_by", "VARCHAR(50)", "Yes", "User who performed the action", "—"],
    ["action_details", "JSON", "No", "Additional context (reason, reject reason, etc.)", "{}"],
    ["ip_address", "VARCHAR(45)", "No", "IP address of the user", "—"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set, immutable", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "action_type", "reference_id", "entity_type", "entity_id", "from_rm_id", "to_rm_id", "action_by"], [
    ["al01", "HANDOVER_INITIATED", "a1b2c3d4", "CLIENT", "080820231", "BASERM2", "WFRM6", "OPSMAKER1"],
    ["al02", "HANDOVER_AUTHORIZED", "a1b2c3d4", "CLIENT", "080820231", "BASERM2", "WFRM6", "OPSCHECKER1"],
    ["al03", "DELEGATION_CREATED", "d1e2f3", "CLIENT", "080820231", "BASERM2", "BASERM3", "OPSMAKER1"],
])

# --- 4.9 HandoverNotification ---
h2(doc, "4.9 HandoverNotification")
para(doc, "Notifications generated by handover/delegation events.")
add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes", "Auto-generated", "uuid_v4()"],
    ["notification_type", "ENUM", "Yes", "'HANDOVER_INITIATED' | 'HANDOVER_AUTHORIZED' | 'HANDOVER_REJECTED' | 'DELEGATION_STARTED' | 'DELEGATION_EXPIRING' | 'DELEGATION_EXPIRED' | 'DELEGATION_EARLY_TERMINATED' | 'DELEGATION_EXTENSION_REQUESTED' | 'DELEGATION_EXTENSION_APPROVED' | 'BULK_UPLOAD_SUPERVISOR_ALERT' | 'BATCH_AUTH_COMPLETE'", "—"],
    ["channel", "ENUM", "Yes", "'EMAIL' | 'IN_APP' | 'BOTH'", "BOTH"],
    ["recipient_user_id", "VARCHAR(50)", "Yes", "Target user", "—"],
    ["recipient_email", "VARCHAR(255)", "No", "For email channel", "—"],
    ["subject", "VARCHAR(500)", "Yes", "Notification subject line", "—"],
    ["body", "TEXT", "Yes", "Notification body content", "—"],
    ["reference_type", "ENUM", "Yes", "'HANDOVER_REQUEST' | 'DELEGATION_REQUEST'", "—"],
    ["reference_id", "UUID", "Yes", "FK to request ID", "—"],
    ["is_read", "BOOLEAN", "Yes", "Read status for in-app", "false"],
    ["sent_at", "TIMESTAMP", "No", "When email was sent", "—"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set", "NOW()"],
])

para(doc, "Sample Data:")
add_styled_table(doc, ["id", "notification_type", "channel", "recipient_user_id", "subject", "is_read"], [
    ["nt01", "HANDOVER_INITIATED", "BOTH", "WFRM6", "New Client Handover: 3 clients assigned to you", "false"],
    ["nt02", "HANDOVER_AUTHORIZED", "BOTH", "OPSMAKER1", "Handover #a1b2c3d4 authorized", "true"],
    ["nt03", "DELEGATION_EXPIRING", "EMAIL", "BASERM2", "Delegation ending in 2 days — 5 clients reverting", "false"],
])

# --- 4.10 Entity Relationships ---
h2(doc, "4.10 Entity Relationships")
add_styled_table(doc, ["Relationship", "Type", "Description"], [
    ["HandoverRequest → HandoverItem", "1:N", "One handover request can transfer multiple leads/prospects/clients"],
    ["HandoverRequest → ScrutinyChecklist", "1:N", "Client handovers have multiple scrutiny checklist items"],
    ["HandoverRequest → HandoverAuditLog", "1:N", "Each request generates multiple audit entries"],
    ["HandoverRequest → HandoverNotification", "1:N", "Each request triggers multiple notifications"],
    ["DelegationRequest → DelegationItem", "1:N", "One delegation can cover multiple entities"],
    ["DelegationRequest → HandoverAuditLog", "1:N", "Each delegation generates audit entries"],
    ["DelegationRequest → HandoverNotification", "1:N", "Each delegation triggers notifications"],
    ["ChecklistItemConfig → ScrutinyChecklist", "1:N", "Config defines templates; checklist instances per request"],
    ["BulkUploadLog → HandoverAuditLog", "1:N", "Each bulk upload generates audit entries"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5: FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "5. Functional Requirements")

# --- FR-001: Lead Handover ---
h2(doc, "5.1 Handover Module")

h3(doc, "FR-001: Lead Handover")
para(doc, "Operations Maker can permanently transfer one or more leads from their current RM to a new incoming RM. The system displays all existing leads available for handover, supports multi-select, and requires a mandatory handover reason before saving.")
para(doc, 'User Story: "As an Operations Maker, I want to select leads from the existing leads list and transfer them to a new RM so that the leads continue to receive relationship coverage when the current RM is departing."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-001.1: System displays Existing Leads grid with columns: Lead ID, Lead Name (English), Lead Name (Local Language), RM Owner Lead ID - RM Owner Lead Name.")
bullet(doc, "AC-001.2: All columns in Existing Leads grid are filterable via text input below column headers.")
bullet(doc, "AC-001.3: Location and Preferred Language dropdowns are available as quick filters above the grid.")
bullet(doc, "AC-001.4: User can select one or more leads using checkboxes; selected leads appear in Selected Lead List section.")
bullet(doc, "AC-001.5: Selected Lead List displays: Lead ID, Lead Name (English), Lead Name (Local Language), RM Owner Lead ID - RM Owner Lead Name.")
bullet(doc, "AC-001.6: Incoming RM/SRM Details section shows: Incoming RM Owner Lead dropdown (searchable), Incoming Supervisor (auto-populated from RM master).")
bullet(doc, "AC-001.7: Handover Reason is a mandatory free-text field (min 5, max 500 characters).")
bullet(doc, "AC-001.8: On Save, system validates all required fields, creates HandoverRequest with status PENDING, and sends to authorization queue.")
bullet(doc, "AC-001.9: Success toast notification: 'Handover Initiated Successfully' with OK button.")
bullet(doc, "AC-001.10: After save, the transferred leads no longer appear in the Existing Leads grid (until authorized/rejected).")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-001.1: Incoming RM cannot be the same as the outgoing RM.")
bullet(doc, "BR-001.2: A lead with an existing PENDING handover request cannot be selected for another handover.")
bullet(doc, "BR-001.3: The Incoming Supervisor is auto-populated based on the selected Incoming RM and is read-only.")
para(doc, "Error Handling:", bold=True)
bullet(doc, "If no leads are selected: 'Please select at least one lead for handover.'")
bullet(doc, "If incoming RM not selected: 'Please select an incoming Relationship Manager.'")
bullet(doc, "If handover reason empty: 'Handover reason is required.'")
bullet(doc, "If server error: 'Unable to process handover. Please try again later.'")

# --- FR-002: Prospect Handover ---
h3(doc, "FR-002: Prospect Handover")
para(doc, "Operations Maker can permanently transfer one or more prospects from their current RM to a new incoming RM. Same workflow as Lead Handover but with prospect-specific data columns.")
para(doc, 'User Story: "As an Operations Maker, I want to select prospects and transfer them to a new RM so that prospect engagement continues without interruption when the current RM changes."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-002.1: System displays Existing Prospects grid with columns: Branch, Prospect ID, Prospect Name, RM ID - RM Name.")
bullet(doc, "AC-002.2: All columns are filterable. Location and Preferred Language quick filters available.")
bullet(doc, "AC-002.3: Selected Prospects List displays: Branch, Prospect ID, Prospect Name, Outgoing RM ID - RM Name.")
bullet(doc, "AC-002.4: Incoming RM section shows: Incoming RM dropdown (searchable), Incoming SRM (auto-populated).")
bullet(doc, "AC-002.5: Handover Reason is mandatory (min 5, max 500 chars).")
bullet(doc, "AC-002.6: On Save, creates PENDING HandoverRequest and displays success toast.")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-002.1: Same as BR-001.1 to BR-001.3 applied to prospects.")

# --- FR-003: Client Handover ---
h3(doc, "FR-003: Client Handover")
para(doc, "Operations Maker can permanently transfer one or more clients to a new RM. Client handover includes additional sections: Selected Client List with base number/name details, Incoming RM/Branch RM/Referring RM section, and a mandatory Scrutiny Checklist that must be completed before saving.")
para(doc, 'User Story: "As an Operations Maker, I want to transfer clients to a new RM with a full scrutiny checklist so that all client commitments, call reports, and contact details are validated before the permanent transfer."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-003.1: Existing Clients grid displays: Branch, Client ID, Client Name, Cust ID, RM ID - RM Name, Referring RM columns.")
bullet(doc, "AC-003.2: All columns filterable. Location and Preferred Language quick filters available.")
bullet(doc, "AC-003.3: Selected Client List displays: Branch, Client ID, Client Name, Cust ID, Outgoing RM ID - RM Name, Referring RM ID - RM Name, Branch RM ID - Branch RM Name, plus expandable Base No and Base Name.")
bullet(doc, "AC-003.4: Incoming RM/Branch RM/Referring RM section shows dropdowns for incoming RM assignment with auto-populated supervisor.")
bullet(doc, "AC-003.5: Scrutiny Check List section displays all active ChecklistItemConfig items with Status dropdown (Not Started / Work In Progress / Completed) and Remarks text field.")
bullet(doc, "AC-003.6: All mandatory checklist items must have status = COMPLETED before save is enabled.")
bullet(doc, "AC-003.7: Handover Reason is mandatory.")
bullet(doc, "AC-003.8: On Save, creates PENDING HandoverRequest with associated HandoverItems and ScrutinyChecklist records.")
bullet(doc, "AC-003.9: System shows AUM summary (total assets under management) for selected clients as portfolio impact indicator.")
bullet(doc, "AC-003.10: System warns if any selected client has pending orders or upcoming maturity dates.")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-003.1: Scrutiny checklist is only required for CLIENT type handovers, not for LEAD or PROSPECT.")
bullet(doc, "BR-003.2: If any mandatory checklist item is not COMPLETED, Save button is disabled with tooltip: 'Complete all mandatory checklist items before saving.'")
bullet(doc, "BR-003.3: AUM impact is informational only and does not block the handover.")

# --- FR-004: Bulk Client Handover ---
h3(doc, "FR-004: Bulk Client Handover via File Upload")
para(doc, "Operations Maker can upload a CSV file to perform mass client handovers. Uploaded records are auto-processed without maker-checker authorization.")
para(doc, 'User Story: "As an Operations Maker, I want to upload a CSV file for bulk client handover so that I can efficiently process mass RM reassignments during branch restructuring."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-004.1: File Upload screen accessible from Relationship menu with upload type 'Client Handover'.")
bullet(doc, "AC-004.2: User selects upload type, browses and selects CSV file, clicks 'Upload File'.")
bullet(doc, "AC-004.3: System validates file format (CSV), size (max 10MB), and row count (max 5000 rows).")
bullet(doc, "AC-004.4: CSV columns: Client ID, Incoming RM ID, Handover Reason.")
bullet(doc, "AC-004.5: System processes rows asynchronously; creates BulkUploadLog record tracking progress.")
bullet(doc, "AC-004.6: Results viewable in error log with success/failure count per row.")
bullet(doc, "AC-004.7: Failed rows include specific error message (e.g., 'Client ID not found', 'RM ID invalid').")
bullet(doc, "AC-004.8: Reset button clears the form. Upload File button is disabled until a file is selected.")
bullet(doc, "AC-004.9: After file selection, system performs a dry-run/preview showing: total rows, affected client count, total AUM impact, list of incoming RMs, and any validation errors — before committing.")
bullet(doc, "AC-004.10: User must click 'Confirm & Process' after reviewing the preview. Cancel returns to the upload form without processing.")
bullet(doc, "AC-004.11: Each row in the bulk upload generates an individual HandoverAuditLog entry with action_type = 'BULK_UPLOAD' and per-row entity details.")
bullet(doc, "AC-004.12: On upload confirmation, system sends email notification to the uploading Maker's supervisor (Branch Manager) with: file name, row count, total AUM impact, and link to upload log.")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-004.1: Bulk uploads bypass full maker-checker authorization but include compensating controls: mandatory preview/dry-run step, per-row audit logging, and supervisor email notification.")
bullet(doc, "BR-004.2: Each row in CSV is processed independently — one row's failure does not affect others.")
bullet(doc, "BR-004.3: Duplicate client IDs in same file are rejected for all but the first occurrence.")
bullet(doc, "BR-004.4: If any client in the upload has an active delegation, the delegation is auto-cancelled with reason 'Superseded by bulk handover #{upload_id}' before processing the handover.")
bullet(doc, "BR-004.5: Preview step must complete within 30 seconds for files up to 5,000 rows.")

# --- FR-005: Handover Authorization ---
h3(doc, "FR-005: Handover Authorization (Maker-Checker)")
para(doc, "Operations Checker can view all pending handover requests and either authorize or reject them. Authorization completes the RM transfer; rejection returns the entities to their original state.")
para(doc, 'User Story: "As an Operations Checker, I want to review pending handover requests with full context (selected entities, incoming RM, scrutiny checklist, AUM impact) so that I can make informed authorize/reject decisions."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-005.1: Checker navigates to Authorization > Authorize and selects handover type (Lead/Prospect/Client Handover).")
bullet(doc, "AC-005.2: Authorization queue displays all PENDING requests with: entity ID, entity name, Cust ID, RM ID - RM Name, Referring RM, and Actions column.")
bullet(doc, "AC-005.3: All columns in authorization queue are filterable.")
bullet(doc, "AC-005.4: Clicking Actions button opens detail view showing: Selected Items List, Incoming RM details, Scrutiny Checklist (for clients), Handover Reason.")
bullet(doc, "AC-005.5: Authorize button — on click, updates HandoverRequest.status to AUTHORIZED, updates RM assignments in master tables, creates audit log, sends notifications.")
bullet(doc, "AC-005.6: Reject button — opens Reject Reason overlay with text input (mandatory, min 5 chars) and Reject/Close buttons. On reject, updates status to REJECTED.")
bullet(doc, "AC-005.7: After authorization, entity disappears from pending queue and authorized entity appears in incoming RM's workspace.")
bullet(doc, "AC-005.8: Toast notification on authorize: 'Record Authorized Successfully'. On reject: 'Record Rejected Successfully'.")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-005.1: Checker cannot authorize their own Maker submissions (segregation of duties).")
bullet(doc, "BR-005.2: A rejected handover frees the entities for a new handover request.")
bullet(doc, "BR-005.3: Rejection requires a mandatory reason (Reject Reason field).")
bullet(doc, "BR-005.4: Concurrent conflict resolution — If any entity in the handover request has an active delegation, upon authorization the delegation is auto-cancelled with status = EARLY_TERMINATED and reason = 'Superseded by handover #{handover_request_id}'. The permanent handover takes precedence over temporary delegation.")
bullet(doc, "BR-005.5: Cross-branch authorization routing — Handover authorization is routed to the Checker pool at the outgoing RM's branch. If no Checker is available at that branch, escalate to the next-level supervisory branch.")
bullet(doc, "BR-005.6: Optimistic locking — Authorize/Reject API must include the HandoverRequest.version field in the request body. If the version does not match the current database version, return HTTP 409 Conflict with message 'This request has been modified since you last viewed it. Please refresh and try again.' This prevents stale-state authorization.")

# --- FR-006 to FR-008: Delegation ---
h2(doc, "5.2 Delegation Module")

h3(doc, "FR-006: Lead Delegation")
para(doc, "Operations Maker can temporarily delegate leads from one RM to another within the same supervisor/branch for a specified date range. Delegations are auto-authorized.")
para(doc, 'User Story: "As an Operations Maker, I want to temporarily delegate leads to a covering RM during the primary RM\'s absence so that lead follow-up continues without gaps."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-006.1: Delegation screen accessible from Relationship > Delegation menu.")
bullet(doc, "AC-006.2: Lead Delegation tab shows Existing Leads with same columns as Lead Handover.")
bullet(doc, "AC-006.3: Selected Lead List, Delegate RM selection (limited to same supervisor/branch), Delegation Reason (mandatory).")
bullet(doc, "AC-006.4: Start Date and End Date fields with date pickers. Start >= today, End <= Start + 90 days.")
bullet(doc, "AC-006.5: On Save, delegation is immediately active (auto-authorized). No maker-checker required.")
bullet(doc, "AC-006.6: Success toast: 'Delegation Initiated Successfully'.")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-006.1: Delegate RM must be under the same supervisor or in the same branch as the outgoing RM.")
bullet(doc, "BR-006.2: Delegation period cannot exceed 90 calendar days.")
bullet(doc, "BR-006.3: An entity cannot have overlapping active delegations.")

h3(doc, "FR-007: Prospect Delegation")
para(doc, "Same workflow as Lead Delegation applied to prospects. Auto-authorized with date range.")
para(doc, 'User Story: "As an Operations Maker, I want to temporarily delegate prospects to a covering RM so that prospect engagement continues during the primary RM\'s absence."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-007.1: Prospect Delegation tab in Delegation screen.")
bullet(doc, "AC-007.2: Same flow as FR-006 with prospect-specific columns (Branch, Prospect ID, Prospect Name, RM ID - RM Name).")
bullet(doc, "AC-007.3: Date range validation and auto-authorization same as FR-006.")

h3(doc, "FR-008: Client Delegation")
para(doc, "Temporary delegation of clients with optional scrutiny checklist. Auto-authorized with date range.")
para(doc, 'User Story: "As an Operations Maker, I want to temporarily delegate clients to a covering RM so that client coverage continues during the primary RM\'s absence, with an optional scrutiny checklist for high-value transfers."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-008.1: Client Delegation tab in Delegation screen.")
bullet(doc, "AC-008.2: Same flow as FR-003 but with delegation-specific fields (Start Date, End Date) instead of permanent transfer.")
bullet(doc, "AC-008.3: Scrutiny Checklist is displayed but configurable — by default NOT required for delegations. System Administrator can enable mandatory checklist for delegations via ChecklistItemConfig.applies_to field ('HANDOVER_ONLY' | 'DELEGATION_ONLY' | 'BOTH').")
bullet(doc, "AC-008.4: Auto-authorized on save. No maker-checker required.")
bullet(doc, "AC-008.5: If a client already has an active delegation that overlaps with the new date range, system blocks the save with error: 'Client {name} already has an active delegation from {start} to {end}.'")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-008.1: Scrutiny checklist is mandatory for Client HANDOVER but optional/configurable for Client DELEGATION. Default configuration: not required for delegations.")
bullet(doc, "BR-008.2: When checklist is configured as mandatory for delegation, same validation rules apply as FR-003 (all mandatory items must be COMPLETED).")

h3(doc, "FR-009: Delegation Auto-Expiry")
para(doc, "System automatically reverts delegated entities to their original RM when the delegation end date is reached.")
para(doc, 'User Story: "As a system operator, I want delegations to automatically expire and revert so that RM assignments are accurate without manual intervention."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-009.1: Scheduled job runs daily at midnight to check for expired delegations (end_date < today).")
bullet(doc, "AC-009.2: For each expired delegation: update DelegationRequest.status to EXPIRED, set auto_revert_completed to true.")
bullet(doc, "AC-009.3: Revert RM assignments in master tables back to original RM.")
bullet(doc, "AC-009.4: Send notification to both original RM and delegate RM that delegation has expired.")
bullet(doc, "AC-009.5: Create audit log entry for each auto-expiry.")
bullet(doc, "AC-009.6: Send warning notification 2 days before expiry to both RMs.")

# --- FR-010 to FR-014: Dashboard & Supporting Features ---
h2(doc, "5.3 Dashboard & Supporting Features")

h3(doc, "FR-010: Handover Dashboard")
para(doc, "Consolidated dashboard showing handover and delegation activity for Operations users and Branch Managers.")
para(doc, 'User Story: "As a Branch Manager, I want a dashboard showing pending authorizations, recent transfers, and delegation coverage so that I can monitor RM capacity and ensure no service gaps."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-010.1: Dashboard shows 4 widgets: Pending Authorizations count, Recent Transfers (last 30 days), Active Delegations count, AUM Impact Summary.")
bullet(doc, "AC-010.2: Pending Authorizations widget shows count by type (Lead/Prospect/Client) with click-through to authorization queue.")
bullet(doc, "AC-010.3: Recent Transfers widget shows last 10 authorized handovers with entity type, count, from/to RM, date.")
bullet(doc, "AC-010.4: Active Delegations widget shows count with expiring-soon (within 7 days) highlighted.")
bullet(doc, "AC-010.5: AUM Impact widget shows total AUM of clients with pending handovers.")

h3(doc, "FR-011: Delegation Calendar")
para(doc, "Visual calendar showing delegation coverage across RMs.")
para(doc, 'User Story: "As an Operations Maker, I want to see a delegation calendar so that I can identify coverage gaps and plan delegations effectively."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-011.1: Calendar view (month/week toggle) showing delegation periods as colored bars.")
bullet(doc, "AC-011.2: Each bar shows: outgoing RM name → delegate RM name, entity count, date range.")
bullet(doc, "AC-011.3: Clicking a bar opens delegation details.")
bullet(doc, "AC-011.4: Filter by branch, RM, delegation type.")

h3(doc, "FR-012: Handover History & Audit Trail")
para(doc, "Searchable history of all handover and delegation actions.")
para(doc, 'User Story: "As a Compliance Officer, I want to search and export handover history so that I can meet regulatory audit requirements for fiduciary relationship transfers."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-012.1: History screen with filters: date range, handover type, status, RM, entity ID.")
bullet(doc, "AC-012.2: Results grid: Date, Action Type, Entity Type, Entity ID, Entity Name, From RM, To RM, Status, Initiated By, Authorized By.")
bullet(doc, "AC-012.3: Export to CSV/Excel functionality.")
bullet(doc, "AC-012.4: Drill-down to full audit trail for a specific handover showing all state transitions.")

h3(doc, "FR-013: Portfolio Impact Assessment")
para(doc, "Display AUM summary for clients being selected for handover to assess financial impact.")
para(doc, 'User Story: "As an Operations Maker, I want to see the AUM impact when selecting clients for handover so that management can assess the financial significance of the transfer."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-013.1: When clients are added to Selected Client List, show total AUM in a summary bar.")
bullet(doc, "AC-013.2: AUM values auto-populated from portfolio/account data.")
bullet(doc, "AC-013.3: AUM summary also visible in authorization detail view.")

h3(doc, "FR-014: Pending Activities Warning")
para(doc, "Alert Operations Maker when selected clients have pending business activities before handover.")
para(doc, 'User Story: "As an Operations Maker, I want to be warned if a client has pending orders or upcoming maturity dates so that I can ensure continuity of service during the handover."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-014.1: On selecting a client for handover, system checks for: open/pending orders, open service requests, maturity dates within next 30 days.")
bullet(doc, "AC-014.2: Warning icon displayed next to affected clients in Selected Client List.")
bullet(doc, "AC-014.3: Clicking warning icon shows details: 'Client has 3 pending orders, 1 open service request, 2 bonds maturing within 30 days.'")
bullet(doc, "AC-014.4: Warning is informational — does not block the handover.")

h3(doc, "FR-015: Notification System")
para(doc, "Email and in-app notifications for all handover and delegation lifecycle events.")
para(doc, 'User Story: "As a Relationship Manager, I want to receive notifications when clients are transferred to me or from me so that I can proactively engage newly assigned clients."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-015.1: Notifications sent for: Handover Initiated, Handover Authorized, Handover Rejected, Delegation Started, Delegation Expiring (2 days before), Delegation Expired.")
bullet(doc, "AC-015.2: Recipients: Outgoing RM, Incoming RM, Initiating Maker, Authorizing Checker (where applicable).")
bullet(doc, "AC-015.3: In-app notification bell with unread count badge.")
bullet(doc, "AC-015.4: Email notification with handover details and link to view in system.")
bullet(doc, "AC-015.5: Notification delivery includes retry mechanism: if email delivery fails, retry up to 3 times with exponential backoff (1 min, 5 min, 30 min). Failed notifications after all retries are logged to a dead-letter queue for manual reconciliation.")
bullet(doc, "AC-015.6: Daily reconciliation report lists any notifications that failed delivery in the last 24 hours, available to System Administrator.")

# --- FR-016: Batch Authorization ---
h2(doc, "5.4 Batch Operations")

h3(doc, "FR-016: Batch Authorization")
para(doc, "Operations Checker can authorize or reject multiple pending handover requests in a single action for high-volume scenarios (e.g., branch restructuring).")
para(doc, 'User Story: "As an Operations Checker, I want to batch-authorize multiple handover requests at once so that I can efficiently process high-volume handover scenarios without reviewing each request individually."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-016.1: Authorization queue displays checkboxes for multi-select of pending requests.")
bullet(doc, "AC-016.2: 'Batch Authorize' and 'Batch Reject' buttons appear when 2 or more requests are selected.")
bullet(doc, "AC-016.3: On Batch Authorize click, confirmation dialog shows: total request count, total entity count, total AUM impact, and list of incoming RMs.")
bullet(doc, "AC-016.4: User must type 'CONFIRM' in a text field within the dialog to proceed (prevents accidental batch authorization).")
bullet(doc, "AC-016.5: On Batch Reject, reject reason overlay appears (same as individual reject) with reason applied to all selected requests.")
bullet(doc, "AC-016.6: Each request in the batch is processed independently — one failure does not roll back others.")
bullet(doc, "AC-016.7: Results summary shown after batch operation: '{X} authorized, {Y} failed' with details for any failures.")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-016.1: Maximum 50 requests per batch operation.")
bullet(doc, "BR-016.2: Segregation of duties still enforced — any request initiated by the current Checker is automatically excluded from the batch with a warning message.")
bullet(doc, "BR-016.3: Optimistic locking applies — if any request's version has changed since page load, that request fails with 409 Conflict and others proceed.")
bullet(doc, "BR-016.4: Batch authorization generates individual audit log entries for each request (not a single batch entry).")

# --- FR-017: Delegation Extension ---
h3(doc, "FR-017: Delegation Extension")
para(doc, "Operations Maker can extend an active delegation by up to 90 additional days when the original delegation period is insufficient (e.g., extended medical leave).")
para(doc, 'User Story: "As an Operations Maker, I want to extend an active delegation beyond its original end date so that client coverage continues when an RM\'s absence is longer than initially planned."')
para(doc, "Acceptance Criteria:", bold=True)
bullet(doc, "AC-017.1: Active delegations in the Delegation Calendar or Active Delegations list show an 'Extend' button.")
bullet(doc, "AC-017.2: Clicking Extend opens a form with: current end date (read-only), new end date (date picker), extension reason (mandatory, min 5 chars).")
bullet(doc, "AC-017.3: New end date must be > current end date and <= current end date + 90 days.")
bullet(doc, "AC-017.4: Extension requires supervisor approval: system sends email to the outgoing RM's supervisor for confirmation. Supervisor receives a link to approve/deny the extension.")
bullet(doc, "AC-017.5: Until supervisor approves, delegation continues with original end date. On approval, end date is updated.")
bullet(doc, "AC-017.6: Maximum 1 extension per delegation. If further extension is needed, a new delegation must be created.")
bullet(doc, "AC-017.7: Audit log entry created for extension request and supervisor approval/denial.")
para(doc, "Business Rules:", bold=True)
bullet(doc, "BR-017.1: A delegation can only be extended once. After extension, the 'Extend' button is replaced with 'Extended until {new_end_date}'.")
bullet(doc, "BR-017.2: If the delegation is within 7 days of expiry, extension can still be requested but the supervisor approval window is reduced to 24 hours with escalation.")
bullet(doc, "BR-017.3: Total delegation period (original + extension) cannot exceed 180 calendar days.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6: USER INTERFACE REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "6. User Interface Requirements")

h2(doc, "6.1 Relationship Menu")
para(doc, "Entry point from the back-office Operations portal sidebar navigation.")
bullet(doc, "Menu item: 'Relationship' in left sidebar with icon.")
bullet(doc, "On click, shows card-based sub-menu with: Family Grouping, Handover, Delegation, FID Grouping.")
bullet(doc, "Each card shows title, description, and info icon.")
bullet(doc, "Search bar at top to filter cards.")

h2(doc, "6.2 Handover Screen")
para(doc, "Main handover screen with collapsible accordion tabs.")
bullet(doc, "Header: 'Handover' with breadcrumb navigation.")
bullet(doc, "Three collapsible accordion sections: Lead Handover, Prospect Handover, Client Handover.")
bullet(doc, "Only one section expanded at a time.")
bullet(doc, "Each section contains: Existing entities grid → Selected entities list → Incoming RM details → (Client only: Scrutiny Checklist) → Back/Save buttons.")
bullet(doc, "Back button at bottom returns to Relationship menu.")

h2(doc, "6.3 Entity Selection Grid (Existing Leads/Prospects/Clients)")
para(doc, "Reusable data grid pattern used across all handover and delegation screens.")
bullet(doc, "Checkbox column for multi-select (header checkbox for select all).")
bullet(doc, "Column headers with filter input fields directly below each header.")
bullet(doc, "Location dropdown filter and Preferred Language dropdown filter above the grid.")
bullet(doc, "Pagination controls at bottom (showing 'X of Y' and page navigation).")
bullet(doc, "Next button to move selected items to Selected List section.")

h2(doc, "6.4 Selected List Section")
para(doc, "Displays items selected for handover/delegation.")
bullet(doc, "Collapsible section with header showing item count.")
bullet(doc, "Grid with checkbox column, entity-specific columns, and outgoing RM details.")
bullet(doc, "For clients: expandable row showing Base No and Base Name sub-details.")

h2(doc, "6.5 Incoming RM/SRM Details Section")
para(doc, "Form section for specifying the destination RM.")
bullet(doc, "Incoming RM: searchable dropdown listing all active RMs (for delegation: limited to same supervisor/branch).")
bullet(doc, "Incoming SRM: auto-populated read-only field based on selected RM's supervisor.")
bullet(doc, "Handover Reason: text input field with required asterisk, 500 char limit.")
bullet(doc, "For delegation: additional Start Date and End Date fields with date pickers.")

h2(doc, "6.6 Scrutiny Checklist Section (Client Handover/Delegation Only)")
para(doc, "Validation checklist that must be completed before saving client transfers.")
bullet(doc, "Collapsible section: 'Scrutiny Check List'.")
bullet(doc, "Table with columns: Validation (label text), Status (dropdown: Not Started / Work In Progress / Completed), Remarks (free text).")
bullet(doc, "Mandatory items highlighted with asterisk.")
bullet(doc, "Save button disabled until all mandatory items are COMPLETED.")

h2(doc, "6.7 Authorization Screen")
para(doc, "Screen for Operations Checker to authorize/reject pending handovers.")
bullet(doc, "Navigation: Authorization > Authorize from sidebar.")
bullet(doc, "Searchable dropdown to select authorization type: Lead Handover, Prospect Handover, Client Handover, Lead Delegation, Prospect Delegation, Client Delegation.")
bullet(doc, "List view of pending records with filterable columns and Actions button (minus icon) per row.")
bullet(doc, "Detail view: read-only display of Selected Items, Incoming RM details, Scrutiny Checklist.")
bullet(doc, "Two buttons: Reject (opens overlay with Reject Reason text input + Reject/Close buttons) and Authorize.")
bullet(doc, "Pagination: 'X of Y' with first/prev/next/last page buttons.")

h2(doc, "6.8 Handover Dashboard Screen")
bullet(doc, "Four KPI cards at top: Pending Authorizations, Recent Transfers, Active Delegations, AUM Impact.")
bullet(doc, "Below cards: Recent Transfers table with entity type, count, from/to RM, date, status.")
bullet(doc, "Side panel: Delegation Calendar (mini calendar with colored bars).")

h2(doc, "6.9 File Upload Screen")
bullet(doc, "Navigation: Relationship > File Upload.")
bullet(doc, "Upload Type dropdown (Client Handover), File Browse button, file name display.")
bullet(doc, "Reset and Upload File buttons at bottom.")
bullet(doc, "After upload: progress indicator, then results summary with success/error count and error detail link.")

h2(doc, "6.10 Responsive Behavior")
bullet(doc, "Desktop-first design (1280px+ primary target for back-office).")
bullet(doc, "Tables horizontally scroll on smaller screens.")
bullet(doc, "Sidebar collapses to hamburger menu on tablet.")
bullet(doc, "Form sections stack vertically on narrow screens.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 7: API & INTEGRATION REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "7. API & Integration Requirements")

h2(doc, "7.1 Authentication")
para(doc, "All API endpoints require valid JWT token via httpOnly cookie. Role-based authorization enforced at endpoint level.")

h2(doc, "7.2 Standardized Error Response Format")
para(doc, 'All error responses follow this format: { "error": { "code": "VALIDATION_ERROR", "message": "Handover reason is required", "field": "handover_reason", "details": [] } }')

h2(doc, "7.3 API Endpoints")
add_styled_table(doc, ["Method", "Path", "Description", "Auth Role"], [
    ["GET", "/api/back-office/handover/leads", "List leads available for handover (filterable, paginated)", "ops-maker, ops-checker"],
    ["GET", "/api/back-office/handover/prospects", "List prospects available for handover", "ops-maker, ops-checker"],
    ["GET", "/api/back-office/handover/clients", "List clients available for handover", "ops-maker, ops-checker"],
    ["POST", "/api/back-office/handover/request", "Create a new handover request", "ops-maker"],
    ["GET", "/api/back-office/handover/request/:id", "Get handover request details", "ops-maker, ops-checker"],
    ["GET", "/api/back-office/handover/pending", "List pending handover requests for authorization", "ops-checker"],
    ["POST", "/api/back-office/handover/authorize/:id", "Authorize a pending handover", "ops-checker"],
    ["POST", "/api/back-office/handover/reject/:id", "Reject a pending handover", "ops-checker"],
    ["GET", "/api/back-office/handover/history", "Search handover history with filters", "ops-maker, ops-checker, branch-mgr"],
    ["POST", "/api/back-office/handover/bulk-upload", "Upload CSV for bulk client handover", "ops-maker"],
    ["GET", "/api/back-office/handover/upload-log/:id", "Get bulk upload processing results", "ops-maker"],
    ["GET", "/api/back-office/delegation/leads", "List leads available for delegation", "ops-maker"],
    ["GET", "/api/back-office/delegation/prospects", "List prospects available for delegation", "ops-maker"],
    ["GET", "/api/back-office/delegation/clients", "List clients available for delegation", "ops-maker"],
    ["POST", "/api/back-office/delegation/request", "Create a new delegation request", "ops-maker"],
    ["GET", "/api/back-office/delegation/active", "List active delegations", "ops-maker, ops-checker, branch-mgr"],
    ["POST", "/api/back-office/delegation/cancel/:id", "Cancel an active delegation", "ops-maker"],
    ["GET", "/api/back-office/delegation/calendar", "Get delegation calendar data", "ops-maker, ops-checker, branch-mgr"],
    ["GET", "/api/back-office/handover/dashboard", "Get dashboard summary data", "ops-maker, ops-checker, branch-mgr"],
    ["GET", "/api/back-office/handover/checklist-config", "Get active scrutiny checklist configuration", "ops-maker"],
    ["GET", "/api/back-office/handover/client-impact/:clientId", "Get client AUM and pending activities", "ops-maker, ops-checker"],
    ["GET", "/api/back-office/handover/rms", "List active RMs for dropdown population (filterable by branch, supervisor)", "ops-maker, ops-checker"],
    ["PATCH", "/api/back-office/handover/request/:id", "Amend a PENDING handover request (add/remove items, change incoming RM)", "ops-maker"],
    ["POST", "/api/back-office/delegation/extend/:id", "Extend an active delegation by up to 90 additional days (requires supervisor approval)", "ops-maker"],
    ["POST", "/api/back-office/handover/batch-authorize", "Batch authorize multiple pending handover requests", "ops-checker"],
    ["POST", "/api/back-office/handover/batch-reject", "Batch reject multiple pending handover requests", "ops-checker"],
    ["POST", "/api/back-office/handover/bulk-upload/preview", "Dry-run preview of bulk upload CSV without committing", "ops-maker"],
])

h2(doc, "7.4 Key Request/Response Examples")

h3(doc, "POST /api/back-office/handover/request — Create Handover")
para(doc, "Request Body:")
para(doc, '{ "handover_type": "CLIENT", "items": [ { "entity_id": "080820231", "entity_type": "CLIENT" }, { "entity_id": "100100231", "entity_type": "CLIENT" } ], "incoming_rm_id": "WFRM6", "handover_reason": "RM leaving organization effective April 30", "scrutiny_checklist": [ { "checklist_item_id": "cfg001", "status": "COMPLETED", "remarks": "All promises documented" }, { "checklist_item_id": "cfg002", "status": "COMPLETED", "remarks": "Final call completed" } ] }')

para(doc, "Success Response (201):")
para(doc, '{ "id": "a1b2c3d4", "handover_type": "CLIENT", "status": "PENDING", "outgoing_rm_id": "BASERM2", "outgoing_rm_name": "BASERMTWO TWO", "incoming_rm_id": "WFRM6", "incoming_rm_name": "Abhay RM", "incoming_srm_name": "Vikas B", "items_count": 2, "total_aum": 4300000.00, "created_at": "2026-04-22T10:30:00Z" }')

h3(doc, "POST /api/back-office/handover/reject/:id — Reject Handover")
para(doc, "Request Body:")
para(doc, '{ "reject_reason": "Incoming RM already at capacity. Please reassign to a different RM." }')

para(doc, "Success Response (200):")
para(doc, '{ "id": "a1b2c3d4", "status": "REJECTED", "reject_reason": "Incoming RM already at capacity. Please reassign to a different RM.", "authorized_by": "OPSCHECKER1", "authorized_at": "2026-04-22T14:15:00Z" }')

h3(doc, "POST /api/back-office/delegation/request — Create Delegation")
para(doc, "Request Body:")
para(doc, '{ "delegation_type": "CLIENT", "items": [ { "entity_id": "080820231", "entity_type": "CLIENT" } ], "delegate_rm_id": "BASERM3", "delegation_reason": "Annual leave", "start_date": "2026-04-25", "end_date": "2026-05-10" }')

para(doc, "Success Response (201):")
para(doc, '{ "id": "d1e2f3", "delegation_type": "CLIENT", "status": "ACTIVE", "outgoing_rm_id": "BASERM2", "delegate_rm_id": "BASERM3", "start_date": "2026-04-25", "end_date": "2026-05-10", "items_count": 1, "created_at": "2026-04-22T10:45:00Z" }')

h3(doc, "POST /api/back-office/handover/batch-authorize — Batch Authorize")
para(doc, "Request Body:")
para(doc, '{ "request_ids": ["a1b2c3d4", "e5f6g7h8", "m3n4o5p6"], "confirmation": "CONFIRM", "versions": {"a1b2c3d4": 1, "e5f6g7h8": 2, "m3n4o5p6": 1} }')
para(doc, "Success Response (200):")
para(doc, '{ "results": [ { "id": "a1b2c3d4", "status": "AUTHORIZED", "success": true }, { "id": "e5f6g7h8", "status": "AUTHORIZED", "success": true }, { "id": "m3n4o5p6", "status": "CONFLICT", "success": false, "error": "Version mismatch — request was modified" } ], "summary": { "total": 3, "authorized": 2, "failed": 1 } }')

h3(doc, "POST /api/back-office/delegation/extend/:id — Extend Delegation")
para(doc, "Request Body:")
para(doc, '{ "new_end_date": "2026-08-10", "extension_reason": "Extended medical leave recovery" }')
para(doc, "Success Response (200):")
para(doc, '{ "id": "d1e2f3", "status": "ACTIVE", "original_end_date": "2026-05-10", "new_end_date": "2026-08-10", "extension_count": 1, "supervisor_approval_status": "PENDING", "supervisor_id": "SRM001", "message": "Extension request sent to supervisor for approval" }')

h3(doc, "POST /api/back-office/handover/bulk-upload/preview — Bulk Upload Preview")
para(doc, "Request Body: multipart/form-data with CSV file")
para(doc, "Success Response (200):")
para(doc, '{ "preview": { "total_rows": 150, "valid_rows": 148, "invalid_rows": 2, "total_aum_impact": 45000000.00, "affected_clients": 148, "incoming_rms": [ {"rm_id": "WFRM6", "rm_name": "Abhay RM", "client_count": 100}, {"rm_id": "BASERM3", "rm_name": "Base RM Three", "client_count": 48} ], "active_delegations_affected": 3, "validation_errors": [ {"row": 45, "field": "client_id", "error": "Client ID not found"}, {"row": 102, "field": "incoming_rm_id", "error": "RM ID is inactive"} ] } }')

h2(doc, "7.5 Rate Limiting")
para(doc, "API rate limits: 100 requests per minute per user for read endpoints. 20 requests per minute per user for write endpoints. Bulk upload: 5 uploads per hour per user.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 8: NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "8. Non-Functional Requirements")

h2(doc, "8.1 Performance")
add_styled_table(doc, ["Metric", "Target"], [
    ["Entity list loading (leads/prospects/clients)", "< 2 seconds for up to 1,000 records"],
    ["Handover request save", "< 3 seconds"],
    ["Authorization action (authorize/reject)", "< 2 seconds"],
    ["Dashboard loading", "< 3 seconds"],
    ["Bulk upload processing", "< 5 minutes for 5,000 rows"],
    ["Audit history search", "< 3 seconds for 1-year date range"],
    ["Concurrent users", "50 simultaneous Operations users"],
])

h2(doc, "8.2 Security")
bullet(doc, "All endpoints require valid JWT token via httpOnly cookies (SEC-07 compliance).")
bullet(doc, "Role-based access control (RBAC) enforced at API gateway and endpoint level.")
bullet(doc, "Maker-Checker segregation: same user cannot both initiate and authorize a handover.")
bullet(doc, "All data transmission over HTTPS/TLS 1.2+.")
bullet(doc, "Audit logs are immutable (append-only, no update/delete).")
bullet(doc, "CSV upload: server-side validation of file content, sanitization against injection attacks.")
bullet(doc, "OWASP Top 10 compliance: input validation, SQL injection prevention, XSS prevention, CSRF tokens.")

h2(doc, "8.3 Scalability")
bullet(doc, "Horizontal scaling: stateless API servers behind load balancer.")
bullet(doc, "Database: connection pooling, read replicas for query-heavy operations (dashboard, history).")
bullet(doc, "Bulk upload processing via background job queue (not blocking API thread).")

h2(doc, "8.4 Availability")
bullet(doc, "Target: 99.9% uptime during business hours (6 AM - 10 PM local time).")
bullet(doc, "Scheduled maintenance window: Sunday 2 AM - 6 AM.")
bullet(doc, "Delegation auto-expiry job: fault-tolerant with retry mechanism.")

h2(doc, "8.5 Data Backup & Recovery")
bullet(doc, "Database backup: daily full backup, hourly incremental.")
bullet(doc, "RPO (Recovery Point Objective): 1 hour.")
bullet(doc, "RTO (Recovery Time Objective): 4 hours.")
bullet(doc, "Audit logs: retained for 7 years per regulatory requirement.")

h2(doc, "8.5.1 Data Archival Strategy")
para(doc, "To manage the growing volume of handover audit data while meeting the 7-year retention requirement:")
bullet(doc, "Partition HandoverAuditLog table by year (e.g., handover_audit_log_2026, handover_audit_log_2027).")
bullet(doc, "Active partition: current year + previous year remain in hot storage (primary database) for fast querying.")
bullet(doc, "Cold storage archival: partitions older than 2 years are archived to cold storage (compressed, read-only) but remain searchable via the Audit History screen with a 'Search archived data' toggle.")
bullet(doc, "Archived data queries may take up to 30 seconds (vs. <3 seconds for hot data). UI displays a loading indicator with message: 'Searching archived records...'")
bullet(doc, "Annual archival job runs on January 1st, moving the partition from 2 years ago to cold storage.")
bullet(doc, "Archived data is immutable and cannot be modified or deleted, even by System Administrators.")
bullet(doc, "Purge policy: data older than 7 years may be permanently purged after Compliance Officer sign-off.")

h2(doc, "8.6 Accessibility")
bullet(doc, "WCAG 2.1 AA compliance.")
bullet(doc, "Keyboard navigation for all interactive elements.")
bullet(doc, "Screen reader compatible table headers and form labels.")
bullet(doc, "Color contrast ratio: minimum 4.5:1.")

h2(doc, "8.7 Browser & Device Support")
add_styled_table(doc, ["Browser/Device", "Support Level"], [
    ["Chrome 90+", "Full support (primary)"],
    ["Firefox 90+", "Full support"],
    ["Safari 15+", "Full support"],
    ["Edge 90+", "Full support"],
    ["Desktop (1280px+)", "Primary target"],
    ["Tablet (768px-1279px)", "Supported with responsive layout"],
    ["Mobile (<768px)", "Not targeted (back-office application)"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 9: WORKFLOW & STATE DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "9. Workflow & State Diagrams")

h2(doc, "9.1 Handover Request Lifecycle")
add_styled_table(doc, ["Current State", "Action", "Next State", "Side Effects"], [
    ["(none)", "Maker saves handover", "PENDING", "HandoverRequest created; HandoverItems created; ScrutinyChecklist created (client only); Audit log: HANDOVER_INITIATED; Notifications sent to Checker, Incoming RM"],
    ["PENDING", "Checker authorizes", "AUTHORIZED", "RM assignments updated in entity master; Audit log: HANDOVER_AUTHORIZED; Notifications sent to Maker, Outgoing RM, Incoming RM; Entities appear in incoming RM workspace"],
    ["PENDING", "Checker rejects", "REJECTED", "Reject reason stored; Entities freed for new handover; Audit log: HANDOVER_REJECTED; Notification sent to Maker with reject reason"],
])

h2(doc, "9.2 Delegation Request Lifecycle")
add_styled_table(doc, ["Current State", "Action", "Next State", "Side Effects"], [
    ["(none)", "Maker saves delegation", "ACTIVE", "DelegationRequest created; DelegationItems created; RM assignments temporarily updated; Audit log: DELEGATION_CREATED; Notifications sent to Outgoing RM, Delegate RM"],
    ["ACTIVE", "End date reached (auto)", "EXPIRED", "RM assignments reverted to original; auto_revert_completed = true; Audit log: DELEGATION_EXPIRED; Notifications sent to both RMs"],
    ["ACTIVE", "Maker cancels", "CANCELLED", "RM assignments reverted immediately; Audit log: DELEGATION_CANCELLED; Notifications sent to both RMs"],
    ["ACTIVE", "Handover authorized for same entity", "EARLY_TERMINATED", "Delegation auto-terminated; reason = 'Superseded by handover #{id}'; Audit log: DELEGATION_EARLY_TERMINATED; No revert (entity now belongs to new RM)"],
    ["ACTIVE", "Maker requests extension", "(stays ACTIVE)", "Extension request created; Supervisor notified; Audit log: DELEGATION_EXTENSION_REQUESTED"],
    ["ACTIVE", "Supervisor approves extension", "(stays ACTIVE)", "end_date updated to new_end_date; extension_count = 1; Audit log: DELEGATION_EXTENDED; Notifications sent to both RMs"],
    ["ACTIVE", "2 days before end_date", "(stays ACTIVE)", "Warning notification sent to both RMs: 'Delegation expiring in 2 days'"],
])

h2(doc, "9.3 Bulk Upload Lifecycle")
add_styled_table(doc, ["Current State", "Action", "Next State", "Side Effects"], [
    ["(none)", "Maker uploads CSV", "PROCESSING", "BulkUploadLog created; File validated; Background job started"],
    ["PROCESSING", "All rows processed", "COMPLETED", "success_count and error_count updated; Audit log: BULK_UPLOAD; Error details stored as JSON"],
    ["PROCESSING", "System error during processing", "FAILED", "Partial results saved; Error details logged; Notification to Maker"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 10: NOTIFICATION & COMMUNICATION REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "10. Notification & Communication Requirements")

add_styled_table(doc, ["Event", "Channel", "Recipients", "Message Template"], [
    ["Handover Initiated", "Email + In-App", "Incoming RM, Ops Checker queue", "Subject: 'Handover Request #{id}: {count} {type}(s) pending your review' Body: Lists entities, outgoing RM, reason"],
    ["Handover Authorized", "Email + In-App", "Maker, Outgoing RM, Incoming RM", "Subject: 'Handover #{id} Authorized — {count} {type}(s) transferred' Body: Confirmation with entity list"],
    ["Handover Rejected", "Email + In-App", "Maker, Outgoing RM", "Subject: 'Handover #{id} Rejected' Body: Reject reason, instruction to re-submit"],
    ["Delegation Started", "Email + In-App", "Outgoing RM, Delegate RM", "Subject: 'Delegation Active: {count} {type}(s) delegated {start} to {end}' Body: Entity list, delegate RM details"],
    ["Delegation Expiring (2 days)", "Email", "Outgoing RM, Delegate RM", "Subject: 'Delegation Expiring: {count} {type}(s) reverting on {end_date}' Body: Reminder to prepare for revert"],
    ["Delegation Expired", "Email + In-App", "Outgoing RM, Delegate RM", "Subject: 'Delegation Expired: {count} {type}(s) reverted to {original_rm}' Body: Confirmation of revert"],
    ["Bulk Upload Completed", "In-App", "Uploading Maker", "Subject: 'Bulk Upload Complete: {success}/{total} processed' Body: Link to error log"],
    ["Bulk Upload Supervisor Alert", "Email", "Maker's Supervisor (Branch Mgr)", "Subject: 'Bulk Client Handover Alert: {count} clients, {aum} AUM' Body: File name, row count, AUM impact, link to upload log"],
    ["Delegation Extension Requested", "Email", "Outgoing RM's Supervisor", "Subject: 'Delegation Extension Request: {delegate_rm} covering for {outgoing_rm}' Body: Current/new end date, reason, approve/deny link"],
    ["Delegation Extension Approved", "Email + In-App", "Outgoing RM, Delegate RM, Maker", "Subject: 'Delegation Extended to {new_end_date}' Body: Updated delegation details"],
    ["Delegation Early Terminated", "Email + In-App", "Delegate RM", "Subject: 'Delegation Terminated: Superseded by Handover #{id}' Body: Explanation that permanent handover takes precedence"],
    ["Batch Authorization Complete", "In-App", "Authorizing Checker", "Subject: 'Batch Authorization: {authorized}/{total} processed' Body: Summary with any failures"],
])

para(doc, "Notification Preferences: Users cannot opt out of handover/delegation notifications as they are operationally critical. Email frequency is per-event (not batched).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 11: REPORTING & ANALYTICS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "11. Reporting & Analytics")

add_styled_table(doc, ["Report Name", "Audience", "Data Sources", "Filters", "Refresh"], [
    ["Handover Activity Report", "Operations Mgr, Compliance", "HandoverRequest, HandoverItem, AuditLog", "Date range, type, status, branch, RM", "Real-time"],
    ["Delegation Coverage Report", "Branch Manager", "DelegationRequest, DelegationItem", "Date range, branch, RM, status", "Real-time"],
    ["AUM Transfer Report", "Management", "HandoverItem, Portfolio data", "Date range, branch, RM, amount range", "Daily"],
    ["Authorization Turnaround Report", "Operations Mgr", "HandoverRequest (PENDING → AUTHORIZED timestamps)", "Date range, checker, type", "Real-time"],
    ["Bulk Upload Summary", "Operations Mgr", "BulkUploadLog", "Date range, uploader, status", "Real-time"],
    ["Regulatory Audit Trail", "Compliance, Auditors", "HandoverAuditLog", "Date range, action type, RM, entity", "Real-time"],
])

h2(doc, "11.1 Key Metric Calculations")
bullet(doc, "Average Authorization Turnaround: AVG(authorized_at - created_at) for AUTHORIZED requests in period.")
bullet(doc, "Delegation Coverage Rate: (days with active delegation / total RM absence days) × 100.")
bullet(doc, "Handover Volume Trend: COUNT(HandoverRequest) grouped by week/month.")
bullet(doc, "AUM Transfer Volume: SUM(HandoverItem.aum_amount) for AUTHORIZED handovers in period.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 12: MIGRATION & LAUNCH PLAN
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "12. Migration & Launch Plan")

h2(doc, "12.1 Data Migration")
bullet(doc, "Import existing RM-Client/Prospect/Lead assignments from current CRM or spreadsheets.")
bullet(doc, "Configure scrutiny checklist items per organizational policy.")
bullet(doc, "Map existing RM hierarchy (RM → Supervisor) in User Maintenance module.")

h2(doc, "12.2 Phased Rollout")
add_styled_table(doc, ["Phase", "Features", "Timeline"], [
    ["Phase 1 (MVP)", "Lead/Prospect/Client Handover with maker-checker authorization (incl. optimistic locking, cross-branch routing, concurrent conflict resolution), Scrutiny Checklist, Handover History", "4 weeks"],
    ["Phase 2", "Delegation module (Lead/Prospect/Client) with auto-expiry, Delegation Calendar, Delegation Extension with supervisor approval", "3 weeks"],
    ["Phase 3", "Bulk Upload with preview/dry-run and supervisor notification, Batch Authorization, Handover Dashboard, Portfolio Impact Assessment, Pending Activities Warning", "3 weeks"],
    ["Phase 4", "Notification System (incl. retry/dead-letter), Reporting & Analytics, Audit Trail Export, Data Archival Setup", "2 weeks"],
])

h2(doc, "12.3 Go-Live Checklist")
bullet(doc, "All RM hierarchy data loaded and verified in User Maintenance.")
bullet(doc, "Scrutiny checklist items configured and reviewed by Compliance.")
bullet(doc, "Operations Maker and Checker users created with correct roles.")
bullet(doc, "Email notification service configured and tested.")
bullet(doc, "Delegation auto-expiry scheduled job deployed and tested.")
bullet(doc, "UAT sign-off from Operations, Compliance, and Branch Management.")
bullet(doc, "Production data backup verified before go-live.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 13: GLOSSARY
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "13. Glossary")

add_styled_table(doc, ["Term", "Definition"], [
    ["RM (Relationship Manager)", "A wealth management professional assigned to manage client relationships, provide investment advice, and handle client portfolios."],
    ["SRM (Senior Relationship Manager) / Supervisor", "The supervisor of an RM, responsible for overseeing a team of RMs and approving certain actions."],
    ["Operations Maker", "A back-office operations user who initiates transactional requests such as handovers and delegations."],
    ["Operations Checker", "A back-office operations user who reviews and authorizes/rejects requests initiated by Makers."],
    ["Handover", "The permanent transfer of Leads, Prospects, or Clients from one RM to another."],
    ["Delegation", "The temporary reassignment of Leads, Prospects, or Clients from one RM to another for a specified date range."],
    ["Lead", "A potential client who has shown interest but has not yet been formally onboarded as a prospect or client."],
    ["Prospect", "An individual or entity that has been identified as a potential client and is in the onboarding pipeline."],
    ["Client", "A fully onboarded customer with active accounts and portfolios managed by the institution."],
    ["Scrutiny Checklist", "A set of validation items that must be completed before a client handover can be saved, ensuring all commitments and documentation are in order."],
    ["AUM (Assets Under Management)", "The total market value of investments managed on behalf of a client by the RM/institution."],
    ["Maker-Checker", "A dual-control authorization pattern where one user (Maker) initiates an action and a different user (Checker) must approve it."],
    ["Bulk Upload", "Mass data import via CSV file for processing multiple records in a single operation."],
    ["Auto-Expiry", "System mechanism that automatically reverts delegated relationships to the original RM when the delegation end date is reached."],
    ["Branch", "A physical office location of the wealth management institution where RMs are based."],
    ["Optimistic Locking", "A concurrency control mechanism where a version field prevents stale-state updates. The API consumer must include the current version; mismatches result in HTTP 409 Conflict."],
    ["Batch Authorization", "The ability for an Operations Checker to authorize or reject multiple pending handover requests in a single action."],
    ["Delegation Extension", "A workflow allowing an active delegation to be extended beyond its original end date, requiring supervisor approval."],
    ["Early Termination", "The automatic cancellation of an active delegation when a permanent handover is authorized for the same entity."],
    ["Cold Storage", "Compressed, read-only archive of historical audit data older than 2 years, queryable but with longer response times."],
    ["Dead-Letter Queue", "A holding queue for failed notification deliveries that have exhausted retry attempts, awaiting manual reconciliation."],
    ["Dry-Run/Preview", "A non-committing validation step for bulk uploads that shows impact (row count, AUM, errors) before actual processing."],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 14: APPENDICES
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "14. Appendices")

h2(doc, "14.1 Bulk Upload CSV Template")
add_styled_table(doc, ["Column", "Type", "Required", "Description"], [
    ["client_id", "VARCHAR(50)", "Yes", "Client ID from client master"],
    ["incoming_rm_id", "VARCHAR(50)", "Yes", "RM ID of the new relationship manager"],
    ["handover_reason", "TEXT", "Yes", "Reason for the handover (min 5 chars)"],
])

para(doc, "Sample CSV:")
para(doc, "client_id,incoming_rm_id,handover_reason")
para(doc, "080820231,WFRM6,Branch restructuring")
para(doc, "100100231,WFRM6,Branch restructuring")
para(doc, "210820231,BASERM3,Client request for RM change")

h2(doc, "14.2 Reference Implementations")
add_styled_table(doc, ["Institution", "Version", "Key Differentiators"], [
    ["Maybank Asset Management", "V12 (Feb 2025)", "Prospect + Client handover; no Lead handover; Agent/SAgent terminology; Maker-Checker authorization"],
    ["HSBC Private Wealth Solutions", "V10 (Nov 2023)", "Lead + Prospect + Client handover; Lead delegation; RM/SRM terminology; Branch RM and Referring RM fields"],
    ["Jio BlackRock (JBR)", "V10 (Nov 2023)", "Lead + Prospect + Client handover; Same structure as HSBC; Base product revision"],
])

h2(doc, "14.3 Regulatory Requirements")
bullet(doc, "Fiduciary transfer records must be retained for minimum 7 years.")
bullet(doc, "Client notification of RM change may be required by local regulations (jurisdiction-specific).")
bullet(doc, "Audit trail must capture: who initiated, who authorized, when, reason, and all affected entities.")
bullet(doc, "Segregation of duties (Maker ≠ Checker) is a regulatory control requirement.")

# ── Save ─────────────────────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(__file__), "Handover_Assignment_Management_BRD_v4_FINAL.docx")
doc.save(out_path)
print(f"BRD saved to: {out_path}")
