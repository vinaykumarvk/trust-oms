#!/usr/bin/env python3
"""Generate Risk Profiling & Proposal Generation Management BRD as .docx"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Helpers ──────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with styled header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Arial'
        set_cell_shading(cell, '1F3864')
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
            if r_idx % 2 == 1:
                set_cell_shading(cell, 'D6E4F0')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)
        run.font.name = 'Arial'
    return p

def h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 56, 100)
        run.font.name = 'Arial'
    return p

def h3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(47, 84, 150)
        run.font.name = 'Arial'
    return p

def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.bold = bold
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 + level * 1.27)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    return p

def numbered(doc, text):
    p = doc.add_paragraph(text, style='List Number')
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    return p

# ── Main Document ────────────────────────────────────────────────────────

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Header
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.text = "CONFIDENTIAL — Trust OMS Risk Profiling & Proposal Generation Management BRD"
hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in hp.runs:
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

# Footer with page numbers
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = fp.add_run("Page ")
run.font.size = Pt(8)
fld = OxmlElement('w:fldSimple')
fld.set(qn('w:instr'), 'PAGE')
fp._p.append(fld)

# ── Title Page ───────────────────────────────────────────────────────────

for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Business Requirements Document")
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(31, 56, 100)
run.font.name = 'Arial'
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Risk Profiling & Proposal Generation Management")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(47, 84, 150)
run.font.name = 'Arial'

proj = doc.add_paragraph()
proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = proj.add_run("Trust OMS — Wealth Management Platform")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(89, 89, 89)
run.font.name = 'Arial'

for _ in range(4):
    doc.add_paragraph()

info_data = [
    ["Document ID", "BRD-TRUSTOMS-RPPG-2026-001"],
    ["Version", "2.0 (Post-Adversarial Review)"],
    ["Date", "2026-04-22"],
    ["Status", "Final"],
    ["Classification", "Confidential"],
    ["Author", "Trust OMS Platform Team"],
]
add_styled_table(doc, ["Property", "Value"], info_data)

doc.add_page_break()

# ── Table of Contents (placeholder) ──────────────────────────────────────

h1(doc, "Table of Contents")
para(doc, "[Table of Contents — auto-generated in Word: Insert > Table of Contents]")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "1. Executive Summary")

h2(doc, "1.1 Project Name")
para(doc, "Trust OMS — Risk Profiling & Proposal Generation Management Module (RP-PGM)")

h2(doc, "1.2 Project Description")
para(doc, "The Risk Profiling & Proposal Generation Management module is a core component of the Trust OMS wealth management platform. It delivers a comprehensive, configurable risk profiling engine and an intelligent investment proposal generation system designed for private banking and wealth management institutions operating across multiple jurisdictions. The module enables Relationship Managers (RMs) to assess client risk appetite through configurable scored and non-scored questionnaires, automatically compute risk categories with model asset allocations, and generate compliant investment proposals that match client suitability profiles. The system supports maker-checker authorization workflows, multi-entity deployments, regulatory compliance across jurisdictions (SEBI, MiFID II, MAS), and provides both back-office administration and client-portal self-service capabilities. The design is validated against five reference implementations: 360One (IIFL Wealth), HSBC Private Wealth Solutions, SBI Wealth, SBI International, and Jio BlackRock (JBR).")

h2(doc, "1.3 Business Objectives")
objectives = [
    "Digitize and standardize risk profiling across all client segments (Individual, Non-Individual, Both) with configurable multi-part questionnaires supporting scored and non-scored assessment types.",
    "Automate risk score computation, risk appetite classification, and model asset allocation recommendation to reduce manual errors and ensure regulatory compliance.",
    "Enable investment proposal generation that enforces suitability rules, concentration limits, and product risk deviation checks in real time.",
    "Support multi-entity, multi-jurisdiction deployments with configurable workflows, disclaimers, and regulatory rules per legal entity.",
    "Provide end-to-end audit trail, maker-checker authorization, and supervisor approval workflows to meet internal control and regulatory requirements.",
]
for o in objectives:
    bullet(doc, o)

h2(doc, "1.4 Target Users and Pain Points")
add_styled_table(doc, ["User Role", "Pain Points Addressed"], [
    ["Operations Administrator", "Manual questionnaire configuration is error-prone; no version control; no maker-checker enforcement on config changes."],
    ["Relationship Manager (RM)", "Paper-based risk profiling is slow and non-compliant; proposal creation is manual with no suitability validation; product-risk mismatches are caught late."],
    ["Supervisor / Team Lead", "No visibility into RM lead pipeline, risk profiling completion rates, or proposal approval backlogs."],
    ["Compliance Officer", "Difficulty auditing suitability assessments; no systematic tracking of risk deviation acknowledgements."],
    ["Client (via Client Portal)", "Lack of self-service risk profiling; cannot view/accept proposals digitally; no transparency into risk-product alignment."],
])

h2(doc, "1.5 Success Metrics (KPIs)")
add_styled_table(doc, ["KPI", "Target", "Measurement Method"], [
    ["Risk profiling completion rate", ">95% of new clients profiled within 48h of onboarding", "System timestamp tracking"],
    ["Questionnaire configuration time", "<30 min per new questionnaire (vs 2h manual)", "Admin activity logs"],
    ["Proposal generation time", "<15 min per proposal (vs 60 min manual)", "RM activity logs"],
    ["Product suitability violation rate", "<2% post-implementation (vs ~12% baseline)", "Deviation alert logs"],
    ["Proposal acceptance rate", ">70% first-pass acceptance", "Proposal lifecycle tracking"],
    ["Supervisor approval SLA", "100% approvals within 24h", "Approval queue timestamps"],
    ["System availability", "99.9% uptime during business hours", "Infrastructure monitoring"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 2: SCOPE & BOUNDARIES
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "2. Scope & Boundaries")

h2(doc, "2.1 In Scope")
in_scope = [
    "Risk Profile Questionnaire Maintenance (CRUD + maker-checker authorization) — Back Office",
    "Risk Appetite Mapping Maintenance (score-to-category mapping with CRUD + authorization)",
    "Asset Allocation Configuration (model portfolios per risk category with CRUD + authorization)",
    "Customer Risk Assessment Journey (3-step wizard: Edit Profile → Assess Risk → Transact)",
    "Risk Score Computation Engine (real-time scoring with normalization for multi-select questions)",
    "Recommended Risk Profile Display (score, category, donut chart, expected return, std deviation)",
    "Risk Profile Deviation handling (optional customer override with configurable enable/disable)",
    "Product Risk Deviation Alerts (popup warning when product risk > client risk appetite)",
    "Supervisor Approval Workflow for risk profile changes",
    "Supervisor Dashboard — Lead Status Widget (Level 1 chart + Level 2 drill-down table)",
    "Investment Proposal Builder (model portfolio auto-suggestion, what-if analysis, product selection)",
    "Proposal Compliance & Suitability Checks (automated validation against risk profile and mandates)",
    "Proposal Document Generation (PDF with branding, charts, disclaimers, digital signature)",
    "Proposal Approval Workflow (multi-level: RM → Team Lead → Compliance → Client)",
    "Model Portfolio Management (define, rebalance, track performance)",
    "Client Portal — Risk Profile view, Proposal view/accept/reject, PDF download",
    "Reporting: Transaction by Product Rating, Product Risk Mismatch, Proposal Pipeline, Risk Distribution",
    "Warning, Acknowledgement, Disclaimer configuration per questionnaire",
    "Biometrics / ID Image Capture and E-Signature integration points",
    "Audit trail for all risk profiling and proposal actions",
]
for s in in_scope:
    bullet(doc, s)

h2(doc, "2.2 Out of Scope")
out_scope = [
    "Customer onboarding / KYC — handled by separate Trust OMS module",
    "Order execution and settlement — downstream from proposal acceptance",
    "Portfolio rebalancing execution — only triggers and recommendations are in scope",
    "Third-party product data feeds (CRISIL, Morningstar) — integration interface defined, but feed ingestion is separate",
    "Regulatory reporting file generation (SEBI/MAS XML formats) — data is provided, format generation is separate",
    "Mobile native app — responsive web only in this phase",
    "AI/ML-based risk profiling or robo-advisory — future phase",
]
for s in out_scope:
    bullet(doc, s)

h2(doc, "2.3 Assumptions")
assumptions = [
    "Users are authenticated via the Trust OMS central authentication service (httpOnly cookie-based JWT).",
    "Role-based access control (RBAC) is enforced by the existing Trust OMS middleware.",
    "Product master data (schemes, NAV, risk ratings, ISIN) is available from existing Trust OMS product service.",
    "Customer master data (demographics, accounts, holdings) is available from existing Trust OMS customer service.",
    "Legal entity and service unit configuration is pre-defined in the Trust OMS organization structure.",
    "PDF generation uses server-side rendering (e.g., Puppeteer or equivalent).",
    "E-signature integration is available via a pluggable adapter interface.",
    "All dates are stored in ISO 8601 format (YYYY-MM-DD); times in UTC.",
    "Each legal entity operates in a single base currency. Proposals are denominated in the entity's base currency. Multi-currency proposals (e.g., NRI accounts with USD/GBP) are out of scope for v1 and will be addressed in a future phase with a currency conversion layer.",
    "Questionnaire content (questions, answers, warnings, disclaimers) is authored in the entity's primary language. Multi-language content support for international deployments will be addressed via a content translation table in Phase 2 (see FR-041).",
]
for a in assumptions:
    bullet(doc, a)

h2(doc, "2.4 Constraints")
constraints = [
    "Must support multi-entity deployments where each entity may have different questionnaire configurations, risk categories, and regulatory requirements.",
    "Must comply with SEBI Suitability Assessment Framework (SAF) for Indian wealth management entities.",
    "All configuration changes (questionnaires, risk mappings, asset allocations) must go through maker-checker authorization.",
    "Risk profiling questionnaire responses must be immutable once submitted (new assessment creates new version).",
    "Must support concurrent usage by 500+ RMs during business hours.",
    "Must integrate with existing Trust OMS tech stack: React/TypeScript frontend, Node.js/Express backend, PostgreSQL database.",
]
for c in constraints:
    bullet(doc, c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 3: USER ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "3. User Roles & Permissions")

para(doc, "The following roles interact with the Risk Profiling & Proposal Generation Management module. Permissions are cumulative where noted.")

h2(doc, "3.1 Role Definitions")

roles = [
    ["Operations Administrator", "Configures and maintains risk profiling questionnaires, risk appetite mappings, and asset allocations in the back office. Acts as Maker in maker-checker workflows."],
    ["Operations Supervisor", "Reviews and authorizes/rejects configuration changes made by Operations Administrators. Acts as Checker in maker-checker workflows."],
    ["Relationship Manager (RM)", "Conducts customer risk profiling, creates and manages investment proposals, selects products for clients."],
    ["RM Supervisor / Team Lead", "Reviews and approves risk profile deviations, monitors RM team performance via dashboard widgets, approves proposals at first level."],
    ["Compliance Officer", "Reviews proposals for regulatory compliance, approves/rejects at compliance level, accesses audit trails and suitability reports."],
    ["Client (Portal User)", "Views own risk profile, responds to risk questionnaires (if self-service enabled), views/accepts/rejects proposals, downloads proposal PDFs."],
]
add_styled_table(doc, ["Role", "Description"], roles)

h2(doc, "3.2 Permissions Matrix")

perm_headers = ["Feature", "Ops Admin", "Ops Supervisor", "RM", "RM Supervisor", "Compliance", "Client"]
perm_rows = [
    ["Questionnaire CRUD", "Create/Edit", "View/Authorize", "No", "No", "View", "No"],
    ["Risk Appetite Mapping CRUD", "Create/Edit", "View/Authorize", "No", "No", "View", "No"],
    ["Asset Allocation CRUD", "Create/Edit", "View/Authorize", "No", "No", "View", "No"],
    ["Conduct Risk Profiling", "No", "No", "Execute", "View/Override", "View", "Self-service (if enabled)"],
    ["Risk Profile Deviation Approval", "No", "No", "Request", "Approve/Reject", "View", "No"],
    ["Create Proposal", "No", "No", "Create/Edit", "View", "View", "No"],
    ["Approve Proposal (L1)", "No", "No", "No", "Approve/Reject", "No", "No"],
    ["Approve Proposal (Compliance)", "No", "No", "No", "No", "Approve/Reject", "No"],
    ["Accept/Reject Proposal", "No", "No", "No", "No", "No", "Accept/Reject"],
    ["View Reports", "Config reports", "Config reports", "Own clients", "Team reports", "All reports", "Own data"],
    ["Supervisor Dashboard", "No", "No", "No", "Full access", "View", "No"],
    ["Audit Trail", "Own actions", "Own + team", "Own actions", "Team actions", "All actions", "No"],
]
add_styled_table(doc, perm_headers, perm_rows)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 4: DATA MODEL
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "4. Data Model")

para(doc, "This section defines every entity required for the Risk Profiling & Proposal Generation Management module. All entities include standard audit columns (created_at, updated_at, created_by, updated_by) and soft-delete flag (is_deleted).")

# 4.1 Questionnaire
h2(doc, "4.1 Entity: Questionnaire")
para(doc, "Stores risk profile questionnaire configurations. Each questionnaire is uniquely identified by the combination of customer category, questionnaire type, and effective date range.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["questionnaire_name", "VARCHAR(200)", "Yes", "Min 3 chars, max 200, unique per active period", "—"],
    ["customer_category", "ENUM", "Yes", "Individual | Non-Individual | Both", "—"],
    ["questionnaire_type", "ENUM", "Yes", "FINANCIAL_PROFILING | INVESTMENT_KNOWLEDGE | SAF | SURVEY | FATCA | PAMM_PRE_INVESTMENT", "—"],
    ["effective_start_date", "DATE", "Yes", "Must be >= today for new records; ISO 8601", "—"],
    ["effective_end_date", "DATE", "Yes", "Must be > effective_start_date", "—"],
    ["valid_period_years", "INTEGER", "Yes", "1-10; determines risk profile expiry date", "2"],
    ["is_score", "BOOLEAN", "Yes", "true = scored questionnaire, false = non-scored", "false"],
    ["warning_text", "TEXT", "No", "Rich text for warning display", "NULL"],
    ["acknowledgement_text", "TEXT", "No", "Rich text for acknowledgement display", "NULL"],
    ["disclaimer_text", "TEXT", "No", "Rich text for disclaimer display", "NULL"],
    ["status", "ENUM", "Yes", "UNAUTHORIZED | MODIFIED | AUTHORIZED | REJECTED", "UNAUTHORIZED"],
    ["maker_id", "UUID (FK→User)", "Yes", "References user who created/modified", "Current user"],
    ["checker_id", "UUID (FK→User)", "No", "References user who authorized/rejected", "NULL"],
    ["authorized_at", "TIMESTAMP", "No", "Set when status → AUTHORIZED", "NULL"],
    ["version", "INTEGER", "Yes", "Incremented on each modification", "1"],
    ["entity_id", "UUID (FK→LegalEntity)", "Yes", "Legal entity this questionnaire belongs to", "—"],
    ["is_deleted", "BOOLEAN", "Yes", "Soft delete flag", "false"],
    ["created_at", "TIMESTAMP", "Yes", "Auto-set on insert", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "Auto-set on update", "NOW()"],
    ["created_by", "UUID (FK→User)", "Yes", "User who created", "Current user"],
    ["updated_by", "UUID (FK→User)", "Yes", "User who last updated", "Current user"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["id", "questionnaire_name", "customer_category", "questionnaire_type", "effective_start_date", "effective_end_date", "valid_period_years", "is_score", "status"], [
    ["a1b2c3d4...", "SAF Questionnaire 2026", "Both", "SAF", "2026-01-01", "2026-12-31", "2", "true", "AUTHORIZED"],
    ["e5f6g7h8...", "Financial Profiling Q1", "Individual", "FINANCIAL_PROFILING", "2026-01-01", "2026-06-30", "1", "false", "AUTHORIZED"],
    ["i9j0k1l2...", "Investment Knowledge", "Both", "INVESTMENT_KNOWLEDGE", "2026-04-01", "2027-03-31", "2", "false", "UNAUTHORIZED"],
])

para(doc, "Uniqueness Constraint: Only one questionnaire per (customer_category, questionnaire_type, overlapping date range) per entity. If customer_category = 'Both', no Individual or Non-Individual questionnaire may exist for overlapping dates and same type.", bold=True)

# 4.2 Question
h2(doc, "4.2 Entity: Question")
para(doc, "Individual questions within a questionnaire. Ordered by question_number.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["questionnaire_id", "UUID (FK→Questionnaire)", "Yes", "Must reference existing questionnaire", "—"],
    ["question_number", "INTEGER", "Yes", "Sequential within questionnaire, >=1", "Auto-increment"],
    ["question_description", "TEXT", "Yes", "Min 10 chars, max 1000", "—"],
    ["is_mandatory", "BOOLEAN", "Yes", "Whether customer must answer", "true"],
    ["is_multi_select", "BOOLEAN", "Yes", "true = checkboxes, false = radio buttons", "false"],
    ["scoring_type", "ENUM", "Yes", "NONE | RANGE", "NONE"],
    ["computation_type", "ENUM", "Yes", "SUM | NONE", "NONE"],
    ["is_deleted", "BOOLEAN", "Yes", "Soft delete", "false"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["id", "questionnaire_id", "question_number", "question_description", "is_mandatory", "is_multi_select", "scoring_type", "computation_type"], [
    ["q1...", "a1b2c3d4...", "1", "What is your annual household income?", "true", "false", "NONE", "SUM"],
    ["q2...", "a1b2c3d4...", "2", "Which products have you invested in before?", "true", "true", "RANGE", "SUM"],
    ["q3...", "a1b2c3d4...", "3", "What is your investment time horizon?", "true", "false", "NONE", "SUM"],
])

para(doc, "Business Rule: When is_multi_select = false, scoring_type is auto-set to NONE and is non-editable.", bold=True)

# 4.3 Answer Option
h2(doc, "4.3 Entity: AnswerOption")
para(doc, "Predefined answer choices for each question with associated weightage scores.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["question_id", "UUID (FK→Question)", "Yes", "Must reference existing question", "—"],
    ["option_number", "INTEGER", "Yes", "Sequential within question, >=1", "Auto-increment"],
    ["answer_description", "VARCHAR(500)", "Yes", "Min 1 char, max 500", "—"],
    ["weightage", "DECIMAL(5,2)", "Yes", ">=0; score assigned when this option is selected", "0"],
    ["is_deleted", "BOOLEAN", "Yes", "Soft delete", "false"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["id", "question_id", "option_number", "answer_description", "weightage"], [
    ["ao1...", "q2...", "1", "Fixed Deposits", "1"],
    ["ao2...", "q2...", "2", "Mutual Funds", "2"],
    ["ao3...", "q2...", "3", "Bonds", "3"],
    ["ao4...", "q2...", "4", "Equity", "4"],
])

# 4.4 Score Normalization Range
h2(doc, "4.4 Entity: ScoreNormalizationRange")
para(doc, "For multi-select questions with scoring_type = RANGE, defines how raw cumulative scores are normalized to a single score value. Example: if customer selects Fixed Deposits (1) + Mutual Funds (2) + Equity (4) = raw score 7, and range 6-10 maps to normalized score 3, the question contributes 3 to the total.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["question_id", "UUID (FK→Question)", "Yes", "Must reference question with scoring_type=RANGE", "—"],
    ["range_from", "DECIMAL(5,2)", "Yes", ">=0; must not overlap with other ranges for same question", "—"],
    ["range_to", "DECIMAL(5,2)", "Yes", "> range_from", "—"],
    ["normalized_score", "DECIMAL(5,2)", "Yes", ">=0", "—"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["id", "question_id", "range_from", "range_to", "normalized_score"], [
    ["snr1...", "q2...", "0", "5", "1"],
    ["snr2...", "q2...", "6", "10", "3"],
    ["snr3...", "q2...", "11", "15", "5"],
])

# 4.5 Risk Appetite Mapping
h2(doc, "4.5 Entity: RiskAppetiteMapping")
para(doc, "Maps computed total risk scores to risk appetite categories. Each mapping defines a score range and the corresponding risk category.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["mapping_name", "VARCHAR(200)", "Yes", "Descriptive name", "—"],
    ["entity_id", "UUID (FK→LegalEntity)", "Yes", "Legal entity", "—"],
    ["effective_start_date", "DATE", "Yes", ">= today for new", "—"],
    ["effective_end_date", "DATE", "Yes", "> start date", "—"],
    ["status", "ENUM", "Yes", "UNAUTHORIZED | MODIFIED | AUTHORIZED | REJECTED", "UNAUTHORIZED"],
    ["maker_id", "UUID (FK→User)", "Yes", "", "Current user"],
    ["checker_id", "UUID (FK→User)", "No", "", "NULL"],
    ["version", "INTEGER", "Yes", "", "1"],
    ["is_deleted", "BOOLEAN", "Yes", "", "false"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["created_by", "UUID (FK→User)", "Yes", "", "Current user"],
    ["updated_by", "UUID (FK→User)", "Yes", "", "Current user"],
])

# 4.6 Risk Appetite Band
h2(doc, "4.6 Entity: RiskAppetiteBand")
para(doc, "Individual bands within a risk appetite mapping defining score-to-category assignments.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["mapping_id", "UUID (FK→RiskAppetiteMapping)", "Yes", "", "—"],
    ["score_from", "DECIMAL(5,2)", "Yes", ">=0; no overlap with other bands", "—"],
    ["score_to", "DECIMAL(5,2)", "Yes", "> score_from", "—"],
    ["risk_category", "VARCHAR(100)", "Yes", "e.g., Conservative, Moderate, Aggressive", "—"],
    ["risk_code", "INTEGER", "Yes", "1-6 numeric code (1=lowest, 6=highest)", "—"],
    ["description", "TEXT", "No", "Category description", "NULL"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["id", "mapping_id", "score_from", "score_to", "risk_category", "risk_code"], [
    ["rab1...", "ram1...", "0", "5", "Conservative", "1"],
    ["rab2...", "ram1...", "6", "10", "Low to Moderate", "2"],
    ["rab3...", "ram1...", "11", "15", "Moderate", "3"],
    ["rab4...", "ram1...", "16", "20", "Moderately High", "4"],
    ["rab5...", "ram1...", "21", "25", "Aggressive", "5"],
    ["rab6...", "ram1...", "26", "30", "Very Aggressive", "6"],
])

# 4.7 Asset Allocation Config
h2(doc, "4.7 Entity: AssetAllocationConfig")
para(doc, "Model portfolio asset allocation per risk appetite category.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["config_name", "VARCHAR(200)", "Yes", "", "—"],
    ["entity_id", "UUID (FK→LegalEntity)", "Yes", "", "—"],
    ["effective_start_date", "DATE", "Yes", "", "—"],
    ["effective_end_date", "DATE", "Yes", "", "—"],
    ["status", "ENUM", "Yes", "UNAUTHORIZED | MODIFIED | AUTHORIZED | REJECTED", "UNAUTHORIZED"],
    ["maker_id", "UUID (FK→User)", "Yes", "", "Current user"],
    ["checker_id", "UUID (FK→User)", "No", "", "NULL"],
    ["version", "INTEGER", "Yes", "", "1"],
    ["is_deleted", "BOOLEAN", "Yes", "", "false"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["created_by", "UUID (FK→User)", "Yes", "", "Current user"],
    ["updated_by", "UUID (FK→User)", "Yes", "", "Current user"],
])

# 4.8 Asset Allocation Line
h2(doc, "4.8 Entity: AssetAllocationLine")
para(doc, "Individual asset class allocation within a model portfolio for a specific risk category.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["config_id", "UUID (FK→AssetAllocationConfig)", "Yes", "", "—"],
    ["risk_category", "VARCHAR(100)", "Yes", "Must match a category from RiskAppetiteBand", "—"],
    ["asset_class", "VARCHAR(100)", "Yes", "e.g., Equity, Fixed Income, Alternatives, Cash", "—"],
    ["allocation_percentage", "DECIMAL(5,2)", "Yes", "0-100; all lines per risk_category must sum to 100", "—"],
    ["expected_return_pct", "DECIMAL(5,2)", "No", "Annualized expected return %", "NULL"],
    ["standard_deviation_pct", "DECIMAL(5,2)", "No", "Annualized std deviation %", "NULL"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["config_id", "risk_category", "asset_class", "allocation_percentage", "expected_return_pct", "standard_deviation_pct"], [
    ["aac1...", "Conservative", "Fixed Income", "70.00", "7.50", "3.20"],
    ["aac1...", "Conservative", "Equity", "15.00", "12.00", "15.00"],
    ["aac1...", "Conservative", "Cash/Money Market", "15.00", "4.00", "0.50"],
    ["aac1...", "Aggressive", "Equity", "60.00", "14.00", "18.00"],
    ["aac1...", "Aggressive", "Fixed Income", "25.00", "8.00", "4.00"],
    ["aac1...", "Aggressive", "Alternatives", "15.00", "16.00", "22.00"],
])

# 4.9 Customer Risk Profile
h2(doc, "4.9 Entity: CustomerRiskProfile")
para(doc, "Stores the result of a customer's risk profiling assessment. A new record is created for each assessment (immutable — no edits, only new versions).")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["customer_id", "UUID (FK→Customer)", "Yes", "Must reference existing customer", "—"],
    ["questionnaire_id", "UUID (FK→Questionnaire)", "Yes", "The questionnaire version used", "—"],
    ["assessment_date", "DATE", "Yes", "Date assessment was completed", "TODAY"],
    ["expiry_date", "DATE", "Yes", "assessment_date + questionnaire.valid_period_years", "Computed"],
    ["total_raw_score", "DECIMAL(7,2)", "No", "Sum of all question scores", "NULL"],
    ["computed_risk_category", "VARCHAR(100)", "Yes", "System-computed category from RiskAppetiteBand", "—"],
    ["computed_risk_code", "INTEGER", "Yes", "1-6 numeric code", "—"],
    ["is_deviated", "BOOLEAN", "Yes", "true if customer deviated from computed profile", "false"],
    ["deviated_risk_category", "VARCHAR(100)", "No", "Customer-chosen category (if deviated)", "NULL"],
    ["deviated_risk_code", "INTEGER", "No", "Customer-chosen code (if deviated)", "NULL"],
    ["deviation_reason", "TEXT", "No", "Free-text reason for deviation", "NULL"],
    ["effective_risk_category", "VARCHAR(100)", "Yes", "= deviated_risk_category if deviated, else computed", "Computed"],
    ["effective_risk_code", "INTEGER", "Yes", "= deviated_risk_code if deviated, else computed", "Computed"],
    ["supervisor_approved", "BOOLEAN", "No", "Whether supervisor approved this profile", "NULL"],
    ["supervisor_id", "UUID (FK→User)", "No", "Approving supervisor", "NULL"],
    ["supervisor_approved_at", "TIMESTAMP", "No", "", "NULL"],
    ["acknowledgement_accepted", "BOOLEAN", "No", "Client acknowledgement checkbox", "NULL"],
    ["disclaimer_accepted", "BOOLEAN", "No", "Client disclaimer checkbox", "NULL"],
    ["is_active", "BOOLEAN", "Yes", "Only one active profile per customer at a time", "true"],
    ["assessed_by", "UUID (FK→User)", "Yes", "RM or system who conducted assessment", "Current user"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["customer_id", "assessment_date", "expiry_date", "total_raw_score", "computed_risk_category", "is_deviated", "effective_risk_category", "is_active"], [
    ["cust001", "2026-04-20", "2028-04-20", "18.5", "Moderately High", "false", "Moderately High", "true"],
    ["cust002", "2026-03-15", "2028-03-15", "8.0", "Low to Moderate", "true", "Moderate", "true"],
    ["cust003", "2025-01-10", "2027-01-10", "24.0", "Aggressive", "false", "Aggressive", "true"],
])

# 4.10 Customer Risk Response
h2(doc, "4.10 Entity: CustomerRiskResponse")
para(doc, "Records each answer selected by a customer during risk profiling.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["risk_profile_id", "UUID (FK→CustomerRiskProfile)", "Yes", "", "—"],
    ["question_id", "UUID (FK→Question)", "Yes", "", "—"],
    ["answer_option_id", "UUID (FK→AnswerOption)", "Yes", "For multi-select, one row per selected option", "—"],
    ["raw_score", "DECIMAL(5,2)", "No", "= answer_option.weightage", "—"],
    ["normalized_score", "DECIMAL(5,2)", "No", "If question uses RANGE normalization", "NULL"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["risk_profile_id", "question_id", "answer_option_id", "raw_score", "normalized_score"], [
    ["crp001", "q1", "ao-income-high", "5", "NULL"],
    ["crp001", "q2", "ao-fd", "1", "NULL"],
    ["crp001", "q2", "ao-mf", "2", "3 (normalized from raw 3, range 0-5→1, but total was 3→range lookup)"],
])

# 4.11 Investment Proposal
h2(doc, "4.11 Entity: InvestmentProposal")
para(doc, "Investment proposal created by an RM for a client.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["proposal_number", "VARCHAR(30)", "Yes", "Auto-generated: PROP-YYYYMMDD-XXXX", "Auto"],
    ["customer_id", "UUID (FK→Customer)", "Yes", "", "—"],
    ["risk_profile_id", "UUID (FK→CustomerRiskProfile)", "Yes", "Must be active, non-expired", "—"],
    ["title", "VARCHAR(300)", "Yes", "Proposal title", "—"],
    ["investment_objective", "ENUM", "Yes", "GROWTH | INCOME | BALANCED | CAPITAL_PRESERVATION | AGGRESSIVE_GROWTH", "—"],
    ["time_horizon_years", "INTEGER", "Yes", "1-30", "—"],
    ["proposed_amount", "DECIMAL(15,2)", "Yes", ">0", "—"],
    ["currency", "VARCHAR(3)", "Yes", "ISO 4217 code", "INR"],
    ["status", "ENUM", "Yes", "DRAFT | SUBMITTED | L1_APPROVED | L1_REJECTED | COMPLIANCE_APPROVED | COMPLIANCE_REJECTED | SENT_TO_CLIENT | CLIENT_ACCEPTED | CLIENT_REJECTED | EXPIRED", "DRAFT"],
    ["suitability_check_passed", "BOOLEAN", "No", "Result of automated suitability check", "NULL"],
    ["suitability_check_details", "JSONB", "No", "Detailed suitability check results", "NULL"],
    ["expected_return_pct", "DECIMAL(5,2)", "No", "Portfolio-weighted expected return", "NULL"],
    ["expected_std_dev_pct", "DECIMAL(5,2)", "No", "Portfolio-weighted std deviation", "NULL"],
    ["sharpe_ratio", "DECIMAL(5,2)", "No", "Computed Sharpe ratio", "NULL"],
    ["max_drawdown_pct", "DECIMAL(5,2)", "No", "Estimated max drawdown %", "NULL"],
    ["proposal_pdf_url", "VARCHAR(500)", "No", "URL to generated PDF", "NULL"],
    ["client_accepted_at", "TIMESTAMP", "No", "", "NULL"],
    ["client_rejected_at", "TIMESTAMP", "No", "", "NULL"],
    ["client_rejection_reason", "TEXT", "No", "", "NULL"],
    ["expires_at", "TIMESTAMP", "No", "Proposal validity end date", "NULL"],
    ["version", "INTEGER", "Yes", "", "1"],
    ["rm_id", "UUID (FK→User)", "Yes", "Creating RM", "Current user"],
    ["entity_id", "UUID (FK→LegalEntity)", "Yes", "", "—"],
    ["is_deleted", "BOOLEAN", "Yes", "", "false"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["created_by", "UUID (FK→User)", "Yes", "", "Current user"],
    ["updated_by", "UUID (FK→User)", "Yes", "", "Current user"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["proposal_number", "customer_id", "title", "investment_objective", "proposed_amount", "status", "suitability_check_passed"], [
    ["PROP-20260420-0001", "cust001", "Growth Portfolio for Mr. Sharma", "GROWTH", "5000000.00", "SENT_TO_CLIENT", "true"],
    ["PROP-20260418-0002", "cust002", "Conservative Income Plan", "INCOME", "2000000.00", "DRAFT", "NULL"],
    ["PROP-20260415-0003", "cust003", "Aggressive Growth Allocation", "AGGRESSIVE_GROWTH", "10000000.00", "CLIENT_ACCEPTED", "true"],
])

# 4.12 Proposal Line Item
h2(doc, "4.12 Entity: ProposalLineItem")
para(doc, "Individual product/asset allocation within an investment proposal.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "", "uuid_generate_v4()"],
    ["proposal_id", "UUID (FK→InvestmentProposal)", "Yes", "", "—"],
    ["asset_class", "VARCHAR(100)", "Yes", "Equity, Fixed Income, Mutual Funds, etc.", "—"],
    ["product_id", "UUID (FK→Product)", "No", "Specific product if selected", "NULL"],
    ["product_name", "VARCHAR(300)", "No", "Product display name", "NULL"],
    ["product_risk_code", "INTEGER", "No", "Product risk rating 1-6", "NULL"],
    ["allocation_percentage", "DECIMAL(5,2)", "Yes", "0-100; all lines must sum to 100", "—"],
    ["allocation_amount", "DECIMAL(15,2)", "Yes", "= proposal.proposed_amount * allocation_pct / 100", "Computed"],
    ["expected_return_pct", "DECIMAL(5,2)", "No", "", "NULL"],
    ["risk_deviation_flagged", "BOOLEAN", "No", "true if product_risk_code > client risk_code", "false"],
    ["deviation_acknowledged", "BOOLEAN", "No", "Client acknowledged the deviation", "false"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["proposal_id", "asset_class", "product_name", "allocation_percentage", "allocation_amount", "product_risk_code", "risk_deviation_flagged"], [
    ["PROP-001", "Equity", "HDFC Top 100 Fund", "40.00", "2000000.00", "4", "false"],
    ["PROP-001", "Fixed Income", "SBI Corporate Bond Fund", "35.00", "1750000.00", "2", "false"],
    ["PROP-001", "Alternatives", "Kotak Real Estate Fund", "25.00", "1250000.00", "5", "true"],
])

# 4.13 Proposal Approval
h2(doc, "4.13 Entity: ProposalApproval")
para(doc, "Audit log of all approval/rejection actions on proposals.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "", "uuid_generate_v4()"],
    ["proposal_id", "UUID (FK→InvestmentProposal)", "Yes", "", "—"],
    ["approval_level", "ENUM", "Yes", "L1_SUPERVISOR | COMPLIANCE | CLIENT", "—"],
    ["action", "ENUM", "Yes", "APPROVED | REJECTED | RETURNED_FOR_REVISION", "—"],
    ["acted_by", "UUID (FK→User)", "Yes", "", "—"],
    ["comments", "TEXT", "No", "Approval/rejection comments", "NULL"],
    ["acted_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["proposal_id", "approval_level", "action", "acted_by", "comments", "acted_at"], [
    ["PROP-001", "L1_SUPERVISOR", "APPROVED", "supervisor01", "Allocation looks appropriate", "2026-04-20 10:30:00"],
    ["PROP-001", "COMPLIANCE", "APPROVED", "compliance01", "Suitability verified", "2026-04-20 14:15:00"],
    ["PROP-002", "L1_SUPERVISOR", "RETURNED_FOR_REVISION", "supervisor01", "Reduce equity allocation for conservative client", "2026-04-19 11:00:00"],
])

# 4.14 Model Portfolio
h2(doc, "4.14 Entity: ModelPortfolio")
para(doc, "Predefined model portfolios linked to risk categories for bulk proposal generation.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "", "uuid_generate_v4()"],
    ["portfolio_name", "VARCHAR(200)", "Yes", "", "—"],
    ["risk_category", "VARCHAR(100)", "Yes", "Must match RiskAppetiteBand category", "—"],
    ["entity_id", "UUID (FK→LegalEntity)", "Yes", "", "—"],
    ["benchmark_index", "VARCHAR(200)", "No", "e.g., NIFTY 50, CRISIL Composite Bond", "NULL"],
    ["rebalance_frequency", "ENUM", "Yes", "MONTHLY | QUARTERLY | SEMI_ANNUAL | ANNUAL", "QUARTERLY"],
    ["drift_threshold_pct", "DECIMAL(5,2)", "Yes", "Trigger rebalance if any class drifts by this %", "5.00"],
    ["is_active", "BOOLEAN", "Yes", "", "true"],
    ["status", "ENUM", "Yes", "UNAUTHORIZED | AUTHORIZED", "UNAUTHORIZED"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["created_by", "UUID (FK→User)", "Yes", "", "Current user"],
    ["updated_by", "UUID (FK→User)", "Yes", "", "Current user"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["portfolio_name", "risk_category", "benchmark_index", "rebalance_frequency", "drift_threshold_pct", "is_active"], [
    ["Conservative Income Model", "Conservative", "CRISIL Composite Bond", "QUARTERLY", "3.00", "true"],
    ["Balanced Growth Model", "Moderate", "NIFTY 50 + CRISIL Short Term", "QUARTERLY", "5.00", "true"],
    ["Aggressive Equity Model", "Aggressive", "NIFTY 50", "SEMI_ANNUAL", "7.00", "true"],
])

# 4.15 ProductRiskDeviation
h2(doc, "4.15 Entity: ProductRiskDeviation")
para(doc, "Records instances where client transacted in a product with higher risk than their risk appetite.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "", "uuid_generate_v4()"],
    ["customer_id", "UUID (FK→Customer)", "Yes", "", "—"],
    ["risk_profile_id", "UUID (FK→CustomerRiskProfile)", "Yes", "", "—"],
    ["product_id", "UUID (FK→Product)", "Yes", "", "—"],
    ["customer_risk_code", "INTEGER", "Yes", "1-6", "—"],
    ["product_risk_code", "INTEGER", "Yes", "1-6", "—"],
    ["deviation_acknowledged", "BOOLEAN", "Yes", "Client clicked Confirm/Agree", "false"],
    ["acknowledged_at", "TIMESTAMP", "No", "", "NULL"],
    ["context", "ENUM", "Yes", "RM_OFFICE | CLIENT_PORTAL", "—"],
    ["order_id", "UUID", "No", "Linked order if applicable", "NULL"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["customer_id", "customer_risk_code", "product_risk_code", "deviation_acknowledged", "context"], [
    ["cust001", "3", "5", "true", "RM_OFFICE"],
    ["cust002", "2", "4", "true", "CLIENT_PORTAL"],
    ["cust003", "1", "3", "false", "RM_OFFICE"],
])

# 4.16 Lead Status (for Supervisor Dashboard)
h2(doc, "4.16 Entity: LeadStatusSummary (View/Materialized View)")
para(doc, "Aggregated lead status data per RM for supervisor dashboard widget. Computed from campaign and lead tables.")

add_styled_table(doc, ["Field", "Type", "Required", "Description", "Default"], [
    ["rm_id", "UUID (FK→User)", "Yes", "Relationship Manager", "—"],
    ["rm_name", "VARCHAR(200)", "Yes", "RM display name", "—"],
    ["total_leads", "INTEGER", "Yes", "Total active campaign leads", "0"],
    ["client_accepted", "INTEGER", "Yes", "Leads with status CLIENT_ACCEPTED", "0"],
    ["client_rejected", "INTEGER", "Yes", "Leads with status CLIENT_REJECTED", "0"],
    ["in_progress", "INTEGER", "Yes", "Leads with status IN_PROGRESS", "0"],
    ["ready_for_followup", "INTEGER", "Yes", "Leads with status READY_FOR_FOLLOWUP", "0"],
    ["new_leads", "INTEGER", "Yes", "Leads with status NEW", "0"],
    ["last_refreshed", "TIMESTAMP", "Yes", "Last computation time", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["rm_name", "total_leads", "client_accepted", "client_rejected", "in_progress", "ready_for_followup", "new_leads"], [
    ["Mariah", "60", "15", "7", "30", "5", "3"],
    ["Winfrid", "80", "10", "12", "35", "14", "9"],
])

# 4.17 Compliance Escalation (Post-Review Addition)
h2(doc, "4.17 Entity: ComplianceEscalation (Post-Review Addition)")
para(doc, "Tracks escalation cases when a customer exceeds the repeat product risk deviation threshold. Added based on adversarial evaluation finding.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["customer_id", "UUID (FK->Customer)", "Yes", "", "—"],
    ["escalation_type", "ENUM", "Yes", "REPEAT_DEVIATION | MANUAL", "REPEAT_DEVIATION"],
    ["deviation_count", "INTEGER", "Yes", "Number of deviations that triggered escalation", "—"],
    ["window_start_date", "DATE", "Yes", "Start of the rolling window", "—"],
    ["window_end_date", "DATE", "Yes", "End of the rolling window (= escalation date)", "—"],
    ["deviation_ids", "UUID[]", "Yes", "Array of ProductRiskDeviation IDs", "—"],
    ["status", "ENUM", "Yes", "OPEN | ACKNOWLEDGED | RESOLVED", "OPEN"],
    ["assigned_to", "UUID (FK->User)", "No", "Compliance Officer assigned", "NULL"],
    ["resolution_action", "ENUM", "No", "NOTED | FLAGGED_FOR_REVIEW | CLIENT_RESTRICTED", "NULL"],
    ["resolution_notes", "TEXT", "No", "Free-text notes from Compliance", "NULL"],
    ["resolved_at", "TIMESTAMP", "No", "", "NULL"],
    ["created_at", "TIMESTAMP", "Yes", "", "NOW()"],
    ["updated_at", "TIMESTAMP", "Yes", "", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["customer_id", "deviation_count", "status", "resolution_action"], [
    ["cust001", "6", "RESOLVED", "NOTED"],
    ["cust004", "8", "OPEN", "NULL"],
])

# 4.18 Risk Profiling Audit Log (Post-Review Addition)
h2(doc, "4.18 Entity: RiskProfilingAuditLog (Post-Review Addition)")
para(doc, "Immutable session-level audit log for every risk profiling assessment. Added based on adversarial evaluation finding for compliance evidence strengthening.")

add_styled_table(doc, ["Field", "Type", "Required", "Validation", "Default"], [
    ["id", "UUID", "Yes (auto)", "System-generated", "uuid_generate_v4()"],
    ["session_id", "UUID", "Yes", "Unique per assessment session", "uuid_generate_v4()"],
    ["customer_id", "UUID (FK->Customer)", "Yes", "", "—"],
    ["initiated_by", "UUID (FK->User)", "Yes", "RM or Client who started assessment", "—"],
    ["initiated_at", "TIMESTAMP", "Yes", "Session start time", "NOW()"],
    ["completed_at", "TIMESTAMP", "No", "Session end time", "NULL"],
    ["duration_seconds", "INTEGER", "No", "completed_at - initiated_at in seconds", "NULL"],
    ["outcome", "ENUM", "No", "COMPLETED | ABANDONED | ERROR", "NULL"],
    ["risk_profile_id", "UUID (FK->CustomerRiskProfile)", "No", "Linked profile if completed", "NULL"],
    ["device_type", "VARCHAR(50)", "No", "desktop | tablet | mobile", "NULL"],
    ["user_agent", "VARCHAR(500)", "No", "Browser user agent string", "NULL"],
    ["ip_address", "VARCHAR(45)", "No", "IPv4 or IPv6", "NULL"],
    ["entity_id", "UUID (FK->LegalEntity)", "Yes", "", "—"],
    ["created_at", "TIMESTAMP", "Yes", "Immutable", "NOW()"],
])

para(doc, "Sample Data:", bold=True)
add_styled_table(doc, ["session_id", "customer_id", "initiated_by", "duration_seconds", "outcome", "device_type"], [
    ["sess001", "cust001", "rm001", "420", "COMPLETED", "desktop"],
    ["sess002", "cust002", "rm002", "NULL", "ABANDONED", "tablet"],
    ["sess003", "cust003", "cust003", "180", "COMPLETED", "mobile"],
])

para(doc, "This entity is write-only (no UPDATE/DELETE operations). Only Compliance and Audit roles have read access.", bold=True)

doc.add_page_break()

# Save progress - we'll continue with remaining sections
output_path = os.path.join(os.path.dirname(__file__), "BRD-RiskProfiling-ProposalGeneration-v2.docx")
doc.save(output_path)
print(f"Phase 1 saved: {output_path}")
print("Continuing with Sections 5-14...")

# ══════════════════════════════════════════════════════════════════════════
# SECTION 5: FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "5. Functional Requirements")

# --- Module A: Risk Profiling Questionnaire Maintenance ---
h2(doc, "5.1 Module A: Risk Profiling Questionnaire Maintenance")

fr_rows_a = [
    ["FR-001", "List Questionnaires", "Display paginated grid of all questionnaire records with columns: Questionnaire Name, Customer Category, Questionnaire Type, Status, Maker ID/Name, Actions (View/Edit/Delete icons). Support column-based search/filter.",
     "As an Operations Administrator, I want to view all configured questionnaires in a searchable grid so that I can quickly find and manage questionnaire records.",
     "1. Grid displays all non-deleted questionnaires for the user's entity.\n2. Columns are sortable and filterable.\n3. Pagination shows 10/25/50 records per page.\n4. Status badges display color-coded: Unauthorized=grey, Modified=amber, Authorized=green, Rejected=red.\n5. Actions column shows View (eye icon), Edit (pencil icon), Delete (trash icon) based on status.",
     "Only UNAUTHORIZED or MODIFIED records can be edited or deleted. AUTHORIZED records can only be viewed."],

    ["FR-002", "Add Questionnaire", "Create a new risk profile questionnaire with header details and dynamic question builder. Questionnaire header includes: name, customer category, questionnaire type, effective start/end dates, valid period, is_score flag.",
     "As an Operations Administrator, I want to create a new risk profile questionnaire with configurable questions and answer options so that the system can assess client risk appetite.",
     "1. Form loads with all mandatory fields marked with *.\n2. Questionnaire Type dropdown shows values based on entity configuration.\n3. Effective Start Date defaults to today; End Date must be after Start.\n4. 'Add Question' button appends a new collapsible question section.\n5. Each question has: Description, Mandatory checkbox, Multi Select dropdown, Scoring Type dropdown, Computation Type dropdown.\n6. 'Add Options' button adds answer description + weightage fields.\n7. 'Save' validates all mandatory fields and creates record with status UNAUTHORIZED.\n8. System shows success toast: 'Record Modified Successfully'.\n9. Grid refreshes showing new record.",
     "Uniqueness: Only 1 questionnaire per (customer_category, questionnaire_type, overlapping date range). If category='Both', blocks Individual/Non-Individual for same period. Validation error shown inline."],

    ["FR-003", "View Questionnaire", "Display questionnaire in read-only mode with all questions and answer options expanded.",
     "As an Operations Administrator, I want to view a questionnaire's complete configuration in read-only mode so that I can review its contents without risk of accidental changes.",
     "1. All fields are displayed but non-editable.\n2. Questions displayed as collapsible sections.\n3. 'Back' button returns to grid.\n4. No Save/Reset buttons shown.",
     "Available for all status values."],

    ["FR-004", "Modify Questionnaire", "Edit an existing questionnaire. All fields editable except Questionnaire Name for AUTHORIZED records. Can add/remove questions and options.",
     "As an Operations Administrator, I want to modify an existing questionnaire's questions, options, and configuration so that I can update the risk assessment criteria.",
     "1. Form loads pre-populated with existing data.\n2. Add/Remove Question and Add/Remove Option buttons functional.\n3. On Save, status changes to MODIFIED (if was AUTHORIZED) or stays UNAUTHORIZED.\n4. Version incremented.\n5. Validation re-applied on save.",
     "Cannot modify if status is REJECTED — must create new. Modified records require re-authorization."],

    ["FR-005", "Delete Questionnaire", "Soft-delete a questionnaire record (set is_deleted=true).",
     "As an Operations Administrator, I want to delete an obsolete questionnaire so that it no longer appears in the active list.",
     "1. Confirmation dialog: 'Are you sure you want to delete this questionnaire?'\n2. Only available for UNAUTHORIZED or MODIFIED records.\n3. Record removed from grid after deletion.\n4. Success toast displayed.",
     "AUTHORIZED questionnaires cannot be deleted. They can only be superseded by creating a new one with overlapping dates."],

    ["FR-006", "Authorize/Reject Questionnaire", "Supervisor reviews UNAUTHORIZED/MODIFIED questionnaires and approves or rejects them.",
     "As an Operations Supervisor, I want to authorize or reject questionnaire changes so that only validated configurations are used for client risk profiling.",
     "1. Authorization queue shows all pending (UNAUTHORIZED/MODIFIED) records.\n2. Supervisor can View full details before deciding.\n3. 'Authorize' sets status=AUTHORIZED, records checker_id and authorized_at.\n4. 'Reject' sets status=REJECTED with mandatory rejection reason.\n5. Maker cannot authorize own records (four-eyes principle).\n6. Email/in-app notification sent to maker on authorize/reject.",
     "Checker must be different from Maker. Rejected records can be modified and resubmitted."],

    ["FR-007", "Score Normalization for Multi-Select", "For multi-select questions with scoring_type=RANGE, define score normalization ranges that map cumulative raw scores to a single normalized score.",
     "As an Operations Administrator, I want to define score normalization ranges for multi-select questions so that the system correctly computes comparable scores regardless of the number of options selected.",
     "1. 'Add Range' button appears when scoring_type=RANGE is selected.\n2. Each range has From, To, and Score fields.\n3. Ranges must not overlap and must cover the full possible score range.\n4. Validation: From < To; Score >=0; no gaps between ranges.",
     "scoring_type is auto-set to NONE and non-editable when is_multi_select=No."],

    ["FR-008", "Warning/Acknowledgement/Disclaimer Config", "Configure rich-text content for Warning, Acknowledgement, and Disclaimer sections displayed during customer risk profiling.",
     "As an Operations Administrator, I want to configure Warning, Acknowledgement, and Disclaimer text per questionnaire so that appropriate legal content is shown to clients during risk profiling.",
     "1. Three collapsible sections at bottom of questionnaire form: Warning, Acknowledgement, Disclaimer.\n2. Each supports rich text input.\n3. Content is optional.\n4. Content is displayed to client during risk profiling in corresponding sections.",
     "Content can differ per questionnaire (and thus per entity/type)."],
]

for row in fr_rows_a:
    h3(doc, f"{row[0]}: {row[1]}")
    para(doc, row[2])
    para(doc, f"User Story: {row[3]}", bold=True)
    para(doc, "Acceptance Criteria:", bold=True)
    for line in row[4].split("\n"):
        bullet(doc, line.strip())
    para(doc, f"Business Rules: {row[5]}", bold=True)

# --- Module B: Risk Appetite & Asset Allocation ---
h2(doc, "5.2 Module B: Risk Appetite Mapping & Asset Allocation Maintenance")

fr_rows_b = [
    ["FR-009", "List Risk Appetite Mappings", "Display grid of all risk appetite mapping configurations with status, effective dates, and actions.",
     "As an Operations Administrator, I want to view all risk appetite mappings so that I can manage score-to-category configurations.",
     "1. Grid displays: Mapping Name, Effective Start/End Date, Status, Maker, Actions.\n2. Sortable, filterable, paginated.\n3. View/Edit/Delete actions based on status.",
     "Same maker-checker rules as questionnaire maintenance."],

    ["FR-010", "Add/Modify Risk Appetite Mapping", "Create or edit a risk appetite mapping with score bands that map total risk scores to risk categories with numeric codes (1-6).",
     "As an Operations Administrator, I want to define score-to-risk-category bands so that the system correctly classifies clients based on their risk profiling scores.",
     "1. Header: Mapping Name, Effective Start/End Date.\n2. Bands table: Score From, Score To, Risk Category (dropdown: Conservative, Low to Moderate, Moderate, Moderately High, Aggressive, Very Aggressive), Risk Code (1-6 auto-assigned).\n3. Bands must be contiguous (no gaps) and non-overlapping.\n4. Save → status UNAUTHORIZED.\n5. Requires authorization.",
     "Risk code 1=lowest risk (Conservative), 6=highest risk (Very Aggressive). Bands must cover entire score range from 0 to max possible score. Band boundary semantics: inclusive-lower, exclusive-upper (score_from <= total_score < score_to). The last band (highest risk) is inclusive on both ends (score_from <= total_score <= score_to)."],

    ["FR-011", "Authorize Risk Appetite Mapping", "Supervisor authorizes or rejects risk appetite mapping changes.",
     "As an Operations Supervisor, I want to authorize risk appetite mappings so that only validated score-category relationships are used.",
     "1. Same authorization workflow as FR-006.\n2. Four-eyes principle enforced.\n3. Notifications sent.",
     "Same as FR-006 business rules."],

    ["FR-012", "List Asset Allocation Configs", "Display grid of all asset allocation configurations.",
     "As an Operations Administrator, I want to view all model asset allocation configurations so that I can manage portfolio templates per risk category.",
     "1. Grid with Config Name, Effective Dates, Status, Maker, Actions.\n2. Same CRUD pattern as other maintenance entities.",
     "Same maker-checker rules."],

    ["FR-013", "Add/Modify Asset Allocation", "Create or edit model asset allocations per risk category with asset class percentages, expected returns, and standard deviation.",
     "As an Operations Administrator, I want to define model asset allocations per risk category so that the system can recommend appropriate portfolio mixes to clients.",
     "1. Per risk category section: add asset class rows with allocation %, expected return %, std deviation %.\n2. Allocation percentages per category must sum to exactly 100%.\n3. Donut chart preview renders real-time as percentages change.\n4. Save → UNAUTHORIZED.\n5. Requires authorization.",
     "Asset classes must match the product universe taxonomy. All six risk categories must have an allocation defined."],

    ["FR-014", "Authorize Asset Allocation", "Supervisor authorizes or rejects asset allocation configuration changes.",
     "As an Operations Supervisor, I want to authorize asset allocation configurations so that only validated portfolio templates are used for proposals.",
     "1. Same authorization workflow as FR-006 and FR-011.",
     "Same business rules."],
]

for row in fr_rows_b:
    h3(doc, f"{row[0]}: {row[1]}")
    para(doc, row[2])
    para(doc, f"User Story: {row[3]}", bold=True)
    para(doc, "Acceptance Criteria:", bold=True)
    for line in row[4].split("\n"):
        bullet(doc, line.strip())
    para(doc, f"Business Rules: {row[5]}", bold=True)

# --- Module C: Customer Risk Assessment ---
h2(doc, "5.3 Module C: Customer Risk Assessment Journey")

fr_rows_c = [
    ["FR-015", "Initiate Risk Profiling", "RM searches for a customer, clicks risk assessment icon, system loads the 3-step wizard (Edit Profile → Assess Risk → Transact/Plan Goals).",
     "As a Relationship Manager, I want to initiate risk profiling for a customer through a guided wizard so that I can efficiently assess and record their risk appetite.",
     "1. Customer search by name/ID/account number.\n2. If active risk profile exists and not expired, show current profile with option to re-assess.\n3. If no active profile or expired, navigate directly to Assess Risk step.\n4. Step indicator shows progress: Edit Profile ✓ → Assess Risk (active) → Transact.\n5. 'Back' button available at each step.",
     "Risk profiling requires customer to have completed at least Step 1 (Edit Profile) with mandatory fields populated."],

    ["FR-016", "Display Risk Questionnaire", "Load the appropriate questionnaire (based on customer category and current date) with 3 collapsible parts: Part A (Financial Profiling), Part B (Investment Knowledge), Part C (Risk Profiling/SAF).",
     "As a Relationship Manager, I want to present the configured risk questionnaire to the client with clear sections so that the client can provide accurate responses.",
     "1. System identifies applicable questionnaire by customer_category + questionnaire_type + current date within effective range.\n2. Three collapsible sections displayed: Part A, Part B, Part C.\n3. First section auto-expanded; others collapsed.\n4. Single-select questions render as radio buttons; multi-select as checkboxes.\n5. Mandatory questions marked with *.\n6. Warning/Acknowledgement/Disclaimer sections displayed at bottom.\n7. 'Confirm' button submits all responses.",
     "If no authorized questionnaire exists for the customer category and current date, display error: 'No active questionnaire configuration found. Please contact Operations.'"],

    ["FR-017", "Compute Risk Score", "Upon questionnaire submission, compute total risk score using configured weightages and normalization rules, then map to risk appetite category.",
     "As a Relationship Manager, I want the system to automatically compute the client's risk score and category so that I can proceed with appropriate investment recommendations.",
     "1. For single-select scored questions: score = selected_option.weightage.\n2. For multi-select with scoring_type=NONE: score = sum of selected weightages.\n3. For multi-select with scoring_type=RANGE: raw_sum = sum of weightages → look up ScoreNormalizationRange → normalized_score.\n4. Total score = sum of all question scores (using normalized where applicable).\n5. Look up RiskAppetiteBand where total_score falls within (score_from, score_to) → risk_category + risk_code.\n6. Non-scored questionnaire parts (Part A, Part B) are recorded but do not contribute to score.",
     "If total score falls outside all defined bands, flag error: 'Score out of range — contact Operations to update Risk Appetite Mapping.' Band matching uses inclusive-lower, exclusive-upper comparison: score_from <= total_score < score_to (last band inclusive on both ends). Non-mandatory questions that are skipped contribute 0 to the total score (they are included with score=0, NOT excluded from the denominator). Validation at questionnaire authorization time must verify all mandatory questions have at least 1 answer option defined."],

    ["FR-018", "Display Recommended Risk Profile", "Show the computed risk profile with score, category, model asset allocation donut chart, expected return, standard deviation, profile date, and expiry date.",
     "As a Relationship Manager, I want to view the system-recommended risk profile with visual portfolio allocation so that I can discuss the results with the client.",
     "1. Risk Score displayed prominently.\n2. Risk Category displayed with color coding.\n3. Interactive donut chart showing model asset allocation per risk category.\n4. Expected Return % and Standard Deviation % displayed.\n5. Risk Profile Date = today; Expiry Date = today + valid_period_years.\n6. 'Continue' button to proceed to deviation check or next step.",
     "Model asset allocation loaded from AssetAllocationConfig linked to the computed risk_category."],

    ["FR-019", "Risk Profile Deviation (Optional)", "Allow customer to optionally deviate from the system-computed risk profile by selecting a different risk category. Configurable per entity (can be disabled).",
     "As a Relationship Manager, I want to allow clients to choose a different risk profile than the system-computed one (when permitted) so that client preferences are accommodated with proper documentation.",
     "1. Dialog: 'Do you agree on your Risk Profile and Risk Class above?' Yes/No.\n2. If No: dropdown to select alternative risk category with mandatory reason text field.\n3. Deviation recorded in CustomerRiskProfile (is_deviated=true, deviated_risk_category, deviation_reason).\n4. Deviation requires supervisor approval.\n5. If deviation feature is disabled for entity, this step is skipped entirely.",
     "For SBI: deviation is disabled (config flag: risk_deviation_enabled=false). For 360One/JBR: deviation is enabled."],

    ["FR-020", "Supervisor Risk Profile Approval", "Supervisor reviews and approves/rejects customer risk profiles, especially deviated ones, via the Customer Onboarding authorization queue.",
     "As an RM Supervisor, I want to review and approve customer risk profiles submitted by my team so that I can ensure risk assessments are accurate and compliant.",
     "1. Supervisor Dashboard → Customer Onboarding dropdown → pending risk profiles.\n2. View full questionnaire responses, computed score, and recommended category.\n3. If deviated: highlight deviation with reason.\n4. Approve or Reject with comments.\n5. On approval: CustomerRiskProfile.supervisor_approved = true.\n6. Notification sent to RM.",
     "Deviated profiles must be approved by supervisor before customer can transact."],
]

for row in fr_rows_c:
    h3(doc, f"{row[0]}: {row[1]}")
    para(doc, row[2])
    para(doc, f"User Story: {row[3]}", bold=True)
    para(doc, "Acceptance Criteria:", bold=True)
    for line in row[4].split("\n"):
        bullet(doc, line.strip())
    para(doc, f"Business Rules: {row[5]}", bold=True)

# --- Module D: Product Risk Deviation ---
h2(doc, "5.4 Module D: Product Risk Deviation & Compliance")

fr_rows_d = [
    ["FR-021", "Product Rating Alert", "When RM or client selects a product with product_risk_code > client effective_risk_code, display a Product Rating Alert popup.",
     "As a Relationship Manager, I want the system to alert me when a selected product has higher risk than the client's risk appetite so that I can ensure suitability compliance.",
     "1. Alert popup shows: 'ALERT! Selected product(s) has (have) Higher Risk than the Customer's Risk Appetite as an Investor.'\n2. Displays: Customer Risk Profile category, table with Product Name and Product Risk Class (numeric code 1-6).\n3. Disclaimer: 'Product rating ranges from 1 to 6, 1 being the lowest risk and 6 being highest risk.'\n4. Buttons: Cancel (remove product) and 'Confirm Notified to Customer' (acknowledge and proceed).\n5. Deviation recorded in ProductRiskDeviation table.",
     "Alert triggers on all product selection screens: Product Selection, Product Cart, Order placement — in both RM Office and Client Portal."],

    ["FR-022", "Risk Rating Filter", "Provide a Risk Rating filter dropdown in product selection screens allowing filtering by risk level.",
     "As a Relationship Manager, I want to filter available products by risk rating so that I can quickly find suitable products for my client's risk profile.",
     "1. Risk Rating dropdown with options: High, Low, Low to Moderate, Moderate, Moderately High, Very High (mapped to codes 1-6).\n2. Multi-select checkboxes.\n3. Products filtered in real-time.\n4. Clear filter option available.",
     "Risk rating codes displayed as numeric 1-6 on product cards, with tooltip showing 'Risk Rating: X' on hover."],

    ["FR-023", "Risk Profile Display Across Screens", "Display product risk profile consistently as numeric codes (1-6) across all screens: Product Selection, Product Selection Summary, Product Cart, Order Book, Scheme Info, Systematic Dashboard.",
     "As a Relationship Manager, I want product risk ratings displayed consistently as numeric codes across all screens so that risk information is clear and comparable.",
     "1. Product risk displayed as badge with numeric code (1-6).\n2. Color coding: 1-2=green, 3-4=amber, 5-6=red.\n3. Tooltip: 'Risk Rating: {code}' on hover.\n4. Consistent across RM Office and Client Portal.",
     "Replaces legacy text-based descriptions (Moderate, Aggressive) with numeric codes per SBICUST-1482/1627 requirements."],
]

for row in fr_rows_d:
    h3(doc, f"{row[0]}: {row[1]}")
    para(doc, row[2])
    para(doc, f"User Story: {row[3]}", bold=True)
    para(doc, "Acceptance Criteria:", bold=True)
    for line in row[4].split("\n"):
        bullet(doc, line.strip())
    para(doc, f"Business Rules: {row[5]}", bold=True)

# --- Module E: Proposal Generation ---
h2(doc, "5.5 Module E: Investment Proposal Generation")

fr_rows_e = [
    ["FR-024", "Create Investment Proposal", "RM creates a new investment proposal for a client by specifying investment objective, time horizon, amount, and selecting products/asset allocations.",
     "As a Relationship Manager, I want to create a personalized investment proposal for my client so that I can recommend a suitable portfolio aligned with their risk profile.",
     "1. Select client from customer search.\n2. System loads client's active risk profile (block creation if no active/non-expired profile).\n3. Auto-suggest model portfolio based on client's risk category from AssetAllocationConfig.\n4. RM can customize: adjust asset class percentages, add/remove specific products.\n5. Real-time validation: allocation percentages must sum to 100%.\n6. Concentration limit checks: max 40% in single asset class, max 10% in single issuer.\n7. Suitability check runs automatically on save.\n8. Save as DRAFT or Submit for approval.",
     "Client must have active, non-expired, supervisor-approved risk profile. Proposal amount > 0. At least one asset class must be allocated."],

    ["FR-025", "What-If Analysis", "RM can adjust allocation percentages and see real-time impact on projected returns, risk metrics (standard deviation, Sharpe ratio, max drawdown).",
     "As a Relationship Manager, I want to perform what-if analysis on portfolio allocations so that I can optimize the proposal for my client's risk-return preferences.",
     "1. Slider or input field per asset class to adjust allocation %.\n2. Real-time recalculation of: Expected Return %, Standard Deviation %, Sharpe Ratio, Max Drawdown %.\n3. Visual bar chart comparing current allocation vs model allocation.\n4. Warning if allocation deviates significantly from model portfolio (>15% drift in any class).\n5. 'Reset to Model' button restores original model allocation.",
     "Sharpe Ratio = (Expected Return - Risk-Free Rate) / Std Deviation. Risk-free rate configured per entity (default: 6.5% for INR)."],

    ["FR-026", "Automated Suitability Check", "System validates proposed products against client risk profile, investment experience, and regulatory requirements.",
     "As a Compliance Officer, I want the system to automatically check proposal suitability so that non-compliant proposals are flagged before reaching the client.",
     "1. Check 1: Each product's risk_code <= client's effective_risk_code (flag but allow with deviation acknowledgement).\n2. Check 2: Product type matches client's investment experience from Part B questionnaire.\n3. Check 3: Concentration limits respected.\n4. Check 4: Investment mandate alignment (advisory vs discretionary constraints).\n5. Result stored in suitability_check_passed and suitability_check_details (JSONB).\n6. Failed checks displayed as itemized warnings with severity (BLOCKER / WARNING).",
     "BLOCKER-level failures prevent submission. WARNING-level allow submission with acknowledgement."],

    ["FR-027", "Proposal Approval Workflow", "Multi-level approval: RM submits → L1 Supervisor reviews → Compliance reviews → Sent to Client.",
     "As a Compliance Officer, I want multi-level proposal approval so that investment recommendations are validated before reaching clients.",
     "1. RM submits proposal → status = SUBMITTED.\n2. L1 Supervisor receives notification → can APPROVE, REJECT, or RETURN_FOR_REVISION.\n3. If L1 approved → Compliance Officer receives notification.\n4. Compliance can APPROVE, REJECT, or RETURN_FOR_REVISION.\n5. If Compliance approved → status = SENT_TO_CLIENT, client notified.\n6. Each action recorded in ProposalApproval table with comments and timestamp.\n7. SLA: 24h per approval level; escalation notification if overdue.",
     "Returned proposals go back to RM as DRAFT for revision. Rejected proposals are terminal. Approval history is immutable."],

    ["FR-028", "Generate Proposal PDF", "Generate a professional PDF document with bank branding, client details, risk profile summary, proposed portfolio, projected performance charts, and disclaimers.",
     "As a Relationship Manager, I want to generate a branded PDF proposal document so that I can present a professional investment recommendation to my client.",
     "1. PDF includes: bank logo, RM contact details, client name/ID, risk profile summary, proposed asset allocation (pie chart), product details table, projected performance (bar chart), disclaimers.\n2. Configurable template per entity/jurisdiction.\n3. Generated server-side (Puppeteer/equivalent).\n4. Stored and linked to proposal record.\n5. Downloadable from both RM Office and Client Portal.",
     "PDF regenerated on any proposal modification. Previous versions archived."],

    ["FR-029", "Client Proposal View & Accept/Reject", "Client views proposals in Client Portal and can accept or reject with digital acknowledgement.",
     "As a Client, I want to view investment proposals in my portal and accept or reject them so that I can make informed investment decisions digitally.",
     "1. Client Portal → Proposals section shows all SENT_TO_CLIENT proposals.\n2. Proposal detail page shows full portfolio with interactive charts.\n3. 'Accept' button with digital acknowledgement checkbox and e-signature.\n4. 'Reject' button with mandatory rejection reason.\n5. Accepted proposals trigger downstream order generation.\n6. PDF downloadable.",
     "Proposals expire after configurable period (default: 30 days). Expired proposals auto-transition to EXPIRED status."],

    ["FR-030", "Model Portfolio Management", "Define, maintain, and track model portfolios per risk category with rebalancing triggers.",
     "As an Operations Administrator, I want to manage model portfolios so that RMs have up-to-date portfolio templates for proposal generation.",
     "1. CRUD for model portfolios with: name, risk category, benchmark, rebalance frequency, drift threshold.\n2. Portfolio line items: asset class, target allocation %, product list.\n3. Performance tracking vs benchmark.\n4. Drift alert when actual allocation deviates from target by > threshold.\n5. Bulk proposal refresh: update all DRAFT proposals using this model when model changes.",
     "Model portfolio changes require authorization. Active proposals using old model are flagged for RM review."],

    ["FR-031", "Proposal Comparison", "Client can compare multiple proposals side-by-side in Client Portal.",
     "As a Client, I want to compare multiple proposals side-by-side so that I can evaluate different investment strategies.",
     "1. Select 2-3 proposals from list.\n2. Side-by-side comparison: asset allocation, expected return, risk metrics, product list.\n3. Visual charts overlay for allocation comparison.\n4. Highlight differences.",
     "Only SENT_TO_CLIENT proposals can be compared."],
]

for row in fr_rows_e:
    h3(doc, f"{row[0]}: {row[1]}")
    para(doc, row[2])
    para(doc, f"User Story: {row[3]}", bold=True)
    para(doc, "Acceptance Criteria:", bold=True)
    for line in row[4].split("\n"):
        bullet(doc, line.strip())
    para(doc, f"Business Rules: {row[5]}", bold=True)

# --- Module F: Supervisor Dashboard ---
h2(doc, "5.6 Module F: Supervisor Dashboard & Reporting")

fr_rows_f = [
    ["FR-032", "Supervisor Leads Widget", "Dashboard widget showing lead status breakdown per RM with Level 1 (bar chart) and Level 2 (detail table) views.",
     "As an RM Supervisor, I want to see a leads status widget in my dashboard so that I can monitor my team's campaign performance at a glance.",
     "1. Level 1: Bar chart with 5 bars (In Progress, New, Client Rejected, Ready for Follow-up, Client Accepted).\n2. Click drill-down icon → Level 2 table.\n3. Level 2 columns: RM Name, Total Leads, Client Accepted, Client Rejected, In Progress, Ready for Follow-up, New.\n4. Only active campaign leads included.\n5. Sorted by Client Accepted ascending; tie-breakers cascade through In Progress, Ready for Follow-up, New, Client Rejected.\n6. Search by RM name.\n7. Pagination.",
     "Only leads from campaigns with status=Active. Supervisor sees only RMs in their reporting hierarchy."],

    ["FR-033", "Risk Profiling Completion Report", "Report showing risk profiling completion rates by RM, branch, and time period.",
     "As a Compliance Officer, I want to view risk profiling completion reports so that I can monitor compliance with profiling requirements.",
     "1. Filters: Date range, RM, Branch, Entity.\n2. Columns: RM Name, Total Clients, Profiled, Pending, Expired, Completion %.\n3. Export to CSV/Excel.\n4. Drill-down to individual client list.",
     "Profiled = active non-expired profile. Expired = profile past expiry date."],

    ["FR-034", "Transaction by Product Rating Report", "Sigma-style report showing all transactions with product rating vs client risk profile.",
     "As a Compliance Officer, I want to view transactions grouped by product rating so that I can identify potential suitability violations.",
     "1. Columns: Primary RM Name, RM Employee ID, Account Name, Client Risk Profile, Product Rating (numeric 1-6), Product Name.\n2. Filters: Date range, RM, Client Risk Profile, Product Rating.\n3. Highlight mismatches (product rating > client risk).\n4. Export to CSV/Excel.",
     "Product Rating column displays numeric codes 1-6."],

    ["FR-035", "Product Risk Mismatch Report", "Report listing all transactions where product risk exceeded client risk appetite.",
     "As a Compliance Officer, I want a dedicated risk mismatch report so that I can audit all deviation instances.",
     "1. Same columns as FR-034 but pre-filtered to only show mismatches.\n2. Additional column: Deviation Acknowledged (Yes/No), Acknowledged Date.\n3. Summary statistics at top: total mismatches, acknowledged %, unacknowledged count.",
     "Report sources from ProductRiskDeviation table."],

    ["FR-036", "Proposal Pipeline Dashboard", "Dashboard showing proposal statistics by status, RM, and time period.",
     "As an RM Supervisor, I want to see a proposal pipeline dashboard so that I can track proposal throughput and bottlenecks.",
     "1. Funnel chart: DRAFT → SUBMITTED → L1_APPROVED → COMPLIANCE_APPROVED → SENT_TO_CLIENT → ACCEPTED.\n2. Cards: Total proposals, Avg time-to-accept, Conversion rate, Proposals pending approval.\n3. Filters: Date range, RM, Entity.\n4. Table view with all proposals and their current status.",
     "Conversion rate = CLIENT_ACCEPTED / SENT_TO_CLIENT * 100."],

    ["FR-037", "Risk Distribution Analytics", "Dashboard showing risk profile distribution across the client base.",
     "As a Compliance Officer, I want to view risk profile distribution analytics so that I can understand the overall risk composition of our client base.",
     "1. Pie chart: % of clients per risk category.\n2. Bar chart: risk category distribution by branch/entity.\n3. Trend line: risk category changes over time.\n4. Filters: Entity, Branch, Date range.",
     "Based on effective_risk_category from active CustomerRiskProfile records."],
]

for row in fr_rows_f:
    h3(doc, f"{row[0]}: {row[1]}")
    para(doc, row[2])
    para(doc, f"User Story: {row[3]}", bold=True)
    para(doc, "Acceptance Criteria:", bold=True)
    for line in row[4].split("\n"):
        bullet(doc, line.strip())
    para(doc, f"Business Rules: {row[5]}", bold=True)

# --- Module G: Post-Review Additions (Adversarial Evaluation Findings) ---
h2(doc, "5.7 Module G: Compliance Hardening & Operational Resilience (Post-Review)")

para(doc, "The following requirements were added based on the adversarial evaluation of the v1.0 BRD to address identified gaps in compliance, concurrency, and operational completeness.", bold=True)

fr_rows_g = [
    ["FR-038", "Repeat Deviation Escalation", "When a customer accumulates more than a configurable threshold of product risk deviations within a configurable time window, automatically escalate to Compliance Officer.",
     "As a Compliance Officer, I want the system to automatically escalate repeat product risk deviations so that I can intervene before systematic suitability violations occur.",
     "1. Configurable threshold per entity: max_deviations (default: 5) within window_days (default: 30).\n2. When threshold exceeded, create compliance escalation notification.\n3. Escalation notification includes: customer name/ID, total deviation count, list of products, dates.\n4. Compliance Officer can: Acknowledge (add notes), Flag for Review (triggers RM meeting), or Restrict (temporarily block further high-risk product transactions for this client).\n5. Escalation status tracked: OPEN → ACKNOWLEDGED → RESOLVED.",
     "Threshold check runs on every new ProductRiskDeviation record insert. Time window is rolling (not calendar-based). Escalation creates entry in a new ComplianceEscalation entity."],

    ["FR-039", "Risk Profiling Audit Log", "Capture detailed session-level audit data for every risk profiling assessment including initiator, duration, device metadata, and completion status.",
     "As a Compliance Officer, I want detailed audit logs for every risk profiling session so that I can demonstrate regulatory compliance during audits.",
     "1. Log entry created when RM initiates risk profiling (FR-015).\n2. Captures: session_id, customer_id, initiated_by (RM user_id), initiated_at (timestamp), device_type (desktop/tablet/mobile), user_agent, IP address.\n3. Updated on completion: completed_at, duration_seconds, outcome (COMPLETED/ABANDONED/ERROR), risk_profile_id (if completed).\n4. If client self-service: captures client_id as initiator.\n5. Logs are immutable (no update/delete). Retained for 7 years per data retention policy.",
     "Audit log is write-only. Only Compliance and Audit roles can read. No UI for editing audit records."],

    ["FR-040", "Cascading Config Validation", "When authorizing a Risk Appetite Mapping, validate that all referenced risk categories have corresponding Asset Allocation Configs. When authorizing an Asset Allocation Config, validate categories match active Risk Appetite Mapping.",
     "As an Operations Supervisor, I want the system to validate cross-entity configuration consistency so that runtime errors from mismatched configs are prevented.",
     "1. On RiskAppetiteMapping authorization: check that for each risk_category in the mapping's bands, an authorized AssetAllocationConfig exists with lines for that category. If missing, show warning (non-blocking): 'Asset Allocation not yet configured for category: {name}. Risk profiling will work but Recommended Risk Profile screen will show incomplete data.'\n2. On AssetAllocationConfig authorization: check that all risk_categories in the config lines match categories in the active RiskAppetiteMapping. Flag any categories not found in mapping.\n3. On Questionnaire authorization: verify all mandatory questions have at least 1 answer option defined.\n4. Validation results shown as warnings (allow authorization) not blockers (except question-option check which is a blocker).",
     "Cross-entity validation is advisory (warning) for mapping/allocation consistency and mandatory (blocker) for question-option completeness."],

    ["FR-041", "Questionnaire Content i18n (Phase 2)", "Support multi-language content for questionnaire questions, answer options, warnings, acknowledgements, and disclaimers via a content translation table.",
     "As an Operations Administrator, I want to configure questionnaire content in multiple languages so that international clients can complete risk profiling in their preferred language.",
     "1. New entity: QuestionnaireContentTranslation with fields: source_entity (QUESTION/ANSWER_OPTION/WARNING/ACKNOWLEDGEMENT/DISCLAIMER), source_id (FK), locale (e.g., en-US, hi-IN, zh-CN), translated_text.\n2. Default locale defined per entity (e.g., en-US for HSBC, en-IN for SBI).\n3. RM Office and Client Portal auto-select locale based on user language preference.\n4. Fallback: if translation not available for requested locale, use default locale.\n5. Admin UI: per-question translation editor with side-by-side original and translated text.",
     "Phase 2 feature. Data model defined now for planning; implementation deferred. Phase 1 operates in single language per entity."],

    ["FR-042", "Data Archival Strategy", "Archive CustomerRiskResponse and RiskProfilingAuditLog records older than the regulatory retention period to cold storage, maintaining query access.",
     "As a System Administrator, I want aged risk profiling data archived to cold storage so that database performance is maintained while meeting retention requirements.",
     "1. Archival job runs monthly.\n2. Records older than 7 years (configurable) are moved to archive schema/table.\n3. Archived records remain queryable via separate archive API endpoints (read-only, slower SLA acceptable).\n4. CustomerRiskProfile records are NEVER archived if is_active=true (regardless of age).\n5. Archival logged in system audit trail.",
     "Archive schema mirrors source schema. No data deletion — only movement. Archive API has separate rate limits."],
]

for row in fr_rows_g:
    h3(doc, f"{row[0]}: {row[1]}")
    para(doc, row[2])
    para(doc, f"User Story: {row[3]}", bold=True)
    para(doc, "Acceptance Criteria:", bold=True)
    for line in row[4].split("\n"):
        bullet(doc, line.strip())
    para(doc, f"Business Rules: {row[5]}", bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 6: UI REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "6. User Interface Requirements")

screens = [
    ["Risk Questionnaire Grid (Back Office)", "List all questionnaires in a data table with search/filter/sort per column. Top-right: 'Add Questionnaire' primary button. Each row: View (eye), Edit (pencil), Delete (trash) action icons. Status badges color-coded. Pagination at bottom-right. 'Back' button at bottom-left returns to Risk Assessor menu.", "Risk Assessor menu → click 'Risk Questionnaire' card", "Risk Assessor menu cards page"],
    ["Questionnaire Add/Edit Form", "Top section: Questionnaire header fields (Name, Category dropdown, Type dropdown, Dates, Valid Period, Is Score). Middle: Dynamic question sections — each question in a collapsible accordion with 'Add Question'/'Remove Question' buttons. Each question has: Description textarea, Mandatory checkbox, Multi Select dropdown, Scoring Type dropdown, Computation Type dropdown, Options sub-section with 'Add Options' button (each option: Description + Weightage fields). For scoring_type=RANGE: 'Add Range' button with From/To/Score fields. Bottom sections: Warning, Acknowledgement, Disclaimer collapsible rich-text areas. Action bar: Back, Reset, Save buttons.", "Grid → Add/Edit action", "Grid page on Back/Save"],
    ["Risk Assessor Menu (Back Office)", "Card-based menu with 3 cards: 'Risk Questionnaire' (settings icon), 'Risk Appetite' (warning icon), 'Asset Allocation' (chart icon). Search bar at top. Each card clickable to navigate to respective grid.", "Operations Office sidebar → Maintenance → Risk Assessor", "Operations Office main menu"],
    ["Customer Risk Assessment Wizard (RM Office)", "Top: Customer info bar (name, ID, current risk profile badge). Step indicator: 3 steps with checkmarks (Edit Profile ✓ → Assess Risk → Transact). Step 2 content: 3 collapsible accordion sections (Part A, Part B, Part C). Each section: questions with radio buttons or checkboxes. Bottom: Warning/Acknowledgement/Disclaimer sections. Action buttons: Back, Confirm.", "Customer search → click risk profiling icon", "Customer search or Recommended Risk Profile screen"],
    ["Recommended Risk Profile Screen", "Top: Risk Score badge (large numeric display) + Risk Category label (color-coded). Center: Interactive donut chart showing model asset allocation with percentages and asset class labels. Below chart: Expected Return % and Standard Deviation % values. Bottom: Risk Profile Date and Expiry Date. Action buttons: Back, Continue.", "After questionnaire submission (FR-017)", "Deviation dialog or Transact step"],
    ["Product Rating Alert Popup (Modal)", "Header: 'Product Rating Alert' with info icon. Body: Alert message text. Customer Risk Profile label. Table: Product Name | Product Risk Class (numeric badge). Footer disclaimer text. Buttons: Cancel (secondary) and 'Confirm Notified to Customer' (primary).", "Triggered when product risk > client risk", "Returns to product selection"],
    ["Investment Proposal Builder", "Left panel: Client info card (name, risk profile, effective category). Center: Asset allocation form with sliders/inputs per asset class, product search and selection per class. Right panel: Live metrics (Expected Return, Std Dev, Sharpe Ratio, Max Drawdown). Top: Proposal header (Title, Objective dropdown, Time Horizon, Amount). Bottom: Submit/Save Draft buttons.", "RM Office → select client → Create Proposal", "Proposal detail or proposal grid"],
    ["Proposal Detail / Client View", "Header: Proposal number, status badge, dates. Portfolio section: Pie chart of asset allocation + product table (Product Name, Asset Class, Allocation %, Amount, Risk Code). Metrics section: Expected Return, Risk, Sharpe, Drawdown. Disclaimers section. Actions: Accept/Reject (client), Approve/Reject (supervisor/compliance), Download PDF.", "Proposal grid → click proposal", "Proposal grid"],
    ["Supervisor Dashboard — Leads Widget", "Level 1: Bar chart with 5 color-coded bars (In Progress=yellow, New=blue, Client Rejected=red, Ready for Follow-up=teal, Client Accepted=pink). Drill-down icon (3 dots). Level 2: Data table with columns (RM Name, Total Leads, Client Accepted, Client Rejected, In Progress, Ready for Follow-up, New). Search bar. Pagination.", "Supervisor Dashboard main page", "Supervisor Dashboard"],
]

for s in screens:
    h2(doc, f"6.x {s[0]}")
    para(doc, f"Layout: {s[1]}")
    para(doc, f"Navigation From: {s[2]}")
    para(doc, f"Navigation To: {s[3]}")

para(doc, "Responsive Behavior: All screens must be responsive. Tables switch to card layout on mobile (<768px). Donut/pie charts resize proportionally. Sidebar collapses to hamburger menu on tablet/mobile.", bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 7: API & INTEGRATION
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "7. API & Integration Requirements")

h2(doc, "7.1 Authentication")
para(doc, "All API endpoints require authentication via httpOnly cookie-based JWT tokens issued by the Trust OMS central auth service. Role-based access control enforced via middleware. Bearer token alternative supported for API-to-API calls.")

h2(doc, "7.2 Standardized Error Response Format")
para(doc, 'All endpoints return errors in this format: { "error": { "code": "VALIDATION_ERROR", "message": "Questionnaire name is required", "field": "questionnaire_name", "details": [] } }. HTTP status codes: 400 (validation), 401 (unauthenticated), 403 (unauthorized), 404 (not found), 409 (conflict/duplicate or version mismatch), 500 (server error).')

h2(doc, "7.2.1 Optimistic Locking (Concurrency Control)")
para(doc, "All PUT/PATCH endpoints for versioned entities (Questionnaire, RiskAppetiteMapping, AssetAllocationConfig, InvestmentProposal, ModelPortfolio) must implement optimistic locking:")
bullet(doc, "Request body must include 'version' field with the expected current version number.")
bullet(doc, "Server validates: WHERE id = :id AND version = :expectedVersion. If no matching row, return HTTP 409 Conflict.")
bullet(doc, "409 Conflict response body: { \"error\": { \"code\": \"VERSION_CONFLICT\", \"message\": \"Record was modified by another user\", \"current_version\": <N>, \"current_record\": {...} } }")
bullet(doc, "UI behavior on 409: Display notification 'This record was modified by another user. Please review the changes.' Reload the record with current data. Highlight fields that differ from the user's in-progress edit. User must manually re-apply their changes and save again.")
bullet(doc, "Authorization queue: When a Checker opens a record for review, no lock is acquired (reads are non-blocking). However, the Authorize/Reject action validates the version to ensure the record hasn't been modified between opening and acting.")

h2(doc, "7.3 Core API Endpoints")

endpoints = [
    ["GET", "/api/back-office/risk-profiling/questionnaires", "List questionnaires", "Query: page, limit, status, customer_category, questionnaire_type, search", "{ items: [...], total, page, limit }"],
    ["POST", "/api/back-office/risk-profiling/questionnaires", "Create questionnaire", "{ questionnaire_name, customer_category, questionnaire_type, effective_start_date, effective_end_date, valid_period_years, is_score, questions: [{ question_description, is_mandatory, is_multi_select, scoring_type, computation_type, options: [{ answer_description, weightage }], ranges: [{ range_from, range_to, normalized_score }] }], warning_text, acknowledgement_text, disclaimer_text }", "{ id, status: 'UNAUTHORIZED', ...full record }"],
    ["PUT", "/api/back-office/risk-profiling/questionnaires/:id", "Update questionnaire", "Same as POST body", "{ id, status: 'MODIFIED', version: N+1 }"],
    ["POST", "/api/back-office/risk-profiling/questionnaires/:id/authorize", "Authorize/Reject", "{ action: 'AUTHORIZE'|'REJECT', reason?: string }", "{ id, status, checker_id, authorized_at }"],
    ["GET", "/api/back-office/risk-profiling/risk-appetite-mappings", "List mappings", "Query params similar", "{ items: [...], total }"],
    ["POST", "/api/back-office/risk-profiling/risk-appetite-mappings", "Create mapping", "{ mapping_name, effective_start_date, effective_end_date, bands: [{ score_from, score_to, risk_category, risk_code }] }", "{ id, status }"],
    ["POST", "/api/back-office/risk-profiling/risk-appetite-mappings/:id/authorize", "Authorize mapping", "{ action, reason }", "{ id, status }"],
    ["GET", "/api/back-office/risk-profiling/asset-allocations", "List configs", "Query params", "{ items: [...], total }"],
    ["POST", "/api/back-office/risk-profiling/asset-allocations", "Create config", "{ config_name, ..., lines: [{ risk_category, asset_class, allocation_percentage, expected_return_pct, standard_deviation_pct }] }", "{ id, status }"],
    ["POST", "/api/risk-profiling/assess/:customerId", "Submit assessment", "{ questionnaire_id, responses: [{ question_id, answer_option_ids: [...] }] }", "{ risk_profile_id, total_raw_score, computed_risk_category, computed_risk_code, model_asset_allocation, expected_return, std_deviation, expiry_date }"],
    ["POST", "/api/risk-profiling/deviate/:riskProfileId", "Submit deviation", "{ deviated_risk_category, deviated_risk_code, deviation_reason }", "{ risk_profile_id, effective_risk_category }"],
    ["POST", "/api/risk-profiling/approve/:riskProfileId", "Supervisor approve", "{ action: 'APPROVE'|'REJECT', comments }", "{ risk_profile_id, supervisor_approved }"],
    ["GET", "/api/risk-profiling/customer/:customerId/profile", "Get active profile", "—", "{ risk_profile, model_allocation, expiry_date }"],
    ["POST", "/api/proposals", "Create proposal", "{ customer_id, title, investment_objective, time_horizon_years, proposed_amount, currency, line_items: [{ asset_class, product_id, allocation_percentage }] }", "{ id, proposal_number, status: 'DRAFT', suitability_check }"],
    ["POST", "/api/proposals/:id/submit", "Submit for approval", "—", "{ id, status: 'SUBMITTED' }"],
    ["POST", "/api/proposals/:id/approve", "Approve proposal", "{ action: 'APPROVED'|'REJECTED'|'RETURNED', comments }", "{ id, status, approval_record }"],
    ["POST", "/api/proposals/:id/client-action", "Client accept/reject", "{ action: 'ACCEPTED'|'REJECTED', reason?, signature? }", "{ id, status }"],
    ["GET", "/api/proposals/:id/pdf", "Download PDF", "—", "PDF binary stream"],
    ["GET", "/api/reports/product-rating-transactions", "Product rating report", "Query: date_from, date_to, rm_id, risk_profile", "{ items: [...], total }"],
    ["GET", "/api/reports/risk-mismatch", "Risk mismatch report", "Query: date_from, date_to", "{ items: [...], summary_stats }"],
    ["GET", "/api/supervisor/leads-summary", "Lead status summary", "Query: supervisor_id", "{ level1_chart_data, level2_table: [...] }"],
]

add_styled_table(doc, ["Method", "Path", "Description", "Request", "Response"], endpoints)

h2(doc, "7.4 External Integrations")
integrations = [
    ["Product Data Service", "Internal", "Fetch product master, NAV, risk ratings, ISIN", "REST API"],
    ["Customer Data Service", "Internal", "Fetch customer demographics, accounts, holdings", "REST API"],
    ["PDF Generation Engine", "Internal", "Server-side PDF rendering (Puppeteer)", "Internal service call"],
    ["E-Signature Service", "External", "Digital signature for proposal acceptance", "REST API (pluggable adapter)"],
    ["Email/Notification Service", "Internal", "Send approval notifications, proposal alerts", "Event bus / REST"],
    ["CRISIL/Morningstar Feed", "External", "Product risk ratings, fund ratings", "File-based / API (interface only)"],
]
add_styled_table(doc, ["Service", "Type", "Purpose", "Protocol"], integrations)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 8: NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "8. Non-Functional Requirements")

nfrs = [
    ["Performance", "API response time <500ms for CRUD operations, <2s for risk score computation, <5s for PDF generation. Support 500 concurrent RM sessions. Dashboard widgets render within 3s."],
    ["Security", "httpOnly JWT cookies, OWASP Top 10 compliance, input sanitization, parameterized queries, RBAC on all endpoints, field-level encryption for PII, TLS 1.3 in transit, AES-256 at rest."],
    ["Scalability", "Horizontal scaling via stateless Node.js instances behind load balancer. Database read replicas for reporting queries. Materialized views for dashboard aggregations."],
    ["Availability", "99.9% uptime during business hours (Mon-Sat 8AM-8PM local). Scheduled maintenance window: Sunday 2AM-6AM."],
    ["Data Backup & Recovery", "Automated daily backups with 30-day retention. RPO: 1 hour. RTO: 4 hours. Point-in-time recovery capability."],
    ["Accessibility", "WCAG 2.1 Level AA compliance. Keyboard navigation for all forms. Screen reader compatible labels. Minimum 4.5:1 contrast ratio."],
    ["Browser Support", "Chrome 90+, Firefox 90+, Safari 15+, Edge 90+. Responsive down to 768px (tablet). Mobile 375px minimum."],
    ["Audit & Compliance", "Complete audit trail for all state changes. Immutable risk profiling records. Maker-checker for all configuration changes. 7-year data retention for regulatory compliance."],
]

for label, desc in nfrs:
    h2(doc, f"8.x {label}")
    para(doc, desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 9: WORKFLOW & STATE DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "9. Workflow & State Diagrams")

h2(doc, "9.1 Questionnaire Authorization Workflow")
add_styled_table(doc, ["Current State", "Action", "Next State", "Actor", "Side Effects"], [
    ["(New)", "Create", "UNAUTHORIZED", "Ops Admin (Maker)", "Record created with version 1"],
    ["UNAUTHORIZED", "Submit for auth", "UNAUTHORIZED (pending)", "Ops Admin", "Appears in Supervisor auth queue"],
    ["UNAUTHORIZED", "Authorize", "AUTHORIZED", "Ops Supervisor (Checker)", "checker_id set, authorized_at set, notification to Maker"],
    ["UNAUTHORIZED", "Reject", "REJECTED", "Ops Supervisor", "Rejection reason recorded, notification to Maker"],
    ["AUTHORIZED", "Modify", "MODIFIED", "Ops Admin", "Version incremented, status changes"],
    ["MODIFIED", "Authorize", "AUTHORIZED", "Ops Supervisor", "Re-authorized with new version"],
    ["MODIFIED", "Reject", "REJECTED", "Ops Supervisor", "Rejection reason recorded"],
    ["REJECTED", "Modify & Resubmit", "UNAUTHORIZED", "Ops Admin", "Can fix and resubmit"],
])

h2(doc, "9.2 Customer Risk Profile Workflow")
add_styled_table(doc, ["Current State", "Action", "Next State", "Actor", "Side Effects"], [
    ["(None)", "Initiate Assessment", "IN_PROGRESS", "RM", "Wizard started at Step 2"],
    ["IN_PROGRESS", "Submit Questionnaire", "COMPUTED", "RM/Client", "Score computed, category assigned"],
    ["COMPUTED", "Accept (No Deviation)", "PENDING_APPROVAL", "RM/Client", "CustomerRiskProfile created, is_deviated=false"],
    ["COMPUTED", "Deviate", "DEVIATED_PENDING", "Client", "is_deviated=true, reason recorded"],
    ["DEVIATED_PENDING", "Supervisor Approve", "ACTIVE", "Supervisor", "supervisor_approved=true, previous profile deactivated"],
    ["DEVIATED_PENDING", "Supervisor Reject", "COMPUTED", "Supervisor", "Deviation rejected, client must accept computed or re-assess"],
    ["PENDING_APPROVAL", "Supervisor Approve", "ACTIVE", "Supervisor", "is_active=true, previous deactivated"],
    ["ACTIVE", "Expiry (automatic)", "EXPIRED", "System", "is_active=false after expiry_date, RM notified"],
    ["ACTIVE", "Re-assess", "IN_PROGRESS", "RM", "New assessment initiated, old remains active until new completed"],
])

h2(doc, "9.3 Investment Proposal Workflow")
add_styled_table(doc, ["Current State", "Action", "Next State", "Actor", "Side Effects"], [
    ["(New)", "Create", "DRAFT", "RM", "Suitability check runs, results stored"],
    ["DRAFT", "Submit", "SUBMITTED", "RM", "L1 Supervisor notified"],
    ["SUBMITTED", "L1 Approve", "L1_APPROVED", "Supervisor", "Compliance Officer notified"],
    ["SUBMITTED", "L1 Reject", "L1_REJECTED", "Supervisor", "RM notified, terminal state"],
    ["SUBMITTED", "L1 Return", "DRAFT", "Supervisor", "RM notified, revision comments attached"],
    ["L1_APPROVED", "Compliance Approve", "COMPLIANCE_APPROVED", "Compliance", "PDF generated, client notified"],
    ["L1_APPROVED", "Compliance Reject", "COMPLIANCE_REJECTED", "Compliance", "RM and Supervisor notified"],
    ["L1_APPROVED", "Compliance Return", "DRAFT", "Compliance", "RM notified"],
    ["COMPLIANCE_APPROVED", "Send to Client", "SENT_TO_CLIENT", "System (auto)", "Client portal notification, email sent"],
    ["SENT_TO_CLIENT", "Client Accept", "CLIENT_ACCEPTED", "Client", "RM notified, downstream order trigger"],
    ["SENT_TO_CLIENT", "Client Reject", "CLIENT_REJECTED", "Client", "RM notified with reason"],
    ["SENT_TO_CLIENT", "Expire", "EXPIRED", "System", "After 30 days (configurable), auto-expire"],
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 10: NOTIFICATIONS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "10. Notification & Communication Requirements")

notifs = [
    ["Questionnaire Authorized", "In-app + Email", "Maker (Ops Admin)", "Checker authorizes questionnaire", "Your questionnaire '{name}' has been authorized by {checker_name}."],
    ["Questionnaire Rejected", "In-app + Email", "Maker (Ops Admin)", "Checker rejects questionnaire", "Your questionnaire '{name}' has been rejected. Reason: {reason}."],
    ["Risk Profile Pending Approval", "In-app", "RM Supervisor", "RM submits customer risk profile", "Risk profile for customer {name} (ID: {id}) is pending your approval."],
    ["Risk Profile Approved", "In-app + Email", "RM", "Supervisor approves risk profile", "Risk profile for {customer_name} has been approved."],
    ["Risk Profile Expiring", "In-app + Email", "RM", "30 days before expiry_date", "Risk profile for {customer_name} expires on {date}. Please schedule re-assessment."],
    ["Proposal Pending L1 Approval", "In-app + Email", "Supervisor", "RM submits proposal", "Proposal {number} for {customer} is pending your review."],
    ["Proposal Pending Compliance", "In-app + Email", "Compliance Officer", "L1 approves proposal", "Proposal {number} requires compliance review."],
    ["Proposal Sent to Client", "In-app + Email + SMS", "Client", "Compliance approves", "A new investment proposal is available in your portal. Please review."],
    ["Proposal Accepted", "In-app + Email", "RM", "Client accepts proposal", "Client {name} has accepted proposal {number}."],
    ["Proposal Rejected", "In-app + Email", "RM", "Client rejects proposal", "Client {name} has rejected proposal {number}. Reason: {reason}."],
    ["Approval SLA Breach", "In-app + Email", "Approver + Escalation", "24h without action", "Proposal {number} approval is overdue. Please take action."],
    ["Product Risk Deviation", "In-app (log)", "Compliance (batch)", "Client transacts in higher-risk product", "Logged for reporting. No real-time notification."],
]
add_styled_table(doc, ["Event", "Channel", "Recipient", "Trigger", "Message Template"], notifs)

para(doc, "Notification Preferences: Users can opt out of email notifications via Settings. In-app notifications cannot be disabled. SMS requires client opt-in.", bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 11: REPORTING & ANALYTICS
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "11. Reporting & Analytics")

reports = [
    ["Risk Profiling Completion Report", "Compliance, Supervisors", "CustomerRiskProfile, Customer, User", "Date range, RM, Branch, Entity", "Daily refresh", "Clients profiled vs pending vs expired by RM/branch"],
    ["Transaction by Product Rating", "Compliance", "Orders, Products, CustomerRiskProfile", "Date range, RM, Product Rating", "Real-time", "All transactions with product risk rating column (1-6)"],
    ["Product Risk Mismatch Report", "Compliance", "ProductRiskDeviation", "Date range, RM", "Real-time", "Transactions where product risk > client risk"],
    ["Proposal Pipeline Dashboard", "Supervisors, Management", "InvestmentProposal, ProposalApproval", "Date range, RM, Entity", "Hourly refresh", "Funnel: Draft→Submitted→Approved→Accepted. Conversion rates."],
    ["Risk Distribution Analytics", "Compliance, Management", "CustomerRiskProfile", "Entity, Branch, Date range", "Daily refresh", "Pie chart of client base by risk category"],
    ["Supervisor Leads Dashboard", "Supervisors", "Leads, Campaigns", "Active campaigns only", "15-min refresh", "Bar chart + detail table of leads by RM by status"],
    ["Risk Profile Deviation Audit", "Compliance, Audit", "CustomerRiskProfile (deviated)", "Date range, Entity", "On-demand", "All deviated risk profiles with reasons and supervisor approvals"],
]
add_styled_table(doc, ["Report Name", "Audience", "Data Sources", "Filters", "Refresh", "Description"], reports)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 12: MIGRATION & LAUNCH
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "12. Migration & Launch Plan")

h2(doc, "12.1 Phased Rollout")
phases = [
    ["Phase 1 (v1.0)", "Core Risk Profiling", "Questionnaire CRUD with maker-checker, Risk Appetite Mapping, Asset Allocation Config, Customer Risk Assessment Wizard, Recommended Risk Profile, Supervisor Approval"],
    ["Phase 2 (v1.1)", "Product Risk Compliance", "Product Rating Alert, Risk Rating Filter, Risk Profile display across screens, Deviation tracking, Sigma reports"],
    ["Phase 3 (v2.0)", "Proposal Generation", "Investment Proposal Builder, What-If Analysis, Suitability Check, Approval Workflow, PDF Generation, Client Portal proposal view"],
    ["Phase 4 (v2.1)", "Advanced Features", "Model Portfolio Management, Proposal Comparison, Rebalancing triggers, Advanced analytics, Bulk proposal generation"],
]
add_styled_table(doc, ["Phase", "Focus Area", "Features"], phases)

h2(doc, "12.2 Data Migration")
bullet(doc, "Migrate existing questionnaire configurations from legacy systems (if any) via SQL scripts with validation.")
bullet(doc, "Import historical risk profile data with status mapping to new schema.")
bullet(doc, "Backfill ProductRiskDeviation records from historical order data where product risk exceeded client risk.")
bullet(doc, "Validate all migrated data against new schema constraints before go-live.")

h2(doc, "12.3 Go-Live Checklist")
checklist = [
    "All questionnaire configurations migrated and AUTHORIZED",
    "Risk Appetite Mappings configured and AUTHORIZED for all entities",
    "Asset Allocation Configs defined for all 6 risk categories per entity",
    "Maker-checker workflow validated with test records",
    "Risk score computation verified against manual calculations (min 20 test cases)",
    "Product Risk Deviation alerts tested for all product types",
    "Supervisor approval workflow tested end-to-end",
    "PDF generation tested with all entity brandings",
    "Client Portal proposal view tested across browsers",
    "Performance testing completed: 500 concurrent users, <2s response times",
    "Security audit passed: OWASP Top 10, penetration testing",
    "User training completed for Ops Admins, RMs, Supervisors, Compliance",
    "Runbook documented for operational support",
]
for c in checklist:
    bullet(doc, c)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 13: GLOSSARY
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "13. Glossary")

glossary = [
    ["SAF", "Suitability Assessment Framework — SEBI-mandated scored questionnaire for assessing investor risk appetite in India."],
    ["Maker-Checker", "Four-eyes authorization principle where the person who creates/modifies a record (Maker) cannot be the same person who authorizes it (Checker)."],
    ["Risk Appetite", "The level of financial risk a customer is willing and able to accept, determined through scored questionnaire responses."],
    ["Risk Code", "Numeric representation (1-6) of risk level. 1 = lowest risk (Conservative), 6 = highest risk (Very Aggressive)."],
    ["Score Normalization", "Process of converting raw cumulative scores from multi-select questions into a single comparable score via predefined ranges."],
    ["Model Portfolio", "A predefined target asset allocation template for a specific risk category, used as the starting point for investment proposals."],
    ["Suitability Check", "Automated validation that a proposed investment aligns with the client's risk profile, investment experience, and regulatory requirements."],
    ["Risk Deviation", "When a client's chosen or transacted risk level differs from the system-computed risk profile, requiring explicit acknowledgement and/or supervisor approval."],
    ["Product Risk Class", "Risk classification (1-6) assigned to investment products based on their inherent risk characteristics."],
    ["Sharpe Ratio", "Risk-adjusted return metric = (Portfolio Return - Risk-Free Rate) / Portfolio Standard Deviation."],
    ["RM", "Relationship Manager — primary wealth management advisor assigned to a client."],
    ["PMS", "Portfolio Management System — the RM-facing application for client management and transactions."],
    ["CRM", "Customer Relationship Management — the back-office maintenance application."],
    ["IIFL / 360One", "India Infoline Wealth Management, now branded as 360One — a reference implementation."],
    ["JBR", "Jio BlackRock — a reference implementation for the Jio Financial Services and BlackRock joint venture."],
    ["Drift Threshold", "The percentage deviation from target allocation that triggers a rebalancing recommendation."],
    ["Effective Risk Category", "The risk category actually used for a client — equals the deviated category if deviation was approved, otherwise the system-computed category."],
]
add_styled_table(doc, ["Term", "Definition"], glossary)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# SECTION 14: APPENDICES
# ══════════════════════════════════════════════════════════════════════════

h1(doc, "14. Appendices")

h2(doc, "14.1 Appendix A: Reference Implementations")
para(doc, "This BRD synthesizes requirements validated across five reference implementations:")
refs = [
    ["360One (IIFL Wealth)", "India", "Full risk profiling with 6 questionnaire types (including FATCA, PAMM Pre-Investment). CRM-based maintenance with maker-checker."],
    ["HSBC Private Wealth Solutions", "Global", "Operations Office-based maintenance. Same 3-card Risk Assessor structure. Survey and SAF questionnaire types."],
    ["SBI Wealth", "India", "3-step wizard (Edit Profile → Assess Risk → Transact). Customer profile with 7 accordion sections. Supervisor approval via Customer Onboarding queue. Deviation feature disabled."],
    ["SBI International", "International", "Operations Portal with ML/OMS/FNC portals. Same Risk Assessor card structure. SAF + non-scored questionnaires."],
    ["Jio BlackRock (JBR)", "India", "Latest implementation (v1.1, 2024). OperationsOffice + PMX + RevenueDesk. Same Risk Assessor pattern with 3 cards."],
]
add_styled_table(doc, ["Entity", "Jurisdiction", "Key Characteristics"], refs)

h2(doc, "14.2 Appendix B: Risk Category Standard Mapping")
add_styled_table(doc, ["Risk Code", "Risk Category", "Description", "Typical Client Profile"], [
    ["1", "Conservative", "Capital preservation focus, minimal risk tolerance", "Retirees, risk-averse investors, short-term goals"],
    ["2", "Low to Moderate", "Low risk with small growth component", "Near-retirement, moderate savings goals"],
    ["3", "Moderate", "Balanced risk-return, diversified portfolio", "Mid-career professionals, medium-term goals"],
    ["4", "Moderately High", "Growth-oriented with some risk tolerance", "Experienced investors, long-term goals"],
    ["5", "Aggressive", "High growth, significant risk tolerance", "High-net-worth, experienced, long horizon"],
    ["6", "Very Aggressive", "Maximum growth, highest risk tolerance", "Ultra-HNI, sophisticated investors, very long horizon"],
])

h2(doc, "14.3 Appendix C: Entity Relationship Summary")
para(doc, "Key relationships between data entities:")
bullet(doc, "Questionnaire 1→N Question 1→N AnswerOption")
bullet(doc, "Question 1→N ScoreNormalizationRange (only when scoring_type=RANGE)")
bullet(doc, "RiskAppetiteMapping 1→N RiskAppetiteBand")
bullet(doc, "AssetAllocationConfig 1→N AssetAllocationLine")
bullet(doc, "Customer 1→N CustomerRiskProfile (only 1 active at a time)")
bullet(doc, "CustomerRiskProfile 1→N CustomerRiskResponse")
bullet(doc, "InvestmentProposal 1→N ProposalLineItem")
bullet(doc, "InvestmentProposal 1→N ProposalApproval")
bullet(doc, "Customer 1→N ProductRiskDeviation")
bullet(doc, "ModelPortfolio N→1 RiskAppetiteBand (via risk_category)")

# ── Save Final Document ──────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(__file__), "BRD-RiskProfiling-ProposalGeneration-v2.docx")
doc.save(output_path)
print(f"BRD saved successfully: {output_path}")
print("14 sections, 37 functional requirements, 16 data entities, complete.")
