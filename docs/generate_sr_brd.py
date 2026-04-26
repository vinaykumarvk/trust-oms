#!/usr/bin/env python3
"""
Generate the Service Request / Task Management Module BRD as a .docx file.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Arial'
    heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    heading_style.font.bold = True
    if level == 1:
        heading_style.font.size = Pt(18)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(14)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    else:
        heading_style.font.size = Pt(12)
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header
        hp = header.paragraphs[0]
        hp.text = "CONFIDENTIAL — TrustOMS Service Request / Task Management BRD v1.0"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.style.font.size = Pt(8)
        hp.style.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.style.font.size = Pt(8)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a table with styled header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = text
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EBF0F7"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    return table


add_header_footer(doc)

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Service Request / Task Management Module")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Business Requirements Document (BRD)")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0x8C)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run("Version 1.0 — April 2026")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
sys_name = doc.add_paragraph()
sys_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sys_name.add_run("TrustOMS — Trust Banking Operations Management System")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = conf.add_run("CONFIDENTIAL")
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# DOCUMENT CONTROL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Document Control", level=1)

doc.add_heading("Version History", level=2)
add_styled_table(doc,
    ["Version", "Date", "Author", "Changes"],
    [
        ["0.1", "2026-04-20", "System Architect", "Initial draft based on reference implementations"],
        ["0.5", "2026-04-21", "System Architect", "Added field-level details, API specs, state diagrams"],
        ["1.0", "2026-04-22", "System Architect", "Final version with all sections complete"],
    ])

doc.add_paragraph()
doc.add_heading("Approvals", level=2)
add_styled_table(doc,
    ["Role", "Name", "Signature", "Date"],
    [
        ["Project Sponsor", "", "", ""],
        ["Business Analyst", "", "", ""],
        ["Technical Lead", "", "", ""],
        ["QA Lead", "", "", ""],
    ])

doc.add_paragraph()
doc.add_heading("References", level=2)
add_styled_table(doc,
    ["#", "Document", "Source", "Relevance"],
    [
        ["1", "HSBC PWS — Journey of Customer App Service Request V10", "Intellect Design", "Full lifecycle, field-level specs, status transitions"],
        ["2", "Maybank DTC — Journey of Customer App Service Request V11", "Intellect Design", "Mobile UX patterns, push notifications, Incomplete status flow"],
        ["3", "Jio BlackRock — Journey of Service Request Customer App V10", "Intellect Design", "Appointed dates, closure date configurability, field editability matrix"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (placeholder)
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. Scope & Boundaries",
    "3. User Roles & Permissions",
    "4. Data Model",
    "5. Functional Requirements",
    "6. User Interface Requirements",
    "7. API & Integration Requirements",
    "8. Non-Functional Requirements",
    "9. Workflow & State Diagrams",
    "10. Notification & Communication Requirements",
    "11. Reporting & Analytics",
    "12. Migration & Launch Plan",
    "13. Glossary",
    "14. Appendices",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("1. Executive Summary", level=1)

doc.add_heading("1.1 Project Overview", level=2)
doc.add_paragraph(
    "The Service Request / Task Management Module (SR-TM) is a core operational module within the "
    "TrustOMS platform that enables end-to-end lifecycle management of client service requests in a "
    "trust banking environment. The module serves two primary audiences: (1) trust banking clients who "
    "initiate, track, and manage service requests through the Client Portal, and (2) back-office staff "
    "(Relationship Managers, Tellers, Operations Officers) who process, verify, and resolve these requests "
    "through the Back-Office application. The system enforces SLA-driven timelines, dual-control "
    "verification workflows, document management, and full audit trails — all aligned with BSP (Bangko "
    "Sentral ng Pilipinas) regulatory requirements for Philippine trust banking operations."
)

doc.add_heading("1.2 Business Objectives", level=2)
objectives = [
    "Digitize the service request lifecycle — eliminate paper-based SR forms and manual tracking spreadsheets, reducing processing time by 60%.",
    "Enforce SLA compliance — auto-compute closure dates based on priority (HIGH=3 days, MEDIUM=5 days, LOW=7 days) and surface overdue SRs in real-time dashboards.",
    "Enable client self-service — allow clients to create, modify, track, and close service requests through the Client Portal without visiting a branch.",
    "Implement dual-control verification — require RM preparation and Teller verification before completing any SR, ensuring segregation of duties per BSP Circular 857.",
    "Provide operational visibility — give operations managers a real-time KPI dashboard with status distribution, SLA breach counts, and aging analysis.",
    "Maintain complete audit trail — log every status change, field modification, and document upload with user ID and timestamp for regulatory examination readiness.",
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading("1.3 Target Users & Pain Points", level=2)
add_styled_table(doc,
    ["User", "Current Pain Point", "How SR-TM Solves It"],
    [
        ["Trust Banking Client", "Must visit branch or call RM to raise service requests; no visibility into request status", "Self-service portal with real-time status tracking, SLA visibility, and document upload"],
        ["Relationship Manager (RM)", "Tracks SRs in spreadsheets; no systematic handoff to teller team; SLA breaches discovered too late", "Structured SR queue with status filters, one-click send-for-verification, overdue SLA alerts"],
        ["Teller / Operations", "Receives SR instructions verbally or via email; no standardized verification workflow", "Dedicated verification queue with complete/incomplete/reject actions and verification notes"],
        ["Operations Manager", "No consolidated view of SR pipeline; manual SLA tracking", "KPI dashboard with real-time status counts, SLA breach metrics, and aging analysis"],
        ["Compliance Officer", "Difficult to reconstruct SR processing history for BSP examinations", "Full audit trail with every state change, user action, and document version logged"],
    ])

doc.add_heading("1.4 Success Metrics (KPIs)", level=2)
add_styled_table(doc,
    ["KPI", "Target", "Measurement Method"],
    [
        ["Average SR Processing Time", "≤ SLA days per priority level", "Median days from creation to COMPLETED status"],
        ["SLA Compliance Rate", "≥ 95% of SRs completed within SLA", "Count of on-time completions / total completions"],
        ["Client Self-Service Adoption", "≥ 70% of SRs created via Client Portal", "Portal-created SRs / total SRs"],
        ["First-Pass Completion Rate", "≥ 85% completed without INCOMPLETE cycle", "SRs going directly READY_FOR_TELLER → COMPLETED"],
        ["Digital Document Upload Rate", "≥ 80% of SRs have documents uploaded digitally", "SRs with documents[] not empty / total SRs"],
        ["Average Response Time (API)", "≤ 500ms for list endpoints", "P95 latency from API monitoring"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 2: SCOPE & BOUNDARIES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("2. Scope & Boundaries", level=1)

doc.add_heading("2.1 In Scope", level=2)
in_scope = [
    "Client Portal: Service Request list view with search, status filters, and pagination",
    "Client Portal: Create New Service Request form with SR type selection, priority, details, remarks, and document upload",
    "Client Portal: Service Request detail view with status-aware field editability",
    "Client Portal: Close Request action with mandatory closure reason dialog",
    "Client Portal: Resubmit for Verification action from INCOMPLETE status",
    "Client Portal: Mobile-responsive design (card-based mobile view, table-based desktop view)",
    "Back-Office: Service Request Workbench with KPI summary cards and filterable data table",
    "Back-Office: Send for Verification workflow (RM fills branch, unit, date, then transitions to READY_FOR_TELLER)",
    "Back-Office: Complete Request action (Teller → COMPLETED)",
    "Back-Office: Mark Incomplete action with verification notes (Teller → INCOMPLETE)",
    "Back-Office: Reject Request action with mandatory reason (Teller → REJECTED)",
    "Back-Office: SR detail update (remarks, documents, closure date)",
    "Backend: RESTful API endpoints for all CRUD and status transition operations",
    "Backend: Role-based access control separating client portal and back-office endpoints",
    "Backend: Paginated list queries with status, priority, and text search filters",
    "Backend: Auto-generated sequential Request IDs (SR-YYYY-NNNNNN format)",
    "Backend: SLA-driven closure date computation",
    "Backend: Full audit trail (created_by, updated_by, timestamps on all records)",
    "Database: Service Request table with all fields, indexes, and relationships",
    "Database: Enum types for status, priority, and SR type",
]
for item in in_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("2.2 Out of Scope", level=2)
out_scope = [
    "SR Type maintenance / configuration UI (SR types are pre-defined as database enums; dynamic type creation is deferred to Phase 2)",
    "Push notification delivery (notification triggers are defined but actual push/SMS delivery requires integration with a notification gateway — deferred)",
    "Email notifications for status changes (will be added when the Notification Service is integrated)",
    "Document storage backend (documents are tracked as JSONB filename references; actual file storage uses the existing upload service)",
    "Service Request analytics / trend reporting beyond the KPI summary dashboard",
    "SLA escalation workflows (auto-escalation to supervisor when SLA is breached — deferred to Phase 2)",
    "Multi-language / i18n support for SR forms and labels",
    "Customer satisfaction surveys post-SR completion",
    "Integration with external CRM systems for SR synchronization",
    "Business-day calendar awareness in SLA computation (SLA currently uses calendar days; business-day support deferred)",
]
for item in out_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("2.3 Assumptions", level=2)
assumptions = [
    "Users are authenticated before accessing SR features; authentication is handled by the existing auth module (JWT-based httpOnly cookies).",
    "Client IDs are pre-existing in the system; the SR module references clients but does not create them.",
    "The upload service at /api/v1/uploads is available for document file storage; SR module stores only file references.",
    "SR types are a fixed set defined in the database enum; adding new types requires a schema migration.",
    "All dates and times are stored in UTC and displayed in the user's local timezone by the frontend.",
    "The back-office application is accessed only from managed corporate devices on internal networks.",
    "The client portal is accessible from any modern browser (Chrome 90+, Safari 14+, Firefox 88+, Edge 90+).",
]
for item in assumptions:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("2.4 Constraints", level=2)
constraints = [
    "Regulatory: All SR state changes must be logged with full audit trail per BSP Circular 857 (Internal Controls).",
    "Technical: Backend must use the existing Express.js + Drizzle ORM + PostgreSQL stack.",
    "Technical: Frontend must use React + React Router + TanStack Query + shadcn/ui component library.",
    "Technical: Document uploads limited to PDF format, maximum 10MB total per SR.",
    "Security: Back-office endpoints must be protected by role-based middleware (requireBackOfficeRole).",
    "Security: Client portal SR endpoints must be scoped to the authenticated client's ID.",
    "Performance: List API endpoints must return results within 500ms for up to 10,000 service requests.",
]
for item in constraints:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 3: USER ROLES & PERMISSIONS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("3. User Roles & Permissions", level=1)

doc.add_heading("3.1 Role Definitions", level=2)
add_styled_table(doc,
    ["Role", "Description", "Application"],
    [
        ["Client (Customer)", "Trust banking client who initiates and tracks their own service requests", "Client Portal"],
        ["Relationship Manager (RM)", "Bank officer assigned to manage a portfolio of clients; prepares SRs for teller verification", "Back-Office"],
        ["Teller", "Operations staff who performs final verification and completes/rejects service requests", "Back-Office"],
        ["Operations Manager", "Supervisor who monitors SR pipeline, SLA compliance, and team performance", "Back-Office"],
        ["Compliance Officer", "Reviews audit trails and ensures SR processing complies with regulatory requirements", "Back-Office"],
        ["System Administrator", "Manages system configuration, user accounts, and SR type definitions", "Back-Office"],
    ])

doc.add_paragraph()
doc.add_heading("3.2 Permissions Matrix", level=2)
add_styled_table(doc,
    ["Action", "Client", "RM", "Teller", "Ops Manager", "Compliance", "Admin"],
    [
        ["Create Service Request", "Yes (own)", "Yes (any client)", "No", "No", "No", "Yes"],
        ["View SR List (own)", "Yes", "—", "—", "—", "—", "—"],
        ["View SR List (all)", "No", "Yes", "Yes", "Yes", "Yes (read-only)", "Yes"],
        ["View SR Detail", "Yes (own)", "Yes", "Yes", "Yes", "Yes (read-only)", "Yes"],
        ["Edit SR Fields (Approved status)", "Yes (own)", "Yes", "No", "No", "No", "Yes"],
        ["Edit SR Fields (Ready for Teller)", "Limited*", "Yes", "No", "No", "No", "Yes"],
        ["Close Request", "Yes (own, open statuses)", "Yes", "No", "No", "No", "Yes"],
        ["Send for Verification", "No", "Yes", "No", "No", "No", "Yes"],
        ["Complete Request", "No", "No", "Yes", "No", "No", "Yes"],
        ["Mark Incomplete", "No", "No", "Yes", "No", "No", "Yes"],
        ["Reject Request", "No", "No", "Yes", "No", "No", "Yes"],
        ["Resubmit for Verification", "Yes (own, INCOMPLETE)", "Yes", "No", "No", "No", "Yes"],
        ["View KPI Summary", "No", "Yes", "Yes", "Yes", "Yes", "Yes"],
        ["View Audit Trail", "No", "No", "No", "Yes", "Yes", "Yes"],
        ["Export SR Data", "No", "No", "No", "Yes", "Yes", "Yes"],
    ])

doc.add_paragraph("* Limited: Customer can edit closure_date only if the SR Type maintenance has 'Closure Date Change Allowed' flag set to Yes.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 4: DATA MODEL
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("4. Data Model", level=1)

doc.add_heading("4.1 Entity: service_requests", level=2)
doc.add_paragraph("Primary table storing all service request records with full lifecycle tracking.")

add_styled_table(doc,
    ["#", "Field Name", "Data Type", "Required", "Default", "Validation Rules", "Description"],
    [
        ["1", "id", "serial (PK)", "Auto", "Auto-increment", "Positive integer", "Internal primary key"],
        ["2", "request_id", "varchar(20)", "Yes", "Auto-generated", "Format: SR-YYYY-NNNNNN, unique", "Human-readable request identifier"],
        ["3", "client_id", "varchar(50)", "Yes", "None", "Must reference existing client", "FK to clients table"],
        ["4", "sr_type", "enum (sr_type_enum)", "Yes", "None", "Must be valid enum value", "Type of service request"],
        ["5", "sr_details", "text", "No", "null", "Max 2000 characters", "Detailed description of the request"],
        ["6", "priority", "enum (sr_priority_enum)", "Yes", "'MEDIUM'", "LOW | MEDIUM | HIGH", "Request priority level"],
        ["7", "sr_status", "enum (sr_status_enum)", "Yes", "'APPROVED'", "Valid status enum value", "Current lifecycle status"],
        ["8", "request_date", "timestamp", "Yes", "now()", "Cannot be future date", "Date/time request was created"],
        ["9", "closure_date", "timestamp", "Yes", "Computed", "Must be >= request_date", "SLA target closure date"],
        ["10", "actual_closure_date", "timestamp", "No", "null", "Set on COMPLETED/REJECTED/CLOSED", "Actual date request was closed"],
        ["11", "appointed_start_date", "timestamp", "No", "null", "Must be >= request_date", "Scheduled start date for service delivery"],
        ["12", "appointed_end_date", "timestamp", "No", "null", "Must be >= appointed_start_date", "Scheduled end date for service delivery"],
        ["13", "remarks", "text", "No", "null", "Max 1000 characters", "Customer or RM remarks"],
        ["14", "closure_reason", "text", "No", "null", "Required when status → CLOSED", "Reason for closing the request"],
        ["15", "rejection_reason", "text", "No", "null", "Required when status → REJECTED", "Reason for rejecting the request"],
        ["16", "verification_notes", "text", "No", "null", "Required when status → INCOMPLETE", "Teller notes explaining why SR is incomplete"],
        ["17", "documents", "jsonb", "No", "[]", "Array of {filename, url, uploadedAt}", "Uploaded document references"],
        ["18", "assigned_rm_id", "integer", "No", "null", "Must reference valid user with RM role", "FK to users table — assigned RM"],
        ["19", "service_branch", "varchar(100)", "No", "null", "Must match known branch code", "Branch handling the request"],
        ["20", "resolution_unit", "varchar(100)", "No", "null", "Free text", "Unit/team responsible for resolution"],
        ["21", "sales_date", "timestamp", "No", "null", "Valid date", "Sales date recorded by RM"],
        ["22", "teller_id", "integer", "No", "null", "Must reference valid user with Teller role", "FK to users table — teller who processed"],
        ["23", "created_by", "varchar(100)", "Yes", "'system'", "Non-empty string", "User who created the record"],
        ["24", "updated_by", "varchar(100)", "Yes", "'system'", "Non-empty string", "User who last updated the record"],
        ["25", "created_at", "timestamp", "Yes", "now()", "Auto-set", "Record creation timestamp"],
        ["26", "updated_at", "timestamp", "Yes", "now()", "Auto-set on every update", "Last modification timestamp"],
        ["27", "is_deleted", "boolean", "Yes", "false", "Soft delete flag", "Whether the record is soft-deleted"],
    ])

doc.add_paragraph()
doc.add_heading("4.2 Enum: sr_status_enum", level=2)
add_styled_table(doc,
    ["Value", "Description", "Terminal?"],
    [
        ["NEW", "Initial status when SR is created (used internally, auto-transitions to APPROVED)", "No"],
        ["APPROVED", "Auto-approved by system; visible to client and RM for processing", "No"],
        ["READY_FOR_TELLER", "RM has prepared the SR and sent for teller verification", "No"],
        ["COMPLETED", "Teller has verified and completed the SR", "Yes"],
        ["INCOMPLETE", "Teller returned the SR with notes; client/RM must address issues", "No"],
        ["REJECTED", "Teller rejected the SR with a reason", "Yes"],
        ["CLOSED", "Manually closed by client or RM before completion", "Yes"],
    ])

doc.add_paragraph()
doc.add_heading("4.3 Enum: sr_priority_enum", level=2)
add_styled_table(doc,
    ["Value", "SLA Days", "Description"],
    [
        ["HIGH", "3", "Urgent requests requiring expedited processing"],
        ["MEDIUM", "5", "Standard priority — default for new requests"],
        ["LOW", "7", "Low-urgency requests with extended processing window"],
    ])

doc.add_paragraph()
doc.add_heading("4.4 Enum: sr_type_enum", level=2)
add_styled_table(doc,
    ["Value", "Display Label", "Description"],
    [
        ["REVIEW_PORTFOLIO", "Review Portfolio", "Client requests a portfolio review meeting with RM"],
        ["MULTIPLE_MANDATE_REGISTRATION", "Multiple Mandate Registration", "Register additional mandates on trust account"],
        ["NOMINEE_UPDATION", "Nominee Update", "Update nominee/beneficiary details on trust account"],
        ["ACCOUNT_CLOSURE", "Account Closure", "Request closure of a trust account"],
        ["STATEMENT_REQUEST", "Statement Request", "Request for specific account statements"],
        ["ADDRESS_CHANGE", "Address Change", "Update client's registered address"],
        ["BENEFICIARY_UPDATE", "Beneficiary Update", "Add, remove, or modify trust beneficiaries"],
        ["GENERAL_INQUIRY", "General Inquiry", "General inquiry not covered by specific SR types"],
    ])

doc.add_paragraph()
doc.add_heading("4.5 Indexes", level=2)
add_styled_table(doc,
    ["Index Name", "Column(s)", "Type", "Purpose"],
    [
        ["idx_sr_request_id", "request_id", "Unique", "Fast lookup by human-readable ID"],
        ["idx_sr_client_id", "client_id", "Non-unique", "Fast filtering by client"],
        ["idx_sr_status", "sr_status", "Non-unique", "Fast filtering by status"],
        ["idx_sr_priority", "priority", "Non-unique", "Fast filtering by priority"],
        ["idx_sr_created_at", "created_at", "Non-unique", "Fast ordering by creation date"],
        ["idx_sr_assigned_rm", "assigned_rm_id", "Non-unique", "Fast filtering by assigned RM"],
    ])

doc.add_paragraph()
doc.add_heading("4.6 Relationships", level=2)
add_styled_table(doc,
    ["From", "To", "Type", "FK Column", "Description"],
    [
        ["service_requests", "clients", "Many-to-One", "client_id", "Each SR belongs to one client"],
        ["service_requests", "users (RM)", "Many-to-One", "assigned_rm_id", "Each SR may be assigned to one RM"],
        ["service_requests", "users (Teller)", "Many-to-One", "teller_id", "Each SR may be processed by one teller"],
    ])

doc.add_paragraph()
doc.add_heading("4.7 Sample Data", level=2)
add_styled_table(doc,
    ["request_id", "client_id", "sr_type", "priority", "sr_status", "request_date", "closure_date", "remarks"],
    [
        ["SR-2026-000001", "CLT-001", "REVIEW_PORTFOLIO", "HIGH", "APPROVED", "2026-04-20 09:30", "2026-04-23 09:30", "Please review my equity allocation"],
        ["SR-2026-000002", "CLT-002", "NOMINEE_UPDATION", "MEDIUM", "READY_FOR_TELLER", "2026-04-18 14:15", "2026-04-23 14:15", "Update nominee to spouse"],
        ["SR-2026-000003", "CLT-001", "STATEMENT_REQUEST", "LOW", "COMPLETED", "2026-04-10 11:00", "2026-04-17 11:00", "Q1 2026 statement needed for tax filing"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 5: FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("5. Functional Requirements", level=1)

# ── Module A: Client Portal ──
doc.add_heading("5.1 Module A: Client Portal — Service Request Management", level=2)

# FR-001
doc.add_heading("FR-001: View Service Request List", level=3)
doc.add_paragraph("As a Client, I want to view a list of all my service requests so that I can track their status and take action on pending items.")
doc.add_paragraph("Description: The Service Request list page displays all SRs belonging to the authenticated client, sorted by creation date (newest first). The page supports search by Request ID and Priority, status tab filters, and pagination.")
p = doc.add_paragraph("Acceptance Criteria:", style='List Bullet')
criteria = [
    "AC-001.1: The list displays columns: Request ID, SR Type, Priority (color-coded badge), Request Date, Closure Date, Request Age (days), Status (color-coded badge).",
    "AC-001.2: Search bar filters results by Request ID (partial match) and Priority text.",
    "AC-001.3: Status tabs filter: All (default), Approved, In Progress (READY_FOR_TELLER), Completed, Rejected. Tab counts reflect filtered totals.",
    "AC-001.4: Desktop displays a data table; mobile (< 768px) displays stacked cards with Request ID as header.",
    "AC-001.5: Clicking a row/card navigates to the Service Request Detail page (/service-requests/:id).",
    "AC-001.6: A '+ New Request' button is visible in the header, linking to /service-requests/new.",
    "AC-001.7: Empty state message 'No service requests found' is displayed when the list is empty.",
    "AC-001.8: Pagination controls appear when results exceed 25 items per page.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

doc.add_paragraph("Business Rules:")
rules = [
    "BR-001.1: Only SRs belonging to the authenticated client (client_id match) are returned.",
    "BR-001.2: Soft-deleted SRs (is_deleted=true) are never displayed.",
    "BR-001.3: Request Age is computed client-side as (today - request_date) in whole days.",
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet 2')

doc.add_paragraph("Edge Cases:")
edges = [
    "EC-001.1: Client with zero SRs sees empty state, not an error.",
    "EC-001.2: If API returns an error, a toast notification with the error message is displayed and the list shows 'Failed to load service requests. Please try again.'",
]
for e in edges:
    doc.add_paragraph(e, style='List Bullet 2')

# FR-002
doc.add_heading("FR-002: Create New Service Request", level=3)
doc.add_paragraph("As a Client, I want to create a new service request so that I can request banking services without visiting a branch.")
doc.add_paragraph("Description: The create form allows the client to select an SR type, enter details, choose priority, add remarks, and upload supporting documents. Upon submission, the system auto-approves the SR and displays a confirmation with the generated Request ID.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-002.1: SR Type dropdown lists all values from sr_type_enum with human-readable labels.",
    "AC-002.2: Service Request Details textarea (max 2000 chars) accepts free-text description.",
    "AC-002.3: Priority toggle shows Low / Medium (default selected) / High as segmented buttons.",
    "AC-002.4: Closure Date (Approx.) is auto-computed and displayed as read-only: request_date + SLA days based on selected priority.",
    "AC-002.5: Remarks textarea (max 1000 chars) is optional.",
    "AC-002.6: Document Upload accepts PDF files only, max 10MB total, with 'Add from files' button.",
    "AC-002.7: 'Send Service Request' button is disabled until SR Type is selected (mandatory field).",
    "AC-002.8: On successful submission, a success dialog displays: '<<Request ID>> is successfully submitted for processing. To view the status or to make modifications, use the Request ID mentioned.'",
    "AC-002.9: After dismissing the success dialog, the user is navigated back to the SR list page.",
    "AC-002.10: The created SR has status 'APPROVED' (auto-approved by system).",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

doc.add_paragraph("Business Rules:")
rules = [
    "BR-002.1: Request ID is auto-generated server-side in format SR-YYYY-NNNNNN (e.g., SR-2026-000001).",
    "BR-002.2: Closure Date = request_date + SLA days (HIGH=3, MEDIUM=5, LOW=7 calendar days).",
    "BR-002.3: The SR is created with status NEW internally, then immediately auto-transitions to APPROVED.",
    "BR-002.4: created_by is set to the authenticated client's user ID.",
    "BR-002.5: client_id is set to the authenticated client's client ID.",
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet 2')

doc.add_paragraph("Validation Rules:")
validations = [
    "VR-002.1: SR Type must be selected (mandatory).",
    "VR-002.2: If documents are uploaded, each file must be PDF and total size must not exceed 10MB.",
    "VR-002.3: Service Request Details, if provided, must not exceed 2000 characters.",
    "VR-002.4: Remarks, if provided, must not exceed 1000 characters.",
]
for v in validations:
    doc.add_paragraph(v, style='List Bullet 2')

# FR-003
doc.add_heading("FR-003: View/Modify Service Request in APPROVED Status", level=3)
doc.add_paragraph("As a Client, I want to view and modify my service request while it is in APPROVED status so that I can add details or correct information before the RM processes it.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-003.1: Header displays Request ID and Status badge (green 'Approved').",
    "AC-003.2: Read-only fields: Request ID, SR Type, Request Date, Closure Date, Priority, Request Age, Status.",
    "AC-003.3: Editable fields: Service Request Details, Remarks.",
    "AC-003.4: Documents section shows uploaded files with option to add more.",
    "AC-003.5: 'Save' button saves edited fields; displays confirmation '<<Request ID>> details saved successfully'.",
    "AC-003.6: 'Close Request' button opens a dialog requiring a closure reason (mandatory textarea).",
    "AC-003.7: After closing, status changes to CLOSED and all fields become read-only.",
    "AC-003.8: 'Back to Service Requests' link navigates to the SR list.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-004
doc.add_heading("FR-004: View/Modify Service Request in READY_FOR_TELLER Status", level=3)
doc.add_paragraph("As a Client, I want to view my service request when it is being verified by the teller so that I can track its progress and make limited edits if allowed.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-004.1: All initiation fields are displayed (SR Type, Details, Priority, Request Date, etc.).",
    "AC-004.2: Additional fields visible: Appointed Start Date, Appointed End Date (set by RM), Service Branch, Resolution Unit.",
    "AC-004.3: Closure Date is editable ONLY if the SR Type has 'Closure Date Change Allowed' flag = Yes.",
    "AC-004.4: Remarks field is editable — client can add additional notes.",
    "AC-004.5: 'Save' button persists changes; 'Close Request' button available with reason dialog.",
    "AC-004.6: Client cannot modify SR Type, Priority, or Service Request Details in this status.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-005
doc.add_heading("FR-005: View/Modify Service Request in INCOMPLETE Status", level=3)
doc.add_paragraph("As a Client, I want to address issues flagged by the teller when my SR is marked incomplete so that I can resubmit it for verification.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-005.1: Verification notes from the teller are displayed prominently (highlighted card or alert).",
    "AC-005.2: Editable fields: Service Request Details, Remarks.",
    "AC-005.3: Document section allows re-uploading corrected documents (replace or add new files).",
    "AC-005.4: 'Save' button saves changes (status remains INCOMPLETE).",
    "AC-005.5: 'Re-send for Verification' button transitions status back to READY_FOR_TELLER.",
    "AC-005.6: 'Close Request' button is available with mandatory reason dialog.",
    "AC-005.7: Confirmation message on resubmit: '<<Request ID>> has been resubmitted for verification.'",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-006
doc.add_heading("FR-006: View Service Request in Terminal Status (COMPLETED / REJECTED / CLOSED)", level=3)
doc.add_paragraph("As a Client, I want to view the final details of a completed, rejected, or closed service request for my records.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-006.1: All fields are displayed in read-only mode — no edit controls visible.",
    "AC-006.2: If REJECTED: rejection_reason is displayed in a red-bordered alert.",
    "AC-006.3: If CLOSED: closure_reason is displayed in a yellow-bordered alert.",
    "AC-006.4: If COMPLETED: actual_closure_date is displayed; no action buttons shown.",
    "AC-006.5: 'Back to Service Requests' link is the only navigation action available.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

doc.add_page_break()

# ── Module B: Back-Office ──
doc.add_heading("5.2 Module B: Back-Office — Service Request Workbench", level=2)

# FR-007
doc.add_heading("FR-007: SR Workbench KPI Dashboard", level=3)
doc.add_paragraph("As an Operations Manager, I want to see a summary dashboard of all service requests so that I can monitor the pipeline and identify SLA breaches.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-007.1: KPI cards displayed at top of workbench: Total SRs, Approved (pending RM action), Ready for Teller, Completed, Overdue SLA (count of open SRs past closure_date).",
    "AC-007.2: Each KPI card shows a count number with a descriptive label.",
    "AC-007.3: Overdue SLA card is highlighted in red/amber when count > 0.",
    "AC-007.4: KPI data refreshes on page load and can be manually refreshed.",
    "AC-007.5: Loading skeleton shown while KPI data is being fetched.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-008
doc.add_heading("FR-008: SR Workbench List View", level=3)
doc.add_paragraph("As an RM, I want to view and filter all service requests so that I can prioritize my work queue and process requests efficiently.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-008.1: Data table columns: Request ID, Client Name, SR Type, Priority (badge), Request Date, Closure Date, Age (days), Status (badge), Actions.",
    "AC-008.2: Status tab filters: All, New, Approved, Ready for Teller, Incomplete, Completed, Rejected, Closed.",
    "AC-008.3: Search bar searches across Request ID, Client Name, SR Type.",
    "AC-008.4: Pagination with configurable page size (10/25/50/100).",
    "AC-008.5: Each row has context-specific action buttons based on status (see FR-009 through FR-013).",
    "AC-008.6: Clicking Request ID opens a detail/edit dialog or navigates to detail view.",
    "AC-008.7: Desktop shows table view; mobile shows card-based list.",
    "AC-008.8: Rows with SLA breached (now > closure_date and status is open) are highlighted with a warning indicator.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-009
doc.add_heading("FR-009: Send for Verification (RM Action)", level=3)
doc.add_paragraph("As an RM, I want to prepare a service request and send it for teller verification so that the request can be processed by the operations team.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-009.1: Action available only for SRs in APPROVED status.",
    "AC-009.2: Clicking 'Send for Verification' opens a dialog with fields: Service Branch (dropdown), Resolution Unit (text input), Sales Date (date picker).",
    "AC-009.3: Service Branch is mandatory; Resolution Unit and Sales Date are optional.",
    "AC-009.4: Submitting the dialog transitions status from APPROVED → READY_FOR_TELLER.",
    "AC-009.5: Success toast: '<<Request ID>> sent for teller verification.'",
    "AC-009.6: The dialog can be cancelled without making changes.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-010
doc.add_heading("FR-010: Complete Request (Teller Action)", level=3)
doc.add_paragraph("As a Teller, I want to mark a service request as completed after verifying and processing it so that the client is notified of successful resolution.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-010.1: Action available only for SRs in READY_FOR_TELLER status.",
    "AC-010.2: Clicking 'Complete' triggers a confirmation dialog: 'Are you sure you want to mark <<Request ID>> as completed?'",
    "AC-010.3: Confirming transitions status from READY_FOR_TELLER → COMPLETED.",
    "AC-010.4: actual_closure_date is set to current timestamp.",
    "AC-010.5: teller_id is set to the authenticated teller's user ID.",
    "AC-010.6: Success toast: '<<Request ID>> completed successfully.'",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-011
doc.add_heading("FR-011: Mark Incomplete (Teller Action)", level=3)
doc.add_paragraph("As a Teller, I want to return a service request to the client/RM as incomplete with notes explaining what needs to be addressed.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-011.1: Action available only for SRs in READY_FOR_TELLER status.",
    "AC-011.2: Clicking 'Mark Incomplete' opens a dialog with a mandatory verification notes textarea.",
    "AC-011.3: Notes textarea requires at least 10 characters (meaningful explanation).",
    "AC-011.4: Submitting transitions status from READY_FOR_TELLER → INCOMPLETE.",
    "AC-011.5: teller_id and verification_notes are saved.",
    "AC-011.6: Success toast: '<<Request ID>> marked as incomplete.'",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-012
doc.add_heading("FR-012: Reject Request (Teller Action)", level=3)
doc.add_paragraph("As a Teller, I want to reject a service request with a documented reason so that invalid or unfulfillable requests are properly dispositioned.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-012.1: Action available only for SRs in READY_FOR_TELLER status.",
    "AC-012.2: Clicking 'Reject' opens a dialog with a mandatory rejection reason textarea.",
    "AC-012.3: Reason textarea requires at least 10 characters.",
    "AC-012.4: Submitting transitions status from READY_FOR_TELLER → REJECTED.",
    "AC-012.5: rejection_reason and actual_closure_date are saved.",
    "AC-012.6: Confirmation dialog warns: 'This action cannot be undone. The service request will be permanently rejected.'",
    "AC-012.7: Success toast: '<<Request ID>> rejected.'",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-013
doc.add_heading("FR-013: Update Service Request Details (RM Action)", level=3)
doc.add_paragraph("As an RM, I want to update service request details, add documents, and modify remarks before sending for verification.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-013.1: Editable fields (for non-terminal statuses): SR Details, Remarks, Documents, Closure Date.",
    "AC-013.2: Changes are persisted via PUT endpoint with optimistic UI update.",
    "AC-013.3: Validation: cannot modify SRs in COMPLETED, REJECTED, or CLOSED status.",
    "AC-013.4: Success toast on save: '<<Request ID>> updated successfully.'",
    "AC-013.5: Error toast if update fails: 'Failed to update <<Request ID>>. [error message]'.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

doc.add_page_break()

# ── Module C: Backend Service ──
doc.add_heading("5.3 Module C: Backend Service & API", level=2)

# FR-014
doc.add_heading("FR-014: Auto-Generate Request ID", level=3)
doc.add_paragraph("As the System, I want to auto-generate unique, sequential request IDs so that each SR has a human-readable, non-colliding identifier.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-014.1: Format: SR-YYYY-NNNNNN (e.g., SR-2026-000001).",
    "AC-014.2: YYYY is the current year at time of creation.",
    "AC-014.3: NNNNNN is a zero-padded 6-digit sequence number, resetting at year boundary.",
    "AC-014.4: Sequence is derived from COUNT of existing SRs for the current year + 1.",
    "AC-014.5: request_id has a UNIQUE database constraint preventing duplicates.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-015
doc.add_heading("FR-015: SLA Closure Date Computation", level=3)
doc.add_paragraph("As the System, I want to compute the closure date based on the priority SLA so that every SR has a target completion date.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-015.1: closure_date = request_date + SLA_DAYS[priority].",
    "AC-015.2: SLA_DAYS mapping: HIGH=3, MEDIUM=5, LOW=7.",
    "AC-015.3: If priority is not recognized, default to 5 days.",
    "AC-015.4: Closure date is set at creation time and not auto-recalculated on priority change.",
    "AC-015.5: Closure date can be manually edited by RM if allowed by SR Type configuration.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-016
doc.add_heading("FR-016: Status Transition Validation", level=3)
doc.add_paragraph("As the System, I want to enforce valid status transitions so that SRs cannot be moved to invalid states.")

doc.add_paragraph("Valid Transitions:")
add_styled_table(doc,
    ["From Status", "To Status", "Triggered By", "Required Fields"],
    [
        ["NEW", "APPROVED", "System (auto)", "None (auto-transition on create)"],
        ["APPROVED", "READY_FOR_TELLER", "RM", "service_branch (mandatory)"],
        ["APPROVED", "CLOSED", "Client / RM", "closure_reason (mandatory)"],
        ["READY_FOR_TELLER", "COMPLETED", "Teller", "teller_id (auto from auth)"],
        ["READY_FOR_TELLER", "INCOMPLETE", "Teller", "teller_id, verification_notes (mandatory)"],
        ["READY_FOR_TELLER", "REJECTED", "Teller", "rejection_reason (mandatory)"],
        ["READY_FOR_TELLER", "CLOSED", "Client / RM", "closure_reason (mandatory)"],
        ["INCOMPLETE", "READY_FOR_TELLER", "Client / RM", "None"],
        ["INCOMPLETE", "CLOSED", "Client / RM", "closure_reason (mandatory)"],
    ])

doc.add_paragraph("Business Rules:")
rules = [
    "BR-016.1: Any transition not in the table above must be rejected with HTTP 400 and message 'Invalid status transition from {current} to {target}'.",
    "BR-016.2: Terminal statuses (COMPLETED, REJECTED, CLOSED) cannot transition to any other status.",
    "BR-016.3: actual_closure_date is set automatically when transitioning to COMPLETED, REJECTED, or CLOSED.",
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet 2')

# FR-017
doc.add_heading("FR-017: Request Age Computation", level=3)
doc.add_paragraph("As the System, I want to compute the request age in days so that users can quickly assess how long a request has been open.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-017.1: request_age = floor((current_timestamp - request_date) / (24 * 60 * 60 * 1000)) in whole days.",
    "AC-017.2: Computed at read time (not stored); included in GET response payloads.",
    "AC-017.3: Age continues counting even for terminal statuses (shows total days from creation).",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

# FR-018
doc.add_heading("FR-018: Soft Delete", level=3)
doc.add_paragraph("As the System, I want to support soft deletion of service requests so that deleted records can be recovered if needed.")

doc.add_paragraph("Acceptance Criteria:")
criteria = [
    "AC-018.1: All list/detail queries include WHERE is_deleted = false.",
    "AC-018.2: Delete operations set is_deleted = true instead of removing the row.",
    "AC-018.3: Soft-deleted SRs do not appear in any list or count query.",
]
for c in criteria:
    doc.add_paragraph(c, style='List Bullet 2')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 6: USER INTERFACE REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("6. User Interface Requirements", level=1)

doc.add_heading("6.1 Client Portal Screens", level=2)

doc.add_heading("6.1.1 Service Request List (/service-requests)", level=3)
doc.add_paragraph("Purpose: Display all service requests belonging to the authenticated client.")
items = [
    "Layout: Full-width page with header (title + '+ New Request' button), search bar, status tab strip, and content area.",
    "Header: 'Service Requests' title left-aligned; '+ New Request' primary button right-aligned.",
    "Search: Single text input with search icon, placeholder 'Search by Request ID or Priority'. Debounced (300ms).",
    "Tabs: Horizontal tab strip — All (default) | Approved | In Progress | Completed | Rejected. Active tab underlined in primary color.",
    "Desktop (≥ 768px): Data table with columns: Request ID, SR Type, Priority, Request Date, Closure Date, Age, Status. Rows clickable.",
    "Mobile (< 768px): Stacked cards — each card shows Request ID (header), SR Type + Closure Date (row 1), Request Age (row 2), Status badge (top-right corner).",
    "Empty State: Centered icon + 'No service requests found' text + '+ Create your first request' link.",
    "Loading State: Skeleton rows/cards matching the expected layout.",
    "Error State: Red alert banner with 'Failed to load service requests' and a 'Retry' button.",
    "Pagination: Page number buttons + Previous/Next at bottom; shows 'Showing 1-25 of 47'.",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("6.1.2 Create Service Request (/service-requests/new)", level=3)
doc.add_paragraph("Purpose: Form for creating a new service request.")
items = [
    "Layout: Single-column centered form card (max-width 600px) with back arrow navigation.",
    "SR Type: Dropdown/Select with placeholder 'Select Service Request Type'. Lists all sr_type_enum values with display labels. Mandatory (asterisk indicator).",
    "Service Request Details: Textarea, 4 rows, placeholder 'Enter Service Request Details'. Character counter showing remaining of 2000.",
    "Priority: Three-button segmented control — Low | Medium (default) | High. Selected button has filled background.",
    "Closure Date (Approx.): Read-only date display, auto-updated when priority changes.",
    "Remarks: Textarea, 2 rows, placeholder 'Enter Remarks If Any'. Optional.",
    "Status: Read-only text displaying 'New'.",
    "Document Upload: Dropzone area with file icon + 'Add from files' button. Note below: 'File format to be in .pdf and the total size of the uploaded document should not exceed 10MB.'",
    "Submit Button: Full-width primary button 'Send Service Request'. Disabled until SR Type is selected. Shows loading spinner during submission.",
    "Success Dialog: Modal with check icon, Request ID in bold, and message text. 'OK' button dismisses and navigates to list.",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("6.1.3 Service Request Detail (/service-requests/:id)", level=3)
doc.add_paragraph("Purpose: View and modify a single service request based on its current status.")
items = [
    "Layout: Card-based detail view with back navigation. Header shows Request ID + Status badge.",
    "Info Grid: Two-column grid showing SR Type, Closure Date (row 1), Request Age, Priority (row 2).",
    "Details Section: 'Service Request Details' text block showing the description.",
    "Remarks Section: Textarea (editable in APPROVED/READY_FOR_TELLER/INCOMPLETE, read-only otherwise).",
    "Documents Section: List of uploaded files with download links. Upload button visible in editable statuses.",
    "Verification Notes: (INCOMPLETE only) Yellow-bordered alert showing teller's notes.",
    "Rejection Reason: (REJECTED only) Red-bordered alert showing rejection reason.",
    "Closure Reason: (CLOSED only) Gray-bordered alert showing closure reason.",
    "Action Buttons: Status-dependent — 'Save' + 'Close Request' (APPROVED/READY_FOR_TELLER), 'Save' + 'Re-send for Verification' + 'Close Request' (INCOMPLETE), None (terminal).",
    "Close Request Dialog: AlertDialog with title 'Close Service Request', mandatory textarea 'Enter reason for closing', Cancel + Confirm buttons.",
    "'Back to Service Requests' link at bottom of page.",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("6.2 Back-Office Screens", level=2)

doc.add_heading("6.2.1 Service Request Workbench (/operations/service-requests)", level=3)
doc.add_paragraph("Purpose: Operations dashboard for processing service requests.")
items = [
    "Layout: Full-width page with KPI cards row at top, tab strip below, then content area with table/cards.",
    "KPI Cards: 5 cards in a horizontal row — Total (blue), Approved (green), Ready for Teller (amber), Completed (teal), Overdue SLA (red when > 0). Each card shows count number + label.",
    "Tabs: All | New | Approved | Ready for Teller | Incomplete | Completed | Rejected | Closed.",
    "Search: Text input searching Request ID, Client Name, SR Type.",
    "Table Columns: Request ID, Client, SR Type, Priority, Request Date, Closure Date, Age, Status, Actions.",
    "Actions Column: Context-specific buttons per row based on status — APPROVED: 'Send for Verification'; READY_FOR_TELLER: 'Complete' | 'Incomplete' | 'Reject'.",
    "Send for Verification Dialog: Fields — Service Branch (dropdown, mandatory), Resolution Unit (text), Sales Date (date picker). Submit + Cancel buttons.",
    "Mark Incomplete Dialog: Verification Notes textarea (mandatory, min 10 chars). Submit + Cancel buttons.",
    "Reject Dialog: Rejection Reason textarea (mandatory, min 10 chars). Warning text 'This action cannot be undone.' Submit + Cancel buttons.",
    "View Toggle: Table view / Card view toggle button (desktop defaults to table, mobile to cards).",
    "Row Highlighting: SRs past SLA (now > closure_date and status is open) have amber background tint.",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 7: API & INTEGRATION REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("7. API & Integration Requirements", level=1)

doc.add_heading("7.1 Authentication", level=2)
doc.add_paragraph(
    "All API endpoints require authentication via JWT tokens stored in httpOnly cookies. "
    "Client portal endpoints are scoped to the authenticated client's client_id. "
    "Back-office endpoints require the 'requireBackOfficeRole()' middleware."
)

doc.add_heading("7.2 Standard Error Response Format", level=2)
doc.add_paragraph("All error responses follow this JSON structure:")
doc.add_paragraph('{\n  "error": {\n    "code": "VALIDATION_ERROR",\n    "message": "SR Type is required",\n    "field": "sr_type"\n  }\n}')
doc.add_paragraph("Error codes: VALIDATION_ERROR, NOT_FOUND, FORBIDDEN, INVALID_TRANSITION, INTERNAL_ERROR.")

doc.add_heading("7.3 Client Portal Endpoints", level=2)
add_styled_table(doc,
    ["Method", "Path", "Description", "Auth", "Request Body", "Success Response"],
    [
        ["GET", "/api/v1/client-portal/service-requests/:clientId", "List SRs for a client", "Client (own)", "Query: ?status=&priority=&search=&page=&pageSize=", '{"data": [...], "total": 47, "page": 1, "pageSize": 25}'],
        ["POST", "/api/v1/client-portal/service-requests", "Create new SR", "Client", '{"client_id", "sr_type", "sr_details?", "priority?", "remarks?", "documents?"}', '{"id": 1, "request_id": "SR-2026-000001", ...}'],
        ["GET", "/api/v1/client-portal/service-requests/detail/:id", "Get SR detail", "Client (own)", "None", '{...full SR object with request_age}'],
        ["PUT", "/api/v1/client-portal/service-requests/:id", "Update SR fields", "Client (own)", '{"sr_details?", "remarks?", "documents?", "closure_date?"}', '{...updated SR object}'],
        ["PUT", "/api/v1/client-portal/service-requests/:id/close", "Close SR", "Client (own)", '{"reason": "No longer needed"}', '{...SR with status CLOSED}'],
        ["PUT", "/api/v1/client-portal/service-requests/:id/resubmit", "Resubmit for verification", "Client (own)", '{"remarks?", "documents?"}', '{...SR with status READY_FOR_TELLER}'],
    ])

doc.add_paragraph()
doc.add_heading("7.4 Back-Office Endpoints", level=2)
add_styled_table(doc,
    ["Method", "Path", "Description", "Auth", "Request Body", "Success Response"],
    [
        ["GET", "/api/v1/service-requests", "List all SRs (paginated)", "Back-office", "Query: ?status=&priority=&search=&page=&pageSize=", '{"data": [...], "total": 150, "page": 1, "pageSize": 25}'],
        ["GET", "/api/v1/service-requests/summary", "KPI summary", "Back-office", "None", '{"byStatus": {...}, "overdueSla": 5, "total": 150}'],
        ["GET", "/api/v1/service-requests/:id", "Get SR detail", "Back-office", "None", '{...full SR object with request_age}'],
        ["PUT", "/api/v1/service-requests/:id", "Update SR fields", "Back-office (RM)", '{"sr_details?", "remarks?", "documents?", "closure_date?"}', '{...updated SR object}'],
        ["PUT", "/api/v1/service-requests/:id/send-for-verification", "RM sends for verification", "Back-office (RM)", '{"service_branch", "resolution_unit?", "sales_date?", "assigned_rm_id?"}', '{...SR with status READY_FOR_TELLER}'],
        ["PUT", "/api/v1/service-requests/:id/complete", "Teller completes", "Back-office (Teller)", '{"teller_id": 42}', '{...SR with status COMPLETED}'],
        ["PUT", "/api/v1/service-requests/:id/incomplete", "Teller marks incomplete", "Back-office (Teller)", '{"teller_id": 42, "notes": "Missing KYC docs"}', '{...SR with status INCOMPLETE}'],
        ["PUT", "/api/v1/service-requests/:id/reject", "Teller rejects", "Back-office (Teller)", '{"reason": "Duplicate request"}', '{...SR with status REJECTED}'],
    ])

doc.add_paragraph()
doc.add_heading("7.5 Example Request/Response Bodies", level=2)

doc.add_heading("7.5.1 Create Service Request (POST)", level=3)
doc.add_paragraph("Request:")
doc.add_paragraph(
    '{\n'
    '  "client_id": "CLT-001",\n'
    '  "sr_type": "REVIEW_PORTFOLIO",\n'
    '  "sr_details": "Please review my equity allocation and suggest rebalancing options",\n'
    '  "priority": "HIGH",\n'
    '  "remarks": "Prefer meeting at Makati branch",\n'
    '  "documents": ["portfolio_summary.pdf"]\n'
    '}'
)
doc.add_paragraph("Response (201 Created):")
doc.add_paragraph(
    '{\n'
    '  "id": 1,\n'
    '  "request_id": "SR-2026-000001",\n'
    '  "client_id": "CLT-001",\n'
    '  "sr_type": "REVIEW_PORTFOLIO",\n'
    '  "sr_details": "Please review my equity allocation and suggest rebalancing options",\n'
    '  "priority": "HIGH",\n'
    '  "sr_status": "APPROVED",\n'
    '  "request_date": "2026-04-22T09:30:00.000Z",\n'
    '  "closure_date": "2026-04-25T09:30:00.000Z",\n'
    '  "actual_closure_date": null,\n'
    '  "remarks": "Prefer meeting at Makati branch",\n'
    '  "documents": ["portfolio_summary.pdf"],\n'
    '  "created_by": "CLT-001",\n'
    '  "created_at": "2026-04-22T09:30:00.000Z"\n'
    '}'
)

doc.add_heading("7.5.2 Send for Verification (PUT)", level=3)
doc.add_paragraph("Request:")
doc.add_paragraph(
    '{\n'
    '  "service_branch": "MAKATI-001",\n'
    '  "resolution_unit": "Trust Operations - Equity",\n'
    '  "sales_date": "2026-04-22",\n'
    '  "assigned_rm_id": 15\n'
    '}'
)
doc.add_paragraph("Response (200 OK):")
doc.add_paragraph(
    '{\n'
    '  "id": 1,\n'
    '  "request_id": "SR-2026-000001",\n'
    '  "sr_status": "READY_FOR_TELLER",\n'
    '  "service_branch": "MAKATI-001",\n'
    '  "resolution_unit": "Trust Operations - Equity",\n'
    '  "sales_date": "2026-04-22T00:00:00.000Z",\n'
    '  "assigned_rm_id": 15,\n'
    '  "updated_at": "2026-04-22T10:15:00.000Z"\n'
    '}'
)

doc.add_heading("7.5.3 List with Filters (GET)", level=3)
doc.add_paragraph("Request: GET /api/v1/service-requests?status=APPROVED&priority=HIGH&page=1&pageSize=10")
doc.add_paragraph("Response (200 OK):")
doc.add_paragraph(
    '{\n'
    '  "data": [\n'
    '    {\n'
    '      "id": 1,\n'
    '      "request_id": "SR-2026-000001",\n'
    '      "client_id": "CLT-001",\n'
    '      "sr_type": "REVIEW_PORTFOLIO",\n'
    '      "priority": "HIGH",\n'
    '      "sr_status": "APPROVED",\n'
    '      "request_date": "2026-04-22T09:30:00.000Z",\n'
    '      "closure_date": "2026-04-25T09:30:00.000Z",\n'
    '      "request_age": 0\n'
    '    }\n'
    '  ],\n'
    '  "total": 1,\n'
    '  "page": 1,\n'
    '  "pageSize": 10\n'
    '}'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 8: NON-FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("8. Non-Functional Requirements", level=1)

doc.add_heading("8.1 Performance", level=2)
add_styled_table(doc,
    ["Metric", "Target", "Notes"],
    [
        ["List API response time (P95)", "≤ 500ms", "For up to 10,000 SRs with filters"],
        ["Detail API response time (P95)", "≤ 200ms", "Single record lookup by ID"],
        ["Create API response time (P95)", "≤ 1000ms", "Includes ID generation and SLA computation"],
        ["UI initial page load", "≤ 2 seconds", "Lazy-loaded page components"],
        ["Concurrent users", "100+", "Back-office users simultaneously using workbench"],
    ])

doc.add_heading("8.2 Security", level=2)
items = [
    "Authentication: JWT tokens in httpOnly cookies (existing auth module).",
    "Authorization: Client portal endpoints scoped to authenticated client_id; back-office endpoints protected by requireBackOfficeRole() middleware.",
    "Input Validation: All user inputs sanitized server-side; SQL injection prevented by Drizzle ORM parameterized queries.",
    "CSRF: Protected by SameSite cookie attributes.",
    "XSS: React's default escaping prevents XSS; no dangerouslySetInnerHTML used.",
    "File Upload: Only PDF files accepted; max 10MB; files scanned for malware by upload service.",
    "Audit: Every state change logged with user ID, timestamp, and previous/new values.",
    "Data Encryption: Data at rest encrypted via PostgreSQL TDE; data in transit via HTTPS/TLS 1.2+.",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("8.3 Scalability", level=2)
doc.add_paragraph("The module is designed for a single-tenant deployment serving up to 50,000 service requests per year. Horizontal scaling is supported via the existing Express.js cluster mode. Database indexes on client_id, sr_status, and created_at ensure query performance scales linearly.")

doc.add_heading("8.4 Availability", level=2)
doc.add_paragraph("Target: 99.5% uptime during business hours (Monday–Friday, 8:00 AM – 6:00 PM PHT). Downtime windows for maintenance: Sunday 2:00 AM – 6:00 AM PHT.")

doc.add_heading("8.5 Accessibility", level=2)
doc.add_paragraph("WCAG 2.1 AA compliance. All form controls have labels, color is not the sole indicator (badges have text + color), keyboard navigation supported for all actions, screen reader announcements for status changes and form submissions.")

doc.add_heading("8.6 Browser & Device Support", level=2)
add_styled_table(doc,
    ["Platform", "Browsers", "Minimum Version"],
    [
        ["Desktop", "Chrome, Firefox, Safari, Edge", "Latest 2 major versions"],
        ["Mobile", "Chrome (Android), Safari (iOS)", "Latest 2 major versions"],
        ["Tablet", "Chrome, Safari", "Latest 2 major versions"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 9: WORKFLOW & STATE DIAGRAMS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("9. Workflow & State Diagrams", level=1)

doc.add_heading("9.1 Service Request Lifecycle State Machine", level=2)
doc.add_paragraph("The following table defines all valid state transitions, the actor triggering each transition, and the side effects of each transition.")

add_styled_table(doc,
    ["#", "Current State", "Action", "Next State", "Actor", "Side Effects"],
    [
        ["T1", "—", "Create SR", "NEW → APPROVED", "Client / RM", "Generate request_id; compute closure_date; set request_date; set created_by"],
        ["T2", "APPROVED", "Send for Verification", "READY_FOR_TELLER", "RM", "Set service_branch, resolution_unit, sales_date, assigned_rm_id"],
        ["T3", "APPROVED", "Close Request", "CLOSED", "Client / RM", "Set closure_reason, actual_closure_date"],
        ["T4", "READY_FOR_TELLER", "Complete", "COMPLETED", "Teller", "Set teller_id, actual_closure_date"],
        ["T5", "READY_FOR_TELLER", "Mark Incomplete", "INCOMPLETE", "Teller", "Set teller_id, verification_notes"],
        ["T6", "READY_FOR_TELLER", "Reject", "REJECTED", "Teller", "Set rejection_reason, actual_closure_date"],
        ["T7", "READY_FOR_TELLER", "Close Request", "CLOSED", "Client / RM", "Set closure_reason, actual_closure_date"],
        ["T8", "INCOMPLETE", "Resubmit", "READY_FOR_TELLER", "Client / RM", "Update remarks, documents"],
        ["T9", "INCOMPLETE", "Close Request", "CLOSED", "Client / RM", "Set closure_reason, actual_closure_date"],
    ])

doc.add_paragraph()
doc.add_heading("9.2 State Diagram (Text Representation)", level=2)
diagram = (
    "                    ┌─────────────────────────────────────────────────────┐\n"
    "                    │                                                     │\n"
    "  [CREATE] ──► NEW ──► APPROVED ──┬──► READY_FOR_TELLER ──┬──► COMPLETED │\n"
    "                         │        │          │             │              │\n"
    "                         │        │          │             ├──► INCOMPLETE│\n"
    "                         │        │          │             │      │       │\n"
    "                         │        │          │             │      │ resubmit\n"
    "                         │        │          │             │      └───────┘\n"
    "                         │        │          │             │              │\n"
    "                         │        │          │             └──► REJECTED  │\n"
    "                         │        │          │                            │\n"
    "                         └────────┴──────────┴─────────────────► CLOSED   │\n"
    "                                                                         │\n"
    "                    └─────────────────────────────────────────────────────┘\n"
)
p = doc.add_paragraph(diagram)
p.style.font.name = 'Courier New'
p.style.font.size = Pt(8)

doc.add_heading("9.3 End-to-End Happy Path Workflow", level=2)
steps = [
    "1. Client logs into Client Portal and navigates to Service Requests.",
    "2. Client clicks '+ New Request' and fills the form: SR Type = Review Portfolio, Priority = HIGH, Details = 'Review equity allocation', uploads portfolio_summary.pdf.",
    "3. System creates SR with status APPROVED, generates Request ID SR-2026-000001, computes Closure Date = request_date + 3 days.",
    "4. Success confirmation displayed to client.",
    "5. RM sees SR-2026-000001 in their Back-Office SR Workbench (status: Approved).",
    "6. RM contacts client, collects additional documents, fills Service Branch = MAKATI-001, Resolution Unit = Trust Ops - Equity, Sales Date = today.",
    "7. RM clicks 'Send for Verification' → status transitions to READY_FOR_TELLER.",
    "8. Teller sees SR-2026-000001 in their queue (status: Ready for Teller).",
    "9. Teller reviews documents, verifies request details, clicks 'Complete'.",
    "10. Status transitions to COMPLETED, actual_closure_date set, teller_id recorded.",
    "11. Client sees status updated to 'Completed' in their Service Request list.",
]
for step in steps:
    doc.add_paragraph(step)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 10: NOTIFICATION & COMMUNICATION
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("10. Notification & Communication Requirements", level=1)

doc.add_paragraph("Note: Actual notification delivery (email/SMS/push) is out of scope for Phase 1. This section defines the notification events and templates for future integration with the Notification Service.")

add_styled_table(doc,
    ["#", "Event", "Channel", "Recipient", "Trigger", "Message Template"],
    [
        ["N1", "SR Created", "In-app, Email", "Client", "New SR created and auto-approved", '"Your service request {request_id} has been submitted successfully. Expected closure date: {closure_date}."'],
        ["N2", "SR Sent for Verification", "In-app", "Client", "RM sends SR to teller", '"Your service request {request_id} is now being processed by our operations team."'],
        ["N3", "SR Completed", "In-app, Email", "Client", "Teller completes SR", '"Your service request {request_id} has been completed successfully."'],
        ["N4", "SR Marked Incomplete", "In-app, Email", "Client", "Teller returns SR", '"Your service request {request_id} requires additional information. Please review and resubmit. Notes: {verification_notes}"'],
        ["N5", "SR Rejected", "In-app, Email", "Client", "Teller rejects SR", '"Your service request {request_id} has been rejected. Reason: {rejection_reason}"'],
        ["N6", "SR Closed", "In-app", "Client, RM", "Client or RM closes SR", '"Service request {request_id} has been closed. Reason: {closure_reason}"'],
        ["N7", "SLA Breach Warning", "In-app, Email", "RM, Ops Manager", "SR is 1 day before SLA breach", '"Service request {request_id} will breach SLA in 24 hours. Current status: {sr_status}."'],
        ["N8", "SLA Breached", "In-app, Email", "RM, Ops Manager", "SR passes closure_date while still open", '"Service request {request_id} has breached its SLA. Priority: {priority}. Age: {request_age} days."'],
        ["N9", "SR Resubmitted", "In-app", "Teller", "Client/RM resubmits from INCOMPLETE", '"Service request {request_id} has been resubmitted for verification."'],
    ])

doc.add_paragraph()
doc.add_paragraph("Notification Preferences: Clients can opt out of email notifications for SR updates via the Preferences page. In-app notifications cannot be disabled.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 11: REPORTING & ANALYTICS
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("11. Reporting & Analytics", level=1)

doc.add_heading("11.1 KPI Summary Dashboard (Real-time)", level=2)
add_styled_table(doc,
    ["Metric", "Calculation", "Audience", "Refresh"],
    [
        ["Total SRs", "COUNT(*) WHERE is_deleted = false", "RM, Ops Manager", "On page load"],
        ["By Status (New/Approved/Ready/Completed/Incomplete/Rejected/Closed)", "COUNT(*) GROUP BY sr_status", "RM, Ops Manager", "On page load"],
        ["Overdue SLA", "COUNT(*) WHERE status IN (open statuses) AND now() > closure_date", "Ops Manager", "On page load"],
        ["Average Processing Time", "AVG(actual_closure_date - request_date) for COMPLETED SRs", "Ops Manager", "Daily batch"],
        ["SLA Compliance Rate", "COMPLETED within SLA / total COMPLETED × 100", "Ops Manager", "Daily batch"],
    ])

doc.add_heading("11.2 Future Reports (Phase 2)", level=2)
reports = [
    "SR Volume Trend: Weekly/monthly SR creation volume chart.",
    "SR Type Distribution: Pie chart showing percentage breakdown by SR type.",
    "Branch Performance: Comparison of average processing time and SLA compliance across branches.",
    "RM Workload: Count of open SRs per RM for load balancing.",
    "Aging Analysis: Histogram of open SRs by age bucket (0-3d, 4-7d, 8-14d, 15+d).",
    "Client Satisfaction: Post-completion survey results (when integrated).",
]
for r in reports:
    doc.add_paragraph(r, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 12: MIGRATION & LAUNCH PLAN
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("12. Migration & Launch Plan", level=1)

doc.add_heading("12.1 Data Migration", level=2)
doc.add_paragraph("No legacy data migration is required. This is a greenfield module. Existing service requests tracked in spreadsheets or email can be manually entered post-launch if needed.")

doc.add_heading("12.2 Phased Rollout", level=2)
add_styled_table(doc,
    ["Phase", "Scope", "Description"],
    [
        ["Phase 1 (MVP)", "Schema + Backend + Client Portal + Back-Office Workbench", "Full SR lifecycle with all status transitions, SLA computation, document upload, KPI dashboard. This BRD covers Phase 1."],
        ["Phase 2", "Notifications + SLA Escalation + SR Type Configuration", "Push/email notifications, auto-escalation on SLA breach, dynamic SR type management UI, business-day calendar for SLA."],
        ["Phase 3", "Analytics + Reporting + Integration", "Trend reports, branch performance analytics, CRM integration, client satisfaction surveys."],
    ])

doc.add_heading("12.3 Go-Live Checklist", level=2)
checklist = [
    "Database migration applied (service_requests table, enums, indexes created).",
    "Backend service passes all unit tests (status transition validation, SLA computation, CRUD operations).",
    "API endpoints return correct responses for all happy path and error scenarios.",
    "Client Portal: Create, List, Detail views render correctly on desktop and mobile.",
    "Back-Office Workbench: KPI cards populate, tab filters work, all action dialogs function.",
    "Role-based access verified: clients cannot access back-office endpoints; tellers cannot RM-only actions.",
    "Document upload/download working end-to-end.",
    "Performance: list endpoint returns < 500ms with 1000+ test records.",
    "Security review passed: no SQL injection, XSS, or CSRF vulnerabilities.",
    "Audit trail verified: all state changes logged with correct user ID and timestamp.",
]
for item in checklist:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 13: GLOSSARY
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("13. Glossary", level=1)

add_styled_table(doc,
    ["Term", "Definition"],
    [
        ["SR", "Service Request — a client-initiated request for banking services"],
        ["SLA", "Service Level Agreement — the agreed-upon timeframe for completing a service request"],
        ["RM", "Relationship Manager — the bank officer assigned to manage a client's accounts"],
        ["Teller", "Operations staff member who performs verification and processing of service requests"],
        ["BSP", "Bangko Sentral ng Pilipinas — the central bank and primary financial regulator of the Philippines"],
        ["KPI", "Key Performance Indicator — a measurable metric used to evaluate operational performance"],
        ["Closure Date", "The SLA-computed target date by which a service request should be completed"],
        ["Request Age", "The number of calendar days since a service request was created"],
        ["Soft Delete", "A deletion method that marks records as deleted without physically removing them from the database"],
        ["Dual Control", "A security principle requiring two different people to perform complementary actions (RM prepares, Teller verifies)"],
        ["UITF", "Unit Investment Trust Fund — a pooled trust fund product common in Philippine trust banking"],
        ["TrustOMS", "Trust Banking Operations Management System — the parent platform containing this module"],
        ["Drizzle ORM", "The TypeScript ORM used for database operations in the TrustOMS backend"],
        ["shadcn/ui", "The React component library used for UI elements in the TrustOMS frontend"],
        ["JWT", "JSON Web Token — the authentication token format used for API authorization"],
        ["JSONB", "PostgreSQL's binary JSON data type used for storing document references"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 14: APPENDICES
# ═══════════════════════════════════════════════════════════════════
doc.add_heading("14. Appendices", level=1)

doc.add_heading("14.1 Appendix A: Field Editability Matrix by Status", level=2)
doc.add_paragraph("This matrix defines which fields are editable by the client in the Client Portal, based on the current SR status.")

add_styled_table(doc,
    ["Field", "Initiate (New)", "Approved", "Ready for Teller", "Incomplete", "Completed", "Rejected", "Closed"],
    [
        ["SR Type", "Select (M)", "Display", "Display", "Display", "Display", "Display", "Display"],
        ["SR Details", "Input", "Display", "Display", "Display", "Display", "Display", "Display"],
        ["Priority", "Select (M)", "Display", "Display", "Display", "Display", "Display", "Display"],
        ["Request Date", "Input (M)", "Display", "Display", "Display", "Display", "Display", "Display"],
        ["Closure Date", "Computed", "Computed", "Editable*", "Editable*", "Display", "Display", "Display"],
        ["Appointed Start", "Input", "Display", "Display", "Display", "Display", "Display", "Display"],
        ["Appointed End", "Input", "Display", "Display", "Display", "Display", "Display", "Display"],
        ["Status", "Display", "Display", "Display", "Display", "Display", "Display", "Display"],
        ["Remarks", "Input", "Input", "Input", "Input", "Display", "Display", "Display"],
        ["Documents", "Upload", "Upload", "Display", "Upload", "Display", "Display", "Display"],
    ])
doc.add_paragraph("(M) = Mandatory. * = Editable only if SR Type maintenance has 'Closure Date Change Allowed' = Yes.")

doc.add_paragraph()
doc.add_heading("14.2 Appendix B: Priority Color Coding", level=2)
add_styled_table(doc,
    ["Priority", "Badge Color", "Hex Code"],
    [
        ["HIGH", "Red", "#EF4444"],
        ["MEDIUM", "Amber", "#F59E0B"],
        ["LOW", "Green", "#10B981"],
    ])

doc.add_paragraph()
doc.add_heading("14.3 Appendix C: Status Badge Color Coding", level=2)
add_styled_table(doc,
    ["Status", "Badge Color", "Hex Code"],
    [
        ["NEW", "Blue", "#3B82F6"],
        ["APPROVED", "Green", "#10B981"],
        ["READY_FOR_TELLER", "Amber", "#F59E0B"],
        ["COMPLETED", "Teal", "#14B8A6"],
        ["INCOMPLETE", "Orange", "#F97316"],
        ["REJECTED", "Red", "#EF4444"],
        ["CLOSED", "Gray", "#6B7280"],
    ])

doc.add_paragraph()
doc.add_heading("14.4 Appendix D: Regulatory References", level=2)
add_styled_table(doc,
    ["Regulation", "Relevance", "Requirement"],
    [
        ["BSP Circular 857 (Internal Controls)", "Audit Trail", "All transactions and status changes must be logged with user ID and timestamp"],
        ["BSP Circular 808 (IT Risk Management)", "Data Security", "Client data must be encrypted at rest and in transit; access must be role-based"],
        ["Data Privacy Act of 2012 (RA 10173)", "Data Protection", "Personal data in SRs must be protected; access logs must be maintained"],
        ["BSP Trust Operations Manual", "Dual Control", "Service request processing must follow dual-control verification (RM + Teller)"],
    ])


# ═══════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════
output_path = "/Users/n15318/Trust OMS/docs/ServiceRequest_TaskManagement_BRD_v1.docx"
doc.save(output_path)
print(f"BRD generated successfully: {output_path}")
