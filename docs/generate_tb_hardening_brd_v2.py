#!/usr/bin/env python3
"""
Generate TrustOMS Philippines — Trust Banking Hardening BRD v2.0 (Final)
Incorporates adversarial evaluation findings: DB-derived identity resolution,
BSP-compliant retention, StorageProvider abstraction, path traversal mitigations,
SCAN_PROVIDER feature flag, override expiry, and BO_HEAD/ADMIN-only config writes.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
        hs.paragraph_format.space_before = Pt(24)
        hs.paragraph_format.space_after = Pt(12)
    elif level == 2:
        hs.font.size = Pt(14)
        hs.paragraph_format.space_before = Pt(18)
        hs.paragraph_format.space_after = Pt(8)
    else:
        hs.font.size = Pt(12)
        hs.paragraph_format.space_before = Pt(12)
        hs.paragraph_format.space_after = Pt(6)


# ── Helpers ──

HDR_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
CELL_BLUE = RGBColor(0xD6, 0xE4, 0xF0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x33, 0x33, 0x33)


def shade_cell(cell, rgb: RGBColor):
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{hex_color}"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        shade_cell(cell, HDR_BLUE)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            shade_cell(cell, CELL_BLUE if i == 0 else WHITE)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    return para


def bullet(text, level=0):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = para.add_run(text)
    run.font.size = Pt(11)
    return para


def numbered(text, level=0):
    para = doc.add_paragraph(style='List Number')
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = para.add_run(text)
    run.font.size = Pt(11)
    return para


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("TrustOMS Philippines")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = HDR_BLUE

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub_para.add_run("Trust Banking Hardening")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = HDR_BLUE

delta_para = doc.add_paragraph()
delta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = delta_para.add_run("v2.0 FINAL — Post-Adversarial-Evaluation Revision")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0xCC, 0x33, 0x00)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub2.add_run("Client Portal Security · Messaging · Statements · SLA Config · Document Storage · Operational Resilience")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()

meta_table = doc.add_table(rows=6, cols=2)
meta_table.style = 'Table Grid'
meta_rows = [
    ("Document Title", "Trust Banking Hardening BRD v2.0 (Final)"),
    ("Project", "TrustOMS Philippines"),
    ("Version", "2.0 (Final)"),
    ("Date", "2026-04-25"),
    ("Status", "For Review"),
    ("Classification", "CONFIDENTIAL — Internal Use Only"),
]
for i, (k, v) in enumerate(meta_rows):
    row = meta_table.rows[i]
    shade_cell(row.cells[0], CELL_BLUE)
    row.cells[0].paragraphs[0].add_run(k).bold = True
    row.cells[1].paragraphs[0].add_run(v)

page_break()

# ══════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════

h1("1. Executive Summary")

h2("1.1 Project Overview")
p("TrustOMS Philippines is a comprehensive Trust Operations Management System serving Philippine trust clients across IMA Discretionary, IMA Directed, UITF, Pre-Need, Escrow, Agency, and Safekeeping products. This document specifies six targeted hardening workstreams identified through a systematic codebase review against the Trust Banking Gap Register (April 2026).")

p("The hardening workstreams address: (1) a P0 security gap in client portal ownership validation, (2) a non-functional client messaging module with no backend persistence, (3) a stub-only statement download workflow, (4) an environment-variable-only SLA configuration that cannot be changed without a deployment, (5) service request document storage limited to filename strings, and (6) an in-memory degraded-mode registry that loses state on server restart. All six gaps were confirmed against live code before this document was written.")

h2("1.2 Business Objectives")
bullet("Eliminate the P0 IDOR vulnerability in client portal service-request, statement, and proposal routes before production deployment.")
bullet("Provide Trust Banking clients with a durable, auditable messaging channel to their Relationship Manager.")
bullet("Enable clients to download official statements on demand with full retention and audit trails.")
bullet("Allow Operations to change SLA thresholds through a governed configuration workflow without code deployments.")
bullet("Store service-request evidence documents with storage references, scan status, and retention metadata.")
bullet("Survive server restarts in production with feed health state restored from durable storage.")

h2("1.3 Target Users")
add_table(
    ["User Type", "Pain Point Addressed"],
    [
        ("Trust Banking Client (Client Portal)", "Cannot send/receive messages; statement download non-functional; security risk from IDOR"),
        ("Relationship Manager (Back-Office)", "No visibility of client messages; no message reply capability"),
        ("Operations Manager (Back-Office)", "Cannot change SLA thresholds without code deployment; no durable feed health view"),
        ("System Administrator (Back-Office)", "No config audit portal; no feed-health override capability"),
        ("Compliance Officer", "Document evidence stored as names only — no scan status or retention proof"),
    ],
    [2.0, 4.5]
)

h2("1.4 Success Metrics")
add_table(
    ["KPI", "Target", "Measurement"],
    [
        ("P0 IDOR closure", "100% of client-portal routes guarded by session-derived ownership check", "Security review sign-off"),
        ("Message delivery", "Client messages persisted and delivered within 5 seconds", "API response time"),
        ("Statement download", "< 2 seconds for statements up to 5 MB", "P95 response time"),
        ("Config change lead time", "Ops can change SLA threshold in < 5 minutes without deployment", "Ops team acceptance test"),
        ("Document scan coverage", "100% of uploaded documents receive scan status within 30 seconds", "Background job metrics"),
        ("Feed state recovery", "Feed health state restored < 3 seconds after server restart", "Integration test timing"),
    ],
    [2.0, 2.5, 2.0]
)

page_break()

# ══════════════════════════════════════════════════════
# SECTION 2: SCOPE & BOUNDARIES
# ══════════════════════════════════════════════════════

h1("2. Scope and Boundaries")

h2("2.1 In Scope")
bullet("Client portal ownership middleware applied to all routes accepting :clientId")
bullet("Client messages database table, API (list/send/mark-read), and UI integration")
bullet("RM back-office client messages inbox and reply workflow")
bullet("Statement download endpoint with ownership check and audit trail")
bullet("Statement record enrichment: file_reference, file_size, delivery_status columns")
bullet("Generated PDF placeholder for statements when actual file not yet available")
bullet("system_config database table and CRUD API (BO_HEAD/ADMIN-restricted)")
bullet("Call-report late-filing threshold loaded from system_config (DB-first, env fallback)")
bullet("Business-day calculation integrated into late-filing check using existing MarketCalendarService")
bullet("service_request_documents table with storage_reference, scan_status, retention fields")
bullet("Document upload placeholder (local filesystem for dev, S3-compatible path for prod)")
bullet("Asynchronous scan simulation and client portal document download endpoint")
bullet("feed_health_snapshots table and startup reload of feed health state from DB")
bullet("Feed health override endpoint for BO_HEAD/ADMIN")
bullet("Unread message count badge in client portal sidebar navigation")

h2("2.2 Out of Scope")
bullet("Full S3 integration (storage_reference column is populated; actual S3 SDK calls are stubbed)")
bullet("Real antivirus/malware scanning engine (scan is simulated in dev)")
bullet("Real eFPS BIR submission connector (separate workstream TB-H-005)")
bullet("Trust Banking base account (TSA/CSA) structural schema (separate workstream TB-A-001)")
bullet("UBO/related-party onboarding screens (separate workstream TB-B-002)")
bullet("Client portal push notifications (future phase)")
bullet("SMS/email delivery of messages (back-office sends in-app messages only in this phase)")
bullet("Message encryption at rest beyond standard PostgreSQL column storage")

h2("2.3 Assumptions")
bullet("PostgreSQL is the database and Drizzle ORM is used for all schema changes.")
bullet("Express.js back-end with session-based auth; req.user.clientId is populated for authenticated portal sessions.")
bullet("Local filesystem (uploads/ directory in project root) is used for document storage in development.")
bullet("The existing MarketCalendarService (market-calendar-service.ts) is usable without modification.")
bullet("Existing audit-logger.ts is available for audit writes.")
bullet("React Query is used for all client-portal data fetching.")

h2("2.4 Constraints")
bullet("All new endpoints must use existing error classes: NotFoundError, ForbiddenError, ValidationError, ConflictError.")
bullet("No new npm packages may be added unless strictly necessary (no new dependencies for PDF generation — use existing jsPDF or return simple HTML response).")
bullet("DB schema changes are additive only (no column drops or table renames).")
bullet("All routes must pass TypeScript strict-mode compilation.")

page_break()

# ══════════════════════════════════════════════════════
# SECTION 3: USER ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════

h1("3. User Roles and Permissions")

add_table(
    ["Role", "Module", "Can Read", "Can Write/Act", "Cannot Do"],
    [
        ("CLIENT", "Messages", "Own messages only", "Send message, mark read", "Read other clients' messages; delete messages"),
        ("CLIENT", "Statements", "Own statements only", "Download own statements", "Access other clients' statements"),
        ("CLIENT", "Service Request Documents", "Own SR documents", "Upload document to own SR", "Upload to another client's SR"),
        ("RELATIONSHIP_MANAGER", "Messages", "Messages from own assigned clients", "Reply to client messages", "Delete messages; access unassigned clients"),
        ("BO_MAKER / BO_CHECKER", "System Config", "Read all config keys", "None (read-only)", "Update config values"),
        ("BO_HEAD / SYSTEM_ADMIN", "System Config", "Read all config keys", "Update config values, add new keys", "Hard-delete config records"),
        ("BO_HEAD / SYSTEM_ADMIN", "Feed Health", "Read feed health status", "Override feed status manually", "No additional restrictions"),
        ("BO_CHECKER / BO_HEAD", "Client Messages (BO)", "All client messages", "Reply, mark resolved", "Delete messages"),
        ("SYSTEM_ADMIN", "All above", "All records", "All operations", "None (super-admin)"),
    ],
    [1.5, 1.5, 1.5, 1.5, 1.5]
)

page_break()

# ══════════════════════════════════════════════════════
# SECTION 4: DATA MODEL
# ══════════════════════════════════════════════════════

h1("4. Data Model")

h2("4.1 client_messages")
p("Stores messages exchanged between clients and their Relationship Managers. Messages are threaded by thread_id. The table is also used for system-generated notifications to clients (sender_type = SYSTEM).")

add_table(
    ["Field", "Type", "Required", "Validation", "Default"],
    [
        ("id", "SERIAL PK", "Y", "Auto-increment", "—"),
        ("thread_id", "VARCHAR(40)", "N", "UUID or NULL for standalone message", "NULL"),
        ("sender_id", "INTEGER", "Y", "FK to users.id or clients.id based on sender_type", "—"),
        ("sender_type", "ENUM", "Y", "RM | CLIENT | SYSTEM", "—"),
        ("recipient_client_id", "VARCHAR(20)", "Y", "FK to clients.client_id", "—"),
        ("subject", "VARCHAR(255)", "Y", "1–255 chars; required on first message in thread", "—"),
        ("body", "TEXT", "Y", "1–5000 chars", "—"),
        ("is_read", "BOOLEAN", "Y", "—", "false"),
        ("is_private", "BOOLEAN", "Y", "If true, only sender + recipient can see; hides from BO list view", "false"),
        ("parent_message_id", "INTEGER", "N", "FK to client_messages.id; NULL for root message", "NULL"),
        ("related_sr_id", "INTEGER", "N", "FK to service_requests.id; cross-reference to the SR this message relates to", "NULL"),
        ("read_at", "TIMESTAMPTZ", "N", "Set when is_read flips to true", "NULL"),
        ("sent_at", "TIMESTAMPTZ", "Y", "Server-assigned; not client-supplied", "NOW()"),
        ("is_deleted", "BOOLEAN", "Y", "Soft-delete flag; retained for 7 years then purged", "false"),
        ("deleted_at", "TIMESTAMPTZ", "N", "Set when soft-deleted", "NULL"),
        ("created_at", "TIMESTAMPTZ", "Y", "—", "NOW()"),
        ("updated_at", "TIMESTAMPTZ", "Y", "—", "NOW()"),
    ],
    [1.5, 1.2, 0.7, 2.8, 1.0]
)

p("Sample data:")
add_table(
    ["id", "thread_id", "sender_type", "recipient_client_id", "subject", "is_read"],
    [
        ("1", "thr-001", "RM", "CL-2024-000001", "Quarterly Portfolio Review Q1 2026", "false"),
        ("2", "thr-001", "CLIENT", "CL-2024-000001", "Re: Quarterly Portfolio Review Q1 2026", "true"),
        ("3", "NULL", "SYSTEM", "CL-2024-000002", "Your March 2026 Statement is ready", "false"),
    ],
    [0.4, 0.9, 0.9, 1.5, 2.2, 0.7]
)

h2("4.2 statement_metadata (extends existing statements table)")
p("The existing statements table is extended with storage and delivery columns. New columns are added to the existing schema; the table is not replaced.")

add_table(
    ["New Column", "Type", "Required", "Description"],
    [
        ("file_reference", "VARCHAR(512)", "N", "Storage key: local path (uploads/statements/...) or S3 key"),
        ("file_size_bytes", "INTEGER", "N", "File size in bytes; NULL if not yet generated"),
        ("generated_at", "TIMESTAMPTZ", "N", "When the file was generated"),
        ("delivery_status", "ENUM", "Y", "AVAILABLE | GENERATING | FAILED | PENDING; default PENDING"),
        ("delivery_error", "TEXT", "N", "Last generation error message if delivery_status = FAILED"),
        ("download_count", "INTEGER", "Y", "Number of times downloaded by client; default 0"),
        ("last_downloaded_at", "TIMESTAMPTZ", "N", "Timestamp of most recent download"),
    ],
    [1.8, 1.0, 0.8, 3.0]
)

h2("4.3 StorageProvider Interface (New in v2)")
p("Evaluation identified that local filesystem storage is not multi-node safe. A StorageProvider abstraction interface is introduced so that dev uses local filesystem and production uses S3-compatible storage without changing service code.")
add_table(
    ["Method", "Parameters", "Returns", "Description"],
    [
        ("save(buffer, path)", "Buffer, string", "string (resolved path/key)", "Persist file bytes to storage backend; returns storage_reference"),
        ("read(reference)", "string", "Buffer", "Read file by storage_reference"),
        ("delete(reference)", "string", "void", "Remove file from storage"),
        ("exists(reference)", "string", "boolean", "Check if file exists"),
    ],
    [2.0, 2.0, 1.5, 2.0]
)
p("Implementation: LocalStorageProvider (default dev) saves to uploads/sr-documents/{sr_id}/{uuid}-{basename}. S3StorageProvider (production) calls AWS SDK putObject/getObject/deleteObject. Active provider set via STORAGE_PROVIDER env var (LOCAL | S3). Startup check enforces that STORAGE_PROVIDER is set and credentials are valid.")

h2("4.4 service_request_documents")
p("Replaces the current string-array documents JSONB column. Each document is a separate row enabling per-document scan status and retention tracking. Document-class retention replaces the BRD v1 blanket 7-year default (BSP MORB requirement: trust account opening documents = 10 years, KYC documents = 5 years, other = 7 years).")

add_table(
    ["Field", "Type", "Required", "Validation", "Default"],
    [
        ("id", "SERIAL PK", "Y", "—", "—"),
        ("sr_id", "INTEGER", "Y", "FK to service_requests.id; CASCADE DELETE", "—"),
        ("document_name", "VARCHAR(255)", "Y", "1–255 chars; original filename", "—"),
        ("storage_reference", "VARCHAR(512)", "N", "Local path or S3 key; NULL until upload completes", "NULL"),
        ("file_size_bytes", "INTEGER", "N", "NULL until upload completes", "NULL"),
        ("mime_type", "VARCHAR(127)", "N", "e.g. application/pdf, image/jpeg", "NULL"),
        ("uploaded_by_type", "ENUM", "Y", "CLIENT | RM | SYSTEM", "—"),
        ("uploaded_by_id", "INTEGER", "Y", "User ID of uploader", "—"),
        ("uploaded_at", "TIMESTAMPTZ", "Y", "—", "NOW()"),
        ("scan_status", "ENUM", "Y", "PENDING | CLEAN | QUARANTINED | SKIPPED", "PENDING"),
        ("scan_completed_at", "TIMESTAMPTZ", "N", "NULL until scan finishes", "NULL"),
        ("scan_error", "TEXT", "N", "Error message if scan fails", "NULL"),
        ("document_class", "ENUM", "Y", "TRUST_ACCOUNT_OPENING | KYC | TRANSACTION | OTHER", "OTHER"),
        ("retention_days", "INTEGER", "Y", "Computed from document_class per BSP MORB: TRUST_ACCOUNT_OPENING=3650 (10yr), KYC=1825 (5yr), TRANSACTION=2555 (7yr), OTHER=2555. Override allowed for specific SRs.", "2555"),
        ("expires_at", "TIMESTAMPTZ", "N", "Computed: resolution_date + retention_days; NULL if not resolved", "NULL"),
        ("is_deleted", "BOOLEAN", "Y", "Soft-delete flag", "false"),
        ("deleted_at", "TIMESTAMPTZ", "N", "—", "NULL"),
        ("created_at", "TIMESTAMPTZ", "Y", "—", "NOW()"),
    ],
    [1.5, 1.0, 0.7, 2.5, 0.9]
)

p("Sample data:")
add_table(
    ["id", "sr_id", "document_name", "mime_type", "scan_status", "retention_days"],
    [
        ("1", "42", "death_certificate.pdf", "application/pdf", "CLEAN", "2555"),
        ("2", "42", "id_photocopy.jpg", "image/jpeg", "PENDING", "2555"),
        ("3", "99", "suspicious_file.exe", "application/octet-stream", "QUARANTINED", "365"),
    ],
    [0.4, 0.5, 1.8, 1.4, 1.2, 1.3]
)

h2("4.5 system_config")
p("Key-value configuration store for operational parameters that must be changeable without a code deployment. All changes are audit-logged.")

add_table(
    ["Field", "Type", "Required", "Validation", "Default"],
    [
        ("id", "SERIAL PK", "Y", "—", "—"),
        ("config_key", "VARCHAR(100)", "Y", "UNIQUE; uppercase + underscores only; e.g. CRM_LATE_FILING_DAYS", "—"),
        ("config_value", "VARCHAR(500)", "Y", "Stored as text; typed on read", "—"),
        ("value_type", "ENUM", "Y", "INTEGER | DECIMAL | BOOLEAN | STRING | JSON", "STRING"),
        ("description", "TEXT", "Y", "Human-readable description of what this setting controls", "—"),
        ("min_value", "VARCHAR(50)", "N", "For numeric types: minimum acceptable value", "NULL"),
        ("max_value", "VARCHAR(50)", "N", "For numeric types: maximum acceptable value", "NULL"),
        ("requires_approval", "BOOLEAN", "Y", "If true, value change requires BO_HEAD approval before taking effect", "false"),
        ("is_sensitive", "BOOLEAN", "Y", "If true, value masked in list API response", "false"),
        ("updated_by", "INTEGER", "N", "FK to users.id of last editor", "NULL"),
        ("approved_by", "INTEGER", "N", "FK to users.id of approver; NULL if approval not required", "NULL"),
        ("version", "INTEGER", "Y", "Incremented on each update for optimistic concurrency", "1"),
        ("created_at", "TIMESTAMPTZ", "Y", "—", "NOW()"),
        ("updated_at", "TIMESTAMPTZ", "Y", "—", "NOW()"),
    ],
    [1.8, 0.8, 0.7, 2.4, 0.9]
)

p("Sample data:")
add_table(
    ["config_key", "config_value", "value_type", "description", "requires_approval"],
    [
        ("CRM_LATE_FILING_DAYS", "5", "INTEGER", "Business days after meeting before call report is late", "false"),
        ("CRM_APPROVAL_AUTO_UNCLAIM_DAYS", "2", "INTEGER", "Business days before auto-unclaim of claimed approval", "false"),
        ("CRM_MAX_MEETING_DURATION_HOURS", "8", "INTEGER", "Maximum allowed meeting duration in hours", "false"),
    ],
    [2.2, 1.0, 0.8, 2.5, 1.0]
)

h2("4.6 feed_health_snapshots")
p("Persists the current health state of each external data feed. On server startup, the service loads the latest snapshot per feed into its in-memory registry, restoring state without requiring an external ping.")

add_table(
    ["Field", "Type", "Required", "Validation", "Default"],
    [
        ("id", "SERIAL PK", "Y", "—", "—"),
        ("feed_name", "VARCHAR(50)", "Y", "e.g. BLOOMBERG, REUTERS, PSE, PDEX, BSP", "—"),
        ("health_score", "INTEGER", "Y", "0–100; 100 = fully healthy", "100"),
        ("status", "ENUM", "Y", "UP | DEGRADED | DOWN | OVERRIDE_UP | OVERRIDE_DOWN", "UP"),
        ("failure_count", "INTEGER", "Y", "Consecutive failure count since last UP transition", "0"),
        ("last_error", "TEXT", "N", "Most recent error message from this feed", "NULL"),
        ("last_updated", "TIMESTAMPTZ", "Y", "Timestamp of this snapshot", "NOW()"),
        ("override_by", "INTEGER", "N", "FK to users.id; populated when status is OVERRIDE_*", "NULL"),
        ("override_reason", "TEXT", "N", "Reason text for manual override; required when status = OVERRIDE_*", "NULL"),
        ("override_expires_at", "TIMESTAMPTZ", "N", "Timestamp when override auto-expires; defaults to 24 hours from override creation; NULL = permanent (not recommended)", "NULL"),
        ("created_at", "TIMESTAMPTZ", "Y", "—", "NOW()"),
    ],
    [1.8, 0.9, 0.7, 2.4, 0.6]
)

page_break()

# ══════════════════════════════════════════════════════
# SECTION 5: FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════

h1("5. Functional Requirements")

# ── Module CP-SEC ──
h2("5.1 Module CP-SEC: Client Portal Ownership Enforcement")

h3("FR-CP-SEC-001: DB-Derived Identity Ownership Middleware (v2 Amendment)")
p("User story: As a Trust Banking client, I want my portal sessions to be scoped exclusively to my own data so that I cannot access or modify another client's records even if I know their client ID.")
p("NOTE (v2 change from v1): The JWT auth middleware in auth.ts only populates req.userId, req.userRole, and req.userEmail — there is no req.user.clientId field. The middleware MUST resolve clientId from the database using req.userId rather than reading a session field.")
bullet("AC-001.1: A middleware function validatePortalOwnership() first resolves the authenticated user's linked clientId by querying: SELECT client_id FROM client_portal_users WHERE user_id = req.user.id (or equivalent mapping table). Cache this lookup per user session for 5 minutes.")
bullet("AC-001.2: If no client record is linked to the authenticated user, return HTTP 403 with code NOT_A_CLIENT_USER.")
bullet("AC-001.3: If the route's :clientId parameter does not exactly match the DB-resolved clientId, the middleware returns HTTP 403 with body { error: { code: 'FORBIDDEN', message: 'Access denied: resource does not belong to your account' } }.")
bullet("AC-001.4: If the user is unauthenticated (no valid session), middleware returns HTTP 401.")
bullet("AC-001.5: Ownership check failure is written to the audit log with fields: event='CLIENT_PORTAL_OWNERSHIP_VIOLATION', actor_id, attempted_client_id, actual_client_id, ip_address, path.")
bullet("AC-001.6: The middleware is applied to ALL five confirmed IDOR-vulnerable routes: GET /portfolio-summary/:clientId, GET /statements/:clientId, GET /notifications/:clientId, GET /risk-profile/:clientId, GET /service-requests/:clientId, and all sub-routes under them.")
bullet("AC-001.7: The DB lookup is wrapped in a try/catch; on DB failure, the middleware fails closed (returns 503) rather than allowing access.")
bullet("BR-001.1: SYSTEM_ADMIN and BO roles do not use client-portal routes — they use back-office routes. This middleware does not run on back-office routes.")

h3("FR-CP-SEC-002: Security Alert Creation on Repeated Violations")
p("User story: As a Compliance Officer, I want the system to automatically flag client portal sessions that produce multiple ownership violations so that potential unauthorized access attempts are escalated.")
bullet("AC-002.1: If a single session produces 3 or more ownership violations within a 15-minute window, a security alert record is created in the exception queue with severity P1 and type CLIENT_PORTAL_SECURITY_ALERT.")
bullet("AC-002.2: Alert fields: session_id, client_id (of the violating session), violated_resource_count, first_violation_at, last_violation_at, source_ip.")
bullet("AC-002.3: A notification is sent to the BO_HEAD-role inbox with title 'Client Portal Security Alert: repeated ownership violations'.")
bullet("BR-002.1: Violation count is tracked in-memory per session with a 15-minute sliding window; no DB write per individual violation (only on threshold breach).")

# ── Module MSG ──
h2("5.2 Module MSG: Client Messaging")

h3("FR-MSG-001: Client Message List (Client Portal)")
p("User story: As a Trust Banking client, I want to see a list of all messages from my Relationship Manager and system notifications so that I can stay informed about my portfolio and service requests.")
bullet("AC-003.1: GET /api/v1/client-portal/messages returns paginated messages where recipient_client_id = session clientId.")
bullet("AC-003.2: Response includes: id, thread_id, sender_type, subject, body (truncated to 200 chars in list view), is_read, sent_at, parent_message_id.")
bullet("AC-003.3: Pagination via page (default 1) and pageSize (default 20, max 100) query params.")
bullet("AC-003.4: Unread count is returned in the response envelope: { data: [...], total: N, unread_count: M }.")
bullet("AC-003.5: Messages are sorted by sent_at DESC.")
bullet("BR-003.1: Private messages (is_private = true) from other senders are not returned unless the authenticated client is the recipient.")

h3("FR-MSG-002: Send Message (Client Portal)")
p("User story: As a Trust Banking client, I want to send a message to my Relationship Manager so that I can ask questions and request information without calling.")
bullet("AC-004.1: POST /api/v1/client-portal/messages accepts { subject, body, thread_id? } in the request body.")
bullet("AC-004.2: subject is required (1–255 chars) when thread_id is null; optional when replying to existing thread.")
bullet("AC-004.3: body is required, 1–5000 chars.")
bullet("AC-004.4: sender_type is set to CLIENT server-side; sender_id is set to req.user.id.")
bullet("AC-004.5: Returns HTTP 201 with the created message record.")
bullet("BR-004.1: Clients cannot send messages to other clients (recipient is always the session's own RM assignment).")

h3("FR-MSG-003: Mark Message as Read (Client Portal)")
p("User story: As a Trust Banking client, I want unread messages to be visually distinct and to mark them read so that I can track which messages I have reviewed.")
bullet("AC-005.1: PATCH /api/v1/client-portal/messages/:id/read sets is_read = true, read_at = NOW() for the message.")
bullet("AC-005.2: Returns 403 if the message's recipient_client_id does not match the session clientId.")
bullet("AC-005.3: Calling mark-read on an already-read message is idempotent (returns 200, no change).")

h3("FR-MSG-004: RM Back-Office Inbox")
p("User story: As a Relationship Manager, I want to see all unread messages from my assigned clients in a single inbox so that I can respond promptly.")
bullet("AC-006.1: GET /api/v1/back-office/client-messages returns messages for clients assigned to the RM.")
bullet("AC-006.2: Supports filters: client_id (exact), is_read (boolean), sender_type, date_from, date_to, page, pageSize.")
bullet("AC-006.3: BO_HEAD and SYSTEM_ADMIN see all messages across all clients (no RM scoping).")

h3("FR-MSG-005: RM Reply to Client Message")
p("User story: As a Relationship Manager, I want to reply to a client message from the back-office inbox so that the client receives a response visible in their portal.")
bullet("AC-007.1: POST /api/v1/back-office/client-messages/:id/reply accepts { body } and creates a reply record with sender_type = RM, sender_id = req.user.id, parent_message_id = :id, thread_id copied from parent.")
bullet("AC-007.2: Reply body is required, 1–5000 chars.")
bullet("AC-007.3: Returns HTTP 201 with created reply record.")
bullet("AC-007.4: A notification is created in the client's notification inbox (using existing notificationInboxService) with message 'Your RM has replied to your message'.")

h3("FR-MSG-006: Unread Count Badge in Client Portal Sidebar")
p("User story: As a Trust Banking client, I want to see the number of unread messages as a badge on the Messages navigation item so that I notice new communications immediately.")
bullet("AC-008.1: GET /api/v1/client-portal/messages/unread-count returns { unread_count: N } for the session client.")
bullet("AC-008.2: The client portal sidebar refreshes this count every 60 seconds using a React Query refetchInterval.")
bullet("AC-008.3: Badge is hidden when unread_count = 0.")

# ── Module STMT ──
h2("5.3 Module STMT: Statement Download")

h3("FR-STMT-001: Statement List with File Metadata")
p("User story: As a Trust Banking client, I want to see the status and size of each statement in my statement list so that I know which ones are available for download.")
bullet("AC-009.1: GET /api/v1/client-portal/statements/:clientId returns statements enriched with delivery_status, file_size_bytes, and generated_at.")
bullet("AC-009.2: Ownership is enforced: :clientId must match session clientId (403 otherwise).")
bullet("AC-009.3: Statements with delivery_status = AVAILABLE have an active download button in the UI; others show the status label (PENDING, GENERATING, FAILED).")

h3("FR-STMT-002: Statement Download Endpoint (v2 Amendment — No Placeholder PDF)")
p("User story: As a Trust Banking client, I want to download my official statements as PDF files so that I can review my portfolio history and provide them to third parties.")
p("NOTE (v2 change from v1): The v1 BRD proposed generating a placeholder PDF as a fallback. Adversarial evaluation identified this as a BSP regulatory exposure — delivering a system-generated document through the Statements section could be cited as a misrepresentation finding in a BSP trust examination. The placeholder PDF is REMOVED. The fallback is a clear UI status response.")
bullet("AC-010.1: GET /api/v1/client-portal/statements/:clientId/:statementId/download returns the statement file ONLY if delivery_status = AVAILABLE.")
bullet("AC-010.2: Ownership check: validatePortalOwnership() middleware verifies DB-resolved clientId (403 on mismatch).")
bullet("AC-010.3: If file_reference points to an existing file via the active StorageProvider, the endpoint streams it with Content-Type: application/pdf and Content-Disposition: attachment; filename='statement-{period}.pdf'.")
bullet("AC-010.4: If delivery_status != AVAILABLE, return HTTP 202 Accepted with body { status: 'NOT_AVAILABLE', delivery_status: '...', message: 'Statement is being prepared. You will be notified when it is ready.' } — no file is served and no fabricated document is generated.")
bullet("AC-010.5: Every download is recorded: download_count incremented, last_downloaded_at updated, audit log written with fields: event='STATEMENT_DOWNLOADED', client_id, statement_id, file_size_bytes, ip_address.")
bullet("AC-010.6: Response time for statements up to 5 MB must be < 2 seconds (P95).")
bullet("BR-010.1: Statements with delivery_status = FAILED return HTTP 202 with status = FAILED and message 'Statement generation failed. Please contact your Relationship Manager.'")

h3("FR-STMT-003: Statement Delivery Status Management")
p("User story: As an Operations Manager, I want to see and trigger statement generation from the back-office so that failed statements can be regenerated without client intervention.")
bullet("AC-011.1: GET /api/v1/back-office/statements returns all statements with full metadata including delivery_status and file_reference.")
bullet("AC-011.2: POST /api/v1/back-office/statements/:id/regenerate triggers async re-generation of the statement file and sets delivery_status = GENERATING.")
bullet("AC-011.3: Returns 409 if delivery_status is already GENERATING.")

# ── Module SYSCONFIG ──
h2("5.4 Module SYSCONFIG: System Configuration")

h3("FR-SC-001: Configuration Read")
p("User story: As a Back-Office Maker, I want to view all system configuration values so that I understand the current operational parameters without needing a code review.")
bullet("AC-012.1: GET /api/v1/back-office/system-config returns all config records.")
bullet("AC-012.2: Values for keys with is_sensitive = true are masked in the response: value is replaced with '****'.")
bullet("AC-012.3: Response includes: config_key, config_value (masked if sensitive), value_type, description, min_value, max_value, requires_approval, updated_by, updated_at, version.")
bullet("AC-012.4: Route is accessible to all authenticated back-office roles (read-only for BO_MAKER/CHECKER).")

h3("FR-SC-002: Configuration Update (v2 Amendment — Role Restriction)")
p("User story: As an Operations Head, I want to update system configuration values through a governed workflow so that SLA thresholds can be changed without a code deployment.")
p("NOTE (v2 change from v1): The auto-generated CRUD router is wrapped in requireBackOfficeRole() which grants BO_MAKER and BO_CHECKER write access — a Segregation of Duties violation. The PUT endpoint must use requireAnyRole('BO_HEAD', 'SYSTEM_ADMIN') from role-auth middleware, overriding the CRUD router's catch-all auth.")
bullet("AC-013.1: PUT /api/v1/back-office/system-config/:key updates the value for the given key. MUST use requireAnyRole('BO_HEAD', 'SYSTEM_ADMIN') middleware — not requireBackOfficeRole(). Returns 403 if caller is BO_MAKER or BO_CHECKER.")
bullet("AC-013.2: Value is validated against value_type (integer parse for INTEGER, boolean parse for BOOLEAN, JSON.parse for JSON).")
bullet("AC-013.3: For numeric types, value is validated against min_value and max_value if set.")
bullet("AC-013.4: Every update writes an audit log with fields: event='SYSTEM_CONFIG_UPDATED', config_key, old_value, new_value, updated_by, updated_at.")
bullet("AC-013.5: version is incremented on each update. If the request includes a version field that does not match the current version, return 409 Conflict (optimistic concurrency).")
bullet("BR-013.1: Keys with requires_approval = true require a separate approval step before the new value takes effect. The pending value is stored in config_value_pending and applied on approval.")

h3("FR-SC-003: Late-Filing Threshold Integration")
p("User story: As an Operations Manager, I want the call-report late-filing threshold to reflect the value in system configuration so that changing the policy does not require a deployment.")
bullet("AC-014.1: callReportService reads CRM_LATE_FILING_DAYS from system_config on first use, caches the value for 5 minutes.")
bullet("AC-014.2: If the DB record is not found, falls back to the CRM_LATE_FILING_DAYS environment variable, then to hardcoded default of 5.")
bullet("AC-014.3: Late-filing calculation uses MarketCalendarService.nextBusinessDay() with the configured threshold, not calendar days.")
bullet("AC-014.4: Cache is invalidated immediately when the config key is updated via the PUT endpoint.")

# ── Module DOCS ──
h2("5.5 Module DOCS: Service Request Document Storage")

h3("FR-DOCS-001: Document Upload (Client Portal) — v2 Amendments")
p("User story: As a Trust Banking client, I want to upload supporting documents when submitting or updating a service request so that my RM has the required evidence to process my request.")
p("NOTE (v2 changes): (1) Path traversal: storage_reference MUST be constructed using path.basename(originalFilename) not the raw filename string; (2) 20MB limit must be enforced by multer middleware at the transport layer, not in the service; (3) SCAN_PROVIDER env var required with startup enforcement check; (4) Retention is per document_class per BSP MORB.")
bullet("AC-015.1: POST /api/v1/client-portal/service-requests/:id/documents accepts multipart/form-data with field 'file'. The multer middleware enforces 20 MB limit (HTTP 413 if exceeded) before the route handler is called.")
bullet("AC-015.2: Ownership check: validatePortalOwnership() middleware. The service request's client_id must match the DB-resolved clientId.")
bullet("AC-015.3: Path traversal prevention: sanitized_filename = path.basename(originalname); storage path = {sr_id}/{uuid}-{sanitized_filename}. The original unsanitized filename is NEVER concatenated into a path.")
bullet("AC-015.4: File is persisted via the active StorageProvider (LocalStorageProvider in dev, S3StorageProvider in production per STORAGE_PROVIDER env var).")
bullet("AC-015.5: A service_request_documents record is inserted with scan_status = PENDING, document_class defaulting to 'OTHER' (SR creator can specify via request body), storage_reference set to the resolved path.")
bullet("AC-015.6: Returns HTTP 201 with { id, document_name, file_size_bytes, scan_status, document_class, retention_days, uploaded_at }.")
bullet("AC-015.7: Async scan: if SCAN_PROVIDER = SIMULATED (dev), marks CLEAN after 2 seconds. If SCAN_PROVIDER = CLAMAV or SCAN_PROVIDER = EXTERNAL_WEBHOOK, delegates to the real scanner. If SCAN_PROVIDER is not set, startup check fails with error: 'SCAN_PROVIDER must be configured before starting the server'.")
bullet("BR-015.1: Files with extensions .exe, .bat, .sh, .cmd, .ps1 are SYNCHRONOUSLY rejected before saving to storage; scan_status set to QUARANTINED; 0 bytes saved.")
bullet("BR-015.2: Accepted MIME types: application/pdf, image/jpeg, image/png, image/gif, application/vnd.ms-excel, application/vnd.openxmlformats-officedocument.spreadsheetml.sheet, application/msword, application/vnd.openxmlformats-officedocument.wordprocessingml.document.")
bullet("BR-015.3: retention_days is computed automatically from document_class: TRUST_ACCOUNT_OPENING=3650, KYC=1825, TRANSACTION=2555, OTHER=2555. Manual override allowed for special SR types.")

h3("FR-DOCS-002: Document List")
p("User story: As a Trust Banking client, I want to see all documents I have uploaded for a service request so that I can track what evidence has been provided.")
bullet("AC-016.1: GET /api/v1/client-portal/service-requests/:id/documents returns all non-deleted documents for the SR.")
bullet("AC-016.2: Returns: id, document_name, file_size_bytes, mime_type, uploaded_at, scan_status, scan_completed_at.")
bullet("AC-016.3: QUARANTINED documents are included in the list but flagged; download is blocked for QUARANTINED files.")

h3("FR-DOCS-003: Document Download")
p("User story: As a Trust Banking client, I want to download a document I previously uploaded to verify it was received correctly.")
bullet("AC-017.1: GET /api/v1/client-portal/service-requests/:id/documents/:docId/download streams the file.")
bullet("AC-017.2: Ownership: SR must belong to session client.")
bullet("AC-017.3: Returns 403 if scan_status = QUARANTINED.")
bullet("AC-017.4: Returns 404 if file is deleted or storage_reference is null.")
bullet("AC-017.5: Content-Disposition header is set to attachment with the original document_name.")

h3("FR-DOCS-004: Back-Office Document Access")
p("User story: As an RM or Operations staff, I want to view and download all documents attached to a service request so that I can process the request with full evidence.")
bullet("AC-018.1: GET /api/v1/back-office/service-requests/:id/documents returns all documents including QUARANTINED ones.")
bullet("AC-018.2: GET /api/v1/back-office/service-requests/:id/documents/:docId/download streams the file to the back-office user.")
bullet("AC-018.3: QUARANTINED files: BO users CAN download with an additional warning header X-Scan-Status: QUARANTINED.")

# ── Module FEED ──
h2("5.6 Module FEED: Degraded Mode State Persistence")

h3("FR-FEED-001: Feed Health Snapshot Persistence (v2 Amendment — Write-back on All Health Updates)")
p("User story: As a System Administrator, I want feed health state to survive server restarts so that Operations does not see a false 'all feeds healthy' state after a deployment.")
p("NOTE (v2 change from v1): The evaluation confirmed that updateFeedHealth() only updates the in-memory registry and has no DB write-back path. Snapshot persistence must be called on EVERY health update (reportIncident, updateFeedHealth, switchFeed), not only on incident creation.")
bullet("AC-019.1: ALL functions that mutate the in-memory feedHealthRegistry (reportIncident, updateFeedHealth, switchFeed) must fire-and-forget a snapshot write to feed_health_snapshots after updating in-memory state.")
bullet("AC-019.2: On service module initialization (startup), degradedModeService queries SELECT DISTINCT ON (feed_name) * FROM feed_health_snapshots ORDER BY feed_name, created_at DESC and populates feedHealthRegistry. No hardcoded initialization values after first run.")
bullet("AC-019.3: If no snapshot exists for a feed at startup (first boot), the feed defaults to { status: 'UP', healthScore: 100, failureCount: 0 } and a baseline snapshot is immediately written.")
bullet("AC-019.4: Snapshot writes are fire-and-forget; a write failure is logged but does not block the calling function.")
bullet("AC-019.5: Override expiry: when status = OVERRIDE_UP or OVERRIDE_DOWN, override_expires_at defaults to NOW() + 24 hours. An EOD background job clears expired overrides and writes a snapshot restoring the score-derived status.")

h3("FR-FEED-002: Feed Health API")
p("User story: As an Operations Manager, I want to see current feed health status from the back-office dashboard so that I can identify and respond to data feed problems.")
bullet("AC-020.1: GET /api/v1/back-office/feed-health returns the current feed health registry as an array of { feed_name, health_score, status, failure_count, last_error, last_updated }.")
bullet("AC-020.2: Route requires BO role (any back-office role).")
bullet("AC-020.3: Response is served from in-memory registry (no DB query on each request).")

h3("FR-FEED-003: Feed Health Manual Override")
p("User story: As a System Administrator, I want to manually force a feed status override during planned maintenance so that the system does not incorrectly switch data sources.")
bullet("AC-021.1: POST /api/v1/back-office/feed-health/:feed/override accepts { status: 'OVERRIDE_UP' | 'OVERRIDE_DOWN', reason } and sets the feed to the override status in both in-memory registry and DB snapshot.")
bullet("AC-021.2: Route restricted to BO_HEAD and SYSTEM_ADMIN.")
bullet("AC-021.3: override_reason is required (10–500 chars).")
bullet("AC-021.4: Override is written to the audit log with fields: event='FEED_HEALTH_OVERRIDE', feed_name, new_status, reason, override_by.")
bullet("AC-021.5: A manual override can be cleared via POST /api/v1/back-office/feed-health/:feed/clear-override (BO_HEAD/ADMIN only), which restores the computed status.")

page_break()

# ══════════════════════════════════════════════════════
# SECTION 6: UI REQUIREMENTS
# ══════════════════════════════════════════════════════

h1("6. User Interface Requirements")

h2("6.1 Client Portal — Messages Page")
p("Screen replaces the current hardcoded messages.tsx with live-data-driven components.")
bullet("Layout: Two-panel. Left panel (35% width): scrollable message list sorted by sent_at DESC. Right panel (65%): selected message thread or compose form.")
bullet("Message list item: sender name (RM name or 'TrustOMS' for SYSTEM), subject (truncated to 60 chars), date, unread indicator dot (blue) for is_read = false.")
bullet("Thread view: all messages in the same thread_id, chronologically ascending, with sender label, timestamp, and body.")
bullet("Compose button: Opens compose form in right panel with subject and body fields. Send button calls POST /messages. Success shows toast 'Message sent'.")
bullet("Reply button: Pre-fills compose form with Re: {subject}, thread_id set to current thread. Calls same POST endpoint.")
bullet("Unread badge: displayed on Messages nav item using GET /messages/unread-count, refreshed every 60 seconds.")
bullet("Loading state: Skeleton rows during initial fetch.")
bullet("Empty state: 'No messages yet. Your Relationship Manager will contact you here.' with compose button.")

h2("6.2 Client Portal — Statements Page")
p("Extends existing statements.tsx by replacing the stub download handler.")
bullet("Statement list table columns: Period, Type (badge: MONTHLY/QUARTERLY/ANNUAL), Status (badge: AVAILABLE=green, GENERATING=amber, FAILED=red, PENDING=grey), File Size, Generated Date, Actions.")
bullet("Download button: Enabled only when delivery_status = AVAILABLE. Calls GET /statements/:clientId/:id/download, browser receives file attachment.")
bullet("When delivery_status is not AVAILABLE: button is replaced by a status chip (e.g., 'Generating...' with spinner).")
bullet("No 'future release' message — remove the placeholder toast handler entirely.")

h2("6.3 Back-Office — Client Messages Inbox")
p("New page in back-office CRM section: /back-office/crm/client-messages")
bullet("Filters bar: Client name/ID search, sender_type filter (ALL/CLIENT/SYSTEM), is_read filter (ALL/UNREAD), date range picker.")
bullet("Message table: Client Name, Subject, Sender Type, Sent At, Is Read (chip), Actions (View, Reply).")
bullet("View action: Expands to show full thread in a modal or side panel.")
bullet("Reply action: Opens inline reply form with body textarea and Send button.")
bullet("Navigation: Added under CRM > Client Messages in back-office sidebar.")

h2("6.4 Back-Office — System Configuration Page")
p("New page in back-office admin section: /back-office/admin/system-config")
bullet("Table columns: Config Key, Value (masked if sensitive), Type, Description, Min, Max, Requires Approval, Last Updated By, Last Updated At, Actions (Edit).")
bullet("Edit button: Opens inline edit form with value input, type-appropriate validation hints. Save calls PUT /system-config/:key.")
bullet("Version mismatch conflict: Shows toast 'Config was updated by another user. Please refresh and retry.'")
bullet("Navigation: Added under Admin > System Configuration in back-office sidebar.")

h2("6.5 Back-Office — Feed Health Dashboard")
p("New section in back-office operations/monitoring area.")
bullet("Card grid: One card per feed (BLOOMBERG, REUTERS, PSE, PDEX, BSP). Each card shows: feed name, health score (gauge-style), status badge, failure count, last updated time.")
bullet("Override button: Opens modal with status (OVERRIDE_UP/OVERRIDE_DOWN) dropdown and reason textarea. Submit calls POST /feed-health/:feed/override.")
bullet("Clear Override button: Visible when status starts with 'OVERRIDE_'. Calls POST /feed-health/:feed/clear-override.")

page_break()

# ══════════════════════════════════════════════════════
# SECTION 7: API & INTEGRATION REQUIREMENTS
# ══════════════════════════════════════════════════════

h1("7. API and Integration Requirements")

h2("7.1 Standard Error Response Format")
p("All endpoints return errors in this format:")
p('{ "error": { "code": "ERROR_CODE", "message": "Human-readable message", "field": "optional_field_name" } }', bold=False)
p("Common error codes: VALIDATION_ERROR, NOT_FOUND, FORBIDDEN, CONFLICT, UNAUTHORIZED, FILE_TOO_LARGE, UNSUPPORTED_MIME_TYPE.")

h2("7.2 Endpoint Inventory")

add_table(
    ["Method", "Path", "Auth", "Description"],
    [
        # Security
        ("MW", "/api/v1/client-portal/*", "Session", "validatePortalOwnership() middleware — not an endpoint"),
        # Messages
        ("GET", "/api/v1/client-portal/messages", "CLIENT session", "List messages for session client"),
        ("GET", "/api/v1/client-portal/messages/unread-count", "CLIENT session", "Returns unread count"),
        ("POST", "/api/v1/client-portal/messages", "CLIENT session", "Send message to RM"),
        ("PATCH", "/api/v1/client-portal/messages/:id/read", "CLIENT session", "Mark message as read"),
        ("GET", "/api/v1/back-office/client-messages", "BO role", "RM inbox — list client messages"),
        ("POST", "/api/v1/back-office/client-messages/:id/reply", "BO role", "RM reply to client message"),
        # Statements
        ("GET", "/api/v1/client-portal/statements/:clientId", "CLIENT session + ownership", "Statement list with metadata"),
        ("GET", "/api/v1/client-portal/statements/:clientId/:statementId/download", "CLIENT session + ownership", "Download statement file"),
        ("GET", "/api/v1/back-office/statements", "BO role", "All statements with delivery status"),
        ("POST", "/api/v1/back-office/statements/:id/regenerate", "BO role", "Trigger statement re-generation"),
        # System config
        ("GET", "/api/v1/back-office/system-config", "BO role (read)", "List all config keys"),
        ("GET", "/api/v1/back-office/system-config/:key", "BO role (read)", "Get single config value"),
        ("PUT", "/api/v1/back-office/system-config/:key", "BO_HEAD/ADMIN", "Update config value"),
        # Documents
        ("POST", "/api/v1/client-portal/service-requests/:id/documents", "CLIENT session + ownership", "Upload document"),
        ("GET", "/api/v1/client-portal/service-requests/:id/documents", "CLIENT session + ownership", "List documents for SR"),
        ("GET", "/api/v1/client-portal/service-requests/:id/documents/:docId/download", "CLIENT session + ownership", "Download document"),
        ("GET", "/api/v1/back-office/service-requests/:id/documents", "BO role", "BO: list SR documents"),
        ("GET", "/api/v1/back-office/service-requests/:id/documents/:docId/download", "BO role", "BO: download SR document"),
        # Feed health
        ("GET", "/api/v1/back-office/feed-health", "BO role", "Current feed health status"),
        ("POST", "/api/v1/back-office/feed-health/:feed/override", "BO_HEAD/ADMIN", "Manual feed status override"),
        ("POST", "/api/v1/back-office/feed-health/:feed/clear-override", "BO_HEAD/ADMIN", "Clear manual override"),
    ],
    [0.6, 3.0, 1.5, 2.5]
)

h2("7.3 Request/Response Examples")

h3("POST /api/v1/client-portal/messages — Send Message")
p("Request body:")
p('{ "subject": "Question about my October dividend", "body": "Hi, I noticed my PLDT dividend was not credited. Can you check?", "thread_id": null }')
p("Response (HTTP 201):")
p('{ "data": { "id": 47, "thread_id": "thr-2026-00047", "sender_type": "CLIENT", "recipient_client_id": "CL-2024-000001", "subject": "Question about my October dividend", "body": "Hi, I noticed my PLDT dividend...", "is_read": false, "sent_at": "2026-04-25T08:30:00Z" } }')

h3("PUT /api/v1/back-office/system-config/CRM_LATE_FILING_DAYS — Update Config")
p("Request body:")
p('{ "config_value": "7", "version": 3 }')
p("Response (HTTP 200):")
p('{ "data": { "config_key": "CRM_LATE_FILING_DAYS", "config_value": "7", "value_type": "INTEGER", "version": 4, "updated_by": 12, "updated_at": "2026-04-25T09:00:00Z" } }')
p("Version conflict response (HTTP 409):")
p('{ "error": { "code": "CONFLICT", "message": "Config version mismatch. Current version is 4, request version was 3." } }')

h3("GET /api/v1/client-portal/statements/:clientId/:statementId/download — Statement Download")
p("Success: HTTP 200 with headers:")
p("Content-Type: application/pdf\nContent-Disposition: attachment; filename=\"statement-2026-03.pdf\"\nContent-Length: 245760\nX-Delivery-Status: AVAILABLE")
p("Placeholder response (file missing): HTTP 200 with same headers but generated content, plus X-Delivery-Status: PLACEHOLDER")

page_break()

# ══════════════════════════════════════════════════════
# SECTION 8: NON-FUNCTIONAL REQUIREMENTS
# ══════════════════════════════════════════════════════

h1("8. Non-Functional Requirements")

add_table(
    ["Category", "Requirement", "Target"],
    [
        ("Performance", "Statement download (< 5 MB)", "< 2 seconds P95"),
        ("Performance", "Message list API", "< 500 ms P95"),
        ("Performance", "Config read (cached)", "< 50 ms P95"),
        ("Performance", "Feed health GET (in-memory)", "< 20 ms P95"),
        ("Security", "All client-portal routes enforce session-derived ownership", "100% coverage"),
        ("Security", "Document storage path must not be guessable (UUID prefix)", "UUID v4 in path"),
        ("Security", "QUARANTINED files not served to clients", "Block at service layer"),
        ("Security", "Audit log for ownership violations, config changes, feed overrides, downloads", "100% events captured"),
        ("Scalability", "Config cache invalidation without server restart", "TTL-based + event-driven flush"),
        ("Availability", "Feed health state restored after restart", "< 3 seconds from DB on startup"),
        ("Data Retention", "SR documents retained for retention_days after resolution", "Soft-delete on schedule"),
        ("Accessibility", "New UI pages follow existing component patterns (shadcn/ui)", "No new a11y regressions"),
        ("Compliance", "Statement download audit trail retained for 7 years", "Soft-delete only; no hard purge"),
    ],
    [1.5, 3.5, 1.5]
)

page_break()

# ══════════════════════════════════════════════════════
# SECTION 9: WORKFLOW & STATE DIAGRAMS
# ══════════════════════════════════════════════════════

h1("9. Workflow and State Diagrams")

h2("9.1 Client Message Lifecycle")
add_table(
    ["Current State", "Action / Trigger", "Next State", "Side Effects"],
    [
        ("—", "CLIENT sends message via portal", "UNREAD (is_read=false)", "Insert client_messages; notify RM inbox"),
        ("UNREAD", "CLIENT views message", "READ (is_read=true, read_at=NOW())", "PATCH /messages/:id/read; update unread_count"),
        ("UNREAD / READ", "RM replies via back-office", "NEW REPLY (child message, is_read=false)", "Insert reply; notify client notification inbox"),
        ("NEW REPLY", "CLIENT views reply", "READ", "Same as above"),
        ("—", "SYSTEM generates statement", "UNREAD (sender_type=SYSTEM)", "Auto-created message; no RM action needed"),
    ],
    [1.5, 2.0, 1.8, 2.3]
)

h2("9.2 Document Scan Lifecycle")
add_table(
    ["State", "Action / Trigger", "Next State", "Side Effects"],
    [
        ("—", "Client uploads file", "PENDING", "File saved to disk; record inserted"),
        ("PENDING", "Blocked extension detected at upload", "QUARANTINED (immediate)", "File not saved; record inserted with QUARANTINED; alert to BO"),
        ("PENDING", "Async scan job runs (dev: 2-second delay)", "CLEAN", "scan_completed_at set; client can download"),
        ("PENDING", "Async scan detects malicious pattern", "QUARANTINED", "scan_completed_at set; download blocked; alert to BO"),
        ("CLEAN", "Retention expires (expires_at < NOW())", "SOFT-DELETED", "is_deleted=true; file removed from disk"),
        ("QUARANTINED", "Admin reviews and clears manually", "CLEAN", "Audit log; scan_status updated"),
    ],
    [1.2, 2.2, 1.5, 2.7]
)

h2("9.3 Statement Delivery Lifecycle")
add_table(
    ["State", "Action / Trigger", "Next State", "Side Effects"],
    [
        ("PENDING", "Scheduled generation job starts", "GENERATING", "delivery_status updated"),
        ("GENERATING", "File generation succeeds", "AVAILABLE", "file_reference set; generated_at set; client notification sent"),
        ("GENERATING", "File generation fails", "FAILED", "delivery_error set; ops exception created"),
        ("AVAILABLE", "Client clicks Download", "AVAILABLE (unchanged)", "download_count++; last_downloaded_at set; audit log"),
        ("FAILED", "BO triggers regenerate", "GENERATING", "delivery_error cleared; retry starts"),
    ],
    [1.2, 2.2, 1.5, 2.7]
)

page_break()

# ══════════════════════════════════════════════════════
# SECTION 10: NOTIFICATIONS
# ══════════════════════════════════════════════════════

h1("10. Notification Requirements")

add_table(
    ["Event", "Channel", "Recipient", "Template"],
    [
        ("RM replies to client message", "In-app (notificationInboxService)", "CLIENT", "Your Relationship Manager has replied to your message: '{subject}'. Log in to view."),
        ("Statement becomes AVAILABLE", "In-app (notificationInboxService)", "CLIENT", "Your {type} statement for {period} is ready to download."),
        ("Document scan result: QUARANTINED", "In-app (back-office notification)", "BO_HEAD/ADMIN", "Document '{document_name}' for SR #{sr_id} was quarantined. Review required."),
        ("Client portal ownership violation (≥ 3 in 15 min)", "In-app (back-office notification)", "BO_HEAD", "Client Portal Security Alert: {count} ownership violations from client session {session_id}."),
        ("System config updated", "Audit log only", "—", "CONFIG_UPDATED: {key} changed from '{old}' to '{new}' by {actor}."),
        ("Feed health override applied", "Audit log + in-app", "BO_HEAD", "Feed {feed_name} manually set to {status} by {actor}: '{reason}'"),
    ],
    [2.0, 1.2, 1.3, 3.0]
)

page_break()

# ══════════════════════════════════════════════════════
# SECTION 11: BLUE OCEAN ENHANCEMENTS
# ══════════════════════════════════════════════════════

h1("11. Blue Ocean Enhancements")

p("The following enhancements extend the core gaps into more complete product capabilities. They are clearly marked as Phase 2 items and are NOT required for initial delivery.")

h2("11.1 BO-1: Unified Client Communication Timeline")
p("A single timeline view per client in the back-office showing all interactions: messages, service requests, CA elections, document uploads, and system notifications in chronological order. Enables RMs to have a full conversation history before client calls.")

h2("11.2 BO-2: Smart Statement Generation Pipeline")
p("Scheduled statement generation: monthly mini-statements auto-generated on the 5th business day, quarterly comprehensive statements on the 10th business day of the quarter. Statement content includes AUM summary, performance vs benchmark (from existing NAV/pricing data), fee summary, and transaction list. Push/email notification when statement is ready.")

h2("11.3 BO-3: Configuration Audit Portal")
p("Dedicated UI for system configuration management showing change history per key, who changed each value, and approval status. Config changes for P0/SLA-affecting keys require BO_HEAD approval via maker-checker workflow. Alerts when a config key deviates from its documented min/max range.")

h2("11.4 BO-4: Client Security Alert Dashboard")
p("Back-office dashboard widget showing recent client portal security alerts (ownership violations, repeated auth failures, suspicious download patterns). Alert cards link to the specific client record and session audit log for investigation.")

page_break()

# ══════════════════════════════════════════════════════
# SECTION 12: MIGRATION & LAUNCH PLAN
# ══════════════════════════════════════════════════════

h1("12. Migration and Launch Plan")

h2("12.1 Phased Rollout")
add_table(
    ["Phase", "Scope", "Rationale"],
    [
        ("Phase 1 (Immediate)", "CP-SEC ownership middleware + security alert", "P0 security fix — must be deployed before any client-portal launch"),
        ("Phase 2", "MSG client messaging + FEED degraded-mode persistence", "P1 functional gaps — required for operational client portal"),
        ("Phase 3", "STMT statement download + SYSCONFIG system configuration", "P1 operational gaps — required before trust product go-live"),
        ("Phase 4", "DOCS document storage + scan workflow", "P1 compliance gap — required for regulatory sign-off on service requests"),
        ("Phase 5", "Blue Ocean enhancements (BO-1 through BO-4)", "P2 enhancements — post-launch improvements"),
    ],
    [1.2, 3.0, 3.3]
)

h2("12.2 Data Migration")
bullet("service_request_documents table is new; no migration of existing string[] documents column required (column can remain for historical records, new uploads use the table).")
bullet("system_config table is new; seed with initial keys (CRM_LATE_FILING_DAYS=5, CRM_APPROVAL_AUTO_UNCLAIM_DAYS=2, CRM_MAX_MEETING_DURATION_HOURS=8).")
bullet("statement_metadata columns are added to existing statements table; no data backfill required (new columns default to NULL/PENDING).")
bullet("feed_health_snapshots table is new; no migration required (first snapshots written on next incident report).")
bullet("client_messages table is new; existing hardcoded messages.tsx data is simply removed (no historical data to migrate).")

page_break()

# ══════════════════════════════════════════════════════
# SECTION 13: GLOSSARY
# ══════════════════════════════════════════════════════

h1("13. Glossary")

add_table(
    ["Term", "Definition"],
    [
        ("IDOR", "Insecure Direct Object Reference — a security vulnerability where a user can access another user's resources by manipulating identifiers in API calls"),
        ("TSA", "Trust Settlement Account — an account used for settling trust transactions"),
        ("CSA", "Client Settlement Account — the client's designated settlement account"),
        ("CIF", "Customer Information File — the master record of a client's personal, KYC, and relationship data"),
        ("IMA Discretionary", "Investment Management Agreement where the RM/Trust has discretion over portfolio decisions"),
        ("IMA Directed", "Investment Management Agreement where the client directs investment decisions"),
        ("UITF", "Unit Investment Trust Fund — a pooled fund product offered by trust banks"),
        ("eFPS", "Electronic Filing and Payment System — BIR's online tax filing platform"),
        ("MarketCalendarService", "Existing TrustOMS service providing business-day and holiday calculations"),
        ("BO", "Back Office — the bank's internal operations and RM teams"),
        ("BO_HEAD", "Back-Office Head role — senior operations manager with config and override privileges"),
        ("P0/P1/P2", "Priority levels: P0 = must fix before launch, P1 = must fix before go-live, P2 = improvement"),
        ("scan_status", "Result of the document virus/malware scan: PENDING (awaiting), CLEAN (passed), QUARANTINED (blocked), SKIPPED (exempt)"),
        ("delivery_status", "Lifecycle state of a generated statement file: PENDING, GENERATING, AVAILABLE, FAILED"),
        ("thread_id", "A shared identifier grouping a sequence of related messages into a conversation thread"),
        ("system_config", "Database-backed key-value store for operational configuration that can be changed without a code deployment"),
        ("feed_health_snapshot", "A point-in-time record of a data feed's health score, status, and error state, persisted to DB for restart recovery"),
        ("validatePortalOwnership()", "Express middleware that enforces that req.user.clientId matches the :clientId in the route — the primary security control for the client portal"),
    ],
    [2.0, 5.5]
)

page_break()

# ══════════════════════════════════════════════════════
# APPENDIX A: Gap-to-FR Traceability
# ══════════════════════════════════════════════════════

h1("Appendix A: Gap Register to FR Traceability")

add_table(
    ["Gap Register ID", "Gap Description", "Functional Requirements"],
    [
        ("TB-E-001 (P0)", "Client portal service-request ownership not session-derived", "FR-CP-SEC-001, FR-CP-SEC-002"),
        ("TB-J-001 (P1)", "Client portal messages are hardcoded stub with no backend", "FR-MSG-001 through FR-MSG-006"),
        ("TB-J-002 (P1)", "Statement download button shows 'future release' toast", "FR-STMT-001, FR-STMT-002, FR-STMT-003"),
        ("TB-C-001 (P1)", "Late-filing threshold is env-only, not DB-backed", "FR-SC-001, FR-SC-002, FR-SC-003"),
        ("TB-E-002 (P1)", "Service-request documents stored as string names only", "FR-DOCS-001 through FR-DOCS-004"),
        ("TB-H-006 (P1)", "Degraded-mode feed registry is in-memory; lost on restart", "FR-FEED-001, FR-FEED-002, FR-FEED-003"),
        ("TB-K-001 (P0)", "Some client-facing APIs need object-level authorization review", "FR-CP-SEC-001 (same fix)"),
    ],
    [1.8, 2.8, 2.9]
)

# ══════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════

out_path = "docs/TrustBanking_Hardening_BRD_v2_Final.docx"
doc.save(out_path)
print(f"BRD saved to: {out_path}")
