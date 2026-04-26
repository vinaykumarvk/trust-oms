#!/usr/bin/env python3
"""Generate Calendar & Call Report Management BRD - Part 1: Setup + Sections 1-3"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Arial'
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1B3A5C')
        shading.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
        if r_idx % 2 == 1:
            for c_idx in range(len(headers)):
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'E8F0FE')
                shading.set(qn('w:val'), 'clear')
                table.rows[r_idx + 1].cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()
    return table

# --- TITLE PAGE ---
for _ in range(6):
    doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Business Requirements Document')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Calendar & Call Report Management Module\nTrustOMS Wealth Management Platform')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run('Version 1.0 | April 2026 | CONFIDENTIAL')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# --- TABLE OF CONTENTS placeholder ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Scope & Boundaries',
    '3. User Roles & Permissions',
    '4. Data Model',
    '5. Functional Requirements',
    '6. User Interface Requirements',
    '7. API & Integration Requirements',
    '8. Non-Functional Requirements',
    '9. Workflow & State Diagrams',
    '10. Notification & Communication Requirements',
    '11. Reporting & Analytics',
    '12. Migration & Launch Plan',
    '13. Glossary',
    '14. Appendices',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
doc.add_page_break()

# ============================================================
# SECTION 1: EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_heading('1.1 Project Name', level=2)
doc.add_paragraph('TrustOMS CRM Calendar & Interaction Management Module (CIM)')

doc.add_heading('1.2 Project Description', level=2)
doc.add_paragraph(
    'The Calendar & Interaction Management (CIM) module extends the TrustOMS wealth management platform '
    'with comprehensive calendar scheduling, call report filing, task management, opportunity tracking, '
    'expense capture, and campaign management capabilities. Designed for Relationship Managers (RMs), '
    'their supervisors (RM Heads), and branch management in a Philippine trust and custody banking environment, '
    'CIM provides end-to-end lifecycle tracking of every client interaction — from initial meeting scheduling '
    'through call report approval, action item closure, and opportunity conversion. The module integrates '
    'seamlessly with existing TrustOMS modules (Client 360, GL, Trust Fee Pro, TTRA, Corporate Actions) '
    'and follows established patterns: Drizzle ORM schema, CRUD factory routes, maker-checker workflows, '
    'hash-chained audit logging, role-based authorization guards, and i18n support (English/Filipino).'
)

doc.add_heading('1.3 Business Objectives', level=2)
objectives = [
    'Centralize all RM-client interactions in a single auditable system, eliminating shadow spreadsheets and manual logs.',
    'Enable supervisors to review, approve, and provide feedback on call reports within defined SLA windows (48 hours for standard, 24 hours for high-value clients).',
    'Capture sales pipeline opportunities during client meetings with structured data (product, expected value, probability) to drive revenue forecasting.',
    'Provide management dashboards showing RM productivity metrics: meetings per week, call report completion rates, opportunity conversion ratios, and average time-to-closure.',
    'Support seamless RM handover by preserving complete interaction history linked to client records, ensuring relationship continuity when RMs transition.',
]
for obj in objectives:
    doc.add_paragraph(obj, style='List Bullet')

doc.add_heading('1.4 Target Users & Pain Points', level=2)
add_table(doc,
    ['Role', 'Pain Point', 'How CIM Solves It'],
    [
        ['Relationship Manager (RM)', 'Tracks meetings in personal notebooks/Excel; no institutional memory', 'Structured calendar with auto-linked call reports and searchable interaction history'],
        ['Senior RM', 'Cannot monitor junior RM activity or coach effectively', 'Dashboard showing team activity, call report quality scores, and opportunity pipeline'],
        ['RM Head / Supervisor', 'No workflow to review/approve call reports; compliance gaps', 'Approval queue with SLA tracking, feedback loop, and late-filing alerts'],
        ['Branch Manager', 'No visibility into branch-level client engagement metrics', 'Branch dashboard with aggregated KPIs: meetings held, reports filed, opportunities won'],
        ['Compliance Officer', 'Cannot audit RM-client interactions for suitability evidence', 'Immutable call report records with hash-chained audit trail and PII access logging'],
        ['Trust Business Head', 'No pipeline visibility across all branches', 'Executive dashboard with cross-branch opportunity funnel and conversion analytics'],
    ]
)

doc.add_heading('1.5 Success Metrics (KPIs)', level=2)
add_table(doc,
    ['KPI', 'Target', 'Measurement Method'],
    [
        ['Call report filing rate', '>=95% of meetings have filed call reports within 48 hours', 'Count(call_reports filed within 48h) / Count(completed meetings) per month'],
        ['Supervisor review completion', '>=90% of submitted call reports reviewed within SLA', 'Count(reviewed within SLA) / Count(submitted) per month'],
        ['Opportunity capture rate', '>=60% of client meetings generate at least one opportunity', 'Count(meetings with opportunities) / Count(total meetings) per quarter'],
        ['Opportunity conversion rate', '>=25% of captured opportunities reach WON status within 90 days', 'Count(WON opportunities) / Count(total opportunities created) per quarter'],
        ['RM calendar adoption', '>=90% of RMs schedule meetings through CIM within 30 days of launch', 'Count(RMs with >=1 meeting/week) / Count(active RMs)'],
        ['Average action item closure', '<=5 business days from creation to closure', 'AVG(closed_at - created_at) for action items per month'],
    ]
)

doc.add_page_break()

# ============================================================
# SECTION 2: SCOPE & BOUNDARIES
# ============================================================
doc.add_heading('2. Scope & Boundaries', level=1)

doc.add_heading('2.1 In Scope', level=2)
in_scope = [
    'Calendar Management: Day/Week/Month/All views, meeting CRUD, required/optional attendees, reminders, reschedule/cancel with notifications, overlapping meeting support',
    'Call Report Management: Scheduled and standalone call reports, meeting particulars capture, meeting information (client status, state of mind, discussion summary), next meeting auto-scheduling, status workflow (Draft -> Submitted -> Under Review -> Approved/Rejected)',
    'Action Item Tracking: Create action items from call reports, assign to RMs, set due dates, track status (Open -> In Progress -> Completed -> Overdue), link to parent call report',
    'Opportunity Management: Capture opportunities during meetings (name, product type, expected value, probability, expected close date), status lifecycle (Open -> Qualified -> Proposal -> Negotiation -> Won/Lost), bulk CSV upload',
    'Expense Tracking: Log client-meeting expenses (type, amount, currency, date), attach receipts, approval workflow',
    'Supervisor Approval Workflow: Call report review queue, approve/reject with comments, conditional rules (late filing >5 days requires supervisor approval), SLA tracking',
    'Feedback & Coaching: Supervisor feedback on call reports, conversation history, quality scoring',
    'Campaign Management: Campaign definition (code, name, type, budget, dates), lead list creation (automatic via filter rules + manual + bulk upload), lead response capture (Interested/Not Interested/Converted), campaign performance metrics',
    'Task Management: Personal to-do list for RMs, supervisor-assigned tasks, reminders, status tracking',
    'Analytics & Reporting: RM productivity dashboard, branch engagement metrics, opportunity pipeline funnel, call report compliance reports, campaign ROI analysis',
    'Integration: Client 360 cross-linking, existing role-based auth (23 roles), maker-checker workflows, hash-chained audit logging, i18n (en/fil), dark mode, WCAG 2.1 AA accessibility',
]
for item in in_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 Out of Scope', level=2)
out_scope = [
    'External calendar sync (Google Calendar, Outlook, iCal) — planned for v2',
    'Video conferencing integration (Zoom, Teams) — meetings tracked by mode but no embedded video',
    'SMS/WhatsApp notification channels — in-app and email only for v1',
    'AI-powered meeting transcription or call recording',
    'Commission/incentive calculation based on opportunities won',
    'Mobile native app — responsive web only for v1',
    'Third-party CRM data import (Salesforce, HubSpot migration)',
    'Automated lead scoring / ML-based propensity models',
    'Document e-signature workflows for meeting outcomes',
]
for item in out_scope:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 Assumptions', level=2)
assumptions = [
    'TrustOMS existing infrastructure (Supabase/PostgreSQL, Express, React, Drizzle ORM) is operational and the CIM module builds on top of it.',
    'Users are already provisioned in the users table with appropriate roles (RELATIONSHIP_MANAGER, SENIOR_RM, BO_HEAD, COMPLIANCE_OFFICER, TRUST_BUSINESS_HEAD, SYSTEM_ADMIN).',
    'The clients table already contains client records that CIM will reference via client_id foreign keys.',
    'Email delivery infrastructure (SMTP or transactional email service) is available for meeting invitations and notification delivery.',
    'Browser support: Chrome 90+, Firefox 90+, Safari 15+, Edge 90+. No IE11 support required.',
    'Maximum concurrent users: 500 RMs, 50 supervisors, 20 branch managers simultaneously.',
    'Average meetings per RM per week: 10-15. Peak: 25 during quarter-end.',
]
for item in assumptions:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.4 Constraints', level=2)
constraints = [
    'Technical: Must use existing TrustOMS stack (React 18, Express, Drizzle ORM, Supabase PostgreSQL). No new runtime dependencies beyond python-docx for report generation.',
    'Regulatory: Call reports containing client PII must comply with Philippine Data Privacy Act (RA 10173). PII fields must be classified and access-logged.',
    'Security: All mutations require role-based authorization guards. Call report content is classified as Sensitive-PII. Audit trail must be immutable (hash-chained).',
    'Performance: Calendar view must load within 2 seconds for a month with 100 meetings. Call report list must support pagination with 10,000+ records.',
    'Naming: All new tables must follow snake_case convention. Entity keys must be kebab-case. API routes must be under /api/v1/ prefix.',
]
for item in constraints:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================
# SECTION 3: USER ROLES & PERMISSIONS
# ============================================================
doc.add_heading('3. User Roles & Permissions', level=1)

doc.add_paragraph(
    'The CIM module leverages the existing TrustOMS 23-role authorization system. '
    'The following roles have CIM-specific permissions. Roles not listed have no CIM access.'
)

doc.add_heading('3.1 Role Descriptions', level=2)
add_table(doc,
    ['Role', 'Description', 'CIM Context'],
    [
        ['RELATIONSHIP_MANAGER', 'Front-office RM managing client relationships', 'Primary user: schedules meetings, files call reports, captures opportunities, logs expenses'],
        ['SENIOR_RM', 'Experienced RM with mentoring responsibilities', 'Same as RM + can view junior RM calendars and call reports within their team'],
        ['BO_HEAD', 'Back-office department head / RM supervisor', 'Reviews and approves call reports, assigns tasks, views team dashboards, manages campaigns'],
        ['TRUST_BUSINESS_HEAD', 'Executive overseeing trust business line', 'Read-only access to all dashboards, cross-branch analytics, campaign performance'],
        ['COMPLIANCE_OFFICER', 'Regulatory compliance staff', 'Read-only access to call reports for audit/suitability review, cannot modify'],
        ['CCO', 'Chief Compliance Officer', 'Same as COMPLIANCE_OFFICER + can flag call reports for regulatory review'],
        ['SYSTEM_ADMIN', 'Technical administrator', 'Configure system settings (reminder intervals, SLA thresholds), manage entity configs, cannot approve call reports'],
        ['BRANCH_ASSOCIATE', 'Junior front-office staff', 'View-only calendar access, can file standalone call reports for walk-in meetings'],
    ]
)

doc.add_heading('3.2 Permissions Matrix', level=2)

doc.add_heading('3.2.1 Calendar Permissions', level=3)
add_table(doc,
    ['Action', 'RM', 'SENIOR_RM', 'BO_HEAD', 'TRUST_BIZ_HEAD', 'COMPLIANCE', 'ADMIN', 'BRANCH_ASSOC'],
    [
        ['View own calendar', 'Yes', 'Yes', 'Yes', 'Yes', 'No', 'No', 'Yes'],
        ['View team calendar', 'No', 'Yes (own team)', 'Yes (all teams)', 'Yes (all)', 'No', 'No', 'No'],
        ['Create meeting', 'Yes', 'Yes', 'Yes', 'No', 'No', 'No', 'No'],
        ['Edit own meeting', 'Yes', 'Yes', 'Yes', 'No', 'No', 'No', 'No'],
        ['Cancel own meeting', 'Yes', 'Yes', 'Yes', 'No', 'No', 'No', 'No'],
        ['Reschedule own meeting', 'Yes', 'Yes', 'Yes', 'No', 'No', 'No', 'No'],
        ['Add attendees to meeting', 'Yes', 'Yes', 'Yes', 'No', 'No', 'No', 'No'],
        ['Configure reminder settings', 'No', 'No', 'No', 'No', 'No', 'Yes', 'No'],
    ]
)

doc.add_heading('3.2.2 Call Report Permissions', level=3)
add_table(doc,
    ['Action', 'RM', 'SENIOR_RM', 'BO_HEAD', 'TRUST_BIZ_HEAD', 'COMPLIANCE', 'CCO', 'BRANCH_ASSOC'],
    [
        ['Create call report (scheduled)', 'Yes', 'Yes', 'No', 'No', 'No', 'No', 'No'],
        ['Create call report (standalone)', 'Yes', 'Yes', 'No', 'No', 'No', 'No', 'Yes'],
        ['Edit own draft call report', 'Yes', 'Yes', 'No', 'No', 'No', 'No', 'Yes'],
        ['Submit call report for review', 'Yes', 'Yes', 'No', 'No', 'No', 'No', 'Yes'],
        ['View own call reports', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes', 'Yes'],
        ['View team call reports', 'No', 'Yes (own team)', 'Yes (all)', 'Yes (all)', 'Yes (all)', 'Yes (all)', 'No'],
        ['Approve/Reject call report', 'No', 'No', 'Yes', 'No', 'No', 'No', 'No'],
        ['Provide feedback on call report', 'No', 'Yes', 'Yes', 'No', 'No', 'Yes', 'No'],
        ['Flag for regulatory review', 'No', 'No', 'No', 'No', 'No', 'Yes', 'No'],
        ['Export call reports (CSV)', 'No', 'No', 'Yes', 'Yes', 'Yes', 'Yes', 'No'],
    ]
)

doc.add_heading('3.2.3 Opportunity & Expense Permissions', level=3)
add_table(doc,
    ['Action', 'RM', 'SENIOR_RM', 'BO_HEAD', 'TRUST_BIZ_HEAD', 'COMPLIANCE'],
    [
        ['Create opportunity', 'Yes', 'Yes', 'No', 'No', 'No'],
        ['Edit own opportunity', 'Yes', 'Yes', 'No', 'No', 'No'],
        ['View team opportunities', 'No', 'Yes (own team)', 'Yes (all)', 'Yes (all)', 'No'],
        ['Bulk upload opportunities', 'Yes', 'Yes', 'Yes', 'No', 'No'],
        ['Create expense entry', 'Yes', 'Yes', 'No', 'No', 'No'],
        ['Approve expenses', 'No', 'No', 'Yes', 'No', 'No'],
        ['View expense reports', 'No', 'No', 'Yes', 'Yes', 'No'],
    ]
)

doc.add_heading('3.2.4 Campaign & Task Permissions', level=3)
add_table(doc,
    ['Action', 'RM', 'SENIOR_RM', 'BO_HEAD', 'TRUST_BIZ_HEAD', 'ADMIN'],
    [
        ['View assigned campaigns', 'Yes', 'Yes', 'Yes', 'Yes', 'No'],
        ['Capture lead responses', 'Yes', 'Yes', 'No', 'No', 'No'],
        ['Create campaign', 'No', 'No', 'Yes', 'No', 'Yes'],
        ['Manage lead lists', 'No', 'No', 'Yes', 'No', 'Yes'],
        ['View campaign analytics', 'No', 'Yes', 'Yes', 'Yes', 'No'],
        ['Create personal task', 'Yes', 'Yes', 'Yes', 'No', 'No'],
        ['Assign task to RM', 'No', 'Yes', 'Yes', 'No', 'No'],
        ['View team tasks', 'No', 'Yes (own team)', 'Yes (all)', 'Yes (all)', 'No'],
    ]
)

# Save intermediate
doc.save('/Users/n15318/Trust OMS/docs/CIM_BRD_Draft.docx')
print("Part 1 complete: Sections 1-3 saved")
