#!/usr/bin/env python3
"""
Patch BRD v1.0 → v1.1: Addresses all gaps from adversarial evaluation.
Adds: FR-018, FR-019, Notification entity, NotificationPreference entity,
AuditLog entity, database indexes appendix, and fixes minor gaps.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import datetime

HEADER_BG = "1B3A5C"
DARK_BLUE = RGBColor(0x1B, 0x3A, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9)
        set_cell_shading(cell, HEADER_BG)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F2F7FB")
    doc.add_paragraph()
    return table

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27)
    return p

# ── Create the addendum document ────────────────────────────────────
doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, size in [('Heading 1', 18), ('Heading 2', 14), ('Heading 3', 12)]:
    s = doc.styles[level]
    s.font.name = 'Arial'
    s.font.size = Pt(size)
    s.font.color.rgb = DARK_BLUE
    s.font.bold = True
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

header = section.header
hp = header.paragraphs[0]
hp.text = "Trust OMS — Calendar & Call Report Management BRD v1.1 Addendum\tConfidential"
hp.style.font.size = Pt(8)

# ── Title ────────────────────────────────────────────────────────────
for _ in range(4):
    doc.add_paragraph()

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = tp.add_run("BRD ADDENDUM — v1.1")
run.font.size = Pt(28)
run.font.color.rgb = DARK_BLUE
run.bold = True

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = tp2.add_run("Calendar & Call Report Management")
run2.font.size = Pt(22)
run2.font.color.rgb = DARK_BLUE

tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = tp3.add_run(f"Addressing Adversarial Evaluation Findings  |  {datetime.date.today().strftime('%B %d, %Y')}")
run3.font.size = Pt(12)
run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── Change Log ───────────────────────────────────────────────────────
doc.add_heading("Change Log", level=1)
add_styled_table(doc,
    ["Version", "Date", "Changes"],
    [
        ["1.0", "2026-04-22", "Initial BRD"],
        ["1.1", datetime.date.today().strftime("%Y-%m-%d"),
         "Addendum: Added FR-018 (Mark Meeting Completed), FR-019 (Meeting No-Show Auto-Transition); "
         "Added Notification, NotificationPreference, AuditLog entities to data model; "
         "Specified polymorphic relationship_id strategy; "
         "Added recommended database indexes; "
         "Added recurring meetings to Out of Scope; "
         "Clarified denormalized name sync strategy, draft uniqueness, timezone for business-day calc; "
         "Added SystemConfig entries for no_show_grace_period_hours and pagination defaults"],
    ])

doc.add_heading("Evaluation Traceability", level=2)
doc.add_paragraph("Each change below references the evaluation gap it addresses.")
add_styled_table(doc,
    ["Gap #", "Severity", "Description", "Resolution Section"],
    [
        ["1", "CRITICAL", "Missing 'Mark Meeting as Completed' FR", "A.1 — FR-018 & FR-019"],
        ["2", "CRITICAL", "Missing Notification data model", "A.2 — Notification & NotificationPreference entities"],
        ["3", "HIGH", "Polymorphic relationship_id ambiguity", "A.3 — Polymorphic FK Strategy"],
        ["4", "MEDIUM", "No AuditLog entity", "A.4 — AuditLog entity"],
        ["5", "MEDIUM", "No database indexing guidance", "A.5 — Recommended Database Indexes"],
        ["6", "LOW", "Timezone edge case in business-day calc", "A.6 — Clarifications"],
        ["7", "LOW", "Draft lifecycle ambiguity", "A.6 — Clarifications"],
        ["8", "LOW", "Denormalized name sync strategy", "A.6 — Clarifications"],
        ["9", "LOW", "No-show grace period not in SystemConfig", "A.6 — Clarifications"],
        ["10", "LOW", "Recurring meetings not explicitly out of scope", "A.6 — Clarifications"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# A.1 — NEW FUNCTIONAL REQUIREMENTS
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("A.1 New Functional Requirements", level=1)

# FR-018
doc.add_heading("FR-018: Mark Meeting as Completed", level=2)
doc.add_paragraph(
    "After a meeting's scheduled end time has passed, the RM can explicitly mark the meeting as completed. "
    "This transitions the meeting status from 'scheduled' to 'completed' and enables the 'File Call Report' button. "
    "The system also displays a 'Mark as Completed' prompt on overdue meetings."
)

p_story = doc.add_paragraph()
run_label = p_story.add_run("User Story: ")
run_label.bold = True
p_story.add_run("As an RM, I want to mark a meeting as completed after it occurs so that I can proceed to file the call report.")

doc.add_heading("Acceptance Criteria", level=3)
criteria_018 = [
    "AC-018-1: A 'Mark as Completed' button appears on meetings where status = 'scheduled' AND end_date_time < NOW().",
    "AC-018-2: The button is NOT visible on meetings whose end_date_time is in the future.",
    "AC-018-3: Clicking the button transitions meeting.status from 'scheduled' to 'completed'.",
    "AC-018-4: On completion, a ConversationHistory entry is created with interaction_type = 'meeting_completed'.",
    "AC-018-5: After marking complete, the 'File Call Report' button becomes visible (per FR-004).",
    "AC-018-6: A confirmation dialog is shown: 'Mark this meeting as completed? You will be able to file a call report after this.'",
    "AC-018-7: The meeting card on the calendar changes color from blue (scheduled) to green (completed).",
]
for ac in criteria_018:
    add_bullet(doc, ac)

doc.add_heading("Business Rules", level=3)
rules_018 = [
    "BR-018-1: Only the meeting creator or a Supervisor can mark a meeting as completed.",
    "BR-018-2: Cancelled meetings cannot be marked as completed.",
    "BR-018-3: If the meeting was marked 'no_show' by the auto-transition (FR-019), it cannot be changed to 'completed'.",
    "BR-018-4: Marking as completed is irreversible — status cannot revert to 'scheduled'.",
]
for r in rules_018:
    add_bullet(doc, r)

doc.add_heading("UI Behavior", level=3)
doc.add_paragraph(
    "The 'Mark as Completed' button appears as a green checkmark icon or a 'Complete' text button on the meeting card/detail view. "
    "On calendar views, past meetings still in 'scheduled' status show a subtle pulsing indicator to remind the RM to mark them complete. "
    "On the All Activities list, a filter for 'Awaiting Completion' (scheduled + end_date < now) is available."
)

doc.add_heading("Edge Cases & Error Handling", level=3)
doc.add_paragraph(
    "If the meeting's end_date_time is more than 90 days in the past and still 'scheduled', show a warning: "
    "'This meeting is over 90 days old. Consider cancelling if it did not occur.' "
    "If two users (RM and supervisor) attempt to mark complete simultaneously, the first succeeds and the second sees 'Meeting already completed.'"
)

# FR-019
doc.add_heading("FR-019: Meeting No-Show Auto-Transition", level=2)
doc.add_paragraph(
    "A scheduled batch job transitions meetings from 'scheduled' to 'no_show' status when the meeting's end time "
    "has passed by more than a configurable grace period and the RM has not marked it as completed."
)

p_story2 = doc.add_paragraph()
run_label2 = p_story2.add_run("User Story: ")
run_label2.bold = True
p_story2.add_run("As a Supervisor, I want meetings that were never completed to be automatically flagged as no-shows so that I can track RM meeting discipline.")

doc.add_heading("Acceptance Criteria", level=3)
criteria_019 = [
    "AC-019-1: A batch job runs every hour (configurable).",
    "AC-019-2: Meetings with status = 'scheduled' AND end_date_time + grace_period < NOW() are transitioned to 'no_show'.",
    "AC-019-3: The grace period is read from SystemConfig key 'no_show_grace_period_hours' (default: 24 hours).",
    "AC-019-4: A notification is sent to the RM: 'Meeting \"{subject}\" has been marked as No Show.'",
    "AC-019-5: A notification is sent to the RM's supervisor.",
    "AC-019-6: No-show meetings appear in red on the calendar.",
    "AC-019-7: A ConversationHistory entry is created with interaction_type = 'meeting_no_show'.",
]
for ac in criteria_019:
    add_bullet(doc, ac)

doc.add_heading("Business Rules", level=3)
rules_019 = [
    "BR-019-1: The batch is idempotent — re-running does not re-process already-transitioned meetings.",
    "BR-019-2: No-show meetings cannot be edited, completed, or have call reports filed against them.",
    "BR-019-3: The RM can appeal a no-show by contacting their supervisor, who can manually override status to 'completed' (creates an audit log entry).",
    "BR-019-4: No-show count per RM is visible on the Supervisor Team Dashboard (FR-017).",
]
for r in rules_019:
    add_bullet(doc, r)

doc.add_heading("API Endpoint", level=3)
doc.add_paragraph("New endpoint for supervisor override:")
add_styled_table(doc,
    ["Method", "Endpoint", "Description", "Status Codes"],
    [
        ["PATCH", "/api/v1/meetings/:id/override-status", "Supervisor overrides no_show → completed. Body: { reason }.", "200, 400, 403, 404"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# A.2 — NEW DATA MODEL ENTITIES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("A.2 New Data Model Entities", level=1)
doc.add_paragraph("These entities supplement Section 4 of the BRD v1.0.")

# Notification
doc.add_heading("4.13 Notification", level=2)
doc.add_paragraph("Stores in-app notifications sent to users. Supports the notification center (bell icon, badge count, slide-out panel).")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["user_id", "UUID", "Yes", "FK → users.id (recipient)", "—"],
        ["event_type", "ENUM", "Yes", "Values: meeting_reminder_24h, meeting_reminder_1h, overdue_call_report, approval_required, approval_decision, meeting_cancelled, action_item_assigned, action_item_overdue, opportunity_expired, bulk_upload_complete, meeting_no_show", "—"],
        ["title", "VARCHAR(255)", "Yes", "Short notification title", "—"],
        ["body", "TEXT", "Yes", "Notification message body. Max 1000 chars.", "—"],
        ["channel", "ENUM", "Yes", "Values: in_app, email, sms, push", "—"],
        ["reference_type", "VARCHAR(50)", "No", "Entity type: meeting, call_report, approval, opportunity, action_item, upload", "NULL"],
        ["reference_id", "UUID", "No", "FK to the referenced entity", "NULL"],
        ["is_read", "BOOLEAN", "Yes", "Whether the user has read/dismissed the notification", "false"],
        ["read_at", "TIMESTAMPTZ", "No", "Set when user marks as read", "NULL"],
        ["email_status", "ENUM", "No", "Values: pending, sent, failed, not_applicable. For email channel tracking.", "not_applicable"],
        ["email_retry_count", "INTEGER", "No", "Number of email send retries (max 3)", "0"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_paragraph("Relationships: Notification N→1 User (user_id).")

doc.add_heading("Sample Data — Notification", level=3)
add_styled_table(doc,
    ["id", "user_id", "event_type", "title", "is_read", "reference_type", "reference_id"],
    [
        ["ntf-001", "rm-user-001", "meeting_reminder_24h", "Meeting tomorrow: Q2 Review – Sharma Family", "false", "meeting", "mtg-abc123"],
        ["ntf-002", "rm-user-001", "overdue_call_report", "Overdue: Branch Visit – Andheri (7 days)", "true", "meeting", "i9j0k1l2-..."],
        ["ntf-003", "sup-user-010", "approval_required", "Call report approval: rm-user-002 – Branch Visit", "false", "approval", "apr-001"],
    ])

# NotificationPreference
doc.add_heading("4.14 NotificationPreference", level=2)
doc.add_paragraph("Per-user per-event-type per-channel notification preferences. Controls which notifications the user receives.")

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["user_id", "UUID", "Yes", "FK → users.id. Composite unique: (user_id, event_type, channel).", "—"],
        ["event_type", "ENUM", "Yes", "Same ENUM as Notification.event_type", "—"],
        ["channel", "ENUM", "Yes", "Values: in_app, email, sms, push", "—"],
        ["enabled", "BOOLEAN", "Yes", "Whether this notification is enabled for this user+channel", "true"],
        ["is_critical", "BOOLEAN", "Yes", "If true, user cannot disable (override enabled = true). Set by system.", "false"],
        ["updated_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_paragraph("Business rule: On user creation, seed all event_type × channel combinations with enabled = true. Critical events (overdue_call_report, approval_required, approval_decision) have is_critical = true and cannot be disabled by the user.")

doc.add_heading("Sample Data — NotificationPreference", level=3)
add_styled_table(doc,
    ["id", "user_id", "event_type", "channel", "enabled", "is_critical"],
    [
        ["np-001", "rm-user-001", "meeting_reminder_24h", "email", "true", "false"],
        ["np-002", "rm-user-001", "meeting_reminder_24h", "in_app", "true", "false"],
        ["np-003", "rm-user-001", "overdue_call_report", "email", "true", "true"],
        ["np-004", "rm-user-001", "meeting_cancelled", "email", "false", "false"],
    ])

# AuditLog
doc.add_heading("4.15 AuditLog", level=2)
doc.add_paragraph(
    "Immutable, append-only audit log for all create/update/delete operations. "
    "Required for regulatory compliance (MiFID II, RBI KYC). "
    "This entity is INSERT-ONLY — no updates or deletes are permitted."
)

add_styled_table(doc,
    ["Field", "Type", "Required", "Validation / Rules", "Default"],
    [
        ["id", "UUID", "Yes", "Auto-generated", "uuid_generate_v4()"],
        ["entity_type", "VARCHAR(50)", "Yes", "Table name: meeting, call_report, opportunity, expense, etc.", "—"],
        ["entity_id", "UUID", "Yes", "PK of the affected entity", "—"],
        ["action", "ENUM", "Yes", "Values: create, update, delete, status_change, approval", "—"],
        ["field_name", "VARCHAR(100)", "No", "For updates: which field changed. NULL for create/delete.", "NULL"],
        ["old_value", "TEXT", "No", "Previous value (JSON-encoded if complex). NULL for create.", "NULL"],
        ["new_value", "TEXT", "No", "New value (JSON-encoded if complex). NULL for delete.", "NULL"],
        ["performed_by", "UUID", "Yes", "FK → users.id", "Current user"],
        ["ip_address", "VARCHAR(45)", "No", "Client IP (IPv4 or IPv6)", "NULL"],
        ["user_agent", "VARCHAR(500)", "No", "Browser user agent string", "NULL"],
        ["created_at", "TIMESTAMPTZ", "Yes", "Auto-set", "NOW()"],
    ])

doc.add_paragraph("Implementation note: Use database triggers or application-level middleware to auto-insert audit records on every mutation. Consider partitioning by created_at (monthly) for large-scale deployments.")

doc.add_heading("Sample Data — AuditLog", level=3)
add_styled_table(doc,
    ["id", "entity_type", "entity_id", "action", "field_name", "old_value", "new_value", "performed_by"],
    [
        ["aud-001", "meeting", "mtg-abc123", "create", "NULL", "NULL", "NULL", "rm-user-001"],
        ["aud-002", "call_report", "cr-002", "status_change", "status", "draft", "pending_approval", "rm-user-002"],
        ["aud-003", "call_report", "cr-002", "approval", "status", "pending_approval", "approved", "sup-user-010"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# A.3 — POLYMORPHIC FK STRATEGY
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("A.3 Polymorphic Relationship ID Strategy", level=1)

doc.add_paragraph(
    "Multiple entities (Meeting, CallReport, Opportunity, Expense, Feedback, ConversationHistory) reference "
    "customers, leads, or prospects via a relationship_id field. This section defines the implementation strategy."
)

doc.add_heading("Chosen Pattern: Discriminator Column", level=2)
doc.add_paragraph(
    "Each entity that references a relationship stores TWO columns:\n\n"
    "1. entity_type (ENUM: 'cif', 'lead', 'prospect') — the discriminator\n"
    "2. relationship_id (UUID) — the FK to the corresponding table\n\n"
    "Referential integrity is enforced at the APPLICATION layer (not database-level FK constraints), "
    "because PostgreSQL does not natively support polymorphic foreign keys. "
    "The application validates on insert/update that the relationship_id exists in the table corresponding to entity_type."
)

doc.add_heading("Rationale", level=3)
reasons = [
    "Aligns with existing meeting_type and entity_type fields already in the BRD (no new columns needed in most entities).",
    "Avoids schema changes to the existing CRM module (no unified relationships table needed).",
    "Avoids nullable column proliferation (three separate FK columns per entity would be verbose).",
    "The existing Trust OMS codebase likely uses this pattern for other cross-entity references.",
]
for r in reasons:
    add_bullet(doc, r)

doc.add_heading("Validation Rule", level=3)
doc.add_paragraph(
    "On every POST/PUT that includes relationship_id:\n"
    "1. Read entity_type from the request.\n"
    "2. Query the corresponding table (customers if 'cif', leads if 'lead', prospects if 'prospect').\n"
    "3. If relationship_id is not found, return 400 VALIDATION_ERROR: 'Relationship not found in {entity_type} table.'\n"
    "4. Auto-populate relationship_name from the source record."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# A.4 — (AuditLog already covered in A.2 above)
# ═══════════════════════════════════════════════════════════════════════

# ═══════════════════════════════════════════════════════════════════════
# A.5 — RECOMMENDED DATABASE INDEXES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("A.5 Recommended Database Indexes", level=1)

doc.add_paragraph(
    "The following indexes are recommended to meet the performance targets in Section 8.1 of the BRD. "
    "All indexes exclude soft-deleted records (WHERE deleted_at IS NULL) where applicable."
)

add_styled_table(doc,
    ["Table", "Index Name", "Columns", "Type", "Justification"],
    [
        ["meetings", "idx_meetings_created_by_start", "created_by, start_date_time", "B-tree", "Calendar view: RM's meetings sorted by date"],
        ["meetings", "idx_meetings_branch_start", "branch_id, start_date_time", "B-tree", "Supervisor/branch calendar view"],
        ["meetings", "idx_meetings_status", "status, call_report_status", "B-tree", "Filter by status; overdue detection query"],
        ["meetings", "idx_meetings_relationship", "relationship_id, entity_type", "B-tree", "Conversation history lookup by client"],
        ["call_reports", "idx_cr_created_by_status", "created_by, status", "B-tree", "RM's call report list filtered by status"],
        ["call_reports", "idx_cr_meeting_id_unique", "meeting_id (UNIQUE, partial: WHERE meeting_id IS NOT NULL AND deleted_at IS NULL)", "B-tree (unique partial)", "Enforce 1:1 meeting→call_report (BR-004-3)"],
        ["call_reports", "idx_cr_branch_status", "branch_id, status", "B-tree", "Supervisor approval queue"],
        ["call_report_approvals", "idx_cra_status_branch", "status, supervisor_id", "B-tree", "Supervisor pending/claimed queries"],
        ["opportunities", "idx_opp_status_due", "status, due_date", "B-tree", "Auto-expiry batch job"],
        ["opportunities", "idx_opp_created_by", "created_by, status", "B-tree", "RM's opportunity list"],
        ["opportunities", "idx_opp_relationship", "relationship_id, entity_type", "B-tree", "Opportunity lookup by client"],
        ["expenses", "idx_exp_created_by_date", "created_by, expense_date", "B-tree", "RM's expense list sorted by date"],
        ["conversation_history", "idx_ch_relationship_date", "relationship_id, interaction_date DESC", "B-tree", "Timeline view (reverse chronological)"],
        ["action_items", "idx_ai_assigned_status", "assigned_to, status, due_date", "B-tree", "My Action Items dashboard"],
        ["notifications", "idx_ntf_user_read", "user_id, is_read, created_at DESC", "B-tree", "Notification center: unread count + list"],
        ["audit_log", "idx_audit_entity", "entity_type, entity_id, created_at DESC", "B-tree", "Audit trail lookup per entity"],
    ])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# A.6 — CLARIFICATIONS & MINOR FIXES
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("A.6 Clarifications & Minor Fixes", level=1)

doc.add_heading("A.6.1 Timezone for Business-Day Calculation (Gap #6)", level=2)
doc.add_paragraph(
    "The 5-day threshold calculation (BR-004-1) uses the RM's configured local timezone to determine "
    "'current date' and 'meeting date.' Specifically:\n\n"
    "1. Convert meeting.end_date_time from UTC to the RM's timezone → meeting_local_date.\n"
    "2. Convert NOW() from UTC to the RM's timezone → current_local_date.\n"
    "3. Count business days between meeting_local_date and current_local_date using the business calendar for the RM's country.\n\n"
    "This ensures that an RM in IST filing at 11:30 PM local time sees the correct day count."
)

doc.add_heading("A.6.2 Draft Lifecycle & Uniqueness (Gap #7)", level=2)
doc.add_paragraph(
    "Clarifications to the call report draft behavior:\n\n"
    "1. A meeting can have at most ONE call report (including drafts). Enforced by a unique partial index: "
    "call_reports(meeting_id) WHERE meeting_id IS NOT NULL AND deleted_at IS NULL.\n"
    "2. If an RM starts a draft and navigates away, the draft is auto-saved (debounced, every 30 seconds of inactivity).\n"
    "3. The RM can resume the draft from the meeting card ('Continue Draft' button) or the Call Reports list.\n"
    "4. Drafts older than 30 days without updates are flagged in the RM's dashboard but NOT auto-deleted.\n"
    "5. Only one draft can exist per meeting. Attempting to create a second returns a 409 CONFLICT error pointing to the existing draft."
)

doc.add_heading("A.6.3 Denormalized Name Sync (Gap #8)", level=2)
doc.add_paragraph(
    "The relationship_name field stored in Meeting, CallReport, Opportunity, Expense, Feedback, and ConversationHistory "
    "is a POINT-IN-TIME SNAPSHOT populated at record creation. It is NOT updated if the source CRM record's name changes. "
    "This is intentional for audit/compliance reasons: the record reflects who the RM interacted with at the time of the interaction. "
    "The UI should display both the snapshot name and (if different) the current name from the CRM with a note: "
    "'Previously known as {snapshot_name}' when names differ."
)

doc.add_heading("A.6.4 New SystemConfig Entries (Gap #9)", level=2)
add_styled_table(doc,
    ["config_key", "config_value", "description"],
    [
        ["no_show_grace_period_hours", "24", "Hours after meeting end_date_time before auto-transitioning to no_show"],
        ["no_show_batch_interval_minutes", "60", "How often the no-show batch job runs (in minutes)"],
        ["draft_auto_save_interval_seconds", "30", "Auto-save debounce interval for call report drafts"],
        ["pagination_default_page_size", "20", "Default number of rows per page for all list endpoints"],
        ["pagination_max_page_size", "100", "Maximum allowed page_size parameter. Requests exceeding this are capped."],
        ["meeting_reminder_hours_before", "24,1", "Comma-separated list of hours before meeting to send reminders"],
        ["email_max_retries", "3", "Maximum retry attempts for failed email notifications"],
    ])

doc.add_heading("A.6.5 Recurring Meetings — Out of Scope (Gap #10)", level=2)
doc.add_paragraph(
    "Added to Section 2.2 (Out of Scope):\n\n"
    "• Recurring meeting patterns (daily, weekly, monthly repeat) — meetings are single-instance only in v1. "
    "Recurring meeting support is planned for v2 alongside Outlook/Google Calendar sync.\n\n"
    "The 'next meeting' feature in the call report form (next_meeting_start, next_meeting_end) creates individual "
    "one-off meetings, NOT a recurrence pattern."
)

doc.add_heading("A.6.6 Pagination Contract (Standardized)", level=2)
doc.add_paragraph(
    "All list/search endpoints (GET with paginated results) follow this contract:\n\n"
    "Query Parameters:\n"
    "• page (integer, default: 1, min: 1)\n"
    "• page_size (integer, default: 20, min: 1, max: 100 — capped silently if exceeded)\n"
    "• sort_by (string, default varies by endpoint — e.g., 'start_date_time' for meetings)\n"
    "• sort_order (string, 'asc' or 'desc', default: 'desc')\n\n"
    "Response envelope:\n"
)
doc.add_paragraph(
    '{\n'
    '  "data": [ ... ],\n'
    '  "pagination": {\n'
    '    "page": 1,\n'
    '    "page_size": 20,\n'
    '    "total_count": 47,\n'
    '    "total_pages": 3,\n'
    '    "has_next": true,\n'
    '    "has_previous": false\n'
    '  }\n'
    '}\n'
)
doc.add_paragraph(
    "If page exceeds total_pages, return an empty data array with the correct pagination metadata (not a 404)."
)

doc.add_heading("A.6.7 Updated Meeting Lifecycle State Diagram", level=2)
doc.add_paragraph("Replaces Section 9.1 of BRD v1.0 to include FR-018 and FR-019:")
add_styled_table(doc,
    ["Current State", "Action", "Next State", "Side Effects"],
    [
        ["(New)", "RM creates meeting", "Scheduled", "ConversationHistory entry; invitee notifications"],
        ["Scheduled", "RM clicks 'Mark as Completed' (FR-018)", "Completed", "ConversationHistory entry; call_report_status stays 'pending'"],
        ["Scheduled", "RM cancels meeting", "Cancelled", "ConversationHistory entry; invitee cancellation notifications"],
        ["Scheduled", "RM edits meeting", "Scheduled", "updated_at set; if date changed, ConversationHistory 'rescheduled' entry"],
        ["Scheduled", "end_date_time + grace_period passes (FR-019 batch)", "No Show", "ConversationHistory entry; RM + supervisor notified"],
        ["No Show", "Supervisor overrides (PATCH /override-status)", "Completed", "AuditLog entry with reason; ConversationHistory entry"],
        ["Completed", "RM files call report within threshold", "Completed (call_report_status → filed)", "CallReport created; ConversationHistory entry"],
        ["Completed", "Threshold days pass without call report", "Completed (call_report_status → overdue)", "Overdue notification sent to RM"],
    ])

doc.add_heading("A.6.8 Updated Entity Relationship Summary", level=2)
doc.add_paragraph("Additional relationships (supplement Section 4.13 of BRD v1.0):")
new_rels = [
    "User 1→N Notification (user_id)",
    "User 1→N NotificationPreference (user_id)",
    "User 1→N AuditLog (performed_by)",
    "Meeting/CallReport/Opportunity/etc. 1→N AuditLog (entity_id via polymorphic entity_type)",
]
for r in new_rels:
    add_bullet(doc, r)

doc.add_page_break()

# ── Final checklist ──────────────────────────────────────────────────
doc.add_heading("Quality Checklist — v1.1 Verification", level=1)

checks = [
    ["Every entity in Section 5 (FRs) has a data model in Section 4?", "YES — Notification and NotificationPreference now defined."],
    ["Every screen in Section 6 references existing entities?", "YES — Notification center backed by Notification entity."],
    ["Meeting lifecycle has explicit completion trigger?", "YES — FR-018 (manual) and FR-019 (auto no-show)."],
    ["Polymorphic FK strategy specified?", "YES — Discriminator column pattern (A.3)."],
    ["Audit trail entity defined?", "YES — AuditLog entity (4.15)."],
    ["Database indexes recommended for perf targets?", "YES — 16 indexes covering all performance-critical queries (A.5)."],
    ["All sample data present for new entities?", "YES — 2-3 rows each for Notification, NotificationPreference, AuditLog."],
    ["SystemConfig complete with all configurable values?", "YES — 7 new entries added (A.6.4)."],
    ["Recurring meetings addressed?", "YES — Explicitly out of scope (A.6.5)."],
    ["Pagination contract standardized?", "YES — Standard envelope with has_next/has_previous (A.6.6)."],
]

add_styled_table(doc,
    ["Check", "Status"],
    checks)

doc.add_paragraph()
p_final = doc.add_paragraph()
run_final = p_final.add_run("BRD v1.1 addresses all 10 gaps identified in the adversarial evaluation. The document is now ready for development handoff.")
run_final.bold = True

# ── Save ─────────────────────────────────────────────────────────────
output_path = "/Users/n15318/Trust OMS/docs/Calendar_CallReport_Management_BRD_v1.1_Addendum.docx"
doc.save(output_path)
print(f"BRD v1.1 Addendum saved to: {output_path}")
